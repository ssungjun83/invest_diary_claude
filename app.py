import sqlite3
from datetime import date, datetime, timedelta
from pathlib import Path
import base64
import time
from difflib import SequenceMatcher
import os
import hashlib
import hmac
import html
from io import BytesIO
from urllib.parse import unquote, quote_plus, urlparse

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import ast
import json
import re
import requests
import streamlit as st

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    HAS_PYTHON_DOCX = True
except Exception:
    HAS_PYTHON_DOCX = False

try:
    import matplotlib  # noqa: F401

    HAS_MATPLOTLIB = True
except Exception:
    HAS_MATPLOTLIB = False

DB_PATH = Path("portfolio.db")
DEFAULT_DATE = date.today()
DEFAULT_EXCEL_PATH = Path("내 주식자산.xlsx")

COL_NAME = "종목명"
COL_QTY = "보유수량"
COL_VALUE = "평가금액"
COL_PNL = "손익금액"
COL_RETURN = "수익률(%)"
COL_CURRENCY = "통화"
COL_FX_RATE = "환율(원화기준)"
COL_VALUE_KRW = "평가금액(원화)"
COL_PNL_KRW = "손익금액(원화)"
COL_PRICE_KRW = "주가(원화환산)"
COL_AVG_BUY_KRW = "평균매수가(원화)"

DEFAULT_USD_KRW = 1350.0

COLUMNS = [COL_NAME, COL_QTY, COL_CURRENCY, COL_FX_RATE, COL_VALUE, COL_PNL, COL_RETURN]

DEFAULT_HOLDINGS = [
    {COL_NAME: "넥스틸", COL_QTY: 200, COL_VALUE: 2349292, COL_PNL: 338292, COL_RETURN: 16.82},
    {COL_NAME: "대신증권", COL_QTY: 55, COL_VALUE: 2179133, COL_PNL: 469632, COL_RETURN: 27.47},
    {COL_NAME: "대창단조", COL_QTY: 300, COL_VALUE: 1916160, COL_PNL: -141840, COL_RETURN: -6.89},
    {COL_NAME: "대한조선", COL_QTY: 100, COL_VALUE: 10079800, COL_PNL: 2866800, COL_RETURN: 39.74},
    {COL_NAME: "동방", COL_QTY: 681, COL_VALUE: 1790628, COL_PNL: -180873, COL_RETURN: -9.17},
    {COL_NAME: "동원개발", COL_QTY: 1360, COL_VALUE: 4031122, COL_PNL: 293722, COL_RETURN: 7.85},
    {COL_NAME: "미창석유", COL_QTY: 18, COL_VALUE: 2094603, COL_PNL: -297597, COL_RETURN: -12.44},
    {COL_NAME: "비에이치아이", COL_QTY: 8, COL_VALUE: 821554, COL_PNL: 156439, COL_RETURN: 23.52},
    {COL_NAME: "삼성생명", COL_QTY: 4, COL_VALUE: 854288, COL_PNL: 21288, COL_RETURN: 2.55},
    {COL_NAME: "삼성화재", COL_QTY: 7, COL_VALUE: 3311364, COL_PNL: -319636, COL_RETURN: -8.80},
    {COL_NAME: "세아제강", COL_QTY: 50, COL_VALUE: 6926120, COL_PNL: 616120, COL_RETURN: 9.76},
    {COL_NAME: "케이씨", COL_QTY: 50, COL_VALUE: 1576840, COL_PNL: -72660, COL_RETURN: -4.40},
    {COL_NAME: "현대제철", COL_QTY: 70, COL_VALUE: 2574341, COL_PNL: -233159, COL_RETURN: -8.30},
    {COL_NAME: "휴스틸", COL_QTY: 500, COL_VALUE: 2492505, COL_PNL: 213895, COL_RETURN: 9.38},
    {COL_NAME: "DL이앤씨", COL_QTY: 80, COL_VALUE: 3912160, COL_PNL: 501160, COL_RETURN: 14.69},
    {COL_NAME: "HD건설기계", COL_QTY: 31, COL_VALUE: 4059063, COL_PNL: 23967, COL_RETURN: 0.59},
]

for item in DEFAULT_HOLDINGS:
    item.setdefault(COL_CURRENCY, "KRW")
    item.setdefault(COL_FX_RATE, 1.0)

SCORE_METRIC_CONFIG = {
    "dividend_yield": {"label": "배당수익률(%)", "min": 0.0, "max": 8.0, "reverse": False},
    "revenue_growth": {"label": "매출성장률(%)", "min": -20.0, "max": 30.0, "reverse": False},
    "eps_growth": {"label": "EPS성장률(%)", "min": -20.0, "max": 30.0, "reverse": False},
    "roe": {"label": "ROE(%)", "min": 0.0, "max": 25.0, "reverse": False},
    "operating_margin": {"label": "영업이익률(%)", "min": 0.0, "max": 30.0, "reverse": False},
    "debt_ratio": {"label": "부채비율(%)", "min": 0.0, "max": 300.0, "reverse": True},
    "current_ratio": {"label": "유동비율(배)", "min": 0.5, "max": 3.0, "reverse": False},
    "per": {"label": "PER", "min": 0.0, "max": 40.0, "reverse": True},
    "pbr": {"label": "PBR", "min": 0.0, "max": 5.0, "reverse": True},
}

DEFAULT_SCORE_WEIGHTS = {
    "dividend": 20.0,
    "growth": 35.0,
    "stability": 25.0,
    "valuation": 20.0,
}

FX_TRACKERS = [
    {"pair": "USD/KRW", "ticker": "KRW=X"},
    {"pair": "EUR/KRW", "ticker": "EURKRW=X"},
    {"pair": "JPY/KRW", "ticker": "JPYKRW=X"},
    {"pair": "CNY/KRW", "ticker": "CNYKRW=X"},
    {"pair": "GBP/KRW", "ticker": "GBPKRW=X"},
    {"pair": "AUD/KRW", "ticker": "AUDKRW=X"},
    {"pair": "CAD/KRW", "ticker": "CADKRW=X"},
    {"pair": "CHF/KRW", "ticker": "CHFKRW=X"},
]

DEFAULT_TICKER_HINTS = {
    "삼성전자": "005930.KS",
    "삼성생명": "032830.KS",
    "삼성화재": "000810.KS",
    "현대제철": "004020.KS",
    "DL이앤씨": "375500.KS",
    "세아제강": "306200.KS",
    "동원개발": "013120.KQ",
    "대신증권": "003540.KS",
    "비에이치아이": "083650.KQ",
    "대한조선": "439260.KS",
    "HD건설기계": "267270.KS",
    "넥스틸": "092790.KS",
    "휴스틸": "005010.KS",
    "TIME 코리아밸류업액티브": "495060.KS",
    "TIME코리아밸류업액티브": "495060.KS",
    "코리아밸류업액티브": "495060.KS",
    "노브": "NOV",
    "노브(주)": "NOV",
    "스타 벌크 캐리어스": "SBLK",
    "스타벌크캐리어스": "SBLK",
    "시드릴": "SDRL",
    "오케아니스 에코 탱커스": "ECO",
    "오케아니스에코탱커스": "ECO",
    "센트러스 에너지": "LEU",
    "센트러스에너지": "LEU",
    "유나이티드헬스 그룹": "UNH",
    "유나이티드헬스그룹": "UNH",
    "차코스 에너지 내비게이션": "TEN",
    "차코스에너지내비게이션": "TEN",
    "프론트라인": "FRO",
    "페트로브라스": "PBR",
    "페트로브라스(ADR)": "PBR",
    "타이드워터": "TDW",
    "텍사스 퍼시픽 랜드": "TPL",
    "텍사스퍼시픽랜드": "TPL",
    "피바디 에너지": "BTU",
    "피바디에너지": "BTU",
    "발레로 에너지": "VLO",
    "발레로에너지": "VLO",
    "터치 브로스": "BROS",
    "터치브로스": "BROS",
    "AGNC 인베스트먼트": "AGNC",
}

AI_PROVIDER_OPTIONS = ["OpenAI", "Claude"]
DEFAULT_OPENAI_MODEL = "gpt-4o-mini"
DEFAULT_CLAUDE_MODEL = "claude-3-5-haiku-latest"
DEFAULT_AI_MODEL = DEFAULT_OPENAI_MODEL
HTTP_HEADERS_COMMON = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/123.0.0.0 Safari/537.36"
    ),
    "Accept": "application/json,text/plain,*/*",
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7",
}


def normalize_ticker_text(value: str) -> str:
    text = str(value or "").strip().upper()
    text = re.sub(r"\s+", "", text)
    text = text.strip("`'\"")
    text = text.strip(".,;:()[]{}")
    return text


def clean_valid_ticker(value: str) -> str:
    ticker = normalize_ticker_text(value)
    if not ticker:
        return ""
    # Reuters 표기(예: AGNC.O, IBM.N)는 Yahoo 표기로 정규화한다.
    m_reuters_us = re.fullmatch(r"([A-Z][A-Z0-9\-]{0,9})\.(O|N)", ticker)
    if m_reuters_us:
        ticker = m_reuters_us.group(1)
    if ticker.startswith("."):
        return ""
    if ticker in {"KS", "KQ"}:
        return ""
    if ticker in {
        "UNKNOWN",
        "NONE",
        "NULL",
        "N/A",
        "NA",
        "STOCK",
        "TICKER",
        "SYMBOL",
        "COMPANY",
        "SHARE",
        "NYSE",
        "NASDAQ",
        "KOSPI",
        "KOSDAQ",
    }:
        return ""
    # ISIN(예: US20080516001)은 티커가 아니다.
    if re.fullmatch(r"[A-Z]{2}[A-Z0-9]{9}\d", ticker):
        return ""
    # 점/구분자 없는 비정상 장문 코드는 오탐 가능성이 높아 차단한다.
    if "." not in ticker and "-" not in ticker and "^" not in ticker and "=" not in ticker and len(ticker) > 7:
        return ""
    if len(ticker) > 24:
        return ""
    if not re.fullmatch(r"[A-Z0-9][A-Z0-9.\-^=]{0,23}", ticker):
        return ""
    return ticker


def get_builtin_ticker_hint(company_name: str) -> str:
    name = (company_name or "").strip()
    hinted = clean_valid_ticker(DEFAULT_TICKER_HINTS.get(name, ""))
    if hinted:
        return hinted
    name_norm = normalize_company_name_for_match(name)
    if not name_norm:
        return ""
    # 공백/기호 차이를 허용하는 정규화 매칭
    for key, value in DEFAULT_TICKER_HINTS.items():
        if normalize_company_name_for_match(key) == name_norm:
            cand = clean_valid_ticker(value)
            if cand:
                return cand
    return ""


def normalize_company_name_for_match(value: str) -> str:
    text = str(value or "").strip().upper()
    if not text:
        return ""
    text = re.sub(r"\([^)]*\)", " ", text)
    text = re.sub(
        r"(주식회사|유한회사|홀딩스|HOLDINGS?|CORP(ORATION)?|INC(ORPORATED)?|CO|COMPANY|LTD|LIMITED|PLC|ADR|CLASS[A-Z]?)",
        " ",
        text,
    )
    text = re.sub(r"[^0-9A-Z가-힣]+", "", text)
    return text.strip()


def _market_pref_normalized(value: str) -> str:
    v = str(value or "").strip().lower()
    if v in {"domestic", "kr", "korea", "국내", "국내주식"}:
        return "domestic"
    if v in {"foreign", "overseas", "global", "해외", "해외주식"}:
        return "foreign"
    return ""


def _ticker_matches_market_preference(ticker: str, market_preference: str) -> bool:
    pref = _market_pref_normalized(market_preference)
    tkr = clean_valid_ticker(ticker)
    if not pref or not tkr:
        return bool(tkr)
    is_kr_ticker = tkr.endswith(".KS") or tkr.endswith(".KQ")
    if pref == "foreign":
        return not is_kr_ticker
    if pref == "domestic":
        return is_kr_ticker
    return True


def _company_name_has_hangul(name: str) -> bool:
    text = str(name or "")
    return bool(re.search(r"[가-힣\u1100-\u11FF\u3130-\u318F]", text))


def _looks_foreign_hangul_name_hint(name: str) -> bool:
    text = str(name or "").strip().upper()
    if not text:
        return False
    # 한글 표기로 자주 입력되는 해외 기업명/키워드 힌트
    foreign_tokens = [
        "스타 벌크",
        "스타벌크",
        "시드릴",
        "오케아니스",
        "센트러스",
        "유나이티드헬스",
        "차코스",
        "페트로브라스",
        "프론트라인",
        "타이드워터",
        "텍사스 퍼시픽",
        "피바디",
        "발레로",
        "터치 브로스",
        "AGNC",
    ]
    return any(tok in text for tok in foreign_tokens)


def _looks_explicit_foreign_company_name(name: str) -> bool:
    text = str(name or "").strip()
    if not text:
        return False
    upper = text.upper()
    if _looks_foreign_hangul_name_hint(text):
        return True
    if "ADR" in upper:
        return True
    domestic_tokens = ["코스피", "코스닥", "KOSPI", "KOSDAQ", "KRX", "스팩"]
    if any(tok in upper for tok in domestic_tokens):
        return False
    if re.search(r"\b[A-Z]{3,6}\b", upper):
        return True
    if any(tok in upper for tok in [" INC", " CORP", " PLC", " LTD", " HOLDINGS", "HLDGS", " S.A", " NV"]):
        return True
    return False


def _looks_domestic_company_name_hint(name: str) -> bool:
    text = str(name or "").strip()
    if not text:
        return False
    builtin = get_builtin_ticker_hint(text)
    if builtin.endswith(".KS") or builtin.endswith(".KQ"):
        return True
    upper = text.upper()
    domestic_tokens = [
        "코리아",
        "코스피",
        "코스닥",
        "KOSPI",
        "KOSDAQ",
        "KRX",
        "ETF",
        "액티브",
        "인버스",
        "레버리지",
        "KODEX",
        "TIGER",
        "KOACT",
        "PLUS",
        "ACE",
        "ARIRANG",
        "SOL",
        "HANARO",
        "RISE",
    ]
    return any(tok in upper for tok in domestic_tokens)


def _is_non_kr_ticker_plausible_for_name(company_name: str, ticker: str) -> bool:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return False
    builtin = get_builtin_ticker_hint(company_name)
    if builtin and tkr == builtin:
        return True
    if tkr.endswith(".KS") or tkr.endswith(".KQ"):
        return True
    name = str(company_name or "").strip()
    if not _company_name_has_hangul(name):
        return True
    upper = name.upper()
    base = tkr.split(".")[0]
    if base and re.search(rf"(?<![A-Z0-9]){re.escape(base)}(?![A-Z0-9])", upper):
        return True
    if "ADR" in upper:
        return True
    if _looks_foreign_hangul_name_hint(name):
        return True
    if _looks_explicit_foreign_company_name(name):
        return True
    return False


def _name_similarity(query: str, candidate: str) -> float:
    q = normalize_company_name_for_match(query)
    c = normalize_company_name_for_match(candidate)
    if not q or not c:
        return 0.0
    if q == c:
        return 1.0
    if q in c or c in q:
        return 0.92
    return float(SequenceMatcher(None, q, c).ratio())


def _is_us_listing_candidate(symbol: str, exchange: str = "", region: str = "") -> bool:
    sym = clean_valid_ticker(symbol)
    ex = str(exchange or "").strip().upper()
    reg = str(region or "").strip().lower()
    if not sym:
        return False

    if "united states" in reg or reg in {"us", "usa"}:
        return True
    if ex in {
        "NMS",
        "NAS",
        "NGM",
        "NCM",
        "NYQ",
        "NYSE",
        "ASE",
        "AMEX",
        "ARCA",
        "BATS",
        "OTC",
        "OTCMKTS",
        "PNK",
        "XNAS",
        "XNYS",
    }:
        return True
    if "." not in sym and re.fullmatch(r"[A-Z][A-Z0-9\-]{0,5}", sym):
        return True
    return False


def choose_best_ticker_candidate(
    company_name: str,
    candidates: list[dict],
    provider_label: str,
    market_preference: str = "",
) -> tuple[str, str]:
    if not candidates:
        return "", f"{provider_label} 검색 결과가 없습니다."

    pref = _market_pref_normalized(market_preference)
    scored = []
    q_has_hangul = bool(re.search(r"[가-힣]", str(company_name or "")))
    q_norm = normalize_company_name_for_match(company_name)
    q_norm_len = len(q_norm)

    # 시장 선호가 명확할 때는 후보군 자체를 우선 필터링한다.
    raw_pool = []
    for cand in candidates:
        symbol = clean_valid_ticker(str(cand.get("symbol") or ""))
        if not symbol:
            continue
        item = dict(cand)
        item["_symbol"] = symbol
        item["_is_kr_ticker"] = symbol.endswith(".KS") or symbol.endswith(".KQ")
        raw_pool.append(item)

    pool = raw_pool
    if pref == "foreign":
        non_kr = [c for c in raw_pool if not bool(c.get("_is_kr_ticker", False))]
        if non_kr:
            pool = non_kr
    elif pref == "domestic":
        only_kr = [c for c in raw_pool if bool(c.get("_is_kr_ticker", False))]
        if only_kr:
            pool = only_kr

    for idx, cand in enumerate(pool):
        symbol = str(cand.get("_symbol") or "").strip()
        if not symbol:
            symbol = clean_valid_ticker(str(cand.get("symbol") or ""))
        if not symbol:
            continue
        display_name = str(cand.get("name") or cand.get("description") or "").strip()
        display_norm = normalize_company_name_for_match(display_name)
        exchange = str(cand.get("exchange") or "").strip()
        region = str(cand.get("region") or "").strip()
        sim = _name_similarity(company_name, display_name or symbol)
        rank_score = max(0.0, 1.0 - idx * 0.10)
        symbol_name_score = 0.0
        symbol_norm = normalize_company_name_for_match(symbol)
        if q_norm and q_norm == symbol_norm:
            symbol_name_score = 1.0
        elif q_norm and q_norm in symbol_norm:
            symbol_name_score = 0.6

        score = (sim * 0.72) + (rank_score * 0.20) + (symbol_name_score * 0.08)
        is_kr_ticker = symbol.endswith(".KS") or symbol.endswith(".KQ")
        is_us_listing = _is_us_listing_candidate(symbol, exchange=exchange, region=region)
        is_other_foreign = ("." in symbol) and (not is_kr_ticker)
        if pref == "foreign" and is_kr_ticker:
            score -= 0.25
        elif pref == "domestic" and is_kr_ticker:
            score += 0.07
        elif pref == "domestic" and not is_kr_ticker:
            score -= 0.07
        elif pref != "domestic":
            # 해외/미분류에서는 미국 상장 티커를 우선 시도
            if is_us_listing:
                score += 0.12
            elif is_other_foreign:
                score -= 0.04
        if pref != "domestic" and q_has_hangul and q_norm_len <= 4 and is_kr_ticker:
            # 짧은 한글명(예: 2~4글자)은 국내 동음이의 종목으로 오탐이 잦아 보수적으로 점수 감점.
            if q_norm and display_norm and q_norm != display_norm:
                score -= 0.26
        score = max(0.0, min(1.0, score))
        scored.append(
            {
                "symbol": symbol,
                "name": display_name,
                "name_norm": display_norm,
                "score": score,
                "sim": sim,
                "is_kr_ticker": is_kr_ticker,
            }
        )

    if not scored:
        return "", f"{provider_label} 유효 티커 결과가 없습니다."

    scored.sort(key=lambda x: (x["score"], x["sim"]), reverse=True)
    best = scored[0]
    second_score = scored[1]["score"] if len(scored) > 1 else 0.0
    threshold = 0.30 if q_has_hangul else 0.42
    if pref:
        threshold = max(threshold, 0.36)
    ambiguous = (best["score"] - second_score) < 0.05 and best["score"] < 0.60
    # 한글 이름 + 시장 미분류에서는 국내(.KS/.KQ)와 해외 후보가 비슷하면 자동확정을 피한다.
    if (not pref) and q_has_hangul and bool(best.get("is_kr_ticker", False)):
        non_kr_sorted = [r for r in scored if not bool(r.get("is_kr_ticker", False))]
        if non_kr_sorted:
            best_non_kr = non_kr_sorted[0]
            if best_non_kr["score"] >= max(0.48, best["score"] - 0.08):
                return "", (
                    f"{provider_label} 후보가 국내/해외로 경합합니다. "
                    f"국내 {best['symbol']}({best['score']:.2f}) vs "
                    f"해외 {best_non_kr['symbol']}({best_non_kr['score']:.2f})"
                )
    if pref != "domestic" and q_has_hangul and q_norm_len <= 4 and bool(best.get("is_kr_ticker", False)):
        best_name_norm = str(best.get("name_norm") or "")
        if q_norm and best_name_norm and q_norm != best_name_norm and best["sim"] < 0.90:
            return "", f"{provider_label} 짧은 한글명 오탐 가능성(국내 {best['symbol']})"
    if best["score"] < threshold or ambiguous:
        reason = f"최고 일치도 {best['score']:.2f}"
        if ambiguous:
            reason += f", 후보 간 점수차 {(best['score'] - second_score):.2f}"
        return "", f"{provider_label} 매칭 신뢰도 낮음({reason})"

    picked_name = best["name"] or best["symbol"]
    return best["symbol"], f"{provider_label} 검색 ({picked_name}, 일치도 {best['score']:.2f})"


def inject_styles() -> None:
    st.markdown(
        """
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@100..900&display=swap');

            :root {
                --ink: #0f172a;
                --muted: #475569;
                --surface: rgba(255, 255, 255, 0.80);
                --surface-strong: rgba(255, 255, 255, 0.93);
                --stroke: rgba(148, 163, 184, 0.35);
                --up: #d92d20;
                --down: #1570ef;
                --flat: #64748b;
            }

            .stApp {
                font-family: 'Noto Sans KR', sans-serif;
                background:
                    radial-gradient(1200px 480px at -10% -25%, #dbeafe 0%, rgba(219, 234, 254, 0) 70%),
                    radial-gradient(1000px 420px at 110% -15%, #ccfbf1 0%, rgba(204, 251, 241, 0) 70%),
                    linear-gradient(180deg, #f8fafc 0%, #edf2f7 55%, #f8fafc 100%);
            }

            .stApp * {
                font-family: 'Noto Sans KR', sans-serif;
            }

            .block-container {
                max-width: 1180px;
                padding-top: 2.1rem;
                padding-bottom: 2.4rem;
            }

            [data-testid="stSidebar"] > div {
                background: linear-gradient(180deg, #0f172a 0%, #1e293b 100%);
                border-right: 1px solid rgba(148, 163, 184, 0.25);
            }

            [data-testid="stSidebar"] * {
                color: #e2e8f0 !important;
            }

            /* Sidebar controls with white surface: force dark text for readability */
            [data-testid="stSidebar"] .stTextInput input,
            [data-testid="stSidebar"] .stDateInput input,
            [data-testid="stSidebar"] .stNumberInput input,
            [data-testid="stSidebar"] .stTextArea textarea,
            [data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] > div {
                color: #0f172a !important;
                background: rgba(255, 255, 255, 0.96) !important;
            }

            [data-testid="stSidebar"] input::placeholder,
            [data-testid="stSidebar"] textarea::placeholder {
                color: #64748b !important;
                opacity: 1 !important;
            }

            [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
                background: #f8fafc !important;
                border: 1px solid rgba(148, 163, 184, 0.45) !important;
            }

            [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] * {
                color: #334155 !important;
            }

            [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button {
                color: #0f172a !important;
                background: #ffffff !important;
                border: 1px solid rgba(148, 163, 184, 0.55) !important;
            }

            [data-testid="stSidebar"] .stCaption {
                color: #cbd5e1 !important;
            }

            .hero {
                background: linear-gradient(120deg, #0f172a 0%, #1d4ed8 58%, #0f766e 100%);
                border-radius: 18px;
                padding: 1.35rem 1.45rem;
                box-shadow: 0 18px 36px rgba(15, 23, 42, 0.26);
                margin-bottom: 1rem;
                color: #fff;
            }

            .hero h1 {
                margin: 0 0 0.3rem 0;
                font-size: 1.72rem;
                line-height: 1.2;
                letter-spacing: -0.01em;
            }

            .hero p {
                margin: 0;
                font-size: 0.97rem;
                opacity: 0.92;
            }

            .section-shell {
                background: transparent;
                border: 0;
                border-top: 1px solid rgba(148, 163, 184, 0.30);
                border-radius: 0;
                padding: 0;
                margin: 0.55rem 0 0.9rem 0;
                box-shadow: none;
                backdrop-filter: none;
                min-height: 0;
                height: 0;
            }

            .section-title {
                font-size: 1.06rem;
                font-weight: 700;
                color: var(--ink);
                margin: 0.1rem 0 0.8rem 0;
            }

            .summary-card {
                background: var(--surface-strong);
                border: 1px solid var(--stroke);
                border-radius: 14px;
                padding: 0.85rem 0.95rem;
                min-height: 104px;
                box-shadow: 0 8px 16px rgba(15, 23, 42, 0.08);
            }

            .summary-label {
                font-size: 0.84rem;
                color: var(--muted);
                margin-bottom: 0.32rem;
            }

            .summary-value {
                font-size: 1.34rem;
                font-weight: 720;
                color: var(--ink);
                line-height: 1.2;
                margin-bottom: 0.2rem;
                letter-spacing: -0.01em;
            }

            .summary-note {
                font-size: 0.86rem;
                font-weight: 600;
            }

            .fx-reference-banner {
                display: flex;
                align-items: center;
                gap: 0.48rem;
                width: fit-content;
                max-width: 100%;
                padding: 0.4rem 0.72rem;
                margin: 0.25rem 0 0.7rem 0;
                border-radius: 999px;
                border: 1px solid rgba(29, 78, 216, 0.32);
                background: linear-gradient(120deg, rgba(219, 234, 254, 0.9) 0%, rgba(204, 251, 241, 0.88) 100%);
                color: #0f172a;
                box-shadow: 0 6px 14px rgba(15, 23, 42, 0.08);
            }

            .fx-reference-banner .fx-label {
                font-size: 0.78rem;
                font-weight: 700;
                color: #1e3a8a;
                opacity: 0.95;
            }

            .fx-reference-banner .fx-date {
                font-size: 0.98rem;
                font-weight: 800;
                letter-spacing: 0.01em;
                color: #0f172a;
            }

            .positive { color: var(--up); }
            .negative { color: var(--down); }
            .neutral { color: var(--flat); }

            [data-testid="stDataEditor"],
            [data-testid="stDataFrame"] {
                border: 1px solid var(--stroke);
                border-radius: 12px;
                overflow: hidden;
                background: rgba(255, 255, 255, 0.86);
            }

            [data-baseweb="tab-list"] {
                gap: 0.35rem;
                margin-bottom: 0.8rem;
            }

            button[data-baseweb="tab"] {
                border: 1px solid rgba(148, 163, 184, 0.4) !important;
                background: rgba(255,255,255,0.55) !important;
                border-radius: 999px !important;
                padding: 0.45rem 1rem !important;
                color: #334155 !important;
                font-weight: 700 !important;
            }

            button[data-baseweb="tab"][aria-selected="true"] {
                background: linear-gradient(120deg, #1d4ed8 0%, #0f766e 100%) !important;
                color: #fff !important;
                border-color: transparent !important;
                box-shadow: 0 10px 20px rgba(15, 23, 42, 0.16);
            }

            .stButton > button,
            .stDownloadButton > button {
                border: 0;
                border-radius: 10px;
                background: linear-gradient(120deg, #1d4ed8 0%, #0f766e 100%);
                color: #fff;
                font-weight: 700;
                box-shadow: 0 8px 16px rgba(15, 23, 42, 0.14);
            }

            .stButton > button:hover,
            .stDownloadButton > button:hover {
                filter: brightness(1.05);
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


def get_conn() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS snapshots (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            snapshot_date TEXT NOT NULL,
            stock_name TEXT NOT NULL,
            quantity REAL NOT NULL,
            market_value REAL NOT NULL,
            pnl_value REAL NOT NULL,
            pnl_pct REAL NOT NULL,
            currency TEXT,
            fx_rate REAL,
            created_at TEXT NOT NULL
        )
        """
    )
    columns = {row[1] for row in conn.execute("PRAGMA table_info(snapshots)").fetchall()}
    if "currency" not in columns:
        conn.execute("ALTER TABLE snapshots ADD COLUMN currency TEXT")
    if "fx_rate" not in columns:
        conn.execute("ALTER TABLE snapshots ADD COLUMN fx_rate REAL")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS fx_rates (
            rate_date TEXT PRIMARY KEY,
            usd_krw REAL NOT NULL,
            market_date TEXT,
            source TEXT,
            fetched_at TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS company_scores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            score_date TEXT NOT NULL,
            stock_name TEXT NOT NULL,
            ticker TEXT,
            dividend_yield REAL,
            revenue_growth REAL,
            eps_growth REAL,
            roe REAL,
            operating_margin REAL,
            debt_ratio REAL,
            current_ratio REAL,
            per REAL,
            pbr REAL,
            dividend_score REAL NOT NULL,
            growth_score REAL NOT NULL,
            stability_score REAL NOT NULL,
            valuation_score REAL NOT NULL,
            total_score REAL NOT NULL,
            weight_dividend REAL NOT NULL,
            weight_growth REAL NOT NULL,
            weight_stability REAL NOT NULL,
            weight_valuation REAL NOT NULL,
            source TEXT,
            note TEXT,
            created_at TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS snapshot_cash (
            snapshot_date TEXT PRIMARY KEY,
            cash_krw REAL NOT NULL DEFAULT 0,
            cash_usd REAL NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )
    snapshot_cash_columns = {row[1] for row in conn.execute("PRAGMA table_info(snapshot_cash)").fetchall()}
    if "cash_krw" not in snapshot_cash_columns:
        conn.execute("ALTER TABLE snapshot_cash ADD COLUMN cash_krw REAL NOT NULL DEFAULT 0")
    if "cash_usd" not in snapshot_cash_columns:
        conn.execute("ALTER TABLE snapshot_cash ADD COLUMN cash_usd REAL NOT NULL DEFAULT 0")
    if "created_at" not in snapshot_cash_columns:
        conn.execute("ALTER TABLE snapshot_cash ADD COLUMN created_at TEXT")
        conn.execute("UPDATE snapshot_cash SET created_at = COALESCE(created_at, datetime('now'))")
    if "updated_at" not in snapshot_cash_columns:
        conn.execute("ALTER TABLE snapshot_cash ADD COLUMN updated_at TEXT")
        conn.execute("UPDATE snapshot_cash SET updated_at = COALESCE(updated_at, datetime('now'))")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS company_analysis (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            analysis_date TEXT NOT NULL,
            stock_name TEXT NOT NULL,
            ticker TEXT,
            company_overview TEXT,
            products_services TEXT,
            raw_materials TEXT,
            profit_up_factors TEXT,
            profit_down_factors TEXT,
            financial_summary_json TEXT,
            source TEXT,
            ai_model TEXT,
            note TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS company_list (
            stock_name TEXT PRIMARY KEY,
            ticker TEXT,
            sector TEXT,
            price_krw REAL,
            price_source TEXT,
            price_updated_at TEXT,
            source TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )
    company_list_columns = {row[1] for row in conn.execute("PRAGMA table_info(company_list)").fetchall()}
    if "sector" not in company_list_columns:
        conn.execute("ALTER TABLE company_list ADD COLUMN sector TEXT")
    if "price_krw" not in company_list_columns:
        conn.execute("ALTER TABLE company_list ADD COLUMN price_krw REAL")
    if "price_source" not in company_list_columns:
        conn.execute("ALTER TABLE company_list ADD COLUMN price_source TEXT")
    if "price_updated_at" not in company_list_columns:
        conn.execute("ALTER TABLE company_list ADD COLUMN price_updated_at TEXT")
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS company_compare_sets (
            set_name TEXT PRIMARY KEY,
            companies_json TEXT NOT NULL,
            metrics_json TEXT,
            weights_json TEXT,
            sector_filter TEXT,
            note TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS app_settings (
            setting_key TEXT PRIMARY KEY,
            setting_value TEXT,
            updated_at TEXT NOT NULL
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS value_chains (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            chain_name TEXT NOT NULL,
            chain_json TEXT NOT NULL,
            matched_companies_json TEXT,
            source_image_hash TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT
        )
        """
    )
    # 성능 인덱스
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_snapshots_date ON snapshots(snapshot_date)"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_company_score_name_date ON company_scores(stock_name, score_date)"
    )
    conn.execute(
        "CREATE INDEX IF NOT EXISTS idx_company_analysis_name_date ON company_analysis(stock_name, analysis_date)"
    )
    conn.commit()
    return conn


# ─────────────────────────────────────────────────────────────────────────────
# 밸류체인 DB 헬퍼
# ─────────────────────────────────────────────────────────────────────────────

def save_value_chain(
    chain_name: str,
    chain_json: dict,
    matched_companies: list | None = None,
    image_hash: str | None = None,
) -> int:
    """밸류체인 추출 결과를 value_chains 테이블에 저장하고 새 id를 반환한다."""
    import json as _json
    conn = get_conn()
    now = datetime.now().isoformat(timespec="seconds")
    try:
        cur = conn.execute(
            """
            INSERT INTO value_chains (chain_name, chain_json, matched_companies_json,
                                      source_image_hash, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                chain_name,
                _json.dumps(chain_json, ensure_ascii=False),
                _json.dumps(matched_companies or [], ensure_ascii=False),
                image_hash,
                now,
                now,
            ),
        )
        conn.commit()
        return cur.lastrowid or 0
    finally:
        conn.close()


def load_value_chains() -> list[dict]:
    """저장된 밸류체인 목록(요약)을 반환한다."""
    conn = get_conn()
    try:
        rows = conn.execute(
            "SELECT id, chain_name, source_image_hash, created_at FROM value_chains ORDER BY created_at DESC"
        ).fetchall()
        return [
            {"id": r[0], "chain_name": r[1], "source_image_hash": r[2], "created_at": r[3]}
            for r in rows
        ]
    finally:
        conn.close()


def load_value_chain_by_id(chain_id: int) -> dict | None:
    """특정 밸류체인의 전체 데이터를 반환한다."""
    import json as _json
    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT id, chain_name, chain_json, matched_companies_json, source_image_hash, created_at "
            "FROM value_chains WHERE id = ?",
            (chain_id,),
        ).fetchone()
        if row is None:
            return None
        return {
            "id": row[0],
            "chain_name": row[1],
            "chain_json": _json.loads(row[2] or "{}"),
            "matched_companies": _json.loads(row[3] or "[]"),
            "source_image_hash": row[4],
            "created_at": row[5],
        }
    finally:
        conn.close()


def delete_value_chain(chain_id: int) -> None:
    """밸류체인을 DB에서 삭제한다."""
    conn = get_conn()
    try:
        conn.execute("DELETE FROM value_chains WHERE id = ?", (chain_id,))
        conn.commit()
    finally:
        conn.close()


def _compute_image_hash(image_bytes: bytes) -> str:
    """이미지 바이트의 SHA-256 해시를 반환한다 (중복 감지용)."""
    import hashlib
    return hashlib.sha256(image_bytes).hexdigest()


def _store_fx_rate(rate_date: date, usd_krw: float, market_date: str | None, source: str) -> None:
    conn = get_conn()
    try:
        conn.execute(
            """
            INSERT INTO fx_rates (rate_date, usd_krw, market_date, source, fetched_at)
            VALUES (?, ?, ?, ?, ?)
            ON CONFLICT(rate_date) DO UPDATE SET
                usd_krw=excluded.usd_krw,
                market_date=excluded.market_date,
                source=excluded.source,
                fetched_at=excluded.fetched_at
            """,
            (
                rate_date.isoformat(),
                float(usd_krw),
                market_date,
                source,
                datetime.now().isoformat(timespec="seconds"),
            ),
        )
        conn.commit()
    finally:
        conn.close()


def _load_cached_fx_rate(rate_date: date) -> tuple[float | None, str, str]:
    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT usd_krw, market_date, source FROM fx_rates WHERE rate_date = ?",
            (rate_date.isoformat(),),
        ).fetchone()
    finally:
        conn.close()

    if not row:
        return None, "", ""

    rate = float(row[0])
    market_date = row[1] or rate_date.isoformat()
    source = row[2] or "cache"
    return rate, source, market_date


def _fetch_usd_krw_rate_from_yfinance(rate_date: date) -> tuple[float | None, str]:
    try:
        import yfinance as yf
    except Exception:
        return None, "yfinance 미설치"

    start = (rate_date - timedelta(days=10)).isoformat()
    end = (rate_date + timedelta(days=1)).isoformat()

    try:
        hist = yf.download("KRW=X", start=start, end=end, interval="1d", progress=False, auto_adjust=False)
    except Exception as exc:
        return None, f"yfinance 조회 실패: {exc}"

    if hist is None or hist.empty:
        return None, "환율 데이터 없음"

    close = hist["Close"]
    if isinstance(close, pd.DataFrame):
        close = close.iloc[:, 0]
    close = close.dropna()
    if close.empty:
        return None, "환율 종가 데이터 없음"

    index = pd.to_datetime(close.index)
    if getattr(index, "tz", None) is not None:
        index = index.tz_convert(None)
    close.index = index

    close = close[close.index.date <= rate_date]
    if close.empty:
        return None, "해당일 이전 환율 데이터 없음"

    market_dt = close.index[-1].date().isoformat()
    rate = float(close.iloc[-1])
    return rate, market_dt


def _fetch_usd_krw_rate_from_frankfurter(rate_date: date) -> tuple[float | None, str]:
    start_date = (rate_date - timedelta(days=10)).isoformat()
    end_date = rate_date.isoformat()
    url = (
        "https://api.frankfurter.app/"
        f"{start_date}..{end_date}?from=USD&to=KRW"
    )
    try:
        resp = requests.get(url, timeout=12)
        resp.raise_for_status()
        payload = resp.json() or {}
    except Exception as exc:
        return None, f"frankfurter 조회 실패: {exc}"

    rates = payload.get("rates") or {}
    if not isinstance(rates, dict) or not rates:
        return None, "frankfurter 환율 데이터 없음"

    latest_dt = ""
    latest_rate = None
    for dt_text in sorted(rates.keys()):
        daily = rates.get(dt_text) or {}
        try:
            krw = float(daily.get("KRW"))
        except Exception:
            continue
        latest_dt = dt_text
        latest_rate = krw

    if latest_rate is None:
        return None, "frankfurter KRW 값 없음"
    return latest_rate, latest_dt


def get_usd_krw_rate_for_date(rate_date: date) -> tuple[float, str]:
    cached_rate, cached_source, cached_market_date = _load_cached_fx_rate(rate_date)
    if cached_rate is not None and cached_source.lower() != "fallback":
        return cached_rate, f"캐시 {cached_source} ({cached_market_date})"

    fetched_rate, market_date_or_msg = _fetch_usd_krw_rate_from_yfinance(rate_date)
    if fetched_rate is not None:
        _store_fx_rate(rate_date, fetched_rate, market_date_or_msg, "yfinance")
        return fetched_rate, f"yfinance ({market_date_or_msg})"

    fetched_rate, market_date_or_msg = _fetch_usd_krw_rate_from_frankfurter(rate_date)
    if fetched_rate is not None:
        _store_fx_rate(rate_date, fetched_rate, market_date_or_msg, "frankfurter")
        return fetched_rate, f"frankfurter ({market_date_or_msg})"

    fallback_rate = DEFAULT_USD_KRW
    fallback_market_date = cached_market_date or rate_date.isoformat()
    if cached_rate is not None and cached_source.lower() == "fallback":
        fallback_rate = float(cached_rate)
    _store_fx_rate(rate_date, fallback_rate, fallback_market_date, "fallback")
    return fallback_rate, f"fallback ({market_date_or_msg})"


@st.cache_data(ttl=60 * 60)
def fetch_benchmark_returns(start_date: str, end_date: str) -> pd.DataFrame:
    """KOSPI(^KS11)·S&P500(^GSPC) 종가를 기준일=100으로 정규화해 반환한다.

    반환 컬럼: date, KOSPI, SP500
    start_date / end_date: 'YYYY-MM-DD' 형식 문자열
    """
    try:
        import yfinance as yf
    except Exception:
        return pd.DataFrame()

    result_df = pd.DataFrame()
    for ticker, col_name in [("^KS11", "KOSPI"), ("^GSPC", "S&P500")]:
        try:
            raw = yf.download(ticker, start=start_date, end=end_date, progress=False, auto_adjust=True)
            if raw is None or raw.empty or "Close" not in raw.columns:
                continue
            close = pd.to_numeric(raw["Close"].squeeze(), errors="coerce").dropna()
            if close.empty:
                continue
            idx = pd.to_datetime(close.index)
            if getattr(idx, "tz", None) is not None:
                idx = idx.tz_convert(None)
            close.index = idx
            normalized = (close / close.iloc[0] * 100).rename(col_name)
            normalized.index = normalized.index.date
            if result_df.empty:
                result_df = normalized.to_frame()
            else:
                result_df = result_df.join(normalized.to_frame(), how="outer")
        except Exception:
            continue

    if not result_df.empty:
        result_df.index.name = "date"
        result_df = result_df.reset_index()
    return result_df


@st.cache_data(ttl=60 * 60)
def fetch_fx_series(ticker: str, start: str, end: str) -> pd.DataFrame:
    try:
        import yfinance as yf
    except Exception:
        return pd.DataFrame()

    try:
        hist = yf.download(
            ticker,
            start=start,
            end=end,
            interval="1d",
            progress=False,
            auto_adjust=False,
            group_by="column",
        )
    except Exception:
        return pd.DataFrame()

    if hist is None or hist.empty or "Close" not in hist.columns:
        return pd.DataFrame()

    close = hist["Close"]
    if isinstance(close, pd.DataFrame):
        close = close.iloc[:, 0]
    close = close.dropna()
    if close.empty:
        return pd.DataFrame()

    s = close.copy()
    idx = pd.to_datetime(s.index)
    if getattr(idx, "tz", None) is not None:
        idx = idx.tz_convert(None)
    s.index = idx

    out = s.reset_index()
    out.columns = ["date", "rate"]
    out["date"] = pd.to_datetime(out["date"])
    out["rate"] = pd.to_numeric(out["rate"], errors="coerce")
    out = out.dropna(subset=["rate"]).sort_values("date")
    return out


def _value_on_or_before(df: pd.DataFrame, base_date: pd.Timestamp) -> float | None:
    subset = df[df["date"] <= base_date]
    if subset.empty:
        return None
    return float(subset.iloc[-1]["rate"])


def get_fx_tracker_summary(series_map: dict[str, pd.DataFrame]) -> pd.DataFrame:
    rows = []
    for pair, df in series_map.items():
        if df.empty:
            continue
        latest = df.iloc[-1]
        latest_date = pd.Timestamp(latest["date"])
        latest_rate = float(latest["rate"])

        prev_rate = _value_on_or_before(df, latest_date - pd.Timedelta(days=1))
        week_rate = _value_on_or_before(df, latest_date - pd.Timedelta(days=7))
        month_rate = _value_on_or_before(df, latest_date - pd.Timedelta(days=30))

        d1 = ((latest_rate / prev_rate - 1) * 100) if prev_rate else None
        w1 = ((latest_rate / week_rate - 1) * 100) if week_rate else None
        m1 = ((latest_rate / month_rate - 1) * 100) if month_rate else None

        rows.append(
            {
                "통화쌍": pair,
                "기준일": latest_date.date(),
                "현재환율": latest_rate,
                "1일변동(%)": d1,
                "1주변동(%)": w1,
                "1개월변동(%)": m1,
            }
        )
    return pd.DataFrame(rows)


def save_snapshot(
    snapshot_date: date,
    df: pd.DataFrame,
    sync_to_github: bool = True,
    sync_reason: str = "",
) -> tuple[bool, str]:
    now_str = datetime.now().isoformat(timespec="seconds")
    date_str = snapshot_date.isoformat()
    usd_krw_rate, _ = get_usd_krw_rate_for_date(snapshot_date)

    conn = get_conn()
    try:
        conn.execute("DELETE FROM snapshots WHERE snapshot_date = ?", (date_str,))
        records = [
            (
                date_str,
                str(row[COL_NAME]),
                float(row[COL_QTY]),
                float(row[COL_VALUE]),
                float(row[COL_PNL]),
                float(row[COL_RETURN]),
                (
                    "USD"
                    if str(row.get(COL_CURRENCY, "KRW")).upper() == "USD"
                    else "KRW"
                ),
                (
                    float(usd_krw_rate)
                    if str(row.get(COL_CURRENCY, "KRW")).upper() == "USD"
                    else 1.0
                ),
                now_str,
            )
            for _, row in df.iterrows()
            if str(row[COL_NAME]).strip()
        ]
        conn.executemany(
            """
            INSERT INTO snapshots (
                snapshot_date,
                stock_name,
                quantity,
                market_value,
                pnl_value,
                pnl_pct,
                currency,
                fx_rate,
                created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            records,
        )
        conn.commit()
    finally:
        conn.close()

    # 스냅샷 저장일의 예수금 행이 없으면 직전 저장값을 승계해 생성한다.
    # (그래프/요약 간 예수금 포함 기준 불일치를 방지)
    try:
        carry_cash_krw, carry_cash_usd = load_snapshot_cash(snapshot_date)
        save_snapshot_cash(snapshot_date, carry_cash_krw, carry_cash_usd)
    except Exception:
        pass

    if not sync_to_github:
        return False, ""

    sync_ok, sync_msg = sync_snapshot_to_github_excel(snapshot_date, df)
    if sync_msg:
        prefix = f"[{sync_reason}] " if sync_reason else ""
        try:
            st.session_state["github_sync_notice"] = prefix + sync_msg
        except Exception:
            pass
    return sync_ok, sync_msg


def delete_snapshot_by_date(
    snapshot_date: date,
    delete_cash: bool = True,
    sync_to_github: bool = True,
) -> tuple[bool, str]:
    date_str = snapshot_date.isoformat()
    conn = get_conn()
    try:
        deleted_snapshot_rows = conn.execute(
            "DELETE FROM snapshots WHERE snapshot_date = ?",
            (date_str,),
        ).rowcount
        deleted_cash_rows = 0
        if delete_cash:
            deleted_cash_rows = conn.execute(
                "DELETE FROM snapshot_cash WHERE snapshot_date = ?",
                (date_str,),
            ).rowcount
        conn.commit()
    finally:
        conn.close()

    if int(deleted_snapshot_rows or 0) <= 0 and int(deleted_cash_rows or 0) <= 0:
        return False, f"{date_str}에 삭제할 기록이 없습니다."

    msg = f"{date_str} 기록 삭제 완료 (보유현황 {int(deleted_snapshot_rows or 0)}건"
    if delete_cash:
        msg += f", 예수금 {int(deleted_cash_rows or 0)}건"
    msg += ")"

    if not sync_to_github:
        return True, msg

    latest_date_text, latest_df = load_latest_snapshot()
    if latest_df.empty:
        target_date = snapshot_date
        latest_df = empty_portfolio_df()
    else:
        target_date = _safe_parse_date(latest_date_text) or snapshot_date
    sync_ok, sync_msg = sync_snapshot_to_github_excel(target_date, latest_df)
    if sync_msg:
        msg += f" / GitHub {'완료' if sync_ok else '경고'}: {sync_msg}"
    return True, msg


def empty_portfolio_df() -> pd.DataFrame:
    return pd.DataFrame(columns=COLUMNS)


def _pick_excel_column(raw_df: pd.DataFrame, aliases: list[str]) -> str:
    if raw_df is None or raw_df.empty:
        return ""
    norm_map = {}
    for col in raw_df.columns:
        key = re.sub(r"[\s_()\-]+", "", str(col or "").strip().lower())
        norm_map[key] = str(col)
    for alias in aliases:
        alias_key = re.sub(r"[\s_()\-]+", "", alias.strip().lower())
        if alias_key in norm_map:
            return norm_map[alias_key]
    return ""


def _to_num_series(series: pd.Series) -> pd.Series:
    s = series.astype(str).str.replace(",", "", regex=False).str.replace("%", "", regex=False).str.strip()
    s = s.replace({"": pd.NA, "nan": pd.NA, "None": pd.NA, "none": pd.NA})
    return pd.to_numeric(s, errors="coerce")


def resolve_excel_path() -> Path | None:
    env_path = str(os.getenv("PORTFOLIO_EXCEL_PATH", "") or "").strip()
    if env_path:
        p = Path(env_path)
        if p.exists() and p.is_file():
            return p
    if DEFAULT_EXCEL_PATH.exists() and DEFAULT_EXCEL_PATH.is_file():
        return DEFAULT_EXCEL_PATH
    return None


def _get_uploaded_excel_bytes() -> tuple[bytes, str]:
    try:
        payload = st.session_state.get("uploaded_portfolio_excel_bytes", b"")
        filename = str(st.session_state.get("uploaded_portfolio_excel_name", "") or "").strip()
    except Exception:
        return b"", ""
    if isinstance(payload, bytearray):
        payload = bytes(payload)
    if isinstance(payload, bytes) and payload:
        return payload, filename
    return b"", ""


def _read_portfolio_excel_source(excel_source) -> pd.DataFrame:
    try:
        xls = pd.ExcelFile(excel_source)
    except Exception:
        return empty_portfolio_df()

    preferred = ["보유현황", "포트폴리오", "입력_오늘", "portfolio", "sheet1", "Sheet1"]
    ordered_sheets = [s for s in preferred if s in xls.sheet_names] + [s for s in xls.sheet_names if s not in preferred]

    for sheet_name in ordered_sheets:
        try:
            raw = pd.read_excel(xls, sheet_name=sheet_name)
        except Exception:
            continue
        if raw is None or raw.empty:
            continue

        name_col = _pick_excel_column(raw, ["종목명", "기업명", "name", "stockname", "stock"])
        qty_col = _pick_excel_column(raw, ["보유수량", "수량", "qty", "quantity", "보유주수", "주수"])
        value_col = _pick_excel_column(raw, ["평가금액", "평가액", "marketvalue", "value", "평가"])
        pnl_col = _pick_excel_column(raw, ["손익금액", "손익", "pnl", "profitloss"])
        ret_col = _pick_excel_column(raw, ["수익률%", "수익률", "pnl%", "return%", "return"])
        currency_col = _pick_excel_column(raw, ["통화", "currency", "cur"])
        fx_col = _pick_excel_column(raw, ["환율원화기준", "환율", "fxrate", "fx"])

        if not name_col or not qty_col or not value_col:
            continue

        view = pd.DataFrame()
        view[COL_NAME] = raw[name_col].astype(str).str.strip()
        view[COL_QTY] = _to_num_series(raw[qty_col])
        view[COL_VALUE] = _to_num_series(raw[value_col])
        view[COL_PNL] = _to_num_series(raw[pnl_col]) if pnl_col else pd.Series([pd.NA] * len(raw))
        view[COL_RETURN] = _to_num_series(raw[ret_col]) if ret_col else pd.Series([pd.NA] * len(raw))
        view[COL_CURRENCY] = raw[currency_col].astype(str).str.strip().str.upper() if currency_col else "KRW"
        view[COL_FX_RATE] = _to_num_series(raw[fx_col]) if fx_col else pd.Series([pd.NA] * len(raw))

        view = view[view[COL_NAME] != ""].copy()
        view = view.dropna(subset=[COL_NAME, COL_QTY, COL_VALUE], how="any")
        if view.empty:
            continue

        # Fill missing PnL/Return from each other when possible.
        if view[COL_PNL].isna().any() and view[COL_RETURN].notna().any():
            principal = view[COL_VALUE] / (1 + (view[COL_RETURN] / 100.0))
            calc_pnl = view[COL_VALUE] - principal
            view.loc[view[COL_PNL].isna(), COL_PNL] = calc_pnl[view[COL_PNL].isna()]
        if view[COL_RETURN].isna().any() and view[COL_PNL].notna().any():
            principal = view[COL_VALUE] - view[COL_PNL]
            calc_ret = (view[COL_PNL] / principal.replace(0, pd.NA)) * 100.0
            view.loc[view[COL_RETURN].isna(), COL_RETURN] = calc_ret[view[COL_RETURN].isna()]

        view[COL_PNL] = pd.to_numeric(view[COL_PNL], errors="coerce").fillna(0.0)
        view[COL_RETURN] = pd.to_numeric(view[COL_RETURN], errors="coerce").fillna(0.0)
        view[COL_CURRENCY] = view[COL_CURRENCY].replace({"": "KRW", "NAN": "KRW", "NONE": "KRW"}).fillna("KRW")
        view[COL_FX_RATE] = pd.to_numeric(view[COL_FX_RATE], errors="coerce")
        view.loc[view[COL_CURRENCY] == "KRW", COL_FX_RATE] = 1.0
        view[COL_FX_RATE] = view[COL_FX_RATE].fillna(1.0)

        return view[COLUMNS]

    return empty_portfolio_df()


def load_portfolio_from_excel() -> pd.DataFrame:
    uploaded_bytes, _ = _get_uploaded_excel_bytes()
    if uploaded_bytes:
        uploaded_df = _read_portfolio_excel_source(BytesIO(uploaded_bytes))
        if not uploaded_df.empty:
            return uploaded_df

    excel_path = resolve_excel_path()
    if excel_path is None:
        return empty_portfolio_df()
    return _read_portfolio_excel_source(excel_path)


def _query_snapshot_for_date(conn: sqlite3.Connection, snapshot_date: date | str) -> pd.DataFrame:
    date_str = snapshot_date if isinstance(snapshot_date, str) else snapshot_date.isoformat()
    query = """
        SELECT
            stock_name AS 종목명,
            quantity AS 보유수량,
            COALESCE(currency, 'KRW') AS 통화,
            COALESCE(fx_rate, 1) AS "환율(원화기준)",
            market_value AS 평가금액,
            pnl_value AS 손익금액,
            pnl_pct AS "수익률(%)"
        FROM snapshots
        WHERE snapshot_date = ?
        ORDER BY market_value DESC
    """
    return pd.read_sql_query(query, conn, params=(date_str,))


def load_snapshot(snapshot_date: date) -> pd.DataFrame:
    conn = get_conn()
    try:
        df = _query_snapshot_for_date(conn, snapshot_date)
        if df.empty:
            latest_row = conn.execute(
                "SELECT MAX(snapshot_date) FROM snapshots WHERE snapshot_date <= ?",
                (snapshot_date.isoformat(),),
            ).fetchone()
            latest_date = str(latest_row[0] or "").strip() if latest_row else ""
    finally:
        conn.close()

    if not df.empty:
        return df

    # 오늘/미래 조회에서는 DB 저장이 없어도 엑셀 최신 입력값을 우선 반영한다.
    if snapshot_date >= date.today():
        excel_df = load_portfolio_from_excel()
        if not excel_df.empty:
            return excel_df

    if latest_date:
        conn = get_conn()
        try:
            df = _query_snapshot_for_date(conn, latest_date)
        finally:
            conn.close()
        if not df.empty:
            return df

    # 과거 날짜 조회에서는 없는 날짜를 엑셀 최신값으로 대체하지 않는다.
    # (선택일 기준 기록 혼선 방지)
    if snapshot_date >= date.today():
        excel_df = load_portfolio_from_excel()
        if not excel_df.empty:
            return excel_df

    return empty_portfolio_df()


def load_latest_snapshot() -> tuple[str | None, pd.DataFrame]:
    conn = get_conn()
    try:
        latest_date_row = conn.execute(
            "SELECT MAX(snapshot_date) AS snapshot_date FROM snapshots"
        ).fetchone()

        latest_date = latest_date_row[0] if latest_date_row and latest_date_row[0] else None
        if latest_date is None:
            return None, pd.DataFrame()

        query = """
            SELECT
                stock_name AS 종목명,
                quantity AS 보유수량,
                COALESCE(currency, 'KRW') AS 통화,
                COALESCE(fx_rate, 1) AS "환율(원화기준)",
                market_value AS 평가금액,
                pnl_value AS 손익금액,
                pnl_pct AS "수익률(%)"
            FROM snapshots
            WHERE snapshot_date = ?
            ORDER BY market_value DESC
        """
        df = pd.read_sql_query(query, conn, params=(latest_date,))
        return latest_date, df
    finally:
        conn.close()


def has_snapshot_on_date(snapshot_date: date) -> bool:
    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT 1 FROM snapshots WHERE snapshot_date = ? LIMIT 1",
            (snapshot_date.isoformat(),),
        ).fetchone()
        return bool(row)
    finally:
        conn.close()


def get_latest_snapshot_date_on_or_before(snapshot_date: date) -> date | None:
    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT MAX(snapshot_date) FROM snapshots WHERE snapshot_date <= ?",
            (snapshot_date.isoformat(),),
        ).fetchone()
    finally:
        conn.close()
    date_text = str(row[0] or "").strip() if row else ""
    return _safe_parse_date(date_text) if date_text else None


def save_snapshot_cash(snapshot_date: date, cash_krw: float | None, cash_usd: float | None) -> None:
    date_str = snapshot_date.isoformat()
    now_str = datetime.now().isoformat(timespec="seconds")
    krw_value = float(cash_krw) if cash_krw is not None and not pd.isna(cash_krw) else 0.0
    usd_value = float(cash_usd) if cash_usd is not None and not pd.isna(cash_usd) else 0.0

    conn = get_conn()
    try:
        conn.execute(
            """
            INSERT INTO snapshot_cash (snapshot_date, cash_krw, cash_usd, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?)
            ON CONFLICT(snapshot_date) DO UPDATE SET
                cash_krw = excluded.cash_krw,
                cash_usd = excluded.cash_usd,
                updated_at = excluded.updated_at
            """,
            (date_str, krw_value, usd_value, now_str, now_str),
        )
        conn.commit()
    finally:
        conn.close()


def load_snapshot_cash(snapshot_date: date) -> tuple[float, float]:
    conn = get_conn()
    try:
        row = conn.execute(
            """
            SELECT COALESCE(cash_krw, 0), COALESCE(cash_usd, 0)
            FROM snapshot_cash
            WHERE snapshot_date = ?
            """,
            (snapshot_date.isoformat(),),
        ).fetchone()
        if not row:
            fallback_date_row = conn.execute(
                """
                SELECT MAX(snapshot_date)
                FROM snapshot_cash
                WHERE snapshot_date <= ?
                """,
                (snapshot_date.isoformat(),),
            ).fetchone()
            fallback_date = str(fallback_date_row[0] or "").strip() if fallback_date_row else ""
            if fallback_date:
                row = conn.execute(
                    """
                    SELECT COALESCE(cash_krw, 0), COALESCE(cash_usd, 0)
                    FROM snapshot_cash
                    WHERE snapshot_date = ?
                    """,
                    (fallback_date,),
                ).fetchone()
    finally:
        conn.close()
    if not row:
        return 0.0, 0.0
    return float(row[0] or 0.0), float(row[1] or 0.0)


def get_snapshot_cash_krw(snapshot_date: date, usd_krw_rate: float | None = None) -> tuple[float, float, float]:
    cash_krw, cash_usd = load_snapshot_cash(snapshot_date)
    fx_rate = float(usd_krw_rate) if usd_krw_rate is not None else float(get_usd_krw_rate_for_date(snapshot_date)[0])
    cash_total_krw = float(cash_krw) + float(cash_usd) * fx_rate
    return cash_total_krw, float(cash_krw), float(cash_usd)


def load_history(as_of_date: date | None = None) -> pd.DataFrame:
    conn = get_conn()
    try:
        query = """
            SELECT
                snapshot_date AS snapshot_date,
                UPPER(COALESCE(currency, 'KRW')) AS currency,
                market_value,
                pnl_value
            FROM snapshots
            ORDER BY snapshot_date
        """
        raw_df = pd.read_sql_query(query, conn)
        cash_df = pd.read_sql_query(
            """
            SELECT
                snapshot_date AS snapshot_date,
                COALESCE(cash_krw, 0) AS cash_krw,
                COALESCE(cash_usd, 0) AS cash_usd
            FROM snapshot_cash
            ORDER BY snapshot_date
            """,
            conn,
        )
    finally:
        conn.close()

    hist_df = pd.DataFrame(columns=["snapshot_date", "total_value", "total_pnl"])
    if not raw_df.empty:
        raw_df["snapshot_date"] = pd.to_datetime(raw_df["snapshot_date"])
        raw_df["fx_effective"] = 1.0
        usd_mask = raw_df["currency"] == "USD"
        if usd_mask.any():
            usd_dates = sorted(raw_df.loc[usd_mask, "snapshot_date"].dt.date.unique().tolist())
            rate_map = {d: get_usd_krw_rate_for_date(d)[0] for d in usd_dates}
            raw_df.loc[usd_mask, "fx_effective"] = raw_df.loc[usd_mask, "snapshot_date"].dt.date.map(rate_map).astype(float)

        raw_df["value_krw"] = raw_df["market_value"] * raw_df["fx_effective"]
        raw_df["pnl_krw"] = raw_df["pnl_value"] * raw_df["fx_effective"]
        hist_df = (
            raw_df.groupby("snapshot_date", as_index=False)
            .agg(total_value=("value_krw", "sum"), total_pnl=("pnl_krw", "sum"))
            .sort_values("snapshot_date")
        )

    if not cash_df.empty:
        cash_df["snapshot_date"] = pd.to_datetime(cash_df["snapshot_date"])
        cash_df = cash_df.sort_values("snapshot_date")
        cash_dates = sorted(cash_df["snapshot_date"].dt.date.unique().tolist())
        cash_rate_map = {d: get_usd_krw_rate_for_date(d)[0] for d in cash_dates}
        cash_df["usd_krw_rate"] = cash_df["snapshot_date"].dt.date.map(cash_rate_map).astype(float)
        cash_df["cash_total_krw"] = cash_df["cash_krw"] + cash_df["cash_usd"] * cash_df["usd_krw_rate"]
        cash_df = cash_df[["snapshot_date", "cash_total_krw", "cash_krw", "cash_usd"]]
        if hist_df.empty:
            hist_df = cash_df.rename(columns={"cash_total_krw": "total_value"})
            hist_df["total_pnl"] = 0.0
            hist_df["cash_krw"] = hist_df["cash_krw"].fillna(0.0)
            hist_df["cash_usd"] = hist_df["cash_usd"].fillna(0.0)
        else:
            # 스냅샷 날짜별 예수금은 동일일이 없더라도 직전값(<=날짜)을 승계 적용한다.
            hist_df = pd.merge_asof(
                hist_df.sort_values("snapshot_date"),
                cash_df.sort_values("snapshot_date"),
                on="snapshot_date",
                direction="backward",
            )
            hist_df["cash_total_krw"] = hist_df["cash_total_krw"].fillna(0.0)
            hist_df["cash_krw"] = hist_df["cash_krw"].fillna(0.0)
            hist_df["cash_usd"] = hist_df["cash_usd"].fillna(0.0)
            hist_df["total_value"] = hist_df["total_value"].fillna(0.0) + hist_df["cash_total_krw"]
            hist_df["total_pnl"] = hist_df["total_pnl"].fillna(0.0)
        if "cash_total_krw" in hist_df.columns:
            hist_df = hist_df.drop(columns=["cash_total_krw"])

    if hist_df.empty:
        return hist_df

    hist_df = hist_df.sort_values("snapshot_date")
    if as_of_date is not None:
        hist_df = hist_df[hist_df["snapshot_date"] <= pd.Timestamp(as_of_date)].copy()
        if hist_df.empty:
            return hist_df
    hist_df["is_carry_forward"] = False
    anchor_date = as_of_date or date.today()
    today_ts = pd.Timestamp(anchor_date)
    last_ts = pd.Timestamp(hist_df["snapshot_date"].max())
    if last_ts.normalize() < today_ts.normalize():
        carry_row = hist_df.iloc[-1].copy()
        carry_row["snapshot_date"] = today_ts
        carry_row["is_carry_forward"] = True
        hist_df = pd.concat([hist_df, pd.DataFrame([carry_row])], ignore_index=True)
        hist_df = hist_df.sort_values("snapshot_date")

    hist_df["total_principal"] = hist_df["total_value"] - hist_df["total_pnl"]
    hist_df["total_return_pct"] = (
        hist_df["total_pnl"] / hist_df["total_principal"].replace(0, pd.NA)
    ) * 100
    return hist_df


def load_stock_history(stock_name: str) -> pd.DataFrame:
    conn = get_conn()
    try:
        query = """
            SELECT
                snapshot_date AS snapshot_date,
                stock_name AS stock_name,
                quantity AS quantity,
                COALESCE(currency, 'KRW') AS currency,
                COALESCE(fx_rate, 1) AS fx_rate,
                market_value AS market_value,
                pnl_value AS pnl_value,
                pnl_pct AS pnl_pct
            FROM snapshots
            WHERE stock_name = ?
            ORDER BY snapshot_date
        """
        stock_df = pd.read_sql_query(query, conn, params=(stock_name,))
    finally:
        conn.close()

    if not stock_df.empty:
        stock_df["snapshot_date"] = pd.to_datetime(stock_df["snapshot_date"])
        stock_df["currency"] = stock_df["currency"].astype(str).str.upper()
        stock_df["fx_effective"] = 1.0
        usd_mask = stock_df["currency"] == "USD"
        if usd_mask.any():
            usd_dates = sorted(stock_df.loc[usd_mask, "snapshot_date"].dt.date.unique().tolist())
            rate_map = {d: get_usd_krw_rate_for_date(d)[0] for d in usd_dates}
            stock_df.loc[usd_mask, "fx_effective"] = stock_df.loc[usd_mask, "snapshot_date"].dt.date.map(rate_map).astype(float)
        stock_df["market_value_krw"] = stock_df["market_value"] * stock_df["fx_effective"]
        stock_df["pnl_value_krw"] = stock_df["pnl_value"] * stock_df["fx_effective"]
    return stock_df


def score_linear(value: float | None, low: float, high: float, reverse: bool = False) -> float:
    if value is None or pd.isna(value):
        return 50.0
    if high <= low:
        return 50.0

    ratio = (float(value) - low) / (high - low)
    ratio = max(0.0, min(1.0, ratio))
    score = (1.0 - ratio) * 100 if reverse else ratio * 100
    return max(0.0, min(100.0, score))


def compute_company_scores(metrics: dict, weights: dict) -> dict:
    dividend_score = score_linear(
        metrics.get("dividend_yield"),
        SCORE_METRIC_CONFIG["dividend_yield"]["min"],
        SCORE_METRIC_CONFIG["dividend_yield"]["max"],
        SCORE_METRIC_CONFIG["dividend_yield"]["reverse"],
    )
    growth_score = (
        score_linear(
            metrics.get("revenue_growth"),
            SCORE_METRIC_CONFIG["revenue_growth"]["min"],
            SCORE_METRIC_CONFIG["revenue_growth"]["max"],
            SCORE_METRIC_CONFIG["revenue_growth"]["reverse"],
        )
        * 0.4
        + score_linear(
            metrics.get("eps_growth"),
            SCORE_METRIC_CONFIG["eps_growth"]["min"],
            SCORE_METRIC_CONFIG["eps_growth"]["max"],
            SCORE_METRIC_CONFIG["eps_growth"]["reverse"],
        )
        * 0.4
        + score_linear(
            metrics.get("roe"),
            SCORE_METRIC_CONFIG["roe"]["min"],
            SCORE_METRIC_CONFIG["roe"]["max"],
            SCORE_METRIC_CONFIG["roe"]["reverse"],
        )
        * 0.2
    )
    stability_score = (
        score_linear(
            metrics.get("debt_ratio"),
            SCORE_METRIC_CONFIG["debt_ratio"]["min"],
            SCORE_METRIC_CONFIG["debt_ratio"]["max"],
            SCORE_METRIC_CONFIG["debt_ratio"]["reverse"],
        )
        * 0.4
        + score_linear(
            metrics.get("current_ratio"),
            SCORE_METRIC_CONFIG["current_ratio"]["min"],
            SCORE_METRIC_CONFIG["current_ratio"]["max"],
            SCORE_METRIC_CONFIG["current_ratio"]["reverse"],
        )
        * 0.3
        + score_linear(
            metrics.get("operating_margin"),
            SCORE_METRIC_CONFIG["operating_margin"]["min"],
            SCORE_METRIC_CONFIG["operating_margin"]["max"],
            SCORE_METRIC_CONFIG["operating_margin"]["reverse"],
        )
        * 0.3
    )
    valuation_score = (
        score_linear(
            metrics.get("per"),
            SCORE_METRIC_CONFIG["per"]["min"],
            SCORE_METRIC_CONFIG["per"]["max"],
            SCORE_METRIC_CONFIG["per"]["reverse"],
        )
        * 0.6
        + score_linear(
            metrics.get("pbr"),
            SCORE_METRIC_CONFIG["pbr"]["min"],
            SCORE_METRIC_CONFIG["pbr"]["max"],
            SCORE_METRIC_CONFIG["pbr"]["reverse"],
        )
        * 0.4
    )

    total_weight = sum(weights.values()) if weights else 0.0
    if total_weight <= 0:
        normalized = DEFAULT_SCORE_WEIGHTS.copy()
    else:
        normalized = {k: v / total_weight * 100.0 for k, v in weights.items()}

    total_score = (
        dividend_score * normalized["dividend"] / 100.0
        + growth_score * normalized["growth"] / 100.0
        + stability_score * normalized["stability"] / 100.0
        + valuation_score * normalized["valuation"] / 100.0
    )

    return {
        "dividend_score": round(dividend_score, 2),
        "growth_score": round(growth_score, 2),
        "stability_score": round(stability_score, 2),
        "valuation_score": round(valuation_score, 2),
        "total_score": round(total_score, 2),
        "weights": normalized,
    }


def compute_company_metric_ranking(
    companies: list[str],
    metric_keys: list[str],
    metric_weights: dict[str, float],
    use_ai_ticker: bool = False,
    ai_provider: str = "openai",
    ai_api_key: str = "",
    ai_model: str = "",
) -> tuple[pd.DataFrame, pd.DataFrame]:
    valid_companies = [str(c).strip() for c in companies if str(c).strip()]
    valid_metrics = [m for m in metric_keys if m in SCORE_METRIC_CONFIG]
    if not valid_companies or not valid_metrics:
        return pd.DataFrame(), pd.DataFrame()

    raw_weights = {m: max(0.0, float(metric_weights.get(m, 0.0))) for m in valid_metrics}
    total_weight = sum(raw_weights.values())
    if total_weight <= 0:
        normalized_weights = {m: 100.0 / len(valid_metrics) for m in valid_metrics}
    else:
        normalized_weights = {m: (raw_weights[m] / total_weight) * 100.0 for m in valid_metrics}

    rows = []
    error_rows = []
    company_list_df = load_company_list()
    sector_hint_map = {}
    if not company_list_df.empty:
        for _, row in company_list_df.iterrows():
            nm = str(row.get("stock_name") or "").strip()
            if not nm:
                continue
            sector_hint_map[nm] = str(row.get("sector") or "").strip()
    for company_name in valid_companies:
        ticker = ""
        ticker_source = ""
        sector_value = sector_hint_map.get(company_name, "").strip() or "미분류"

        list_ticker = get_company_list_ticker(company_name)
        if list_ticker:
            ticker = list_ticker
            ticker_source = "기업 리스트 저장값"
        else:
            ticker, ticker_source = resolve_ticker_auto_with_retry(
                company_name,
                use_ai=use_ai_ticker,
                api_key=ai_api_key,
                model=ai_model,
                provider=ai_provider,
            )

        if not ticker:
            row = {
                "기업명": company_name,
                "산업섹터": sector_value,
                "티커": "",
                "티커소스": ticker_source or "티커 없음",
                "데이터소스": "-",
                "상태": "티커 없음",
                "총점": None,
            }
            for m in valid_metrics:
                label = SCORE_METRIC_CONFIG[m]["label"]
                row[f"{label} 값"] = None
                row[f"{label} 점수"] = None
            rows.append(row)
            error_rows.append({"기업명": company_name, "티커": "", "오류": ticker_source or "티커를 찾지 못했습니다."})
            continue

        metrics, _, err_msg, metric_source = fetch_company_metrics_multi_source(ticker)
        if err_msg:
            row = {
                "기업명": company_name,
                "산업섹터": sector_value,
                "티커": ticker,
                "티커소스": ticker_source,
                "데이터소스": metric_source or "-",
                "상태": "데이터 오류",
                "총점": None,
            }
            for m in valid_metrics:
                label = SCORE_METRIC_CONFIG[m]["label"]
                row[f"{label} 값"] = None
                row[f"{label} 점수"] = None
            rows.append(row)
            error_rows.append({"기업명": company_name, "티커": ticker, "오류": err_msg})
            continue

        upsert_company_list_entry(company_name, ticker, source="analysis_compare")

        row = {
            "기업명": company_name,
            "산업섹터": sector_value,
            "티커": ticker,
            "티커소스": ticker_source,
            "데이터소스": metric_source or "-",
            "상태": "정상",
        }
        weighted_sum = 0.0
        missing_metric_labels = []
        available_metric_count = 0
        for m in valid_metrics:
            cfg = SCORE_METRIC_CONFIG[m]
            label = cfg["label"]
            metric_value = metrics.get(m)
            has_value = metric_value is not None and not pd.isna(metric_value)
            if has_value:
                metric_value_float = float(metric_value)
                metric_score = score_linear(metric_value_float, cfg["min"], cfg["max"], cfg["reverse"])
                row[f"{label} 값"] = metric_value_float
                row[f"{label} 점수"] = round(metric_score, 2)
                weighted_sum += metric_score * normalized_weights[m]
                available_metric_count += 1
            else:
                row[f"{label} 값"] = None
                row[f"{label} 점수"] = None
                missing_metric_labels.append(label)

        if available_metric_count < len(valid_metrics):
            row["상태"] = "지표 부족"
            row["총점"] = None
            rows.append(row)
            error_rows.append(
                {
                    "기업명": company_name,
                    "티커": ticker,
                    "오류": f"누락 지표: {', '.join(missing_metric_labels)}",
                }
            )
            continue

        row["총점"] = round(weighted_sum / 100.0, 2)
        rows.append(row)

    result_df = pd.DataFrame(rows)
    error_df = pd.DataFrame(error_rows)
    if result_df.empty:
        return result_df, error_df

    status_rank = {"정상": 0, "지표 부족": 1, "데이터 오류": 2, "티커 없음": 3}
    result_df["__status_rank"] = result_df["상태"].map(status_rank).fillna(9)
    result_df = result_df.sort_values(["__status_rank", "총점"], ascending=[True, False]).drop(columns="__status_rank")
    result_df["순위"] = 0
    normal_idx = result_df[result_df["상태"] == "정상"].index.tolist()
    for idx, ridx in enumerate(normal_idx, start=1):
        result_df.loc[ridx, "순위"] = idx
    return result_df, error_df


def sanitize_compare_result_df(result_df: pd.DataFrame, metric_keys: list[str]) -> pd.DataFrame:
    if result_df is None or result_df.empty:
        return result_df

    fixed = result_df.copy()
    if "상태" not in fixed.columns:
        return fixed

    value_cols = []
    metric_to_score_col = {}
    for m in metric_keys or []:
        cfg = SCORE_METRIC_CONFIG.get(m)
        if not cfg:
            continue
        value_col = f"{cfg['label']} 값"
        score_col = f"{cfg['label']} 점수"
        if value_col in fixed.columns:
            value_cols.append(value_col)
        metric_to_score_col[value_col] = score_col

    if not value_cols:
        return fixed

    normal_mask = fixed["상태"].astype(str).eq("정상")
    if not normal_mask.any():
        return fixed

    missing_any_mask = fixed[value_cols].isna().any(axis=1)
    downgrade_mask = normal_mask & missing_any_mask
    if not downgrade_mask.any():
        return fixed

    fixed.loc[downgrade_mask, "상태"] = "지표 부족"
    if "총점" in fixed.columns:
        fixed.loc[downgrade_mask, "총점"] = None

    for value_col in value_cols:
        score_col = metric_to_score_col.get(value_col, "")
        if score_col and score_col in fixed.columns:
            fixed.loc[downgrade_mask & fixed[value_col].isna(), score_col] = None

    if "순위" in fixed.columns:
        fixed["순위"] = 0
        normal_sorted = fixed[fixed["상태"] == "정상"].sort_values("총점", ascending=False).index.tolist()
        for idx, ridx in enumerate(normal_sorted, start=1):
            fixed.loc[ridx, "순위"] = idx

    status_rank = {"정상": 0, "지표 부족": 1, "데이터 오류": 2, "티커 없음": 3}
    fixed["__status_rank"] = fixed["상태"].map(status_rank).fillna(9)
    fixed = fixed.sort_values(["__status_rank", "총점"], ascending=[True, False]).drop(columns="__status_rank")
    return fixed


def get_saved_ticker_hint(company_name: str) -> str:
    name = (company_name or "").strip()
    if not name:
        return ""
    conn = get_conn()
    try:
        row = conn.execute(
            """
            SELECT ticker
            FROM company_scores
            WHERE stock_name = ? AND ticker IS NOT NULL AND ticker != ''
            ORDER BY score_date DESC, created_at DESC
            LIMIT 1
            """,
            (name,),
        ).fetchone()
    finally:
        conn.close()
    raw = (row[0] or "").strip().upper() if row else ""
    return clean_valid_ticker(raw)


def _extract_ticker_candidates_from_web_text(text: str) -> list[str]:
    raw = str(text or "")
    if not raw:
        return []
    decoded = unquote(unquote(raw))
    corpus = f"{raw}\n{decoded}"
    candidates = []

    # 야후 검색 결과 제목(h3)에서 티커를 우선 추출한다.
    for m in re.finditer(r"<h3[^>]*>(.*?)</h3>", corpus, re.I | re.S):
        title_html = m.group(1) or ""
        title = re.sub(r"<.*?>", " ", title_html)
        title = html.unescape(title)
        title = re.sub(r"\s+", " ", title).strip()
        if not title:
            continue

        title_patterns = [
            r"\(([A-Za-z0-9.\-^=]{1,24})\)\s*(?:Stock|주가|Quote|개요)",
            r"\b([A-Za-z0-9.\-^=]{1,24})\s*-\s*\([^)]*?(?:NYSE|NASDAQ|AMEX|KOSPI|KOSDAQ)",
        ]
        for pat in title_patterns:
            for hit in re.finditer(pat, title, re.I):
                cand = clean_valid_ticker(hit.group(1))
                if cand:
                    candidates.append(cand)

    # 링크 안에 finance.yahoo.com/quote/{TICKER}가 있으면 추출한다.
    for m in re.finditer(r"finance\.yahoo\.com/quote/([A-Za-z0-9.\-^=]{1,24})", corpus, re.I):
        cand = clean_valid_ticker(m.group(1))
        if cand:
            candidates.append(cand)

    # 일반적인 시장 접두 패턴도 보조로 수집한다.
    for m in re.finditer(r"(?:NASDAQ|NYSE|AMEX|OTC|KOSPI|KOSDAQ)\s*[:\-]\s*([A-Za-z0-9.\-^=]{1,24})", corpus, re.I):
        cand = clean_valid_ticker(m.group(1))
        if cand:
            candidates.append(cand)

    # 최후 보조: "ticker/symbol" 컨텍스트 근처
    for m in re.finditer(r"(?:stock symbol|ticker)\s*[:\-]?\s*([A-Za-z0-9.\-^=]{1,24})", corpus, re.I):
        cand = clean_valid_ticker(m.group(1))
        if cand:
            candidates.append(cand)

    # RU=... redirect 안에 숨어있는 실제 URL도 한 번 더 펼쳐본다.
    for m in re.finditer(r"/RU=([^/]+)/", raw):
        ru = unquote(unquote(m.group(1)))
        for hit in re.finditer(r"finance\.yahoo\.com/quote/([A-Za-z0-9.\-^=]{1,24})", ru, re.I):
            cand = clean_valid_ticker(hit.group(1))
            if cand:
                candidates.append(cand)

    blocked = {
        "YAHOO",
        "SYMBOL",
        "PRICE",
        "TODAY",
        "LIVE",
        "TAPE",
        "FOR",
        "CODE",
        "USD",
        "KRW",
        "JPY",
        "EUR",
        "CNY",
        "GBP",
        "ETF",
        "ADR",
        "NYSE",
        "NASDAQ",
        "AMEX",
        "OTC",
        "KOSPI",
        "KOSDAQ",
    }
    cleaned = []
    seen = set()
    for cand in candidates:
        cand = clean_valid_ticker(cand)
        if not cand:
            continue
        if cand in blocked:
            continue
        if cand in seen:
            continue
        seen.add(cand)
        cleaned.append(cand)
    return cleaned


@st.cache_data(ttl=60 * 30, show_spinner=False)
def _web_search_ticker_candidates(company_name: str) -> tuple[list[dict], str]:
    name = (company_name or "").strip()
    if not name:
        return [], "기업명이 비어 있습니다."

    queries = [
        f"{name} ticker",
        f"{name} stock ticker",
        f"{name} 주식 티커",
    ]
    score_map: dict[str, float] = {}
    info_map: dict[str, dict] = {}
    errs = []

    for q_idx, q in enumerate(queries):
        try:
            resp = requests.get(
                "https://search.yahoo.com/search",
                params={"p": q},
                headers=HTTP_HEADERS_COMMON,
                timeout=12,
            )
            resp.raise_for_status()
            body = resp.text or ""
            hits = _extract_ticker_candidates_from_web_text(body)
            for rank, ticker in enumerate(hits):
                score = 6.0 - q_idx - (rank * 0.18)
                if ticker not in score_map or score > score_map[ticker]:
                    score_map[ticker] = score
                    info_map[ticker] = {
                        "symbol": ticker,
                        "name": name,
                        "exchange": "WEB",
                        "region": "",
                    }
        except Exception as exc:
            errs.append(str(exc))

    if not score_map:
        err_text = " / ".join(errs[:2]) if errs else "웹 검색 결과가 없습니다."
        return [], err_text

    ordered = sorted(score_map.keys(), key=lambda t: score_map[t], reverse=True)
    return [info_map[t] for t in ordered], ""


def search_ticker_web_first(company_name: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."
    candidates, err = _web_search_ticker_candidates(name)
    if not candidates:
        return "", f"웹 검색 실패: {err}"
    return choose_best_ticker_candidate(name, candidates, "웹검색", market_preference=market_preference)


@st.cache_data(ttl=60 * 30, show_spinner=False)
def _fetch_yahoo_search_quotes_cached(company_name: str) -> tuple[list[dict], str]:
    name = (company_name or "").strip()
    if not name:
        return [], "기업명이 비어 있습니다."

    endpoints = [
        ("https://query2.finance.yahoo.com/v1/finance/search", "query2"),
        ("https://query1.finance.yahoo.com/v1/finance/search", "query1"),
    ]
    last_err = ""
    for endpoint, endpoint_label in endpoints:
        for attempt in range(3):
            try:
                resp = requests.get(
                    endpoint,
                    params={"q": name, "quotesCount": 12, "newsCount": 0},
                    headers=HTTP_HEADERS_COMMON,
                    timeout=10,
                )
                if resp.status_code == 429:
                    last_err = f"429 Too Many Requests ({endpoint_label})"
                    if attempt < 2:
                        time.sleep(0.8 * (attempt + 1))
                        continue
                    break
                resp.raise_for_status()
                payload = resp.json() or {}
                quotes = payload.get("quotes") or []
                if quotes:
                    return quotes, endpoint_label
                break
            except Exception as exc:
                last_err = f"{endpoint_label} 조회 실패: {exc}"
                if attempt < 2:
                    time.sleep(0.7 * (attempt + 1))
                    continue
                break

    return [], last_err or "검색 결과가 없습니다."


@st.cache_data(ttl=60 * 60 * 12, show_spinner=False)
def _load_sec_company_ticker_dataset() -> tuple[list[dict], str]:
    try:
        resp = requests.get(
            "https://www.sec.gov/files/company_tickers.json",
            headers={
                "User-Agent": "invest-diary/1.0 (public app)",
                "Accept": "application/json",
            },
            timeout=16,
        )
        resp.raise_for_status()
        payload = resp.json() or {}
    except Exception as exc:
        return [], f"SEC 목록 조회 실패: {exc}"

    rows = []
    if isinstance(payload, dict):
        iterable = payload.values()
    else:
        iterable = payload
    for item in iterable:
        if not isinstance(item, dict):
            continue
        symbol = clean_valid_ticker(str(item.get("ticker") or ""))
        title = str(item.get("title") or "").strip()
        if not symbol or not title:
            continue
        rows.append(
            {
                "symbol": symbol,
                "name": title,
                "exchange": "SEC",
                "region": "United States",
            }
        )
    return rows, ""


def search_ticker_sec_dataset(company_name: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."
    if _market_pref_normalized(market_preference) == "domestic":
        return "", "SEC 검색은 해외 종목용입니다."

    rows, err = _load_sec_company_ticker_dataset()
    if err:
        return "", err
    if not rows:
        return "", "SEC 목록 데이터가 비어 있습니다."

    q_norm = normalize_company_name_for_match(name)
    direct = []
    fuzzy = []
    for row in rows:
        title = str(row.get("name") or "").strip()
        title_norm = normalize_company_name_for_match(title)
        if q_norm and title_norm and (q_norm in title_norm or title_norm in q_norm):
            direct.append(row)
            continue
        sim = _name_similarity(name, title)
        if sim >= 0.50:
            item = dict(row)
            item["_sim"] = sim
            fuzzy.append(item)

    candidates = direct
    if not candidates and fuzzy:
        fuzzy.sort(key=lambda x: float(x.get("_sim", 0.0)), reverse=True)
        candidates = fuzzy[:18]
    if not candidates:
        return "", "SEC 목록에서 일치 기업을 찾지 못했습니다."

    return choose_best_ticker_candidate(name, candidates, "SEC 목록", market_preference="foreign")


def search_ticker_yfinance(company_name: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."

    quotes, query_source = _fetch_yahoo_search_quotes_cached(name)
    if not quotes:
        if "429" in str(query_source):
            return "", f"yfinance 검색 제한: {query_source}"
        return "", f"yfinance 검색 실패: {query_source}"

    def is_equity(q):
        return str(q.get("quoteType", "")).upper() == "EQUITY"

    candidates = []
    for q in quotes:
        if not is_equity(q):
            continue
        symbol = clean_valid_ticker(str(q.get("symbol", "")))
        if not symbol:
            continue
        longname = q.get("longname") or q.get("shortname") or q.get("name") or ""
        candidates.append(
            {
                "symbol": symbol,
                "name": str(longname).strip(),
                "exchange": str(q.get("exchange") or q.get("exchDisp") or "").strip(),
                "region": str(q.get("region") or "").strip(),
            }
        )
    if not candidates:
        return "", "주식 티커 결과가 없습니다."
    ticker, msg = choose_best_ticker_candidate(name, candidates, "yfinance", market_preference=market_preference)
    if ticker:
        return ticker, f"{msg}, {query_source}"
    return "", msg


def _build_ticker_from_naver_item(item: dict) -> str:
    if not isinstance(item, dict):
        return ""
    code = clean_valid_ticker(str(item.get("code") or ""))
    reuters = clean_valid_ticker(str(item.get("reutersCode") or ""))
    type_code = str(item.get("typeCode") or "").strip().upper()
    nation_code = str(item.get("nationCode") or "").strip().upper()
    url = str(item.get("url") or "").strip().lower()

    is_domestic = nation_code == "KOR" or "/domestic/stock/" in url
    if is_domestic:
        base = clean_valid_ticker(code or reuters)
        if re.fullmatch(r"[0-9A-Z]{6}", base or ""):
            suffix = ".KQ" if type_code == "KOSDAQ" else ".KS"
            return f"{base}{suffix}"
    if nation_code == "USA":
        return clean_valid_ticker(code or reuters)
    return clean_valid_ticker(reuters or code)


@st.cache_data(ttl=60 * 30, show_spinner=False)
def _fetch_naver_stock_candidates(company_name: str) -> tuple[list[dict], str]:
    name = (company_name or "").strip()
    if not name:
        return [], "기업명이 비어 있습니다."

    try:
        resp = requests.get(
            "https://ac.stock.naver.com/ac",
            params={"q": name, "target": "stock,ipo,index,marketindicator"},
            headers=HTTP_HEADERS_COMMON,
            timeout=12,
        )
        resp.raise_for_status()
        payload = resp.json() or {}
    except Exception as exc:
        return [], f"Naver 검색 실패: {exc}"

    items = payload.get("items") or []
    if not isinstance(items, list) or not items:
        return [], "Naver 검색 결과가 없습니다."

    candidates = []
    for item in items:
        if not isinstance(item, dict):
            continue
        category = str(item.get("category") or "").strip().lower()
        if category and category != "stock":
            continue
        symbol = _build_ticker_from_naver_item(item)
        if not symbol:
            continue
        candidates.append(
            {
                "symbol": symbol,
                "name": str(item.get("name") or "").strip(),
                "exchange": str(item.get("typeCode") or "").strip(),
                "region": str(item.get("nationCode") or "").strip(),
            }
        )
    if not candidates:
        return [], "Naver 유효 티커 결과가 없습니다."
    return candidates, ""


def search_ticker_naver(company_name: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."
    candidates, err = _fetch_naver_stock_candidates(name)
    if not candidates:
        return "", err
    return choose_best_ticker_candidate(name, candidates, "Naver증권", market_preference=market_preference)


@st.cache_data(ttl=60 * 60 * 6, show_spinner=False)
def _resolve_krx_code_to_yahoo_ticker(code: str) -> str:
    raw = normalize_ticker_text(code)
    if not re.fullmatch(r"[0-9A-Z]{6}", raw or ""):
        return ""
    try:
        resp = requests.get(
            "https://ac.stock.naver.com/ac",
            params={"q": raw, "target": "stock"},
            headers=HTTP_HEADERS_COMMON,
            timeout=10,
        )
        resp.raise_for_status()
        payload = resp.json() or {}
    except Exception:
        return f"{raw}.KS"
    items = payload.get("items") or []
    for item in items:
        if not isinstance(item, dict):
            continue
        nation = str(item.get("nationCode") or "").strip().upper()
        code_item = normalize_ticker_text(str(item.get("code") or ""))
        reuters = normalize_ticker_text(str(item.get("reutersCode") or ""))
        if nation != "KOR":
            continue
        if raw not in {code_item, reuters}:
            continue
        type_code = str(item.get("typeCode") or "").strip().upper()
        suffix = ".KQ" if type_code == "KOSDAQ" else ".KS"
        return f"{raw}{suffix}"
    return f"{raw}.KS"


def _normalize_external_ticker_candidate(raw_symbol: str) -> str:
    raw = normalize_ticker_text(raw_symbol)
    if not raw:
        return ""
    # Reuters 접미(.O/.N)를 Yahoo 표기로 변환
    m_reuters_us = re.fullmatch(r"([A-Z][A-Z0-9\-]{0,9})\.(O|N)", raw)
    if m_reuters_us:
        raw = m_reuters_us.group(1)
    # 국내 6자리(숫자+문자 포함) 코드는 Yahoo 접미를 보강
    if "." not in raw and re.fullmatch(r"[0-9A-Z]{6}", raw) and any(ch.isdigit() for ch in raw):
        return _resolve_krx_code_to_yahoo_ticker(raw)
    return clean_valid_ticker(raw)


def _extract_toss_ticker_candidates_from_web_text(text: str) -> list[str]:
    raw = str(text or "")
    if not raw:
        return []
    decoded = unquote(unquote(raw))
    corpus = f"{raw}\n{decoded}"
    candidates = []
    for m in re.finditer(r"(?:www\.)?tossinvest\.com/stocks/([A-Za-z0-9.\-]{1,24})", corpus, re.I):
        cand = _normalize_external_ticker_candidate(m.group(1))
        if cand:
            candidates.append(cand)
    uniq = []
    seen = set()
    for c in candidates:
        if c in seen:
            continue
        seen.add(c)
        uniq.append(c)
    return uniq


@st.cache_data(ttl=60 * 30, show_spinner=False)
def _search_ticker_toss_candidates(company_name: str) -> tuple[list[dict], str]:
    name = (company_name or "").strip()
    if not name:
        return [], "기업명이 비어 있습니다."
    queries = [
        f"site:tossinvest.com/stocks {name}",
        f"{name} tossinvest stock",
    ]
    score_map: dict[str, float] = {}
    info_map: dict[str, dict] = {}
    errs = []
    for q_idx, q in enumerate(queries):
        try:
            resp = requests.get(
                "https://search.yahoo.com/search",
                params={"p": q},
                headers=HTTP_HEADERS_COMMON,
                timeout=12,
            )
            resp.raise_for_status()
            hits = _extract_toss_ticker_candidates_from_web_text(resp.text or "")
            for rank, ticker in enumerate(hits):
                score = 5.8 - q_idx - (rank * 0.2)
                if ticker not in score_map or score > score_map[ticker]:
                    score_map[ticker] = score
                    info_map[ticker] = {
                        "symbol": ticker,
                        "name": "",
                        "exchange": "TOSS",
                        "region": "",
                    }
        except Exception as exc:
            errs.append(str(exc))
    if not score_map:
        return [], (" / ".join(errs[:2]) if errs else "Toss 웹검색 결과가 없습니다.")
    ordered = sorted(score_map.keys(), key=lambda t: score_map[t], reverse=True)
    return [info_map[t] for t in ordered], ""


def search_ticker_toss(company_name: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."
    candidates, err = _search_ticker_toss_candidates(name)
    if not candidates:
        return "", f"Toss 검색 실패: {err}"
    return choose_best_ticker_candidate(name, candidates, "Toss웹검색", market_preference=market_preference)


def _extract_google_finance_ticker_candidates_from_web_text(text: str) -> list[str]:
    raw = str(text or "")
    if not raw:
        return []
    decoded = unquote(unquote(raw))
    corpus = f"{raw}\n{decoded}"
    candidates = []
    for m in re.finditer(
        r"google\.com/finance/quote/([A-Za-z0-9.\-]{1,24}):([A-Za-z0-9.\-]{1,20})",
        corpus,
        re.I,
    ):
        symbol = normalize_ticker_text(m.group(1))
        exch = normalize_ticker_text(m.group(2))
        picked = ""
        if exch in {"NASDAQ", "NYSE", "AMEX", "OTCMKTS", "OTC"}:
            picked = _normalize_external_ticker_candidate(symbol)
        elif exch in {"KOSDAQ"} and re.fullmatch(r"[0-9A-Z]{6}", symbol):
            picked = f"{symbol}.KQ"
        elif exch in {"KRX", "KOSPI"} and re.fullmatch(r"[0-9A-Z]{6}", symbol):
            picked = _resolve_krx_code_to_yahoo_ticker(symbol)
        if not picked:
            picked = _normalize_external_ticker_candidate(symbol)
        if picked:
            candidates.append(picked)
    uniq = []
    seen = set()
    for c in candidates:
        if c in seen:
            continue
        seen.add(c)
        uniq.append(c)
    return uniq


@st.cache_data(ttl=60 * 30, show_spinner=False)
def _search_ticker_google_candidates(company_name: str) -> tuple[list[dict], str]:
    name = (company_name or "").strip()
    if not name:
        return [], "기업명이 비어 있습니다."
    queries = [
        f"site:google.com/finance/quote {name}",
        f"{name} google finance quote",
    ]
    score_map: dict[str, float] = {}
    info_map: dict[str, dict] = {}
    errs = []
    for q_idx, q in enumerate(queries):
        try:
            resp = requests.get(
                "https://search.yahoo.com/search",
                params={"p": q},
                headers=HTTP_HEADERS_COMMON,
                timeout=12,
            )
            resp.raise_for_status()
            hits = _extract_google_finance_ticker_candidates_from_web_text(resp.text or "")
            for rank, ticker in enumerate(hits):
                score = 5.5 - q_idx - (rank * 0.22)
                if ticker not in score_map or score > score_map[ticker]:
                    score_map[ticker] = score
                    info_map[ticker] = {
                        "symbol": ticker,
                        "name": "",
                        "exchange": "GOOGLE",
                        "region": "",
                    }
        except Exception as exc:
            errs.append(str(exc))
    if not score_map:
        return [], (" / ".join(errs[:2]) if errs else "Google 웹검색 결과가 없습니다.")
    ordered = sorted(score_map.keys(), key=lambda t: score_map[t], reverse=True)
    return [info_map[t] for t in ordered], ""


def search_ticker_google(company_name: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."
    candidates, err = _search_ticker_google_candidates(name)
    if not candidates:
        return "", f"Google 검색 실패: {err}"
    return choose_best_ticker_candidate(name, candidates, "GoogleFinance웹검색", market_preference=market_preference)


def search_ticker_alpha_vantage(company_name: str, api_key: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    key = (api_key or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."
    if not key:
        return "", "Alpha Vantage API Key 없음"

    try:
        resp = requests.get(
            "https://www.alphavantage.co/query",
            params={"function": "SYMBOL_SEARCH", "keywords": name, "apikey": key},
            timeout=12,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        return "", f"Alpha Vantage 검색 실패: {exc}"

    matches = data.get("bestMatches") or []
    if not matches:
        note = str(data.get("Note") or "").strip()
        if note:
            return "", f"Alpha Vantage 제한: {note}"
        return "", "Alpha Vantage 검색 결과가 없습니다."

    candidates = []
    for m in matches:
        symbol = clean_valid_ticker(str(m.get("1. symbol", "")))
        if not symbol:
            continue
        longname = str(m.get("2. name", "")).strip()
        candidates.append(
            {
                "symbol": symbol,
                "name": longname,
                "exchange": str(m.get("4. region", "")).strip(),
                "region": str(m.get("4. region", "")).strip(),
            }
        )
    if not candidates:
        return "", "Alpha Vantage 티커 결과가 없습니다."
    return choose_best_ticker_candidate(company_name, candidates, "Alpha Vantage", market_preference=market_preference)


def search_ticker_finnhub(company_name: str, api_key: str, market_preference: str = "") -> tuple[str, str]:
    name = (company_name or "").strip()
    key = (api_key or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."
    if not key:
        return "", "Finnhub API Key 없음"

    try:
        resp = requests.get(
            "https://finnhub.io/api/v1/search",
            params={"q": name, "token": key},
            timeout=12,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        return "", f"Finnhub 검색 실패: {exc}"

    results = data.get("result") or []
    if not results:
        return "", "Finnhub 검색 결과가 없습니다."

    def type_of(item):
        return str(item.get("type", "")).strip().lower()

    filtered = []
    for r in results:
        symbol = clean_valid_ticker(str(r.get("symbol", "")))
        if not symbol:
            continue
        filtered.append(
            {
                "symbol": symbol,
                "name": str(r.get("description", "")).strip(),
                "type": type_of(r),
                "exchange": str(r.get("mic") or "").strip(),
                "region": str(r.get("currency") or "").strip(),
            }
        )
    if not filtered:
        return "", "Finnhub 티커 결과가 없습니다."
    equity = [r for r in filtered if r.get("type", "") in {"common stock", "equity"}]
    pool = equity if equity else filtered
    return choose_best_ticker_candidate(company_name, pool, "Finnhub", market_preference=market_preference)


def normalize_ai_provider(provider: str) -> str:
    p = (provider or "").strip().lower()
    if p in {"claude", "anthropic"}:
        return "claude"
    if p in {"openai", "gpt"}:
        return "openai"
    return "claude"


def ai_provider_label(provider: str) -> str:
    p = (provider or "").strip().lower()
    if p == "template":
        return "Template"
    return "Claude" if normalize_ai_provider(provider) == "claude" else "OpenAI"


def _extract_openai_output_text(data: dict) -> str:
    text = (data.get("output_text") or "").strip()
    if text:
        return text
    outputs = data.get("output") or []
    chunks = []
    for item in outputs:
        for content in item.get("content", []):
            if content.get("type") == "output_text":
                chunks.append(content.get("text", ""))
    return "\n".join([c for c in chunks if c]).strip()


def call_ai_text(
    provider: str,
    api_key: str,
    model: str,
    system_prompt: str,
    user_prompt: str,
    temperature: float,
    max_output_tokens: int,
    timeout_sec: int = 20,
) -> tuple[str, str]:
    key = (api_key or "").strip()
    if not key:
        return "", "AI API 키를 입력해 주세요."

    normalized = normalize_ai_provider(provider)
    selected_model = (
        model
        or (DEFAULT_CLAUDE_MODEL if normalized == "claude" else DEFAULT_OPENAI_MODEL)
    )
    selected_model = str(selected_model).strip()
    label = ai_provider_label(normalized)

    last_exc = None
    max_attempts = 3
    for attempt in range(max_attempts):
        request_timeout = int(timeout_sec + attempt * 12)
        try:
            if normalized == "claude":
                body = {
                    "model": selected_model,
                    "max_tokens": int(max_output_tokens),
                    "temperature": float(temperature),
                    "system": system_prompt,
                    "messages": [{"role": "user", "content": user_prompt}],
                }
                resp = requests.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": key,
                        "anthropic-version": "2023-06-01",
                        "Content-Type": "application/json",
                    },
                    json=body,
                    timeout=request_timeout,
                )
                resp.raise_for_status()
                data = resp.json()
                pieces = []
                for block in data.get("content", []):
                    if block.get("type") == "text":
                        pieces.append(str(block.get("text", "")))
                text = "\n".join([p for p in pieces if p]).strip()
            else:
                body = {
                    "model": selected_model,
                    "input": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "temperature": float(temperature),
                    "max_output_tokens": int(max_output_tokens),
                }
                resp = requests.post(
                    "https://api.openai.com/v1/responses",
                    headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                    json=body,
                    timeout=request_timeout,
                )
                resp.raise_for_status()
                data = resp.json()
                text = _extract_openai_output_text(data)

            if text:
                return text, ""
            raise RuntimeError(f"{label} 응답이 비어 있습니다.")
        except Exception as exc:
            last_exc = exc
            msg = str(exc).lower()
            retryable = any(token in msg for token in ["timeout", "timed out", "429", "502", "503", "504", "connection"])
            if attempt < (max_attempts - 1) and retryable:
                time.sleep(1.2 * (attempt + 1))
                continue
            break

    return "", f"{label} API 호출 실패: {last_exc}"


def infer_ticker_with_ai(
    company_name: str,
    api_key: str,
    model: str,
    provider: str = "openai",
    market_preference: str = "",
) -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."

    pref = _market_pref_normalized(market_preference)
    market_hint = ""
    if pref == "foreign":
        market_hint = "Market hint: This is likely an overseas company (US listing preferred, then non-Korea listing). "
    elif pref == "domestic":
        market_hint = "Market hint: This is likely a Korea-listed company (.KS or .KQ preferred). "

    user_prompt = (
        "Return the single most likely Yahoo Finance stock ticker symbol for this company. "
        "Output only one ticker symbol, no explanation. "
        "Use .KS or .KQ only when the company is actually listed in Korea. "
        "If it is a US/overseas company, return its primary listing ticker without .KS/.KQ. "
        "If multiple overseas listings exist, prefer a US-listed ticker (NYSE/NASDAQ) when available. "
        "If uncertain, return UNKNOWN. "
        + market_hint
        + f"Company name: {name}"
    )
    system_prompt = "You return only one valid Yahoo Finance ticker symbol."
    text, err = call_ai_text(
        provider=provider,
        api_key=api_key,
        model=model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        temperature=0.0,
        max_output_tokens=40,
        timeout_sec=12,
    )
    if err:
        return "", f"AI 추론 실패: {err}"

    candidates = re.findall(r"[A-Za-z0-9][A-Za-z0-9.^_=\\-]{0,23}", str(text or ""))
    invalid_tokens = {"UNKNOWN", "N/A", "NA", "NONE", "NULL"}
    for cand in candidates:
        if str(cand or "").strip().upper() in invalid_tokens:
            continue
        cleaned = clean_valid_ticker(cand)
        if cleaned:
            return cleaned, f"AI 추론 ({ai_provider_label(provider)})"
    return "", f"AI 응답에서 유효 티커를 찾지 못했습니다: {text}"


def infer_sector_with_ai(
    company_name: str,
    ticker: str,
    api_key: str,
    model: str,
    provider: str = "openai",
) -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명이 비어 있습니다."

    user_prompt = (
        "아래 기업의 산업 섹터를 한 단어 또는 짧은 구로만 답하세요. 설명 금지.\n"
        f"기업명: {name}\n"
        f"티커: {(ticker or '').strip().upper()}"
    )
    system_prompt = "너는 기업 분류기다. 산업 섹터만 간결하게 답한다."
    text, err = call_ai_text(
        provider=provider,
        api_key=api_key,
        model=model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        temperature=0.0,
        max_output_tokens=30,
        timeout_sec=12,
    )
    if err:
        return "", f"AI 섹터 추론 실패: {err}"

    first = str(text or "").strip().splitlines()[0] if text else ""
    first = re.sub(r"^[\\-\\*\\d\\.)\\s]+", "", first).strip()
    first = re.sub(r"[\"'`]", "", first).strip()
    if not first:
        return "", "AI 응답에서 섹터를 찾지 못했습니다."
    return first[:40], f"AI 추론 ({ai_provider_label(provider)})"


@st.cache_data(ttl=60 * 60 * 6, show_spinner=False)
def fetch_sector_from_yahoo_asset_profile(ticker: str) -> tuple[str, str]:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return "", "티커가 비어 있습니다."

    url = f"https://query2.finance.yahoo.com/v10/finance/quoteSummary/{tkr}"
    params = {"modules": "assetProfile"}
    last_err = ""
    for attempt in range(3):
        try:
            resp = requests.get(
                url,
                params=params,
                headers=HTTP_HEADERS_COMMON,
                timeout=12,
            )
            if resp.status_code == 429:
                last_err = "429 Too Many Requests"
                if attempt < 2:
                    time.sleep(0.8 * (attempt + 1))
                    continue
                break
            resp.raise_for_status()
            payload = resp.json() or {}
            result_arr = ((payload.get("quoteSummary") or {}).get("result") or [])
            profile = result_arr[0].get("assetProfile") if result_arr and isinstance(result_arr[0], dict) else {}
            if not isinstance(profile, dict):
                profile = {}
            sector = str(profile.get("sector") or "").strip()
            industry = str(profile.get("industry") or "").strip()
            picked = sector or industry
            if picked:
                return picked, "yahoo_asset_profile"
            return "", "assetProfile에 섹터 정보 없음"
        except Exception as exc:
            last_err = str(exc)
            if attempt < 2:
                time.sleep(0.7 * (attempt + 1))
                continue
            break
    return "", f"yahoo assetProfile 조회 실패: {last_err}"


@st.cache_data(ttl=60 * 60 * 6, show_spinner=False)
def fetch_sector_from_naver_domestic(ticker: str) -> tuple[str, str]:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return "", "티커가 비어 있습니다."

    code = ""
    if re.fullmatch(r"\d{6}\.(KS|KQ)", tkr):
        code = tkr.split(".")[0]
    elif re.fullmatch(r"\d{6}", tkr):
        code = tkr
    if not code:
        return "", "Naver 섹터 조회는 국내 티커(.KS/.KQ)만 지원합니다."

    try:
        resp = requests.get(
            "https://finance.naver.com/item/main.naver",
            params={"code": code},
            headers=HTTP_HEADERS_COMMON,
            timeout=12,
        )
        resp.raise_for_status()
        body = resp.text or ""
    except Exception as exc:
        return "", f"Naver 종목 페이지 조회 실패: {exc}"

    matches = re.findall(r"sise_group_detail\.naver\?type=upjong[^\">']*[\"']>\s*([^<]+)\s*<", body, re.I)
    for raw in matches:
        sector = html.unescape(str(raw or "")).strip()
        sector = re.sub(r"\s+", " ", sector)
        if sector:
            return sector, "naver_upjong"
    return "", "Naver 업종 정보를 찾지 못했습니다."


def resolve_sector_auto(
    company_name: str,
    ticker: str,
    ai_api_key: str = "",
    ai_model: str = "",
    ai_provider: str = "claude",
) -> tuple[str, str]:
    name = (company_name or "").strip()
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        heuristic = infer_sector_from_name_heuristic(name, "")
        return (heuristic, "name_heuristic") if heuristic else ("", "티커 없음")

    # 국내 종목은 네이버 업종을 우선 사용한다.
    if _extract_domestic_code_from_ticker(tkr):
        naver_sector, naver_src = fetch_sector_from_naver_domestic(tkr)
        if naver_sector:
            return naver_sector, naver_src

    fast_sector, fast_src = fetch_sector_from_yahoo_asset_profile(tkr)
    if fast_sector:
        return fast_sector, fast_src

    if not _extract_domestic_code_from_ticker(tkr):
        naver_sector, naver_src = fetch_sector_from_naver_domestic(tkr)
        if naver_sector:
            return naver_sector, naver_src

    fetched_summary, _, fetched_source = fetch_company_financial_summary_multi_source(tkr)
    if isinstance(fetched_summary, dict) and fetched_summary:
        summary_sector = str(fetched_summary.get("sector") or fetched_summary.get("industry") or "").strip()
        if summary_sector:
            return summary_sector, fetched_source or "financial_summary"

    if ai_api_key:
        ai_sector, ai_src = infer_sector_with_ai(
            name,
            ticker=tkr,
            api_key=ai_api_key,
            model=ai_model,
            provider=ai_provider,
        )
        if ai_sector:
            return ai_sector, ai_src

    heuristic = infer_sector_from_name_heuristic(name, tkr)
    if heuristic:
        return heuristic, "name_heuristic"
    return "", "섹터 자동 탐색 실패"


def infer_sector_from_name_heuristic(company_name: str, ticker: str = "") -> str:
    text = f"{str(company_name or '').strip()} {str(ticker or '').strip().upper()}".strip().upper()
    if not text:
        return ""
    if "ETF" in text or text.endswith(".KS") and "KODEX" in text:
        return "ETF"
    if any(token in text for token in ["REIT", "리츠"]):
        return "REIT"
    if any(token in text for token in ["BANK", "은행", "FINANCE", "금융", "증권", "INSURANCE", "보험"]):
        return "Finance"
    if any(token in text for token in ["BIO", "PHARMA", "헬스", "제약", "바이오", "MEDICAL"]):
        return "Healthcare"
    if any(token in text for token in ["SEMICON", "반도체", "SOFTWARE", "CLOUD", "DATA", "TECH"]):
        return "Technology"
    if any(token in text for token in ["SHIP", "조선", "해운", "물류", "MARINE"]):
        return "해운/조선"
    if any(token in text for token in ["STEEL", "철강", "METAL", "금속", "소재", "화학", "CHEM"]):
        return "소재/산업재"
    if any(token in text for token in ["AUTO", "자동차", "EV", "BATTERY", "배터리"]):
        return "자동차/배터리"
    if any(token in text for token in ["CONSTRUCT", "건설", "INFRA", "ENGINEERING", "토목"]):
        return "건설/인프라"
    if any(token in text for token in ["OIL", "GAS", "에너지", "ENERGY", "POWER", "전력"]):
        return "에너지"
    if any(token in text for token in ["CONSUMER", "RETAIL", "유통", "식품", "FOOD", "BEVERAGE"]):
        return "소비재"
    return ""


def resolve_ticker_auto(
    company_name: str,
    use_ai: bool,
    api_key: str,
    model: str,
    provider: str = "openai",
    market_preference: str = "",
) -> tuple[str, str]:
    name = (company_name or "").strip()
    if not name:
        return "", "기업명을 입력해 주세요."
    pref = _market_pref_normalized(market_preference)
    has_hangul_name = bool(re.search(r"[가-힣]", name))

    builtin = get_builtin_ticker_hint(name)
    if builtin and _ticker_matches_market_preference(builtin, pref):
        return builtin, "내장 힌트"

    nv_ticker = ""
    nv_source = ""
    naver_pref = pref or ("domestic" if has_hangul_name else "")
    run_naver_first = has_hangul_name or _market_pref_normalized(pref) == "domestic" or not pref
    if run_naver_first:
        # 한국기업 탐색 정확도를 위해 네이버증권을 우선 시도한다.
        # 시장 선호가 비어 있고 한글 기업명이면 국내를 우선 시도한 뒤, 실패 시 일반 탐색으로 재시도한다.
        nv_ticker, nv_source = search_ticker_naver(name, market_preference=naver_pref)
        if (
            not nv_ticker
            and not pref
            and naver_pref == "domestic"
        ):
            nv_ticker_any, nv_source_any = search_ticker_naver(name, market_preference="")
            if nv_ticker_any:
                nv_ticker, nv_source = nv_ticker_any, nv_source_any
            elif nv_source_any:
                nv_source = " | ".join([m for m in [nv_source, nv_source_any] if m])
        if nv_ticker and _ticker_matches_market_preference(nv_ticker, naver_pref or pref):
            return nv_ticker, nv_source

    saved = clean_valid_ticker(get_saved_ticker_hint(name))
    if saved and _ticker_matches_market_preference(saved, pref):
        return saved, "기존 저장 이력"

    yf_ticker, yf_source = search_ticker_yfinance(name, market_preference=market_preference)
    if yf_ticker and _ticker_matches_market_preference(yf_ticker, pref):
        return yf_ticker, yf_source

    if not run_naver_first:
        nv_ticker, nv_source = search_ticker_naver(name, market_preference=market_preference)
        if nv_ticker and _ticker_matches_market_preference(nv_ticker, pref):
            return nv_ticker, nv_source

    google_ticker, google_source = search_ticker_google(name, market_preference=market_preference)
    if google_ticker and _ticker_matches_market_preference(google_ticker, pref):
        return google_ticker, google_source

    toss_ticker, toss_source = search_ticker_toss(name, market_preference=market_preference)
    if toss_ticker and _ticker_matches_market_preference(toss_ticker, pref):
        return toss_ticker, toss_source

    web_ticker, web_source = search_ticker_web_first(name, market_preference=market_preference)
    if web_ticker and _ticker_matches_market_preference(web_ticker, pref):
        return web_ticker, web_source

    sec_ticker = ""
    sec_source = ""
    if pref != "domestic":
        sec_ticker, sec_source = search_ticker_sec_dataset(name, market_preference=market_preference)
        if sec_ticker and _ticker_matches_market_preference(sec_ticker, "foreign"):
            return sec_ticker, sec_source

    alpha_key, finnhub_key = get_market_data_api_keys()
    alpha_source = ""
    fin_source = ""
    if alpha_key:
        alpha_ticker, alpha_source = search_ticker_alpha_vantage(name, alpha_key, market_preference=market_preference)
        if alpha_ticker and _ticker_matches_market_preference(alpha_ticker, pref):
            return alpha_ticker, alpha_source
    if finnhub_key:
        fin_ticker, fin_source = search_ticker_finnhub(name, finnhub_key, market_preference=market_preference)
        if fin_ticker and _ticker_matches_market_preference(fin_ticker, pref):
            return fin_ticker, fin_source

    ai_source_msg = ""
    yf_429 = "429" in str(yf_source or "")
    force_ai = bool(api_key) and (pref == "foreign" or yf_429)
    if (use_ai or force_ai) and api_key:
        ai_ticker, ai_source = infer_ticker_with_ai(
            name,
            api_key,
            model,
            provider=provider,
            market_preference=market_preference,
        )
        if ai_ticker and _ticker_matches_market_preference(ai_ticker, pref):
            if use_ai:
                return ai_ticker, ai_source
            return ai_ticker, f"{ai_source} (자동 보조)"
        ai_source_msg = ai_source or "AI 시장선호 조건 불일치"

    fallback_msgs = [
        msg
        for msg in [
            ai_source_msg,
            nv_source,
            yf_source,
            google_source,
            toss_source,
            web_source,
            sec_source,
            alpha_source,
            fin_source,
        ]
        if msg
    ]
    return "", " | ".join(fallback_msgs) if fallback_msgs else "티커 자동 탐색에 실패했습니다."


def resolve_ticker_auto_with_retry(
    company_name: str,
    use_ai: bool,
    api_key: str,
    model: str,
    provider: str = "openai",
    market_preference: str = "",
) -> tuple[str, str]:
    name = (company_name or "").strip()
    has_hangul_name = _company_name_has_hangul(name)
    foreign_name_hint = _looks_explicit_foreign_company_name(name)

    ticker, source = resolve_ticker_auto(
        company_name=company_name,
        use_ai=use_ai,
        api_key=api_key,
        model=model,
        provider=provider,
        market_preference=market_preference,
    )
    ticker = clean_valid_ticker(ticker)
    pref = _market_pref_normalized(market_preference)

    if ticker and pref == "" and has_hangul_name and not _is_non_kr_ticker_plausible_for_name(name, ticker):
        source = f"{source} | 한글명 대비 해외티커 신뢰도 낮음"
        ticker = ""

    is_kr = ticker.endswith(".KS") or ticker.endswith(".KQ")

    if not ticker and pref == "" and has_hangul_name and not foreign_name_hint:
        retry_kr, retry_kr_source = resolve_ticker_auto(
            company_name=company_name,
            use_ai=False,
            api_key=api_key,
            model=model,
            provider=provider,
            market_preference="domestic",
        )
        retry_kr = clean_valid_ticker(retry_kr)
        if retry_kr.endswith(".KS") or retry_kr.endswith(".KQ"):
            return retry_kr, f"{retry_kr_source} (국내 우선 재시도)"

    # 해외 우선이 명확하거나, 한글명이어도 국내 힌트가 없는 경우 해외 재시도를 허용한다.
    has_domestic_hint = _looks_domestic_company_name_hint(name)
    need_retry_foreign = (pref == "foreign" and (not ticker or is_kr)) or (
        (not ticker)
        and (
            (not has_hangul_name)
            or foreign_name_hint
            or (has_hangul_name and not has_domestic_hint)
        )
    )
    if need_retry_foreign:
        retry_ticker, retry_source = resolve_ticker_auto(
            company_name=company_name,
            use_ai=use_ai,
            api_key=api_key,
            model=model,
            provider=provider,
            market_preference="foreign",
        )
        retry_ticker = clean_valid_ticker(retry_ticker)
        retry_is_kr = str(retry_ticker or "").endswith(".KS") or str(retry_ticker or "").endswith(".KQ")
        if retry_ticker and not retry_is_kr:
            return retry_ticker, f"{retry_source} (해외 우선 재시도)"

    return ticker, source


def _fetch_currency_to_krw_rate(currency: str, rate_date: date) -> tuple[float | None, str]:
    curr = str(currency or "").strip().upper()
    if not curr:
        return None, "통화 정보 없음"
    if curr == "KRW":
        return 1.0, "KRW"
    if curr == "USD":
        rate, src = get_usd_krw_rate_for_date(rate_date)
        return float(rate), f"USD/KRW:{src}"

    fx_ticker_map = {
        "EUR": "EURKRW=X",
        "JPY": "JPYKRW=X",
        "CNY": "CNYKRW=X",
        "GBP": "GBPKRW=X",
        "AUD": "AUDKRW=X",
        "CAD": "CADKRW=X",
        "CHF": "CHFKRW=X",
    }
    fx_ticker = fx_ticker_map.get(curr, "")
    if not fx_ticker:
        return None, f"지원되지 않는 통화: {curr}"

    try:
        import yfinance as yf
    except Exception:
        return None, "yfinance 미설치"

    start = (rate_date - timedelta(days=10)).isoformat()
    end = (rate_date + timedelta(days=1)).isoformat()
    try:
        hist = yf.download(fx_ticker, start=start, end=end, interval="1d", progress=False, auto_adjust=False)
    except Exception as exc:
        return None, f"{fx_ticker} 조회 실패: {exc}"
    if hist is None or hist.empty or "Close" not in hist.columns:
        return None, f"{fx_ticker} 환율 데이터 없음"

    close = hist["Close"]
    if isinstance(close, pd.DataFrame):
        close = close.iloc[:, 0]
    close = close.dropna()
    if close.empty:
        return None, f"{fx_ticker} 종가 데이터 없음"

    idx = pd.to_datetime(close.index)
    if getattr(idx, "tz", None) is not None:
        idx = idx.tz_convert(None)
    close.index = idx
    close = close[close.index.date <= rate_date]
    if close.empty:
        return None, f"{fx_ticker} 해당일 이전 데이터 없음"

    market_dt = close.index[-1].date().isoformat()
    return float(close.iloc[-1]), f"{fx_ticker}:{market_dt}"


def fetch_current_price_krw_from_ticker(ticker: str, rate_date: date | None = None) -> tuple[float | None, str]:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return None, "티커 없음"
    target_date = rate_date or date.today()

    try:
        import yfinance as yf
    except Exception:
        return None, "yfinance 미설치"

    obj = yf.Ticker(tkr)
    price_native = None
    currency = ""
    source_label = ""

    try:
        fast = obj.fast_info or {}
    except Exception:
        fast = {}

    if isinstance(fast, dict):
        for key in ["lastPrice", "regularMarketPrice", "previousClose", "last_price"]:
            val = _safe_to_float(fast.get(key))
            if val is not None and val > 0:
                price_native = float(val)
                source_label = f"yfinance.fast_info.{key}"
                break
        currency = str(fast.get("currency") or "").strip().upper()

    if price_native is None:
        try:
            hist = obj.history(period="5d", interval="1d", auto_adjust=False)
        except Exception:
            hist = pd.DataFrame()
        if hist is not None and not hist.empty and "Close" in hist.columns:
            close = pd.to_numeric(hist["Close"], errors="coerce").dropna()
            if not close.empty:
                price_native = float(close.iloc[-1])
                source_label = "yfinance.history.close"

    if not currency:
        try:
            info = obj.info or {}
        except Exception:
            info = {}
        currency = str(info.get("currency") or "").strip().upper()

    if not currency:
        if tkr.endswith(".KS") or tkr.endswith(".KQ"):
            currency = "KRW"
        else:
            currency = "USD"

    if price_native is None or price_native <= 0:
        return None, "현재 주가를 찾지 못했습니다."

    fx_rate, fx_src = _fetch_currency_to_krw_rate(currency, target_date)
    if fx_rate is None or fx_rate <= 0:
        return None, f"환산 실패({currency}): {fx_src}"

    price_krw = float(price_native) * float(fx_rate)
    return price_krw, f"{source_label}/{currency}→KRW({fx_src})"


def _safe_to_float(value) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        try:
            return float(value)
        except Exception:
            return None

    text = str(value).strip()
    if not text:
        return None
    text = (
        text.replace(",", "")
        .replace("원", "")
        .replace("주", "")
        .replace("%", "")
        .replace("KRW", "")
        .replace("USD", "")
        .replace("$", "")
    )
    text = re.sub(r"[^0-9.+\-]", "", text)
    if text in {"", "+", "-", ".", "+.", "-."}:
        return None
    try:
        return float(text)
    except Exception:
        return None


def _safe_parse_date(value) -> date | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    try:
        parsed = pd.to_datetime(text, errors="coerce")
    except Exception:
        return None
    if pd.isna(parsed):
        return None
    return parsed.date()


def _normalize_currency(value) -> str:
    text = str(value or "").strip().upper()
    if text in {"USD", "US$", "$"}:
        return "USD"
    if "USD" in text or "$" in text:
        return "USD"
    return "KRW"


def _classify_cash_bucket(name_text: str, currency_text: str = "") -> str:
    name = str(name_text or "").strip().upper()
    curr = _normalize_currency(currency_text)
    compact = re.sub(r"\s+", "", name)
    if not compact:
        return ""
    if any(token in compact for token in ["USD", "US$", "$", "달러", "외화"]):
        return "USD"
    if any(token in compact for token in ["KRW", "원화"]):
        return "KRW"
    if compact in {"예수금", "현금", "CASH", "MONEY"}:
        return curr if curr in {"KRW", "USD"} else "KRW"
    return ""


def extract_cash_from_ai_payload(parsed_payload: dict) -> tuple[float | None, float | None]:
    if not isinstance(parsed_payload, dict):
        return None, None

    cash_krw = None
    cash_usd = None

    def _pick_num(source: dict, keys: list[str]) -> float | None:
        if not isinstance(source, dict):
            return None
        for key in keys:
            if key in source:
                val = _safe_to_float(source.get(key))
                if val is not None:
                    return float(val)
        return None

    cash_krw = _pick_num(parsed_payload, ["cash_krw", "krw_cash", "원화예수금", "원화", "krw"])
    cash_usd = _pick_num(parsed_payload, ["cash_usd", "usd_cash", "달러예수금", "달러", "usd"])

    cash_obj = parsed_payload.get("cash")
    if isinstance(cash_obj, dict):
        if cash_krw is None:
            cash_krw = _pick_num(cash_obj, ["krw", "KRW", "원화", "원화예수금", "cash_krw", "krw_cash"])
        if cash_usd is None:
            cash_usd = _pick_num(cash_obj, ["usd", "USD", "달러", "달러예수금", "cash_usd", "usd_cash"])

    balances_obj = parsed_payload.get("balances")
    if isinstance(balances_obj, dict):
        if cash_krw is None:
            cash_krw = _pick_num(balances_obj, ["krw", "KRW", "원화", "원화예수금", "cash_krw", "krw_cash"])
        if cash_usd is None:
            cash_usd = _pick_num(balances_obj, ["usd", "USD", "달러", "달러예수금", "cash_usd", "usd_cash"])

    holdings = parsed_payload.get("holdings")
    if isinstance(holdings, list):
        for row in holdings:
            if not isinstance(row, dict):
                continue
            nm = str(row.get("stock_name") or row.get("name") or row.get("기업명") or "").strip()
            bucket = _classify_cash_bucket(nm, str(row.get("currency") or ""))
            if not bucket:
                continue
            amount = _safe_to_float(row.get("market_value"))
            if amount is None:
                amount = _safe_to_float(row.get("value"))
            if amount is None:
                continue
            if bucket == "USD":
                cash_usd = float(amount)
            else:
                cash_krw = float(amount)

    return cash_krw, cash_usd


def build_holdings_df_from_ai_rows(rows: list[dict], usd_krw_rate: float) -> pd.DataFrame:
    def _pick_num(item: dict, keys: list[str]) -> float | None:
        for k in keys:
            if k in item:
                v = _safe_to_float(item.get(k))
                if v is not None:
                    return float(v)
        return None

    normalized_rows = []
    for item in rows or []:
        if not isinstance(item, dict):
            continue
        stock_name = str(item.get("stock_name") or item.get("name") or item.get("기업명") or "").strip()
        if not stock_name:
            continue
        if _classify_cash_bucket(stock_name, str(item.get("currency") or "")):
            continue

        quantity = _pick_num(item, ["quantity", "qty", "shares", "보유수량", "수량", "보유주수", "주수"])
        market_value = _pick_num(item, ["market_value", "value", "평가금액", "평가금액(원화)", "평가액"])
        pnl_value = _pick_num(item, ["pnl_value", "pnl", "손익금액", "손익", "손익금액(원화)"])
        return_pct = _pick_num(item, ["return_pct", "return", "수익률", "수익률(%)", "pnl_pct"])
        price_krw = _pick_num(item, ["price_krw", "current_price", "price", "주가", "주가(원)", "현재가"])
        invest_won = _pick_num(item, ["invest_won", "principal", "cost", "투자금액", "투자금액(원)", "원금"])

        currency = _normalize_currency(item.get("currency"))
        if market_value is None and quantity is not None and quantity > 0 and price_krw is not None and price_krw > 0:
            market_value = float(quantity) * float(price_krw)
        if invest_won is None and market_value is not None and pnl_value is not None:
            invest_won = float(market_value) - float(pnl_value)
        if pnl_value is None and market_value is not None and invest_won is not None:
            pnl_value = float(market_value) - float(invest_won)
        if return_pct is None and market_value is not None and pnl_value is not None:
            principal = float(invest_won) if invest_won is not None else (float(market_value) - float(pnl_value))
            return_pct = (float(pnl_value) / principal * 100.0) if principal != 0 else 0.0
        if pnl_value is None and return_pct is not None and market_value is not None:
            denom = 1.0 + float(return_pct) / 100.0
            if denom != 0:
                principal = float(market_value) / denom
                pnl_value = float(market_value) - principal

        normalized_rows.append(
            {
                COL_NAME: stock_name,
                COL_QTY: float(quantity) if quantity is not None else 0.0,
                COL_CURRENCY: currency,
                COL_FX_RATE: float(usd_krw_rate) if currency == "USD" else 1.0,
                COL_VALUE: float(market_value) if market_value is not None else 0.0,
                COL_PNL: float(pnl_value) if pnl_value is not None else 0.0,
                COL_RETURN: float(return_pct) if return_pct is not None else 0.0,
            }
        )

    if not normalized_rows:
        return pd.DataFrame(columns=COLUMNS)

    df = pd.DataFrame(normalized_rows)
    df[COL_NAME] = df[COL_NAME].astype(str).str.strip()
    df = df[df[COL_NAME] != ""]
    df = df.drop_duplicates(subset=[COL_NAME], keep="last")
    df = ensure_portfolio_columns(df, usd_krw_rate, force_usd_rate=True)
    return ensure_numeric(df, usd_krw_rate)


def merge_holdings_overwrite(base_df: pd.DataFrame, incoming_df: pd.DataFrame, usd_krw_rate: float) -> pd.DataFrame:
    base = ensure_portfolio_columns(base_df, usd_krw_rate, force_usd_rate=True).copy()
    incoming = ensure_portfolio_columns(incoming_df, usd_krw_rate, force_usd_rate=True).copy()
    if incoming.empty:
        return ensure_numeric(base, usd_krw_rate) if not base.empty else incoming
    if base.empty:
        return ensure_numeric(incoming, usd_krw_rate)

    base[COL_NAME] = base[COL_NAME].astype(str).str.strip()
    incoming[COL_NAME] = incoming[COL_NAME].astype(str).str.strip()
    incoming_names = set(incoming[COL_NAME].tolist())
    kept_base = base[~base[COL_NAME].isin(incoming_names)]
    merged = pd.concat([kept_base, incoming], ignore_index=True)
    merged = merged.drop_duplicates(subset=[COL_NAME], keep="last")
    return ensure_numeric(merged, usd_krw_rate)


def extract_holdings_from_image_with_ai(
    image_bytes: bytes,
    mime_type: str,
    provider: str,
    api_key: str,
    model: str,
) -> tuple[dict, str]:
    key = (api_key or "").strip()
    if not key:
        return {}, "API 설정 탭에서 AI API 키를 먼저 저장해 주세요."
    if not image_bytes:
        return {}, "이미지 데이터가 비어 있습니다."

    normalized_provider = normalize_ai_provider(provider)
    selected_model = (
        str(model).strip()
        or (DEFAULT_CLAUDE_MODEL if normalized_provider == "claude" else DEFAULT_OPENAI_MODEL)
    )
    media_type = str(mime_type or "image/png").strip().lower()
    if not media_type.startswith("image/"):
        media_type = "image/png"

    system_prompt = (
        "너는 포트폴리오 이미지 OCR/구조화 엔진이다. "
        "이미지에서 보이는 종목 보유 정보를 정확히 읽고 JSON만 출력한다."
    )
    user_prompt = """
다음 형식 JSON만 출력:
{
  "as_of_date": "YYYY-MM-DD 또는 빈문자열",
  "cash_krw": 숫자 또는 null,
  "cash_usd": 숫자 또는 null,
  "holdings": [
    {
      "stock_name": "종목명",
      "ticker": "티커(모르면 빈문자열)",
      "sector": "산업섹터(모르면 빈문자열)",
      "quantity": 숫자,
      "market_value": 숫자,
      "pnl_value": 숫자,
      "return_pct": 숫자,
      "currency": "KRW 또는 USD"
    }
  ]
}

규칙:
- 화면에 원화/달러 예수금(또는 현금, cash balance)이 보이면 cash_krw/cash_usd에 반드시 넣기
- 보이는 행마다 holdings 배열에 추가
- 숫자는 쉼표/원/주/% 제거 후 숫자만
- 손실은 음수
- 통화 표기가 없으면 KRW
- 설명 문장 없이 JSON만
""".strip()

    image_b64 = base64.b64encode(image_bytes).decode("utf-8")

    try:
        if normalized_provider == "claude":
            body = {
                "model": selected_model,
                "max_tokens": 2200,
                "temperature": 0.0,
                "system": system_prompt,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {"type": "base64", "media_type": media_type, "data": image_b64},
                            },
                            {"type": "text", "text": user_prompt},
                        ],
                    }
                ],
            }
            resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                },
                json=body,
                timeout=35,
            )
            resp.raise_for_status()
            data = resp.json()
            pieces = []
            for block in data.get("content", []):
                if block.get("type") == "text":
                    pieces.append(str(block.get("text", "")))
            text = "\n".join([p for p in pieces if p]).strip()
        else:
            data_url = f"data:{media_type};base64,{image_b64}"
            body = {
                "model": selected_model,
                "input": [
                    {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": user_prompt},
                            {"type": "input_image", "image_url": data_url},
                        ],
                    },
                ],
                "temperature": 0.0,
                "max_output_tokens": 2200,
            }
            resp = requests.post(
                "https://api.openai.com/v1/responses",
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json=body,
                timeout=35,
            )
            resp.raise_for_status()
            data = resp.json()
            text = _extract_openai_output_text(data)
    except Exception as exc:
        return {}, f"{ai_provider_label(normalized_provider)} 이미지 분석 실패: {exc}"

    if not text:
        return {}, "AI 응답이 비어 있습니다."

    parsed = _extract_json_from_text(text)
    if not parsed:
        return {}, f"AI 응답에서 JSON을 파싱하지 못했습니다. 원인: {_json_parse_failure_reason(text)}"
    if isinstance(parsed, list):
        parsed = {"as_of_date": "", "holdings": parsed}
    return parsed, ""


def extract_company_watchlist_from_image_with_ai(
    image_bytes: bytes,
    mime_type: str,
    provider: str,
    api_key: str,
    model: str,
) -> tuple[list[dict], str]:
    key = (api_key or "").strip()
    if not key:
        return [], "API 설정 탭 또는 기업정보의 AI 설정에 API 키를 먼저 입력해 주세요."
    if not image_bytes:
        return [], "이미지 데이터가 비어 있습니다."

    normalized_provider = normalize_ai_provider(provider)
    selected_model = (
        str(model).strip()
        or (DEFAULT_CLAUDE_MODEL if normalized_provider == "claude" else DEFAULT_OPENAI_MODEL)
    )
    media_type = str(mime_type or "image/png").strip().lower()
    if not media_type.startswith("image/"):
        media_type = "image/png"

    system_prompt = (
        "너는 기업명/티커/섹터 목록 OCR 구조화 엔진이다. "
        "보이는 기업 리스트를 정확히 읽고 JSON만 출력한다."
    )
    user_prompt = """
다음 형식 JSON만 출력:
{
  "companies": [
    {
      "stock_name": "기업명",
      "ticker": "야후 티커(없으면 빈문자열)",
      "sector": "산업섹터(없으면 빈문자열)"
    }
  ]
}

규칙:
- 관심종목/기업리스트 영역에서 보이는 기업행만 추출
- 숫자(수량/평가금/손익)는 무시
- 기업명은 최대한 정확히, 티커는 대문자/점(.) 포함 원문 유지
- 확실하지 않은 값은 빈문자열
- 설명 문장 없이 JSON만
""".strip()

    image_b64 = base64.b64encode(image_bytes).decode("utf-8")

    try:
        if normalized_provider == "claude":
            body = {
                "model": selected_model,
                "max_tokens": 2000,
                "temperature": 0.0,
                "system": system_prompt,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {"type": "base64", "media_type": media_type, "data": image_b64},
                            },
                            {"type": "text", "text": user_prompt},
                        ],
                    }
                ],
            }
            resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                },
                json=body,
                timeout=35,
            )
            resp.raise_for_status()
            data = resp.json()
            pieces = []
            for block in data.get("content", []):
                if block.get("type") == "text":
                    pieces.append(str(block.get("text", "")))
            text = "\n".join([p for p in pieces if p]).strip()
        else:
            data_url = f"data:{media_type};base64,{image_b64}"
            body = {
                "model": selected_model,
                "input": [
                    {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": user_prompt},
                            {"type": "input_image", "image_url": data_url},
                        ],
                    },
                ],
                "temperature": 0.0,
                "max_output_tokens": 2000,
            }
            resp = requests.post(
                "https://api.openai.com/v1/responses",
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json=body,
                timeout=35,
            )
            resp.raise_for_status()
            data = resp.json()
            text = _extract_openai_output_text(data)
    except Exception as exc:
        return [], f"{ai_provider_label(normalized_provider)} 이미지 분석 실패: {exc}"

    if not text:
        return [], "AI 응답이 비어 있습니다."

    parsed = _extract_json_from_text(text)
    if not parsed:
        return [], f"AI 응답에서 JSON을 파싱하지 못했습니다. 원인: {_json_parse_failure_reason(text)}"

    raw_items = []
    if isinstance(parsed, list):
        raw_items = parsed
    elif isinstance(parsed, dict):
        for key_name in ["companies", "watchlist", "company_list", "items", "rows", "holdings"]:
            candidate = parsed.get(key_name)
            if isinstance(candidate, list):
                raw_items = candidate
                break
        if not raw_items:
            single = parsed.get("company")
            if isinstance(single, dict):
                raw_items = [single]

    normalized_rows = []
    seen_names = set()
    for item in raw_items:
        if not isinstance(item, dict):
            continue
        stock_name = str(
            item.get("stock_name")
            or item.get("company_name")
            or item.get("name")
            or item.get("기업명")
            or item.get("종목명")
            or ""
        ).strip()
        if not stock_name:
            continue
        name_key = normalize_company_name_for_match(stock_name)
        if not name_key or name_key in seen_names:
            continue
        seen_names.add(name_key)
        ticker = clean_valid_ticker(
            str(item.get("ticker") or item.get("symbol") or item.get("티커") or "").strip()
        )
        sector = str(item.get("sector") or item.get("industry") or item.get("산업섹터") or "").strip()
        normalized_rows.append({"stock_name": stock_name, "ticker": ticker, "sector": sector})

    return normalized_rows, ""


def _normalize_value_chain_stage_name(value: str) -> str:
    text = str(value or "").strip()
    low = text.lower()
    if ("업스트림" in text) or ("upstream" in low):
        return "업스트림"
    if ("미드스트림" in text) or ("midstream" in low):
        return "미드스트림"
    if ("다운스트림" in text) or ("downstream" in low):
        return "다운스트림"
    return text or "기타"


def _normalize_value_chain_payload(parsed) -> dict:
    chain_name = ""
    rows: list[dict] = []

    def _to_company_list(value) -> list[str]:
        if isinstance(value, list):
            out = []
            for item in value:
                nm = str(item or "").strip()
                if nm:
                    out.append(nm)
            return out
        if isinstance(value, str):
            parts = re.split(r"[,/·\n|]", value)
            return [p.strip() for p in parts if p.strip()]
        return []

    def _append_row(stage: str, segment: str, companies) -> None:
        stg = _normalize_value_chain_stage_name(stage)
        seg = str(segment or "").strip() or "기타"
        comps = []
        seen = set()
        for nm in _to_company_list(companies):
            key = normalize_company_name_for_match(nm)
            if not key or key in seen:
                continue
            seen.add(key)
            comps.append(nm)
        if comps:
            rows.append({"stage": stg, "segment": seg, "companies": comps})

    if isinstance(parsed, dict):
        chain_name = str(parsed.get("chain_name") or parsed.get("value_chain_name") or parsed.get("name") or "").strip()
        stage_items = parsed.get("stages")
        if isinstance(stage_items, list):
            for item in stage_items:
                if not isinstance(item, dict):
                    continue
                stage = item.get("stage") or item.get("stage_name") or item.get("단계")
                segment = item.get("segment") or item.get("sub_stage") or item.get("세부단계") or item.get("category")
                if isinstance(item.get("segments"), list):
                    for seg in item.get("segments"):
                        if not isinstance(seg, dict):
                            continue
                        _append_row(
                            stage=stage,
                            segment=seg.get("segment") or seg.get("name") or seg.get("세부단계") or segment,
                            companies=seg.get("companies") or seg.get("기업들") or [],
                        )
                else:
                    _append_row(
                        stage=stage,
                        segment=segment,
                        companies=item.get("companies") or item.get("company_names") or item.get("기업들") or [],
                    )

        stage_key_map = {"upstream": "업스트림", "midstream": "미드스트림", "downstream": "다운스트림"}
        for raw_key, stage_label in stage_key_map.items():
            bucket = parsed.get(raw_key)
            if not bucket:
                continue
            if isinstance(bucket, dict):
                for seg_name, seg_companies in bucket.items():
                    _append_row(stage_label, str(seg_name), seg_companies)
            elif isinstance(bucket, list):
                for item in bucket:
                    if isinstance(item, dict):
                        _append_row(
                            stage_label,
                            item.get("segment") or item.get("name") or item.get("세부단계") or "기타",
                            item.get("companies") or item.get("기업들") or [],
                        )
                    elif isinstance(item, str):
                        _append_row(stage_label, "기타", [item])
            elif isinstance(bucket, str):
                _append_row(stage_label, "기타", bucket)

    elif isinstance(parsed, list):
        for item in parsed:
            if not isinstance(item, dict):
                continue
            _append_row(
                stage=item.get("stage") or item.get("단계") or "기타",
                segment=item.get("segment") or item.get("세부단계") or item.get("category") or "기타",
                companies=item.get("companies") or item.get("기업들") or item.get("company") or [],
            )

    merged: dict[tuple[str, str], list[str]] = {}
    for row in rows:
        key = (str(row.get("stage") or "기타"), str(row.get("segment") or "기타"))
        if key not in merged:
            merged[key] = []
        existing_norm = {normalize_company_name_for_match(v) for v in merged[key]}
        for nm in row.get("companies") or []:
            nkey = normalize_company_name_for_match(nm)
            if not nkey or nkey in existing_norm:
                continue
            merged[key].append(str(nm).strip())
            existing_norm.add(nkey)

    norm_rows = [{"stage": k[0], "segment": k[1], "companies": v} for k, v in merged.items() if v]
    chain_name = chain_name or "밸류체인"
    return {"chain_name": chain_name, "rows": norm_rows}


def extract_value_chain_from_image_with_ai(
    image_bytes: bytes,
    mime_type: str,
    provider: str,
    api_key: str,
    model: str,
) -> tuple[dict, str]:
    key = (api_key or "").strip()
    if not key:
        return {}, "API 설정 탭에서 AI API 키를 먼저 저장해 주세요."
    if not image_bytes:
        return {}, "이미지 데이터가 비어 있습니다."

    normalized_provider = normalize_ai_provider(provider)
    selected_model = (
        str(model).strip()
        or (DEFAULT_CLAUDE_MODEL if normalized_provider == "claude" else DEFAULT_OPENAI_MODEL)
    )
    media_type = str(mime_type or "image/png").strip().lower()
    if not media_type.startswith("image/"):
        media_type = "image/png"

    system_prompt = (
        "너는 산업 밸류체인 이미지 OCR/구조화 엔진이다. "
        "이미지의 단계-세부공정-기업명을 읽고 JSON만 출력한다."
    )
    user_prompt = """
다음 형식 JSON만 출력:
{
  "chain_name": "예: LNG 밸류체인",
  "stages": [
    {
      "stage": "업스트림/미드스트림/다운스트림",
      "segment": "세부 공정명(예: 운반, 피팅밸브, 강관)",
      "companies": ["기업명1", "기업명2"]
    }
  ]
}

규칙:
- 이미지에서 읽은 텍스트만 사용하고 추측 금지
- 기업명은 가능한 정확히
- 단계/세부공정이 불명확하면 "기타" 사용
- 설명문 없이 JSON만 출력
""".strip()

    image_b64 = base64.b64encode(image_bytes).decode("utf-8")
    try:
        if normalized_provider == "claude":
            body = {
                "model": selected_model,
                "max_tokens": 2200,
                "temperature": 0.0,
                "system": system_prompt,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {"type": "base64", "media_type": media_type, "data": image_b64},
                            },
                            {"type": "text", "text": user_prompt},
                        ],
                    }
                ],
            }
            resp = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json",
                },
                json=body,
                timeout=35,
            )
            resp.raise_for_status()
            data = resp.json()
            pieces = []
            for block in data.get("content", []):
                if block.get("type") == "text":
                    pieces.append(str(block.get("text", "")))
            text = "\n".join([p for p in pieces if p]).strip()
        else:
            data_url = f"data:{media_type};base64,{image_b64}"
            body = {
                "model": selected_model,
                "input": [
                    {"role": "system", "content": [{"type": "input_text", "text": system_prompt}]},
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": user_prompt},
                            {"type": "input_image", "image_url": data_url},
                        ],
                    },
                ],
                "temperature": 0.0,
                "max_output_tokens": 2200,
            }
            resp = requests.post(
                "https://api.openai.com/v1/responses",
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json=body,
                timeout=35,
            )
            resp.raise_for_status()
            data = resp.json()
            text = _extract_openai_output_text(data)
    except Exception as exc:
        return {}, f"{ai_provider_label(normalized_provider)} 이미지 분석 실패: {exc}"

    if not text:
        return {}, "AI 응답이 비어 있습니다."
    parsed = _extract_json_from_text(text)
    if not parsed:
        return {}, f"AI 응답에서 JSON을 파싱하지 못했습니다. 원인: {_json_parse_failure_reason(text)}"
    normalized = _normalize_value_chain_payload(parsed)
    if not normalized.get("rows"):
        return {}, "이미지에서 밸류체인 단계/기업 정보를 충분히 읽지 못했습니다."
    return normalized, ""


def _match_value_chain_company_name(
    raw_name: str,
    watch_names: list[str],
    watch_ticker_map: dict[str, str] | None = None,
) -> tuple[str, float]:
    name = str(raw_name or "").strip()
    if not name:
        return "", 0.0
    ticker_map = watch_ticker_map or {}
    nkey = normalize_company_name_for_match(name)

    # 티커 단독 입력인 경우 우선 매칭
    maybe_ticker = clean_valid_ticker(name)
    if maybe_ticker:
        for wn, wt in ticker_map.items():
            if clean_valid_ticker(wt) == maybe_ticker:
                return wn, 1.0

    if nkey:
        exact = []
        for wn in watch_names:
            if normalize_company_name_for_match(wn) == nkey:
                exact.append(wn)
        if len(exact) == 1:
            return exact[0], 1.0
        if len(exact) > 1:
            return exact[0], 0.95

    best_name = ""
    best_score = 0.0
    for wn in watch_names:
        sim = _name_similarity(name, wn)
        if nkey and normalize_company_name_for_match(wn).find(nkey) >= 0:
            sim = max(sim, 0.9)
        if sim > best_score:
            best_score = float(sim)
            best_name = wn
    if best_score >= 0.74:
        return best_name, best_score
    return "", best_score


def _build_value_chain_match_rows(chain_payload: dict, company_list_df: pd.DataFrame) -> list[dict]:
    rows = chain_payload.get("rows") if isinstance(chain_payload, dict) else []
    if not isinstance(rows, list):
        return []
    cl_df = company_list_df if isinstance(company_list_df, pd.DataFrame) else pd.DataFrame()
    watch_names = cl_df["stock_name"].dropna().astype(str).str.strip().tolist() if not cl_df.empty else []
    watch_names = [v for v in watch_names if v]
    watch_ticker_map = {}
    watch_sector_map = {}
    if not cl_df.empty:
        for _, rr in cl_df.iterrows():
            nm = str(rr.get("stock_name") or "").strip()
            if not nm:
                continue
            watch_ticker_map[nm] = clean_valid_ticker(str(rr.get("ticker") or ""))
            watch_sector_map[nm] = str(rr.get("sector") or "").strip()

    matched_rows: list[dict] = []
    for row in rows:
        stage = str(row.get("stage") or "기타").strip() or "기타"
        segment = str(row.get("segment") or "기타").strip() or "기타"
        companies = row.get("companies") if isinstance(row.get("companies"), list) else []
        for raw_name in companies:
            name = str(raw_name or "").strip()
            if not name:
                continue
            matched_name, score = _match_value_chain_company_name(name, watch_names, watch_ticker_map=watch_ticker_map)
            matched_rows.append(
                {
                    "stage": _normalize_value_chain_stage_name(stage),
                    "segment": segment,
                    "input_company": name,
                    "matched_company": matched_name,
                    "matched": bool(matched_name),
                    "match_score": float(score),
                    "ticker": watch_ticker_map.get(matched_name, "") if matched_name else "",
                    "sector": watch_sector_map.get(matched_name, "") if matched_name else "",
                }
            )
    return matched_rows


def _build_value_chain_sankey_figure(match_rows: list[dict], chain_name: str):
    if not match_rows:
        return None
    df = pd.DataFrame(match_rows)
    if df.empty:
        return None

    stage_order = ["업스트림", "미드스트림", "다운스트림"]
    df["stage_order"] = df["stage"].apply(lambda s: stage_order.index(s) if s in stage_order else 99)
    df = df.sort_values(["stage_order", "segment", "input_company"]).copy()

    nodes: list[str] = []
    node_index: dict[str, int] = {}

    def _idx(name: str) -> int:
        if name not in node_index:
            node_index[name] = len(nodes)
            nodes.append(name)
        return node_index[name]

    links_source = []
    links_target = []
    links_value = []
    links_color = []

    stage_color = {"업스트림": "#2563eb", "미드스트림": "#0f766e", "다운스트림": "#ea580c", "기타": "#64748b"}

    stage_seg_counts = df.groupby(["stage", "segment"]).size().reset_index(name="cnt")
    for _, r in stage_seg_counts.iterrows():
        s = str(r["stage"])
        seg = f"{s} | {str(r['segment'])}"
        links_source.append(_idx(s))
        links_target.append(_idx(seg))
        links_value.append(float(r["cnt"]))
        links_color.append("rgba(148,163,184,0.35)")

    seg_company_counts = df.groupby(["stage", "segment", "input_company", "matched"]).size().reset_index(name="cnt")
    for _, r in seg_company_counts.iterrows():
        s = str(r["stage"])
        seg = f"{s} | {str(r['segment'])}"
        cname = str(r["input_company"])
        label = f"{cname} {'(매칭)' if bool(r['matched']) else '(미매칭)'}"
        links_source.append(_idx(seg))
        links_target.append(_idx(label))
        links_value.append(float(r["cnt"]))
        links_color.append("rgba(16,185,129,0.35)" if bool(r["matched"]) else "rgba(239,68,68,0.30)")

    node_colors = []
    for n in nodes:
        if n in stage_color:
            node_colors.append(stage_color.get(n, "#64748b"))
        elif " | " in n:
            stage_name = n.split(" | ", 1)[0]
            base = stage_color.get(stage_name, "#64748b")
            node_colors.append(base)
        elif "(매칭)" in n:
            node_colors.append("#16a34a")
        else:
            node_colors.append("#dc2626")

    fig = go.Figure(
        data=[
            go.Sankey(
                arrangement="snap",
                node=dict(label=nodes, pad=15, thickness=16, color=node_colors),
                link=dict(source=links_source, target=links_target, value=links_value, color=links_color),
            )
        ]
    )
    fig.update_layout(
        title=f"{chain_name} 매칭 흐름",
        font=dict(size=11, family="Noto Sans KR"),
        paper_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=10, r=10, t=45, b=10),
    )
    return fig


def _build_value_chain_rows_brief_text(match_rows: list[dict], max_stage_segments: int = 24) -> str:
    if not match_rows:
        return ""
    df = pd.DataFrame(match_rows)
    if df.empty:
        return ""

    stage_order = {"업스트림": 0, "미드스트림": 1, "다운스트림": 2}
    df["stage_order"] = df["stage"].map(stage_order).fillna(99)
    grouped = (
        df.groupby(["stage", "segment"], dropna=False)["input_company"]
        .apply(lambda s: sorted({str(v).strip() for v in s if str(v).strip()}))
        .reset_index()
        .sort_values(["stage_order", "segment"])
    )
    lines = []
    for idx, row in grouped.head(max_stage_segments).iterrows():
        stage = str(row.get("stage") or "기타").strip() or "기타"
        segment = str(row.get("segment") or "기타").strip() or "기타"
        comps = row.get("input_company") if isinstance(row.get("input_company"), list) else []
        comp_text = ", ".join(comps[:12]) if comps else "-"
        lines.append(f"- {stage} > {segment}: {comp_text}")
    return "\n".join(lines)


def generate_value_chain_overview_with_ai(
    chain_name: str,
    match_rows: list[dict],
    provider: str,
    api_key: str,
    model: str,
) -> tuple[str, str, str]:
    key = (api_key or "").strip()
    if not key:
        return "", "API 설정 탭에서 AI API 키를 먼저 저장해 주세요.", ""

    rows_df = pd.DataFrame(match_rows) if isinstance(match_rows, list) else pd.DataFrame()
    if rows_df.empty:
        return "", "밸류체인 매칭 데이터가 비어 있습니다.", ""

    matched_names = []
    for v in rows_df["matched_company"].dropna().astype(str).tolist():
        vv = v.strip()
        if vv and vv not in matched_names:
            matched_names.append(vv)
    unmatched_names = []
    for v in rows_df.loc[rows_df["matched"] != True, "input_company"].dropna().astype(str).tolist():  # noqa: E712
        vv = v.strip()
        if vv and vv not in unmatched_names:
            unmatched_names.append(vv)

    stage_brief = _build_value_chain_rows_brief_text(match_rows)
    topic_queries = [
        f"{chain_name} 업스트림 미드스트림 다운스트림 설명",
        f"{chain_name} 수요 공급 가격 변수",
        f"{chain_name} 관련 상장사 투자 포인트 리스크",
    ]
    google_context_text, google_note = fetch_google_topic_research_context(
        topic=chain_name,
        extra_queries=topic_queries,
        max_items=12,
    )

    matched_text = ", ".join(matched_names[:40]) if matched_names else "없음"
    unmatched_text = ", ".join(unmatched_names[:40]) if unmatched_names else "없음"
    user_prompt = f"""
주제: {chain_name}

내 밸류체인 추출 결과:
{stage_brief}

내 관심기업 매칭:
- 매칭됨: {matched_text}
- 미매칭: {unmatched_text}

구글 검색 발췌:
{google_context_text or "검색 컨텍스트 없음"}

요청:
- 한국어 Markdown으로만 작성
- 아래 섹션을 정확히 포함
  1) 밸류체인 총 설명
  2) 단계별 핵심 포인트(업스트림/미드스트림/다운스트림)
  3) 가격/수요/원가/규제 민감도
  4) 내 관심기업 포지셔닝(매칭된 기업 중심)
  5) 체크해야 할 리스크와 데이터 포인트
- 일반론보다 실제 산업 키워드/공정/제품명을 구체적으로 쓸 것
- 확실하지 않은 내용은 '(확인필요)' 표기
""".strip()

    text, err = call_ai_text(
        provider=provider,
        api_key=key,
        model=model,
        system_prompt="너는 산업 밸류체인 분석 어시스턴트다. 사실 기반으로 한국어 Markdown을 작성한다.",
        user_prompt=user_prompt,
        temperature=0.2,
        max_output_tokens=1400,
        timeout_sec=35,
    )
    if err:
        return "", f"밸류체인 설명 생성 실패: {err}", google_note
    out = str(text or "").strip()
    if not out:
        return "", "밸류체인 설명 생성 결과가 비어 있습니다.", google_note
    return out, "", google_note


def _analysis_preview_lines(text: str, limit: int = 5) -> list[str]:
    lines = _split_report_lines(text)
    if not lines:
        raw = str(text or "").strip()
        return [raw] if raw else []
    return lines[: max(1, int(limit))]


def _render_value_chain_company_detail(
    input_company: str,
    matched_company: str,
    company_list_df: pd.DataFrame,
    latest_analysis_map: dict[str, dict],
) -> None:
    input_name = str(input_company or "").strip()
    matched_name = str(matched_company or "").strip()
    target_name = matched_name or input_name
    st.markdown(f"**기업명:** {target_name or '-'}")
    if input_name and matched_name and input_name != matched_name:
        st.caption(f"이미지 추출명: {input_name}")

    cl_df = company_list_df if isinstance(company_list_df, pd.DataFrame) else pd.DataFrame()
    company_row = None
    if not cl_df.empty and target_name:
        picked = cl_df[cl_df["stock_name"].astype(str) == target_name]
        if not picked.empty:
            company_row = picked.iloc[0]

    if company_row is not None:
        ticker = clean_valid_ticker(str(company_row.get("ticker") or ""))
        sector = str(company_row.get("sector") or "").strip()
        price_krw = _safe_to_float(company_row.get("price_krw"))
        price_source = str(company_row.get("price_source") or "").strip()
        created_at = str(company_row.get("created_at") or "").strip()
        updated_at = str(company_row.get("updated_at") or "").strip()
        c1, c2 = st.columns(2)
        with c1:
            st.caption(f"티커: {ticker or '-'}")
            st.caption(f"산업섹터: {sector or '-'}")
            st.caption(f"리스트소스: {str(company_row.get('source') or '-').strip() or '-'}")
        with c2:
            if price_krw is not None and price_krw > 0:
                st.caption(f"현재주가(원화환산): {price_krw:,.0f}원")
            else:
                st.caption("현재주가(원화환산): -")
            st.caption(f"주가소스: {price_source or '-'}")
            st.caption(f"등록/수정: {created_at or '-'} / {updated_at or '-'}")
    else:
        st.caption("기업정보 리스트에서 상세 정보를 찾지 못했습니다.")

    latest = latest_analysis_map.get(target_name) if target_name else None
    if not isinstance(latest, dict):
        st.caption("저장된 기업분석 이력이 없습니다.")
        return

    st.markdown("---")
    st.caption(
        "최근 기업분석: "
        f"{str(latest.get('analysis_date') or '-')} / "
        f"티커 {str(latest.get('ticker') or '-')} / "
        f"소스 {str(latest.get('source') or '-')}"
    )

    overview_lines = _analysis_preview_lines(str(latest.get("company_overview") or ""), limit=4)
    if overview_lines:
        st.markdown("**기업 개요**")
        for ln in overview_lines:
            st.write(f"- {ln}")

    product_lines = _analysis_preview_lines(str(latest.get("products_services") or ""), limit=5)
    if product_lines:
        st.markdown("**핵심 제품/서비스**")
        for ln in product_lines:
            st.write(f"- {ln}")

    raw_lines = _analysis_preview_lines(str(latest.get("raw_materials") or ""), limit=4)
    if raw_lines:
        st.markdown("**핵심 원재료/투입요소**")
        for ln in raw_lines:
            st.write(f"- {ln}")

    fs = parse_financial_summary_json(latest.get("financial_summary_json"))
    if isinstance(fs, dict) and fs:
        st.markdown("**핵심 재무 요약**")
        st.caption(
            f"시가총액 {_fmt_num_brief(fs.get('market_cap'))} / "
            f"매출 {_fmt_num_brief(fs.get('total_revenue'))} / "
            f"ROE {_fmt_pct_brief(fs.get('roe_pct'))}"
        )


def fetch_company_metrics_from_yfinance(ticker: str) -> tuple[dict, str, str]:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return {}, "", "티커를 입력해 주세요."

    try:
        import yfinance as yf
    except Exception:
        return {}, "", "yfinance 패키지가 없어 자동 불러오기를 사용할 수 없습니다."

    try:
        info = yf.Ticker(tkr).info or {}
    except Exception as exc:
        return {}, "", f"데이터 불러오기 실패: {exc}"

    if not info:
        return {}, "", "해당 티커에서 데이터를 찾지 못했습니다."

    def to_pct(value):
        if value is None:
            return None
        try:
            value = float(value)
        except Exception:
            return None
        return value * 100.0 if -1.5 <= value <= 1.5 else value

    metrics = {
        "dividend_yield": to_pct(info.get("dividendYield")),
        "revenue_growth": to_pct(info.get("revenueGrowth")),
        "eps_growth": to_pct(info.get("earningsGrowth")),
        "roe": to_pct(info.get("returnOnEquity")),
        "operating_margin": to_pct(info.get("operatingMargins")),
        "debt_ratio": info.get("debtToEquity"),
        "current_ratio": info.get("currentRatio"),
        "per": info.get("trailingPE") if info.get("trailingPE") is not None else info.get("forwardPE"),
        "pbr": info.get("priceToBook"),
    }

    name = info.get("shortName") or info.get("longName") or ""
    cleaned = {}
    for key, value in metrics.items():
        try:
            cleaned[key] = None if value is None else float(value)
        except Exception:
            cleaned[key] = None

    return cleaned, name, ""


def _to_pct_value(value) -> float | None:
    fv = _safe_to_float(value)
    if fv is None:
        return None
    return fv * 100.0 if -1.5 <= fv <= 1.5 else fv


def _first_number(data: dict, keys: list[str]) -> float | None:
    for key in keys:
        if key in data:
            value = _safe_to_float(data.get(key))
            if value is not None:
                return value
    return None


def fetch_company_metrics_from_alpha_vantage(ticker: str, api_key: str) -> tuple[dict, str, str]:
    tkr = clean_valid_ticker(ticker)
    key = (api_key or "").strip()
    if not tkr:
        return {}, "", "티커를 입력해 주세요."
    if not key:
        return {}, "", "Alpha Vantage API Key가 없습니다."

    try:
        resp = requests.get(
            "https://www.alphavantage.co/query",
            params={"function": "OVERVIEW", "symbol": tkr, "apikey": key},
            timeout=14,
        )
        resp.raise_for_status()
        info = resp.json()
    except Exception as exc:
        return {}, "", f"Alpha Vantage 조회 실패: {exc}"

    if not info or not isinstance(info, dict):
        return {}, "", "Alpha Vantage 응답이 비어 있습니다."
    if str(info.get("Note") or "").strip():
        return {}, "", f"Alpha Vantage 제한: {info.get('Note')}"
    if str(info.get("Information") or "").strip():
        return {}, "", f"Alpha Vantage 안내: {info.get('Information')}"

    metrics = {
        "dividend_yield": _to_pct_value(info.get("DividendYield")),
        "revenue_growth": _to_pct_value(info.get("QuarterlyRevenueGrowthYOY")),
        "eps_growth": _to_pct_value(info.get("QuarterlyEarningsGrowthYOY")),
        "roe": _to_pct_value(info.get("ReturnOnEquityTTM")),
        "operating_margin": _to_pct_value(info.get("OperatingMarginTTM")),
        "debt_ratio": _safe_to_float(info.get("DebtToEquity")),
        "current_ratio": _safe_to_float(info.get("CurrentRatio")),
        "per": _safe_to_float(info.get("PERatio")),
        "pbr": _safe_to_float(info.get("PriceToBookRatio")),
    }
    name = str(info.get("Name") or "").strip()
    return metrics, name, ""


def fetch_company_metrics_from_finnhub(ticker: str, api_key: str) -> tuple[dict, str, str]:
    tkr = clean_valid_ticker(ticker)
    key = (api_key or "").strip()
    if not tkr:
        return {}, "", "티커를 입력해 주세요."
    if not key:
        return {}, "", "Finnhub API Key가 없습니다."

    try:
        profile_resp = requests.get(
            "https://finnhub.io/api/v1/stock/profile2",
            params={"symbol": tkr, "token": key},
            timeout=14,
        )
        profile_resp.raise_for_status()
        profile = profile_resp.json() or {}

        metric_resp = requests.get(
            "https://finnhub.io/api/v1/stock/metric",
            params={"symbol": tkr, "metric": "all", "token": key},
            timeout=14,
        )
        metric_resp.raise_for_status()
        metric_payload = metric_resp.json() or {}
    except Exception as exc:
        return {}, "", f"Finnhub 조회 실패: {exc}"

    metric = metric_payload.get("metric") or {}
    if not metric and not profile:
        return {}, "", "Finnhub에서 데이터를 찾지 못했습니다."

    metrics = {
        "dividend_yield": _to_pct_value(_first_number(metric, ["dividendYieldIndicatedAnnual", "dividendYield5Y"])),
        "revenue_growth": _to_pct_value(_first_number(metric, ["revenueGrowthTTMYoy", "revenueGrowth3Y"])),
        "eps_growth": _to_pct_value(_first_number(metric, ["epsGrowthTTMYoy", "epsGrowth5Y"])),
        "roe": _to_pct_value(_first_number(metric, ["roeTTM", "roeRfy"])),
        "operating_margin": _to_pct_value(_first_number(metric, ["operatingMarginTTM", "operatingMarginAnnual"])),
        "debt_ratio": _first_number(metric, ["debtEquityAnnual", "totalDebt/totalEquityAnnual"]),
        "current_ratio": _first_number(metric, ["currentRatioAnnual", "currentRatioQuarterly"]),
        "per": _first_number(metric, ["peTTM", "peBasicExclExtraTTM"]),
        "pbr": _first_number(metric, ["pbAnnual", "pbQuarterly"]),
    }
    name = str(profile.get("name") or "").strip()
    return metrics, name, ""


def fetch_company_metrics_multi_source(ticker: str) -> tuple[dict, str, str, str]:
    metrics, name, err = fetch_company_metrics_from_yfinance(ticker)
    if not err and metrics:
        return metrics, name, "", "yfinance"

    alpha_key, finnhub_key = get_market_data_api_keys()
    alpha_err = ""
    fin_err = ""
    if alpha_key:
        alpha_metrics, alpha_name, alpha_err = fetch_company_metrics_from_alpha_vantage(ticker, alpha_key)
        if not alpha_err and alpha_metrics:
            return alpha_metrics, alpha_name, "", "alpha_vantage"

    if finnhub_key:
        fin_metrics, fin_name, fin_err = fetch_company_metrics_from_finnhub(ticker, finnhub_key)
        if not fin_err and fin_metrics:
            return fin_metrics, fin_name, "", "finnhub"

    errs = [e for e in [err, alpha_err, fin_err] if e]
    return {}, "", " | ".join(errs) if errs else "기업 지표를 불러오지 못했습니다.", ""


def save_company_score(
    score_date: date,
    stock_name: str,
    ticker: str,
    metrics: dict,
    scores: dict,
    note: str,
    source: str,
) -> None:
    now_str = datetime.now().isoformat(timespec="seconds")
    conn = get_conn()
    try:
        conn.execute(
            """
            INSERT INTO company_scores (
                score_date,
                stock_name,
                ticker,
                dividend_yield,
                revenue_growth,
                eps_growth,
                roe,
                operating_margin,
                debt_ratio,
                current_ratio,
                per,
                pbr,
                dividend_score,
                growth_score,
                stability_score,
                valuation_score,
                total_score,
                weight_dividend,
                weight_growth,
                weight_stability,
                weight_valuation,
                source,
                note,
                created_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                score_date.isoformat(),
                stock_name,
                ticker or None,
                metrics.get("dividend_yield"),
                metrics.get("revenue_growth"),
                metrics.get("eps_growth"),
                metrics.get("roe"),
                metrics.get("operating_margin"),
                metrics.get("debt_ratio"),
                metrics.get("current_ratio"),
                metrics.get("per"),
                metrics.get("pbr"),
                scores["dividend_score"],
                scores["growth_score"],
                scores["stability_score"],
                scores["valuation_score"],
                scores["total_score"],
                scores["weights"]["dividend"],
                scores["weights"]["growth"],
                scores["weights"]["stability"],
                scores["weights"]["valuation"],
                source or "manual",
                note or None,
                now_str,
            ),
        )
        conn.commit()
    finally:
        conn.close()


def load_company_score_history() -> pd.DataFrame:
    conn = get_conn()
    try:
        query = """
            SELECT
                score_date,
                stock_name,
                ticker,
                dividend_yield,
                revenue_growth,
                eps_growth,
                roe,
                operating_margin,
                debt_ratio,
                current_ratio,
                per,
                pbr,
                dividend_score,
                growth_score,
                stability_score,
                valuation_score,
                total_score,
                weight_dividend,
                weight_growth,
                weight_stability,
                weight_valuation,
                source,
                note,
                created_at
            FROM company_scores
            ORDER BY score_date, stock_name
        """
        df = pd.read_sql_query(query, conn)
    finally:
        conn.close()

    if not df.empty:
        df["score_date"] = pd.to_datetime(df["score_date"])
    return df


def _json_parse_failure_reason(text: str) -> str:
    raw = (text or "").strip()
    if not raw:
        return "AI 응답이 비어 있습니다."
    if ("```" in raw and ("{" in raw or "[" in raw)) or ("\n" in raw and not raw.startswith("{")):
        cause = "코드블록/설명문이 JSON과 함께 섞여 있습니다."
    elif ("{" not in raw) and ("[" not in raw):
        cause = "JSON 시작 괄호({ 또는 [)가 없습니다."
    else:
        cause = "따옴표/쉼표/괄호 형식이 JSON 규격과 다릅니다."
    preview = re.sub(r"\s+", " ", raw)[:220]
    return f"{cause} 응답 일부: {preview}"


def _extract_json_from_text(text: str) -> dict | list | None:
    raw = (text or "").strip()
    if not raw:
        return None

    def _balanced_chunks(src: str, open_ch: str, close_ch: str) -> list[str]:
        out: list[str] = []
        in_str = False
        escaped = False
        depth = 0
        start = -1
        for i, ch in enumerate(src):
            if in_str:
                if escaped:
                    escaped = False
                elif ch == "\\":
                    escaped = True
                elif ch == '"':
                    in_str = False
                continue
            if ch == '"':
                in_str = True
                continue
            if ch == open_ch:
                if depth == 0:
                    start = i
                depth += 1
            elif ch == close_ch and depth > 0:
                depth -= 1
                if depth == 0 and start >= 0:
                    out.append(src[start : i + 1].strip())
                    start = -1
        return out

    candidates: list[str] = [raw]
    for m in re.finditer(r"```(?:json)?\s*([\s\S]*?)```", raw, re.I):
        block = (m.group(1) or "").strip()
        if block:
            candidates.append(block)

    candidates.extend(_balanced_chunks(raw, "{", "}"))
    candidates.extend(_balanced_chunks(raw, "[", "]"))

    seen: set[str] = set()
    for cand in candidates:
        c = str(cand or "").strip()
        if not c or c in seen:
            continue
        seen.add(c)

        variants = [
            c,
            c.replace("“", '"').replace("”", '"').replace("’", "'").replace("‘", "'"),
            re.sub(r",\s*([}\]])", r"\1", c),
        ]
        for v in variants:
            try:
                parsed = json.loads(v)
                if isinstance(parsed, (dict, list)):
                    return parsed
            except Exception:
                pass
            try:
                parsed = ast.literal_eval(v)
                if isinstance(parsed, (dict, list)):
                    return parsed
            except Exception:
                pass
    return None


def _normalize_research_url(url: str) -> str:
    u = str(url or "").strip()
    if not u:
        return ""
    try:
        parsed = urlparse(u)
        if not parsed.scheme or not parsed.netloc:
            return ""
        return f"{parsed.scheme.lower()}://{parsed.netloc.lower()}{parsed.path}"
    except Exception:
        return u


def _is_valid_google_search_result_url(url: str) -> bool:
    u = str(url or "").strip()
    if not u.startswith(("http://", "https://")):
        return False
    try:
        host = (urlparse(u).netloc or "").lower()
    except Exception:
        host = ""
    blocked_hosts = {
        "www.google.com",
        "google.com",
        "support.google.com",
        "accounts.google.com",
        "maps.google.com",
        "news.google.com",
        "encrypted-tbn0.gstatic.com",
        "encrypted-tbn1.gstatic.com",
        "encrypted-tbn2.gstatic.com",
        "encrypted-tbn3.gstatic.com",
    }
    if host in blocked_hosts:
        return False
    if "google." in host and "youtube.com" not in host:
        return False
    return True


def _parse_google_search_markdown_results(markdown_text: str, limit: int = 10) -> list[dict]:
    lines = (markdown_text or "").splitlines()
    rows: list[dict] = []
    seen: set[str] = set()
    n = len(lines)

    i = 0
    while i < n and len(rows) < limit:
        line = str(lines[i] or "").strip()
        m_empty_anchor = re.match(r"^\[\]\((https?://[^)]+)\)\s*$", line)
        m_labeled_anchor = re.match(r"^\[([^\]]+)\]\((https?://[^)]+)\)\s*$", line)
        url = ""
        title = ""
        if m_empty_anchor:
            url = html.unescape(m_empty_anchor.group(1)).strip()
        elif m_labeled_anchor:
            title = html.unescape(m_labeled_anchor.group(1)).strip()
            url = html.unescape(m_labeled_anchor.group(2)).strip()
        else:
            i += 1
            continue

        if not _is_valid_google_search_result_url(url):
            i += 1
            continue

        url_key = _normalize_research_url(url)
        if not url_key or url_key in seen:
            i += 1
            continue

        if not title:
            for j in range(i + 1, min(i + 8, n)):
                cand = str(lines[j] or "").strip()
                if not cand:
                    continue
                if cand.startswith(("![", "[](", "[Skip ", "# ", "* ")):
                    continue
                if re.match(r"^\[[^\]]+\]\(https?://", cand):
                    continue
                if cand.lower() in {"google", "people also ask"}:
                    continue
                title = cand
                break

        snippet_parts: list[str] = []
        for k in range(i + 1, min(i + 12, n)):
            cand = str(lines[k] or "").strip()
            if not cand:
                continue
            if cand.startswith(("![", "[](", "# ", "[Skip ", "* ")):
                continue
            if re.match(r"^\[[^\]]+\]\(https?://", cand):
                continue
            if title and cand == title:
                continue
            if len(cand) < 18:
                continue
            if "faviconV2" in cand or "gstatic.com" in cand:
                continue
            snippet_parts.append(cand)
            if len(" ".join(snippet_parts)) >= 260:
                break
        snippet = " ".join(snippet_parts).strip()
        if not title and not snippet:
            i += 1
            continue

        try:
            domain = (urlparse(url).netloc or "").lower()
        except Exception:
            domain = ""
        rows.append(
            {
                "title": title or "(제목 없음)",
                "snippet": snippet,
                "url": url,
                "domain": domain,
            }
        )
        seen.add(url_key)
        i += 1

    return rows


def fetch_google_company_research_context(
    company_name: str,
    ticker: str,
    max_items: int = 8,
) -> tuple[str, str]:
    name = str(company_name or "").strip()
    tkr = clean_valid_ticker(str(ticker or "").strip())
    if not name:
        return "", "구글 검색 컨텍스트 생성 실패: 기업명이 비어 있습니다."

    query_parts = [name]
    if tkr:
        query_parts.append(tkr)

    queries = [
        " ".join(query_parts + ["products", "services", "business model", "investor presentation"]),
        " ".join(query_parts + ["raw materials", "suppliers", "cost structure", "operations"]),
        " ".join(query_parts + ["growth strategy", "new projects", "pipeline", "capex"]),
    ]

    all_rows: list[dict] = []
    errs: list[str] = []
    for q in queries:
        url = f"https://r.jina.ai/http://www.google.com/search?q={quote_plus(q)}&hl=en&num=8"
        try:
            resp = requests.get(
                url,
                headers={**HTTP_HEADERS_COMMON, "Accept-Language": "en-US,en;q=0.9"},
                timeout=18,
            )
            resp.raise_for_status()
            parsed_rows = _parse_google_search_markdown_results(resp.text, limit=10)
            all_rows.extend(parsed_rows)
        except Exception as exc:
            errs.append(str(exc))

    deduped: list[dict] = []
    seen_urls: set[str] = set()
    for row in all_rows:
        url_key = _normalize_research_url(row.get("url"))
        if not url_key or url_key in seen_urls:
            continue
        seen_urls.add(url_key)
        deduped.append(row)

    if not deduped:
        msg = "구글 검색 결과를 확보하지 못했습니다."
        if errs:
            msg = f"{msg} ({'; '.join(errs[:2])})"
        return "", msg

    context_lines = []
    for idx, row in enumerate(deduped[: max(1, int(max_items))], start=1):
        title = str(row.get("title") or "").strip()
        snippet = str(row.get("snippet") or "").strip()
        snippet = re.sub(r"\s+", " ", snippet)[:360]
        domain = str(row.get("domain") or "").strip()
        url = str(row.get("url") or "").strip()
        context_lines.append(
            f"{idx}. title={title} | snippet={snippet} | source={domain} | url={url}"
        )
    context_text = "\n".join(context_lines)
    return context_text, f"google_web ({len(context_lines)}건)"


def fetch_google_topic_research_context(
    topic: str,
    extra_queries: list[str] | None = None,
    max_items: int = 10,
) -> tuple[str, str]:
    topic_text = str(topic or "").strip()
    if not topic_text:
        return "", "구글 검색 컨텍스트 생성 실패: 주제가 비어 있습니다."

    queries = [
        f"{topic_text} value chain upstream midstream downstream",
        f"{topic_text} industry structure demand supply price driver",
        f"{topic_text} 주요 기업 경쟁구도 투자 포인트 리스크",
    ]
    for q in (extra_queries or []):
        qq = str(q or "").strip()
        if qq:
            queries.append(qq)

    all_rows: list[dict] = []
    errs: list[str] = []
    for q in queries:
        url = f"https://r.jina.ai/http://www.google.com/search?q={quote_plus(q)}&hl=ko&num=8"
        try:
            resp = requests.get(
                url,
                headers={**HTTP_HEADERS_COMMON, "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7"},
                timeout=18,
            )
            resp.raise_for_status()
            all_rows.extend(_parse_google_search_markdown_results(resp.text, limit=12))
        except Exception as exc:
            errs.append(str(exc))

    deduped: list[dict] = []
    seen_urls: set[str] = set()
    for row in all_rows:
        url_key = _normalize_research_url(row.get("url"))
        if not url_key or url_key in seen_urls:
            continue
        seen_urls.add(url_key)
        deduped.append(row)

    if not deduped:
        msg = "구글 검색 결과를 확보하지 못했습니다."
        if errs:
            msg = f"{msg} ({'; '.join(errs[:2])})"
        return "", msg

    lines = []
    for idx, row in enumerate(deduped[: max(1, int(max_items))], start=1):
        title = str(row.get("title") or "").strip()
        snippet = re.sub(r"\s+", " ", str(row.get("snippet") or "").strip())[:360]
        domain = str(row.get("domain") or "").strip()
        url = str(row.get("url") or "").strip()
        lines.append(f"{idx}. title={title} | snippet={snippet} | source={domain} | url={url}")
    return "\n".join(lines), f"google_web ({len(lines)}건)"


def _is_missing_summary_value(value) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return not value.strip()
    if isinstance(value, (list, tuple, dict)):
        return len(value) == 0
    return False


def _merge_financial_summary_dicts(base: dict, incoming: dict) -> tuple[dict, bool]:
    merged = dict(base or {})
    changed = False
    for key, val in (incoming or {}).items():
        if key not in merged or _is_missing_summary_value(merged.get(key)):
            if not _is_missing_summary_value(val):
                merged[key] = val
                changed = True
    return merged, changed


def _extract_domestic_code_from_ticker(ticker: str) -> str:
    tkr = clean_valid_ticker(ticker)
    if re.fullmatch(r"\d{6}\.(KS|KQ)", tkr):
        return tkr.split(".")[0]
    if re.fullmatch(r"\d{6}", tkr):
        return tkr
    return ""


def _parse_number_with_short_unit(value_text: str) -> float | None:
    text = str(value_text or "").strip().replace(",", "")
    if not text:
        return None
    m = re.search(r"([+\-]?\d+(?:\.\d+)?)\s*([TMBK]?)", text, re.I)
    if not m:
        return _safe_to_float(text)
    num = _safe_to_float(m.group(1))
    unit = (m.group(2) or "").upper()
    if num is None:
        return None
    factor_map = {"": 1.0, "K": 1e3, "M": 1e6, "B": 1e9, "T": 1e12}
    return float(num) * float(factor_map.get(unit, 1.0))


def _extract_google_finance_metric_value(page_html: str, label: str) -> str:
    if not page_html:
        return ""
    pattern = rf">{re.escape(label)}</div>.*?<div class=\"P6K39c\">(.*?)</div>"
    m = re.search(pattern, page_html, re.I | re.S)
    if not m:
        return ""
    raw = re.sub(r"<.*?>", " ", m.group(1) or "")
    raw = html.unescape(raw)
    return re.sub(r"\s+", " ", raw).strip()


def _google_quote_symbol_candidates(ticker: str) -> list[str]:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return []
    code = _extract_domestic_code_from_ticker(tkr)
    if code:
        return [f"{code}:KRX"]

    if "." in tkr:
        # 국제 접미사가 있는 티커는 점을 하이픈으로 바꾼 후보도 함께 시도.
        no_dot = tkr.replace(".", "-")
        return [tkr, no_dot]
    return [f"{tkr}:NASDAQ", f"{tkr}:NYSE", f"{tkr}:AMEX", tkr]


def fetch_company_financial_summary_from_naver(ticker: str) -> tuple[dict, str]:
    code = _extract_domestic_code_from_ticker(ticker)
    if not code:
        return {}, "Naver 재무요약은 국내 티커(.KS/.KQ)만 지원합니다."

    try:
        resp = requests.get(
            "https://finance.naver.com/item/main.naver",
            params={"code": code},
            headers=HTTP_HEADERS_COMMON,
            timeout=14,
        )
        resp.raise_for_status()
        body = resp.text or ""
    except Exception as exc:
        return {}, f"Naver 기업 페이지 조회 실패: {exc}"

    title_match = re.search(r"<title>(.*?)</title>", body, re.I | re.S)
    raw_title = html.unescape(title_match.group(1)).strip() if title_match else ""
    name = raw_title.split(":")[0].strip() if ":" in raw_title else raw_title
    sector, _ = fetch_sector_from_naver_domestic(f"{code}.KS")

    def cell_metric(label: str) -> float | None:
        pat = rf"<th[^>]*>\s*<strong>\s*{label}[^<]*</strong>\s*</th>\s*<td[^>]*>(.*?)</td>"
        m = re.search(pat, body, re.I | re.S)
        if not m:
            return None
        cell = re.sub(r"<.*?>", " ", m.group(1) or "")
        cell = html.unescape(re.sub(r"\s+", " ", cell)).strip()
        return _safe_to_float(cell)

    summary = {
        "name": name or "",
        "sector": sector or "",
        "industry": sector or "",
        "country": "대한민국",
        "market_cap": None,
        "enterprise_value": None,
        "total_revenue": None,
        "ebitda": None,
        "net_income_to_common": None,
        "operating_cashflow": None,
        "free_cashflow": None,
        "dividend_yield_pct": _to_pct_value(cell_metric("배당수익률")),
        "revenue_growth_pct": None,
        "earnings_growth_pct": None,
        "roe_pct": _to_pct_value(cell_metric("ROE")),
        "operating_margin_pct": None,
        "gross_margin_pct": None,
        "debt_to_equity": None,
        "current_ratio": None,
        "trailing_pe": cell_metric("PER"),
        "forward_pe": None,
        "price_to_book": cell_metric("PBR"),
        "beta": None,
        "website": "",
        "business_summary": "",
        "income_statement_annual": [],
        "balance_sheet_annual": [],
        "cashflow_annual": [],
    }
    has_core = bool(summary.get("name") or summary.get("sector") or summary.get("trailing_pe") or summary.get("price_to_book"))
    if not has_core:
        return {}, "Naver 요약 정보를 찾지 못했습니다."
    return summary, ""


def fetch_company_financial_summary_from_google_finance(ticker: str) -> tuple[dict, str]:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return {}, "티커가 비어 있습니다."

    page_html = ""
    final_symbol = ""
    last_err = ""
    for symbol in _google_quote_symbol_candidates(tkr):
        url = f"https://www.google.com/finance/quote/{symbol}"
        try:
            resp = requests.get(
                url,
                headers={**HTTP_HEADERS_COMMON, "Accept-Language": "en-US,en;q=0.9"},
                timeout=14,
            )
            if resp.status_code >= 400:
                last_err = f"{symbol} 상태코드 {resp.status_code}"
                continue
            text = resp.text or ""
            if "google.com/finance/quote" not in text:
                last_err = f"{symbol} 페이지 구조 불일치"
                continue
            page_html = text
            final_symbol = symbol
            break
        except Exception as exc:
            last_err = str(exc)
            continue

    if not page_html:
        return {}, f"Google Finance 조회 실패: {last_err or '결과 없음'}"

    title_match = re.search(r"<title>(.*?)</title>", page_html, re.I | re.S)
    title_text = html.unescape(title_match.group(1)).strip() if title_match else ""
    name = title_text.split("(")[0].strip() if "(" in title_text else title_text
    meta_match = re.search(r'<meta name=\"description\" content=\"([^\"]+)\"', page_html, re.I)
    business_summary = html.unescape(meta_match.group(1)).strip() if meta_match else ""

    market_cap_text = _extract_google_finance_metric_value(page_html, "Market cap")
    pe_text = _extract_google_finance_metric_value(page_html, "P/E ratio")
    div_text = _extract_google_finance_metric_value(page_html, "Dividend yield")
    pbr_text = _extract_google_finance_metric_value(page_html, "Price to book")
    de_text = _extract_google_finance_metric_value(page_html, "D/E ratio")

    summary = {
        "name": name or "",
        "sector": "",
        "industry": "",
        "country": "대한민국" if ":KRX" in final_symbol else "",
        "market_cap": _parse_number_with_short_unit(market_cap_text),
        "enterprise_value": None,
        "total_revenue": None,
        "ebitda": None,
        "net_income_to_common": None,
        "operating_cashflow": None,
        "free_cashflow": None,
        "dividend_yield_pct": _to_pct_value(div_text),
        "revenue_growth_pct": None,
        "earnings_growth_pct": None,
        "roe_pct": None,
        "operating_margin_pct": None,
        "gross_margin_pct": None,
        "debt_to_equity": _safe_to_float(de_text),
        "current_ratio": None,
        "trailing_pe": _safe_to_float(pe_text),
        "forward_pe": None,
        "price_to_book": _safe_to_float(pbr_text),
        "beta": None,
        "website": "",
        "business_summary": business_summary,
        "income_statement_annual": [],
        "balance_sheet_annual": [],
        "cashflow_annual": [],
    }
    has_core = bool(summary.get("name") or summary.get("market_cap") or summary.get("trailing_pe"))
    if not has_core:
        return {}, "Google Finance 핵심 요약값을 찾지 못했습니다."
    return summary, ""


def fetch_company_financial_summary_from_yfinance(ticker: str) -> tuple[dict, str]:
    tkr = clean_valid_ticker(ticker)
    if not tkr:
        return {}, "티커가 비어 있습니다."

    try:
        import yfinance as yf
    except Exception:
        return {}, "yfinance 패키지를 찾을 수 없습니다."

    try:
        obj = yf.Ticker(tkr)
        info = obj.info or {}
    except Exception as exc:
        return {}, f"yfinance 조회 실패: {exc}"

    if not info:
        return {}, "기업 기본정보를 불러오지 못했습니다."

    def pct(v):
        if v is None:
            return None
        try:
            fv = float(v)
        except Exception:
            return None
        return fv * 100 if -1.5 <= fv <= 1.5 else fv

    summary = {
        "name": info.get("shortName") or info.get("longName") or "",
        "sector": info.get("sector"),
        "industry": info.get("industry"),
        "country": info.get("country"),
        "market_cap": info.get("marketCap"),
        "enterprise_value": info.get("enterpriseValue"),
        "total_revenue": info.get("totalRevenue"),
        "ebitda": info.get("ebitda"),
        "net_income_to_common": info.get("netIncomeToCommon"),
        "operating_cashflow": info.get("operatingCashflow"),
        "free_cashflow": info.get("freeCashflow"),
        "dividend_yield_pct": pct(info.get("dividendYield")),
        "revenue_growth_pct": pct(info.get("revenueGrowth")),
        "earnings_growth_pct": pct(info.get("earningsGrowth")),
        "roe_pct": pct(info.get("returnOnEquity")),
        "operating_margin_pct": pct(info.get("operatingMargins")),
        "gross_margin_pct": pct(info.get("grossMargins")),
        "debt_to_equity": info.get("debtToEquity"),
        "current_ratio": info.get("currentRatio"),
        "trailing_pe": info.get("trailingPE"),
        "forward_pe": info.get("forwardPE"),
        "price_to_book": info.get("priceToBook"),
        "beta": info.get("beta"),
        "website": info.get("website"),
        "business_summary": info.get("longBusinessSummary"),
    }

    # Annual income statement snapshot (recent years)
    try:
        fin = obj.financials
        if fin is None or fin.empty:
            fin = obj.income_stmt
    except Exception:
        fin = pd.DataFrame()
    try:
        bal = obj.balance_sheet
    except Exception:
        bal = pd.DataFrame()
    try:
        cash = obj.cashflow
    except Exception:
        cash = pd.DataFrame()

    def table_to_records(df: pd.DataFrame, targets: list[str], max_cols: int = 4) -> list[dict]:
        if df is None or df.empty:
            return []
        cols = list(df.columns)[:max_cols]
        years = []
        for c in cols:
            try:
                years.append(str(pd.to_datetime(c).year))
            except Exception:
                years.append(str(c))
        records = []
        for row_name in targets:
            if row_name not in df.index:
                continue
            values = []
            for c in cols:
                v = df.loc[row_name, c]
                try:
                    v = float(v)
                except Exception:
                    v = None
                values.append(v)
            records.append({"item": row_name, **{years[i]: values[i] for i in range(len(years))}})
        return records

    summary["income_statement_annual"] = table_to_records(
        fin,
        ["Total Revenue", "Gross Profit", "Operating Income", "Net Income", "EBITDA"],
        max_cols=4,
    )
    summary["balance_sheet_annual"] = table_to_records(
        bal,
        ["Total Assets", "Total Liabilities Net Minority Interest", "Stockholders Equity", "Total Debt", "Cash And Cash Equivalents"],
        max_cols=4,
    )
    summary["cashflow_annual"] = table_to_records(
        cash,
        ["Operating Cash Flow", "Free Cash Flow", "Capital Expenditure", "Investing Cash Flow", "Financing Cash Flow"],
        max_cols=4,
    )

    for key, value in list(summary.items()):
        if isinstance(value, (int, float)):
            summary[key] = float(value)
    return summary, ""


def _alpha_table_records(reports: list[dict], mapping: dict[str, list[str]], max_years: int = 4) -> list[dict]:
    if not reports:
        return []
    records = []
    rows = []
    for rep in reports[:max_years]:
        if not isinstance(rep, dict):
            continue
        year = str(rep.get("fiscalDateEnding", ""))[:4]
        if not year:
            continue
        rows.append((year, rep))
    if not rows:
        return []

    for item_name, keys in mapping.items():
        row = {"item": item_name}
        has_value = False
        for year, rep in rows:
            val = None
            for key in keys:
                v = _safe_to_float(rep.get(key))
                if v is not None:
                    val = v
                    break
            row[year] = val
            if val is not None:
                has_value = True
        if has_value:
            records.append(row)
    return records


def fetch_company_financial_summary_from_alpha_vantage(ticker: str, api_key: str) -> tuple[dict, str]:
    tkr = clean_valid_ticker(ticker)
    key = (api_key or "").strip()
    if not tkr:
        return {}, "티커가 비어 있습니다."
    if not key:
        return {}, "Alpha Vantage API Key가 없습니다."

    base_url = "https://www.alphavantage.co/query"
    try:
        overview_resp = requests.get(
            base_url,
            params={"function": "OVERVIEW", "symbol": tkr, "apikey": key},
            timeout=16,
        )
        overview_resp.raise_for_status()
        overview = overview_resp.json() or {}
    except Exception as exc:
        return {}, f"Alpha Vantage 기업개요 조회 실패: {exc}"

    note = str(overview.get("Note") or "").strip()
    if note:
        return {}, f"Alpha Vantage 제한: {note}"
    if not overview or not overview.get("Symbol"):
        return {}, "Alpha Vantage 기업개요 데이터가 없습니다."

    def pct(v):
        return _to_pct_value(v)

    summary = {
        "name": overview.get("Name") or "",
        "sector": overview.get("Sector"),
        "industry": overview.get("Industry"),
        "country": overview.get("Country"),
        "market_cap": _safe_to_float(overview.get("MarketCapitalization")),
        "enterprise_value": _safe_to_float(overview.get("EVToRevenue")),
        "total_revenue": _safe_to_float(overview.get("RevenueTTM")),
        "ebitda": _safe_to_float(overview.get("EBITDA")),
        "net_income_to_common": None,
        "operating_cashflow": None,
        "free_cashflow": None,
        "dividend_yield_pct": pct(overview.get("DividendYield")),
        "revenue_growth_pct": pct(overview.get("QuarterlyRevenueGrowthYOY")),
        "earnings_growth_pct": pct(overview.get("QuarterlyEarningsGrowthYOY")),
        "roe_pct": pct(overview.get("ReturnOnEquityTTM")),
        "operating_margin_pct": pct(overview.get("OperatingMarginTTM")),
        "gross_margin_pct": pct(overview.get("ProfitMargin")),
        "debt_to_equity": _safe_to_float(overview.get("DebtToEquity")),
        "current_ratio": _safe_to_float(overview.get("CurrentRatio")),
        "trailing_pe": _safe_to_float(overview.get("PERatio")),
        "forward_pe": None,
        "price_to_book": _safe_to_float(overview.get("PriceToBookRatio")),
        "beta": _safe_to_float(overview.get("Beta")),
        "website": overview.get("OfficialSite"),
        "business_summary": overview.get("Description"),
    }

    try:
        income_resp = requests.get(
            base_url,
            params={"function": "INCOME_STATEMENT", "symbol": tkr, "apikey": key},
            timeout=16,
        )
        income_resp.raise_for_status()
        income_data = income_resp.json() or {}
        income_reports = income_data.get("annualReports") or []
    except Exception:
        income_reports = []

    try:
        balance_resp = requests.get(
            base_url,
            params={"function": "BALANCE_SHEET", "symbol": tkr, "apikey": key},
            timeout=16,
        )
        balance_resp.raise_for_status()
        balance_data = balance_resp.json() or {}
        balance_reports = balance_data.get("annualReports") or []
    except Exception:
        balance_reports = []

    try:
        cash_resp = requests.get(
            base_url,
            params={"function": "CASH_FLOW", "symbol": tkr, "apikey": key},
            timeout=16,
        )
        cash_resp.raise_for_status()
        cash_data = cash_resp.json() or {}
        cash_reports = cash_data.get("annualReports") or []
    except Exception:
        cash_reports = []

    summary["income_statement_annual"] = _alpha_table_records(
        income_reports,
        {
            "Total Revenue": ["totalRevenue"],
            "Gross Profit": ["grossProfit"],
            "Operating Income": ["operatingIncome"],
            "Net Income": ["netIncome"],
            "EBITDA": ["ebitda"],
        },
    )
    summary["balance_sheet_annual"] = _alpha_table_records(
        balance_reports,
        {
            "Total Assets": ["totalAssets"],
            "Total Liabilities Net Minority Interest": ["totalLiabilities"],
            "Stockholders Equity": ["totalShareholderEquity"],
            "Total Debt": ["shortLongTermDebtTotal", "longTermDebt", "shortTermDebt"],
            "Cash And Cash Equivalents": ["cashAndCashEquivalentsAtCarryingValue"],
        },
    )
    summary["cashflow_annual"] = _alpha_table_records(
        cash_reports,
        {
            "Operating Cash Flow": ["operatingCashflow"],
            "Free Cash Flow": ["freeCashFlow"],
            "Capital Expenditure": ["capitalExpenditures"],
            "Investing Cash Flow": ["cashflowFromInvestment"],
            "Financing Cash Flow": ["cashflowFromFinancing"],
        },
    )

    if not summary["cashflow_annual"] and cash_reports:
        # freeCashFlow 필드가 없을 때 대체 계산
        rows = []
        for rep in cash_reports[:4]:
            year = str(rep.get("fiscalDateEnding", ""))[:4]
            if not year:
                continue
            op = _safe_to_float(rep.get("operatingCashflow"))
            capex = _safe_to_float(rep.get("capitalExpenditures"))
            fcf = (op - capex) if op is not None and capex is not None else None
            rows.append({"item": "Free Cash Flow", year: fcf})
        if rows:
            summary["cashflow_annual"] = rows

    for key_name, value in list(summary.items()):
        if isinstance(value, (int, float)):
            summary[key_name] = float(value)
    return summary, ""


def fetch_company_financial_summary_from_finnhub(ticker: str, api_key: str) -> tuple[dict, str]:
    tkr = clean_valid_ticker(ticker)
    key = (api_key or "").strip()
    if not tkr:
        return {}, "티커가 비어 있습니다."
    if not key:
        return {}, "Finnhub API Key가 없습니다."

    try:
        profile_resp = requests.get(
            "https://finnhub.io/api/v1/stock/profile2",
            params={"symbol": tkr, "token": key},
            timeout=14,
        )
        profile_resp.raise_for_status()
        profile = profile_resp.json() or {}

        metric_resp = requests.get(
            "https://finnhub.io/api/v1/stock/metric",
            params={"symbol": tkr, "metric": "all", "token": key},
            timeout=14,
        )
        metric_resp.raise_for_status()
        metric_payload = metric_resp.json() or {}
    except Exception as exc:
        return {}, f"Finnhub 조회 실패: {exc}"

    metric = metric_payload.get("metric") or {}
    if not profile and not metric:
        return {}, "Finnhub에서 기업 정보를 찾지 못했습니다."

    summary = {
        "name": profile.get("name") or "",
        "sector": profile.get("finnhubIndustry"),
        "industry": profile.get("finnhubIndustry"),
        "country": profile.get("country"),
        "market_cap": _safe_to_float(profile.get("marketCapitalization")),
        "enterprise_value": None,
        "total_revenue": None,
        "ebitda": None,
        "net_income_to_common": None,
        "operating_cashflow": None,
        "free_cashflow": None,
        "dividend_yield_pct": _to_pct_value(_first_number(metric, ["dividendYieldIndicatedAnnual", "dividendYield5Y"])),
        "revenue_growth_pct": _to_pct_value(_first_number(metric, ["revenueGrowthTTMYoy", "revenueGrowth3Y"])),
        "earnings_growth_pct": _to_pct_value(_first_number(metric, ["epsGrowthTTMYoy", "epsGrowth5Y"])),
        "roe_pct": _to_pct_value(_first_number(metric, ["roeTTM", "roeRfy"])),
        "operating_margin_pct": _to_pct_value(_first_number(metric, ["operatingMarginTTM", "operatingMarginAnnual"])),
        "gross_margin_pct": None,
        "debt_to_equity": _first_number(metric, ["debtEquityAnnual", "totalDebt/totalEquityAnnual"]),
        "current_ratio": _first_number(metric, ["currentRatioAnnual", "currentRatioQuarterly"]),
        "trailing_pe": _first_number(metric, ["peTTM", "peBasicExclExtraTTM"]),
        "forward_pe": None,
        "price_to_book": _first_number(metric, ["pbAnnual", "pbQuarterly"]),
        "beta": _safe_to_float(profile.get("beta")),
        "website": profile.get("weburl"),
        "business_summary": "",
        "income_statement_annual": [],
        "balance_sheet_annual": [],
        "cashflow_annual": [],
    }
    return summary, ""


def fetch_company_financial_summary_multi_source(ticker: str) -> tuple[dict, str, str]:
    source_chain: list[str] = []
    errs: list[str] = []
    merged_summary: dict = {}

    summary, err = fetch_company_financial_summary_from_yfinance(ticker)
    if not err and summary:
        merged_summary = dict(summary)
        source_chain.append("yfinance")
    elif err:
        errs.append(err)

    naver_summary, naver_err = fetch_company_financial_summary_from_naver(ticker)
    if not naver_err and naver_summary:
        merged_summary, _ = _merge_financial_summary_dicts(merged_summary, naver_summary)
        if "naver" not in source_chain:
            source_chain.append("naver")
    elif naver_err:
        errs.append(naver_err)

    google_summary, google_err = fetch_company_financial_summary_from_google_finance(ticker)
    if not google_err and google_summary:
        merged_summary, _ = _merge_financial_summary_dicts(merged_summary, google_summary)
        if "google_finance" not in source_chain:
            source_chain.append("google_finance")
    elif google_err:
        errs.append(google_err)

    alpha_key, finnhub_key = get_market_data_api_keys()
    alpha_err = ""
    fin_err = ""
    if alpha_key:
        alpha_summary, alpha_err = fetch_company_financial_summary_from_alpha_vantage(ticker, alpha_key)
        if not alpha_err and alpha_summary:
            merged_summary, _ = _merge_financial_summary_dicts(merged_summary, alpha_summary)
            if "alpha_vantage" not in source_chain:
                source_chain.append("alpha_vantage")
        elif alpha_err:
            errs.append(alpha_err)

    if finnhub_key:
        fin_summary, fin_err = fetch_company_financial_summary_from_finnhub(ticker, finnhub_key)
        if not fin_err and fin_summary:
            merged_summary, _ = _merge_financial_summary_dicts(merged_summary, fin_summary)
            if "finnhub" not in source_chain:
                source_chain.append("finnhub")
        elif fin_err:
            errs.append(fin_err)

    if merged_summary:
        if source_chain:
            merged_summary["data_sources"] = source_chain
        source_label = "+".join(source_chain) if source_chain else "multi_source"
        return merged_summary, "", source_label

    uniq_errs = []
    seen_err = set()
    for msg in errs:
        m = str(msg or "").strip()
        if not m or m in seen_err:
            continue
        uniq_errs.append(m)
        seen_err.add(m)
    return {}, " | ".join(uniq_errs) if uniq_errs else "기업 재무 정보를 불러오지 못했습니다.", ""


def _financial_summary_quality_score(summary: dict) -> int:
    if not isinstance(summary, dict) or not summary:
        return 0
    score = 0
    core_keys = [
        "name",
        "sector",
        "industry",
        "country",
        "market_cap",
        "enterprise_value",
        "total_revenue",
        "ebitda",
        "trailing_pe",
        "price_to_book",
        "revenue_growth_pct",
        "roe_pct",
        "business_summary",
    ]
    for key in core_keys:
        if not _is_missing_summary_value(summary.get(key)):
            score += 1
    for key in ["income_statement_annual", "balance_sheet_annual", "cashflow_annual"]:
        value = summary.get(key)
        if isinstance(value, list) and value:
            score += 2
    return score


def _collect_analysis_ticker_candidates(
    company_name: str,
    ticker_input: str,
    use_ai_ticker: bool,
    ai_provider: str,
    ai_api_key: str,
    ai_model: str,
    market_preference: str = "",
) -> tuple[list[str], list[str]]:
    name = (company_name or "").strip()
    raw_ticker = (ticker_input or "").strip()
    pref = _market_pref_normalized(market_preference)
    candidates: list[str] = []
    notes: list[str] = []

    def add_candidate(value: str, reason: str) -> None:
        tkr = clean_valid_ticker(value)
        if pref and not _ticker_matches_market_preference(tkr, pref):
            return
        if not tkr or tkr in candidates:
            return
        candidates.append(tkr)
        notes.append(f"{reason}:{tkr}")

    add_candidate(raw_ticker, "입력")
    normalized_external = _normalize_external_ticker_candidate(raw_ticker)
    if normalized_external and normalized_external != clean_valid_ticker(raw_ticker):
        add_candidate(normalized_external, "입력정규화")

    if name:
        saved_hint = clean_valid_ticker(get_company_list_ticker(name))
        if saved_hint:
            add_candidate(saved_hint, "리스트저장값")

        auto_ticker, auto_src = resolve_ticker_auto_with_retry(
            name,
            use_ai=bool(use_ai_ticker),
            api_key=ai_api_key,
            model=ai_model,
            provider=ai_provider,
            market_preference=market_preference,
        )
        if auto_ticker:
            add_candidate(auto_ticker, f"자동탐색({auto_src or 'source'})")
    return candidates, notes


def fetch_company_financial_summary_for_analysis(
    company_name: str,
    ticker_input: str,
    use_ai_ticker: bool,
    ai_provider: str,
    ai_api_key: str,
    ai_model: str,
    market_preference: str = "",
) -> tuple[str, dict, str, str]:
    candidates, notes = _collect_analysis_ticker_candidates(
        company_name=company_name,
        ticker_input=ticker_input,
        use_ai_ticker=use_ai_ticker,
        ai_provider=ai_provider,
        ai_api_key=ai_api_key,
        ai_model=ai_model,
        market_preference=market_preference,
    )
    if not candidates:
        return "", {}, "", "유효한 티커를 찾지 못했습니다."

    best_ticker = ""
    best_summary: dict = {}
    best_source = ""
    best_score = -1.0
    errs: list[str] = []

    for cand in candidates:
        summary, err, source = fetch_company_financial_summary_multi_source(cand)
        score = float(_financial_summary_quality_score(summary))
        if summary:
            ref_name = str(summary.get("name") or cand).strip()
            sim = _name_similarity(company_name, ref_name) if company_name else 0.0
            score += sim * 3.0
        if summary and score > best_score:
            best_score = score
            best_ticker = cand
            best_summary = dict(summary)
            best_source = source or "multi_source"
        if err:
            errs.append(f"{cand}: {err}")

    tried_text = ", ".join(candidates)
    route_text = " -> ".join(notes) if notes else tried_text
    if best_summary:
        source_text = f"{best_source} | 시도: {tried_text}"
        if route_text:
            source_text = f"{source_text} | 경로: {route_text}"
        return best_ticker, best_summary, source_text, ""

    uniq_errs = []
    seen = set()
    for msg in errs:
        m = str(msg or "").strip()
        if not m or m in seen:
            continue
        seen.add(m)
        uniq_errs.append(m)
    err_text = " | ".join(uniq_errs[:6]) if uniq_errs else "기업 재무 정보를 불러오지 못했습니다."
    if tried_text:
        err_text = f"{err_text} | 시도 티커: {tried_text}"
    return candidates[0], {}, "", err_text


def _lines_to_text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        lines = [str(v).strip() for v in value if str(v).strip()]
        return "\n".join([f"- {line}" for line in lines])
    return str(value).strip()


def _split_report_lines(text: str) -> list[str]:
    raw = str(text or "").strip()
    if not raw:
        return []
    lines = []
    for line in raw.splitlines():
        ln = str(line or "").strip()
        if not ln:
            continue
        ln = re.sub(r"^[\-•\*\u2022]\s*", "", ln).strip()
        if ln:
            lines.append(ln)
    return lines


def _safe_report_filename(company_name: str, ticker: str, analysis_date_value) -> str:
    date_text = ""
    try:
        if hasattr(analysis_date_value, "isoformat"):
            date_text = str(analysis_date_value.isoformat())
        else:
            date_text = str(analysis_date_value or "").strip()
    except Exception:
        date_text = ""
    base = f"company_report_{company_name}_{ticker}_{date_text}".strip("_")
    safe = re.sub(r"[^0-9A-Za-z가-힣._-]+", "_", base)
    safe = re.sub(r"_+", "_", safe).strip("_.")
    if not safe:
        safe = f"company_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    return f"{safe}.docx"


def _fmt_ratio_brief(value, digits: int = 2) -> str:
    num = _safe_to_float(value)
    if num is None:
        return "데이터 없음"
    return f"{num:,.{digits}f}"


def _set_docx_run_noto(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    if run is None:
        return
    run.font.name = "Noto Sans KR"
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bool(bold)
    try:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Noto Sans KR")
        rfonts.set(qn("w:hAnsi"), "Noto Sans KR")
        rfonts.set(qn("w:eastAsia"), "Noto Sans KR")
        rfonts.set(qn("w:cs"), "Noto Sans KR")
    except Exception:
        pass


def _set_docx_paragraph_noto(paragraph, size_pt: float | None = None, bold: bool | None = None) -> None:
    if paragraph is None:
        return
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        _set_docx_run_noto(run, size_pt=size_pt, bold=bold)
    try:
        if size_pt is not None:
            paragraph.paragraph_format.space_after = Pt(4)
    except Exception:
        pass


def _set_docx_style_noto(doc, style_name: str, size_pt: float, bold: bool = False) -> None:
    try:
        style = doc.styles[style_name]
    except Exception:
        return
    style.font.name = "Noto Sans KR"
    style.font.size = Pt(size_pt)
    style.font.bold = bool(bold)
    try:
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), "Noto Sans KR")
        rfonts.set(qn("w:hAnsi"), "Noto Sans KR")
        rfonts.set(qn("w:eastAsia"), "Noto Sans KR")
        rfonts.set(qn("w:cs"), "Noto Sans KR")
    except Exception:
        pass


def _docx_add_heading(doc, text: str, level: int = 1):
    style_name = "Title" if level == 0 else ("Heading 1" if level == 1 else "Heading 2")
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(str(text or "").strip())
    _set_docx_run_noto(run, size_pt=(24 if level == 0 else 15 if level == 1 else 12), bold=True)
    return p


def _docx_add_paragraph(doc, text: str, size_pt: float = 11, bold: bool = False, style_name: str = ""):
    p = doc.add_paragraph(style=style_name) if style_name else doc.add_paragraph()
    run = p.add_run(str(text or ""))
    _set_docx_run_noto(run, size_pt=size_pt, bold=bold)
    return p


def _docx_add_bullet(doc, text: str, size_pt: float = 11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(str(text or ""))
    _set_docx_run_noto(run, size_pt=size_pt, bold=False)
    return p


def _fmt_statement_value(value) -> str:
    num = _safe_to_float(value)
    if num is None:
        return "-"
    return f"{num:,.0f}"


def _docx_add_financial_statement_table(doc, title: str, records: list[dict], max_rows: int = 12) -> bool:
    rows = [r for r in (records or []) if isinstance(r, dict) and str(r.get("item") or "").strip()]
    if not rows:
        return False

    year_cols: list[str] = []
    for row in rows:
        for key in row.keys():
            k = str(key or "").strip()
            if not k or k == "item":
                continue
            if k not in year_cols:
                year_cols.append(k)
    year_cols = year_cols[:4]
    if not year_cols:
        return False

    _docx_add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=1 + len(year_cols))
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "항목"
    for idx, year in enumerate(year_cols, start=1):
        header[idx].text = str(year)

    for row in rows[:max_rows]:
        tr = table.add_row().cells
        tr[0].text = str(row.get("item") or "").strip()
        for idx, year in enumerate(year_cols, start=1):
            tr[idx].text = _fmt_statement_value(row.get(year))

    for ridx, tr in enumerate(table.rows):
        for cell in tr.cells:
            for p in cell.paragraphs:
                _set_docx_paragraph_noto(p, size_pt=10.5, bold=(ridx == 0))
    return True


def build_company_analysis_docx_bytes(
    company_name: str,
    ticker: str,
    analysis_date_value,
    analysis_fields: dict,
    financial_summary: dict | None = None,
) -> bytes:
    if not HAS_PYTHON_DOCX:
        raise RuntimeError("python-docx 패키지가 설치되어 있지 않습니다.")

    doc = Document()
    _set_docx_style_noto(doc, "Normal", size_pt=11, bold=False)
    _set_docx_style_noto(doc, "Title", size_pt=24, bold=True)
    _set_docx_style_noto(doc, "Heading 1", size_pt=15, bold=True)
    _set_docx_style_noto(doc, "Heading 2", size_pt=12, bold=True)

    title_company = str(company_name or "기업명 미입력").strip()
    title_ticker = clean_valid_ticker(str(ticker or "").strip())
    title_date = ""
    try:
        if hasattr(analysis_date_value, "strftime"):
            title_date = analysis_date_value.strftime("%Y-%m-%d")
        elif hasattr(analysis_date_value, "isoformat"):
            title_date = str(analysis_date_value.isoformat())
        else:
            title_date = str(analysis_date_value or "").strip()
    except Exception:
        title_date = str(analysis_date_value or "").strip()

    _docx_add_heading(doc, f"{title_company} 기업 분석 보고서", level=0)
    subtitle = f"기준일: {title_date or '미입력'}"
    if title_ticker:
        subtitle = f"{subtitle} | 티커: {title_ticker}"
    _docx_add_paragraph(doc, subtitle, size_pt=11)

    fs = financial_summary if isinstance(financial_summary, dict) else {}
    if fs:
        _docx_add_heading(doc, "기초 정보", level=1)
        name = str(fs.get("name") or "").strip()
        sector = str(fs.get("sector") or fs.get("industry") or "").strip()
        country = str(fs.get("country") or "").strip()
        market_cap = _fmt_num_brief(fs.get("market_cap"))
        revenue = _fmt_num_brief(fs.get("total_revenue"))
        if name:
            _docx_add_paragraph(doc, f"회사명: {name}", size_pt=11)
        if sector:
            _docx_add_paragraph(doc, f"산업섹터: {sector}", size_pt=11)
        if country:
            _docx_add_paragraph(doc, f"국가: {country}", size_pt=11)
        _docx_add_paragraph(doc, f"시가총액(추정): {market_cap}", size_pt=11)
        _docx_add_paragraph(doc, f"매출 규모(최근): {revenue}", size_pt=11)

        _docx_add_heading(doc, "핵심 재무지표", level=1)
        metric_rows = [
            ("시가총액", _fmt_num_brief(fs.get("market_cap"))),
            ("매출", _fmt_num_brief(fs.get("total_revenue"))),
            ("EBITDA", _fmt_num_brief(fs.get("ebitda"))),
            ("순이익", _fmt_num_brief(fs.get("net_income_to_common"))),
            ("영업현금흐름", _fmt_num_brief(fs.get("operating_cashflow"))),
            ("잉여현금흐름", _fmt_num_brief(fs.get("free_cashflow"))),
            ("매출성장률(%)", _fmt_pct_brief(fs.get("revenue_growth_pct"))),
            ("이익성장률(%)", _fmt_pct_brief(fs.get("earnings_growth_pct"))),
            ("ROE(%)", _fmt_pct_brief(fs.get("roe_pct"))),
            ("영업이익률(%)", _fmt_pct_brief(fs.get("operating_margin_pct"))),
            ("매출총이익률(%)", _fmt_pct_brief(fs.get("gross_margin_pct"))),
            ("배당수익률(%)", _fmt_pct_brief(fs.get("dividend_yield_pct"))),
            ("부채비율(D/E)", _fmt_ratio_brief(fs.get("debt_to_equity"))),
            ("유동비율", _fmt_ratio_brief(fs.get("current_ratio"))),
            ("PER", _fmt_ratio_brief(fs.get("trailing_pe"))),
            ("Forward PER", _fmt_ratio_brief(fs.get("forward_pe"))),
            ("PBR", _fmt_ratio_brief(fs.get("price_to_book"))),
            ("Beta", _fmt_ratio_brief(fs.get("beta"), digits=3)),
        ]
        metric_table = doc.add_table(rows=1, cols=2)
        metric_table.style = "Table Grid"
        metric_table.rows[0].cells[0].text = "지표"
        metric_table.rows[0].cells[1].text = "값"
        for label, val in metric_rows:
            row_cells = metric_table.add_row().cells
            row_cells[0].text = str(label)
            row_cells[1].text = str(val)
        for ridx, row in enumerate(metric_table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    _set_docx_paragraph_noto(p, size_pt=10.5, bold=(ridx == 0))

        _docx_add_heading(doc, "재무제표 요약(연간)", level=1)
        has_stmt = False
        has_stmt = _docx_add_financial_statement_table(
            doc,
            "손익계산서(최근 연도)",
            fs.get("income_statement_annual") if isinstance(fs.get("income_statement_annual"), list) else [],
            max_rows=10,
        ) or has_stmt
        has_stmt = _docx_add_financial_statement_table(
            doc,
            "재무상태표(최근 연도)",
            fs.get("balance_sheet_annual") if isinstance(fs.get("balance_sheet_annual"), list) else [],
            max_rows=10,
        ) or has_stmt
        has_stmt = _docx_add_financial_statement_table(
            doc,
            "현금흐름표(최근 연도)",
            fs.get("cashflow_annual") if isinstance(fs.get("cashflow_annual"), list) else [],
            max_rows=10,
        ) or has_stmt
        if not has_stmt:
            _docx_add_paragraph(doc, "재무제표 연간 데이터가 없어 핵심 재무지표만 표시했습니다.", size_pt=10.5)

    sections = [
        ("기업 개요", analysis_fields.get("company_overview", "")),
        ("핵심 제품/서비스·돈 버는 방식", analysis_fields.get("products_services", "")),
        ("핵심 원재료/투입요소", analysis_fields.get("raw_materials", "")),
        ("이익 증가 요인·좋은 변화(촉매)", analysis_fields.get("profit_up_factors", "")),
        ("이익 감소 요인(리스크)", analysis_fields.get("profit_down_factors", "")),
        ("요약 메모", analysis_fields.get("key_takeaway", "")),
    ]
    for section_title, section_text in sections:
        _docx_add_heading(doc, section_title, level=1)
        lines = _split_report_lines(section_text)
        if lines:
            for ln in lines:
                _docx_add_bullet(doc, ln, size_pt=11)
        else:
            fallback_text = str(section_text or "").strip()
            _docx_add_paragraph(doc, fallback_text or "내용 없음", size_pt=11)

    for p in doc.paragraphs:
        style_name = str(getattr(getattr(p, "style", None), "name", "") or "")
        if style_name == "Title":
            _set_docx_paragraph_noto(p, size_pt=24, bold=True)
        elif style_name == "Heading 1":
            _set_docx_paragraph_noto(p, size_pt=15, bold=True)
        elif style_name == "Heading 2":
            _set_docx_paragraph_noto(p, size_pt=12, bold=True)
        else:
            _set_docx_paragraph_noto(p, size_pt=11, bold=None)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def generate_company_analysis_with_ai(
    company_name: str,
    ticker: str,
    financial_summary: dict,
    api_key: str,
    model: str,
    provider: str = "openai",
    google_context_text: str = "",
    google_research_note: str = "",
) -> tuple[dict, str, str]:
    if not str(google_context_text or "").strip():
        google_context_text, google_research_note = fetch_google_company_research_context(
            company_name=company_name,
            ticker=ticker,
            max_items=8,
        )
    user_prompt = f"""
회사명: {company_name}
티커: {ticker}
재무 데이터(JSON):
{json.dumps(financial_summary, ensure_ascii=False)}

구글 웹검색 발췌(기업 실체 정보, 우선 참고):
{google_context_text or "구글 검색 컨텍스트 없음"}

위 정보를 바탕으로 장기투자 관점의 기업 분석을 한국어로 작성해 주세요.
반드시 JSON 객체만 출력하세요. 키는 아래와 같습니다.
- company_overview: string (4~7문장)
- products_services: array[string] (실제 제품/서비스/브랜드/고객군/지역을 구체명으로 6~12개)
- raw_materials: array[string] (실제 원재료/부품/연료/물류/설비/외주요소를 구체명으로 6~12개)
- profit_up_factors: array[string] (이익 증가 조건 + 좋은 변화/촉매 6~12개)
- profit_down_factors: array[string] (이익 감소 조건 + 핵심 리스크 6~12개)
- key_takeaway: string (핵심 결론 + 체크포인트 3~6문장)

반드시 포함할 내용:
1) 회사가 정확히 무엇을 하는지(사업모델/고객/주요 지역/판매 채널)
2) 실제 판매 제품/서비스명(고유명사)을 구체적으로 제시
3) 실제 원재료/투입요소(고유명사)와 원가 민감도
4) 어떤 상황에서 이익이 늘고/줄어드는지(구체 트리거)
5) 기업에 유리한 변화(촉매)와 불리한 리스크
6) 미래 아이템/신규 사업/신규 투자 파이프라인 최소 3개 이상

주의:
- 추정 표현은 명시(예: 추정/가정)
- 일반론 문장(예: "판매량×단가×마진")만 반복하지 말 것
- 구글 발췌에서 제품명/원재료명/프로젝트명 등 고유명사를 최소 10개 이상 반영
- 확인이 불확실한 항목은 '(확인필요)'를 붙일 것
"""
    system_prompt = "너는 재무 데이터 기반 기업분석 어시스턴트다. 반드시 JSON만 출력한다."
    text, err = call_ai_text(
        provider=provider,
        api_key=api_key,
        model=model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        temperature=0.2,
        max_output_tokens=1400,
        timeout_sec=35,
    )
    if err:
        return {}, f"AI 생성 실패: {err}", google_research_note

    parsed = _extract_json_from_text(text)
    if not parsed:
        return {}, f"AI 응답에서 JSON을 파싱하지 못했습니다. 원인: {_json_parse_failure_reason(text)}", google_research_note

    analysis_keys = {
        "company_overview",
        "products_services",
        "raw_materials",
        "profit_up_factors",
        "profit_down_factors",
        "key_takeaway",
    }

    def _extract_analysis_obj(value) -> dict:
        if isinstance(value, dict):
            nested_candidates = [
                value,
                value.get("analysis"),
                value.get("result"),
                value.get("data"),
                value.get("output"),
                value.get("response"),
            ]
            for cand in nested_candidates:
                if isinstance(cand, dict):
                    if analysis_keys.intersection(set(cand.keys())):
                        return cand
            for cand in nested_candidates:
                if isinstance(cand, dict):
                    return cand
            return {}
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict) and analysis_keys.intersection(set(item.keys())):
                    return item
            for item in value:
                if isinstance(item, dict):
                    return item
        return {}

    parsed_obj = _extract_analysis_obj(parsed)
    if not parsed_obj:
        return {}, "AI 응답 JSON 구조가 예상과 달라 분석 항목을 찾지 못했습니다.", google_research_note

    def _pick(obj: dict, *keys):
        for key in keys:
            val = obj.get(key)
            if val is None:
                continue
            if isinstance(val, str) and not val.strip():
                continue
            if isinstance(val, list) and not val:
                continue
            return val
        return None

    analysis = {
        "company_overview": _lines_to_text(
            _pick(parsed_obj, "company_overview", "overview", "companyOverview", "기업개요")
        ),
        "products_services": _lines_to_text(
            _pick(parsed_obj, "products_services", "products", "services", "productsServices", "핵심제품서비스")
        ),
        "raw_materials": _lines_to_text(
            _pick(parsed_obj, "raw_materials", "inputs", "cost_drivers", "rawMaterials", "핵심원재료투입요소")
        ),
        "profit_up_factors": _lines_to_text(
            _pick(parsed_obj, "profit_up_factors", "upside", "profitUpFactors", "이익증가요인")
        ),
        "profit_down_factors": _lines_to_text(
            _pick(parsed_obj, "profit_down_factors", "risks", "downside", "profitDownFactors", "이익감소요인")
        ),
        "key_takeaway": _lines_to_text(
            _pick(parsed_obj, "key_takeaway", "takeaway", "summary", "keyTakeaway", "요약메모")
        ),
    }
    if not any((analysis.get(k) or "").strip() for k in analysis.keys()):
        return {}, "AI 응답 JSON에 분석 본문이 비어 있습니다.", google_research_note
    return analysis, "", google_research_note


def _fmt_num_brief(value) -> str:
    num = _safe_to_float(value)
    if num is None:
        return "데이터 없음"
    abs_num = abs(num)
    if abs_num >= 1_000_000_000_000:
        return f"{num / 1_000_000_000_000:.2f}T"
    if abs_num >= 1_000_000_000:
        return f"{num / 1_000_000_000:.2f}B"
    if abs_num >= 1_000_000:
        return f"{num / 1_000_000:.2f}M"
    return f"{num:,.0f}"


def _fmt_pct_brief(value) -> str:
    num = _safe_to_float(value)
    if num is None:
        return "데이터 없음"
    return f"{num:.2f}%"


def _extract_facts_from_google_context(google_context_text: str, limit: int = 8) -> list[str]:
    text = str(google_context_text or "").strip()
    if not text:
        return []
    facts: list[str] = []
    for line in text.splitlines():
        ln = str(line or "").strip()
        if not ln:
            continue
        ln = re.sub(r"^\d+\.\s*", "", ln)
        ln = re.sub(r"\s*\|\s*url=https?://\S+\s*$", "", ln)
        ln = re.sub(r"\s*\|\s*source=[^|]+", "", ln)
        ln = re.sub(r"\s*title=", "", ln)
        ln = re.sub(r"\s*snippet=", " / ", ln)
        ln = re.sub(r"\s+", " ", ln).strip(" -")
        if len(ln) < 18:
            continue
        facts.append(ln)
        if len(facts) >= max(1, int(limit)):
            break
    return facts


def generate_company_analysis_template(
    company_name: str,
    ticker: str,
    financial_summary: dict,
    google_context_text: str = "",
) -> dict:
    summary = financial_summary if isinstance(financial_summary, dict) else {}
    name = str(summary.get("name") or company_name or ticker or "해당 기업").strip()
    tkr = clean_valid_ticker(ticker)
    sector = str(summary.get("sector") or summary.get("industry") or "미분류").strip()
    country = str(summary.get("country") or "").strip()

    market_cap_text = _fmt_num_brief(summary.get("market_cap"))
    revenue_text = _fmt_num_brief(summary.get("total_revenue"))
    rev_growth_text = _fmt_pct_brief(summary.get("revenue_growth_pct"))
    roe_text = _fmt_pct_brief(summary.get("roe_pct"))
    pe_text = _fmt_num_brief(summary.get("trailing_pe"))
    pb_text = _fmt_num_brief(summary.get("price_to_book"))
    debt_text = _fmt_num_brief(summary.get("debt_to_equity"))
    google_facts = _extract_facts_from_google_context(google_context_text, limit=8)

    base_overview = [
        f"{name}({tkr or '티커 미확정'})는 {sector} 섹터에서 사업을 운영하는 기업입니다.",
        "핵심은 어떤 상품/서비스를 누구에게 판매하는지와 해당 시장의 수급 구조입니다.",
    ]
    if country:
        base_overview.append(f"주요 상장/사업 국가는 {country}입니다.")
    base_overview.append(f"시가총액 추정치는 {market_cap_text}, 최근 매출 규모는 {revenue_text}입니다.")
    if not _is_missing_summary_value(summary.get("business_summary")):
        base_overview.append("사업 개요는 확보된 재무 데이터의 회사 설명을 반영했습니다.")
    else:
        base_overview.append("세부 사업설명은 추가 공시/IR 자료 확인이 필요합니다.")
    if google_facts:
        base_overview.append(f"구글 검색 기반 핵심 사실: {google_facts[0]}")

    products = [
        f"주요 매출원: {sector} 관련 제품/서비스 판매",
        "돈 버는 구조: 판매량(물동량) × 단가(가격/운임) × 마진",
        "수익성 레버: 고정비 레버리지, 가동률, 단위당 원가",
        "고객/시장 구조: 주요 고객군·계약 방식·지역 노출 점검",
        "실적 확인 포인트: 매출 성장률, 영업이익률, 현금흐름의 동행 여부",
    ]
    if google_facts:
        for fact in google_facts[:4]:
            products.append(f"구글 검색 기반 제품/사업 사실: {fact}")
    raw_materials = [
        "원재료/부품 조달 단가 및 가격 전가 가능성",
        "인건비·에너지비·운영비 고정비 부담",
        "물류/운송비와 공급망 병목",
        "환율·금리·차입비용 변동",
        "규제·환경비용·유지보수(CAPEX) 부담",
    ]
    if google_facts:
        for fact in google_facts[2:6]:
            raw_materials.append(f"구글 검색 기반 원가/투입요소 단서: {fact}")
    up_factors = [
        f"매출 성장률 개선 시 이익 증가 가능 ({rev_growth_text})",
        f"수익성 지표 개선 여지 확대 (ROE {roe_text})",
        "수요 회복/가격(운임) 상승/스프레드 확대",
        "원가 하락 또는 가격 전가 성공",
        "가동률 상승 및 고정비 레버리지 효과",
        "신규 고객·시장 진입, 제품 믹스 고도화",
        "비핵심 자산 정리·재무구조 개선",
        "주주환원 확대/정책 수혜 등 긍정 촉매",
    ]
    if google_facts:
        up_factors.append(f"미래 아이템/확장 단서(구글): {google_facts[-1]}")
    down_factors = [
        f"밸류에이션 부담 가능성 (PER {pe_text}, PBR {pb_text})",
        f"재무 레버리지 부담 (부채비율/유사 지표 {debt_text})",
        "수요 둔화·가격(운임) 하락·스프레드 축소",
        "원자재/환율/금리 변동성 확대",
        "가동률 하락 및 고정비 부담 심화",
        "공급과잉·경쟁 심화·점유율 하락",
        "규제/정책/환경 리스크",
        "대규모 CAPEX 집행 실패 또는 현금흐름 악화",
    ]
    if google_facts:
        down_factors.append("구글 검색 문구 중 불확실 정보는 반드시 확인 필요(오탐 가능성).")
    takeaway = [
        f"{name}는 {sector} 업황과 가격/비용 사이클에 실적 민감도가 큰 편입니다.",
        "이익은 수요·단가·원가·가동률 조합에서 결정되므로 분기별 지표를 함께 추적해야 합니다.",
        "좋은 변화(촉매): 수요 회복, 가격 전가, 원가 안정, 재무구조 개선, 정책 수혜.",
        "핵심 리스크: 수요 둔화, 가격 하락, 비용 상승, 레버리지 부담, 규제 변화.",
        "체크포인트: 매출/마진/현금흐름/가이던스의 동행 여부.",
    ]

    return {
        "company_overview": _lines_to_text(base_overview),
        "products_services": _lines_to_text(products),
        "raw_materials": _lines_to_text(raw_materials),
        "profit_up_factors": _lines_to_text(up_factors),
        "profit_down_factors": _lines_to_text(down_factors),
        "key_takeaway": _lines_to_text(takeaway),
    }


def save_company_analysis(
    analysis_date: date,
    stock_name: str,
    ticker: str,
    financial_summary: dict,
    analysis: dict,
    source: str,
    ai_model: str,
    note: str,
) -> None:
    date_str = analysis_date.isoformat()
    now_str = datetime.now().isoformat(timespec="seconds")
    conn = get_conn()
    try:
        financial_payload = financial_summary or {}
        if not financial_payload:
            same_row = conn.execute(
                """
                SELECT financial_summary_json
                FROM company_analysis
                WHERE analysis_date = ? AND stock_name = ?
                LIMIT 1
                """,
                (date_str, stock_name),
            ).fetchone()
            if same_row and same_row[0]:
                financial_payload = parse_financial_summary_json(same_row[0])
            else:
                latest_row = conn.execute(
                    """
                    SELECT financial_summary_json
                    FROM company_analysis
                    WHERE stock_name = ? AND financial_summary_json IS NOT NULL AND financial_summary_json != ''
                    ORDER BY analysis_date DESC, updated_at DESC
                    LIMIT 1
                    """,
                    (stock_name,),
                ).fetchone()
                if latest_row and latest_row[0]:
                    financial_payload = parse_financial_summary_json(latest_row[0])

        conn.execute(
            "DELETE FROM company_analysis WHERE analysis_date = ? AND stock_name = ?",
            (date_str, stock_name),
        )
        conn.execute(
            """
            INSERT INTO company_analysis (
                analysis_date,
                stock_name,
                ticker,
                company_overview,
                products_services,
                raw_materials,
                profit_up_factors,
                profit_down_factors,
                financial_summary_json,
                source,
                ai_model,
                note,
                created_at,
                updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                date_str,
                stock_name,
                ticker or None,
                analysis.get("company_overview") or None,
                analysis.get("products_services") or None,
                analysis.get("raw_materials") or None,
                analysis.get("profit_up_factors") or None,
                analysis.get("profit_down_factors") or None,
                json.dumps(financial_payload, ensure_ascii=False) if financial_payload else None,
                source or "ai",
                ai_model or DEFAULT_AI_MODEL,
                note or None,
                now_str,
                now_str,
            ),
        )
        conn.commit()
    finally:
        conn.close()


def load_company_analysis_history(stock_name: str | None = None) -> pd.DataFrame:
    conn = get_conn()
    try:
        if stock_name:
            df = pd.read_sql_query(
                """
                SELECT *
                FROM company_analysis
                WHERE stock_name = ?
                ORDER BY analysis_date DESC, updated_at DESC
                """,
                conn,
                params=(stock_name,),
            )
        else:
            df = pd.read_sql_query(
                """
                SELECT *
                FROM company_analysis
                ORDER BY analysis_date DESC, stock_name
                """,
                conn,
            )
    finally:
        conn.close()

    if not df.empty:
        df["analysis_date"] = pd.to_datetime(df["analysis_date"])
    return df


def parse_financial_summary_json(value: str | None) -> dict:
    if not value:
        return {}
    try:
        return json.loads(value)
    except Exception:
        return {}


def load_company_list() -> pd.DataFrame:
    conn = get_conn()
    try:
        df = pd.read_sql_query(
            """
            SELECT stock_name, ticker, sector, price_krw, price_source, price_updated_at, source, created_at, updated_at
            FROM company_list
            ORDER BY stock_name
            """,
            conn,
        )
    finally:
        conn.close()
    return df


def get_company_list_ticker(stock_name: str) -> str:
    name = (stock_name or "").strip()
    if not name:
        return ""
    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT ticker FROM company_list WHERE stock_name = ?",
            (name,),
        ).fetchone()
    finally:
        conn.close()
    raw = (row[0] or "").strip().upper() if row and row[0] else ""
    saved = clean_valid_ticker(raw)
    if saved:
        return saved
    return get_builtin_ticker_hint(name)


def get_company_list_sector(stock_name: str) -> str:
    name = (stock_name or "").strip()
    if not name:
        return ""
    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT sector FROM company_list WHERE stock_name = ?",
            (name,),
        ).fetchone()
    finally:
        conn.close()
    return (row[0] or "").strip() if row and row[0] else ""


def build_company_price_krw_maps(company_list_df: pd.DataFrame | None = None) -> tuple[dict[str, float], dict[str, float]]:
    exact_map: dict[str, float] = {}
    norm_last_price: dict[str, float] = {}
    norm_count: dict[str, int] = {}

    src_df = company_list_df if isinstance(company_list_df, pd.DataFrame) else load_company_list()
    if src_df is None or src_df.empty:
        return exact_map, {}

    for _, row in src_df.iterrows():
        stock_name = str(row.get("stock_name") or "").strip()
        if not stock_name:
            continue
        price_val = _safe_to_float(row.get("price_krw"))
        if price_val is None or price_val <= 0:
            continue
        price = float(price_val)
        exact_map[stock_name] = price
        norm_name = normalize_company_name_for_match(stock_name)
        if norm_name:
            norm_last_price[norm_name] = price
            norm_count[norm_name] = int(norm_count.get(norm_name, 0)) + 1

    norm_unique_map = {k: v for k, v in norm_last_price.items() if int(norm_count.get(k, 0)) == 1}
    return exact_map, norm_unique_map


def lookup_company_price_krw(stock_name: str, exact_map: dict[str, float], norm_map: dict[str, float]) -> float | None:
    name = str(stock_name or "").strip()
    if not name:
        return None
    exact = _safe_to_float(exact_map.get(name))
    if exact is not None and exact > 0:
        return float(exact)

    norm_name = normalize_company_name_for_match(name)
    if not norm_name:
        return None
    norm_price = _safe_to_float(norm_map.get(norm_name))
    if norm_price is None or norm_price <= 0:
        return None
    return float(norm_price)


def build_price_series_with_company_fallback(
    names: pd.Series,
    qty: pd.Series,
    value_krw: pd.Series,
    company_price_exact: dict[str, float],
    company_price_norm: dict[str, float],
) -> pd.Series:
    qty_num = pd.to_numeric(qty, errors="coerce")
    value_num = pd.to_numeric(value_krw, errors="coerce")
    calc_series = value_num / qty_num.replace(0, pd.NA)
    mapped_series = names.astype(str).apply(
        lambda nm: lookup_company_price_krw(nm, company_price_exact, company_price_norm)
    )
    mapped_series = pd.to_numeric(mapped_series, errors="coerce")
    merged = mapped_series.where(mapped_series > 0, calc_series)
    return merged.fillna(0.0)


def recalculate_portfolio_from_price_and_avg_buy(
    df: pd.DataFrame,
    usd_krw_rate: float,
    company_price_exact: dict[str, float],
    company_price_norm: dict[str, float],
) -> pd.DataFrame:
    if df is None or df.empty:
        return ensure_portfolio_columns(pd.DataFrame(columns=COLUMNS), usd_krw_rate, force_usd_rate=True)

    base = ensure_portfolio_columns(df.copy(), usd_krw_rate, force_usd_rate=True)
    base[COL_NAME] = base[COL_NAME].astype(str).str.strip()
    qty = pd.to_numeric(base[COL_QTY], errors="coerce")
    fx = pd.to_numeric(base[COL_FX_RATE], errors="coerce").replace(0, pd.NA)

    view = to_krw_view(base, usd_krw_rate, force_usd_rate=True)
    current_price_krw = build_price_series_with_company_fallback(
        names=base[COL_NAME],
        qty=qty,
        value_krw=view[COL_VALUE_KRW],
        company_price_exact=company_price_exact,
        company_price_norm=company_price_norm,
    )

    invest_krw_existing = view[COL_VALUE_KRW] - view[COL_PNL_KRW]
    avg_buy_krw = invest_krw_existing / qty.replace(0, pd.NA)

    recalc_value_krw = qty * current_price_krw
    recalc_invest_krw = qty * avg_buy_krw
    recalc_pnl_krw = recalc_value_krw - recalc_invest_krw
    recalc_return = (recalc_pnl_krw / recalc_invest_krw.replace(0, pd.NA)) * 100.0

    valid = (qty > 0) & (pd.to_numeric(current_price_krw, errors="coerce") > 0) & avg_buy_krw.notna() & fx.notna()

    value_num = pd.to_numeric(base[COL_VALUE], errors="coerce")
    pnl_num = pd.to_numeric(base[COL_PNL], errors="coerce")
    ret_num = pd.to_numeric(base[COL_RETURN], errors="coerce")

    recalc_value = recalc_value_krw / fx
    recalc_pnl = recalc_pnl_krw / fx

    base[COL_VALUE] = value_num.where(~valid, recalc_value)
    base[COL_PNL] = pnl_num.where(~valid, recalc_pnl)
    base[COL_RETURN] = ret_num.where(~valid, recalc_return)
    base[COL_PRICE_KRW] = pd.to_numeric(current_price_krw, errors="coerce")
    base[COL_AVG_BUY_KRW] = pd.to_numeric(avg_buy_krw, errors="coerce")
    return base


def upsert_company_list_entry(
    stock_name: str,
    ticker: str = "",
    sector: str = "",
    source: str | None = "manual",
    price_krw: float | None = None,
    price_source: str | None = None,
) -> None:
    name = (stock_name or "").strip()
    if not name:
        return

    tkr = clean_valid_ticker(ticker)
    sec = (sector or "").strip()
    price_val = _safe_to_float(price_krw)
    if price_val is not None and price_val <= 0:
        price_val = None
    p_source = (price_source or "").strip()
    source_text = None if source is None else str(source).strip()
    builtin_hint = get_builtin_ticker_hint(name)
    # 자동/일괄 경로에서는 내장 티커 힌트를 우선 적용해 오탐 저장을 방지한다.
    # 수동 편집(manual_edit)에서만 사용자의 명시 입력을 그대로 존중한다.
    if builtin_hint and not (source_text == "manual_edit" and tkr):
        tkr = builtin_hint
    now_str = datetime.now().isoformat(timespec="seconds")

    conn = get_conn()
    try:
        row = conn.execute(
            "SELECT ticker, sector, price_krw, price_source FROM company_list WHERE stock_name = ?",
            (name,),
        ).fetchone()
        existing_ticker = clean_valid_ticker((row[0] or "").strip().upper()) if row and row[0] else ""
        existing_sector = (row[1] or "").strip() if row and len(row) > 1 and row[1] else ""
        existing_price = _safe_to_float(row[2]) if row and len(row) > 2 else None
        existing_price_source = (row[3] or "").strip() if row and len(row) > 3 and row[3] else ""
        next_ticker = tkr or existing_ticker
        next_sector = sec or existing_sector
        next_price = price_val if price_val is not None else existing_price
        next_price_source = p_source or existing_price_source
        conn.execute(
            """
            INSERT INTO company_list (stock_name, ticker, sector, price_krw, price_source, price_updated_at, source, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(stock_name) DO UPDATE SET
                ticker = CASE
                    WHEN excluded.ticker IS NULL OR excluded.ticker = '' THEN company_list.ticker
                    ELSE excluded.ticker
                END,
                sector = CASE
                    WHEN excluded.sector IS NULL OR excluded.sector = '' THEN company_list.sector
                    ELSE excluded.sector
                END,
                price_krw = CASE
                    WHEN excluded.price_krw IS NULL OR excluded.price_krw <= 0 THEN company_list.price_krw
                    ELSE excluded.price_krw
                END,
                price_source = CASE
                    WHEN excluded.price_source IS NULL OR excluded.price_source = '' THEN company_list.price_source
                    ELSE excluded.price_source
                END,
                price_updated_at = CASE
                    WHEN excluded.price_krw IS NULL OR excluded.price_krw <= 0 THEN company_list.price_updated_at
                    ELSE excluded.price_updated_at
                END,
                source = CASE
                    WHEN excluded.source IS NULL OR excluded.source = '' THEN company_list.source
                    ELSE excluded.source
                END,
                updated_at = excluded.updated_at
            """,
            (
                name,
                next_ticker or None,
                next_sector or None,
                float(next_price) if next_price is not None else None,
                next_price_source or None,
                now_str if next_price is not None else None,
                source_text if source_text is not None else None,
                now_str,
                now_str,
            ),
        )
        conn.commit()
    finally:
        conn.close()


def delete_company_list_entry(stock_name: str) -> None:
    name = (stock_name or "").strip()
    if not name:
        return
    conn = get_conn()
    try:
        conn.execute("DELETE FROM company_list WHERE stock_name = ?", (name,))
        conn.commit()
    finally:
        conn.close()


def clear_company_list_ticker(stock_name: str, source: str = "manual_edit_clear") -> bool:
    name = (stock_name or "").strip()
    if not name:
        return False
    now_str = datetime.now().isoformat(timespec="seconds")
    conn = get_conn()
    try:
        row = conn.execute("SELECT stock_name FROM company_list WHERE stock_name = ?", (name,)).fetchone()
        if not row:
            return False
        conn.execute(
            """
            UPDATE company_list
            SET ticker = NULL,
                price_krw = NULL,
                price_source = NULL,
                price_updated_at = NULL,
                source = ?,
                updated_at = ?
            WHERE stock_name = ?
            """,
            ((source or "manual_edit_clear").strip(), now_str, name),
        )
        conn.commit()
        return True
    finally:
        conn.close()


def reconcile_builtin_ticker_overrides() -> int:
    df = load_company_list()
    if df is None or df.empty:
        return 0
    updated = 0
    for _, row in df.iterrows():
        name = str(row.get("stock_name") or "").strip()
        if not name:
            continue
        hint = get_builtin_ticker_hint(name)
        if not hint:
            continue
        current = clean_valid_ticker(str(row.get("ticker") or ""))
        if current == hint:
            continue
        upsert_company_list_entry(name, ticker=hint, source="builtin_override")
        updated += 1
    return updated


def clear_company_list_meta_all(source: str = "manual_meta_reset") -> int:
    now_str = datetime.now().isoformat(timespec="seconds")
    conn = get_conn()
    try:
        row = conn.execute("SELECT COUNT(*) FROM company_list").fetchone()
        total = int(row[0]) if row and row[0] is not None else 0
        if total <= 0:
            return 0
        conn.execute(
            """
            UPDATE company_list
            SET ticker = NULL,
                sector = NULL,
                price_krw = NULL,
                price_source = NULL,
                price_updated_at = NULL,
                source = ?,
                updated_at = ?
            """,
            ((source or "manual_meta_reset").strip(), now_str),
        )
        conn.commit()
        return total
    finally:
        conn.close()


def clear_company_list_meta_by_names(stock_names: list[str], source: str = "manual_meta_reset_selected") -> int:
    names = []
    seen = set()
    for raw in stock_names or []:
        name = str(raw or "").strip()
        if not name or name in seen:
            continue
        names.append(name)
        seen.add(name)
    if not names:
        return 0

    now_str = datetime.now().isoformat(timespec="seconds")
    placeholders = ",".join(["?"] * len(names))
    conn = get_conn()
    try:
        params = tuple(names)
        row = conn.execute(
            f"SELECT COUNT(*) FROM company_list WHERE stock_name IN ({placeholders})",
            params,
        ).fetchone()
        total = int(row[0]) if row and row[0] is not None else 0
        if total <= 0:
            return 0

        conn.execute(
            f"""
            UPDATE company_list
            SET ticker = NULL,
                sector = NULL,
                price_krw = NULL,
                price_source = NULL,
                price_updated_at = NULL,
                source = ?,
                updated_at = ?
            WHERE stock_name IN ({placeholders})
            """,
            ((source or "manual_meta_reset_selected").strip(), now_str, *names),
        )
        conn.commit()
        return total
    finally:
        conn.close()


def load_company_compare_sets() -> pd.DataFrame:
    conn = get_conn()
    try:
        df = pd.read_sql_query(
            """
            SELECT set_name, companies_json, metrics_json, weights_json, sector_filter, note, created_at, updated_at
            FROM company_compare_sets
            ORDER BY updated_at DESC, set_name
            """,
            conn,
        )
    finally:
        conn.close()
    return df


def save_company_compare_set(
    set_name: str,
    companies: list[str],
    metrics: list[str],
    weights: dict[str, float],
    sector_filter: str = "전체",
    note: str = "",
) -> None:
    name = (set_name or "").strip()
    if not name:
        return

    saved_companies = [str(v).strip() for v in companies if str(v).strip()]
    saved_metrics = [str(v).strip() for v in metrics if str(v).strip()]
    if not saved_companies or not saved_metrics:
        return

    cleaned_weights = {}
    for key, value in (weights or {}).items():
        k = str(key).strip()
        if not k:
            continue
        try:
            cleaned_weights[k] = float(value)
        except Exception:
            cleaned_weights[k] = 0.0

    now_str = datetime.now().isoformat(timespec="seconds")
    conn = get_conn()
    try:
        conn.execute(
            """
            INSERT INTO company_compare_sets (
                set_name, companies_json, metrics_json, weights_json, sector_filter, note, created_at, updated_at
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(set_name) DO UPDATE SET
                companies_json=excluded.companies_json,
                metrics_json=excluded.metrics_json,
                weights_json=excluded.weights_json,
                sector_filter=excluded.sector_filter,
                note=excluded.note,
                updated_at=excluded.updated_at
            """,
            (
                name,
                json.dumps(saved_companies, ensure_ascii=False),
                json.dumps(saved_metrics, ensure_ascii=False),
                json.dumps(cleaned_weights, ensure_ascii=False),
                (sector_filter or "전체").strip(),
                (note or "").strip() or None,
                now_str,
                now_str,
            ),
        )
        conn.commit()
    finally:
        conn.close()


def delete_company_compare_set(set_name: str) -> None:
    name = (set_name or "").strip()
    if not name:
        return
    conn = get_conn()
    try:
        conn.execute("DELETE FROM company_compare_sets WHERE set_name = ?", (name,))
        conn.commit()
    finally:
        conn.close()


def get_ai_settings_from_session(prefix: str) -> tuple[str, str, str]:
    _ = prefix
    provider = normalize_ai_provider(st.session_state.get("global_ai_provider", "claude"))
    openai_key = (st.session_state.get("global_openai_api_key", "") or "").strip()
    claude_key = (st.session_state.get("global_claude_api_key", "") or "").strip()
    openai_model = (st.session_state.get("global_openai_model", DEFAULT_OPENAI_MODEL) or DEFAULT_OPENAI_MODEL).strip()
    claude_model = (st.session_state.get("global_claude_model", DEFAULT_CLAUDE_MODEL) or DEFAULT_CLAUDE_MODEL).strip()

    if provider == "claude":
        api_key = claude_key
        model = claude_model
    else:
        api_key = openai_key
        model = openai_model
    return provider, api_key, model


def get_market_data_api_keys() -> tuple[str, str]:
    try:
        alpha_key = (st.session_state.get("global_alpha_vantage_api_key", "") or "").strip()
        finnhub_key = (st.session_state.get("global_finnhub_api_key", "") or "").strip()
    except Exception:
        return "", ""
    return alpha_key, finnhub_key


def _to_bool_flag(value) -> bool:
    return str(value or "").strip().lower() in {"1", "true", "t", "yes", "y", "on"}


def _coerce_string_list(value) -> list[str]:
    if value is None:
        return []
    if isinstance(value, str):
        text = value.strip()
        return [text] if text else []

    if isinstance(value, dict):
        return []

    items = []
    try:
        for v in value:
            text = str(v or "").strip()
            if text:
                items.append(text)
    except Exception:
        text = str(value or "").strip()
        if text:
            items.append(text)
    return items


def _coerce_choice(value, allowed: set[str], default: str) -> str:
    text = str(value or "").strip()
    return text if text in allowed else default


def _sanitize_widget_text(value, default: str = "") -> str:
    if isinstance(value, (list, tuple, set, dict)):
        return default
    text = str(value or "").strip()
    if not text:
        return default
    lowered = text.lower()
    if lowered in {"[object object]", "nan", "none", "null", "undefined"}:
        return default
    if re.match(r"^[\._]*arr\d+", lowered):
        return default
    return text


def _read_first_secret_or_env(keys: list[str]) -> str:
    for key in keys:
        value = ""
        key_variants = [str(key or "").strip()]
        upper_key = key_variants[0].upper()
        lower_key = key_variants[0].lower()
        if upper_key and upper_key not in key_variants:
            key_variants.append(upper_key)
        if lower_key and lower_key not in key_variants:
            key_variants.append(lower_key)
        for key_variant in key_variants:
            try:
                value = str(st.secrets.get(key_variant, "") or "").strip()
            except Exception:
                value = ""
            if value:
                break
        if not value:
            for key_variant in key_variants:
                value = str(os.getenv(key_variant, "") or "").strip()
                if value:
                    break
        if value:
            return value
    return ""


def _normalize_github_repo_value(repo_text: str) -> str:
    text = str(repo_text or "").strip()
    if not text:
        return ""
    # URL 입력(https://github.com/owner/repo(.git))도 owner/repo로 정규화
    m = re.search(r"github\.com[:/]+([^/\s]+)/([^/\s]+?)(?:\.git)?/?$", text, re.I)
    if m:
        return f"{m.group(1)}/{m.group(2)}"
    return text


def _load_github_settings_from_secrets() -> dict[str, str]:
    result = {
        "repo": "",
        "branch": "",
        "excel_path": "",
        "token": "",
        "sync_enabled": "",
        "sync_on_change": "",
    }

    # 1) 최상위 키
    repo = _read_first_secret_or_env(
        ["GITHUB_REPO", "GLOBAL_GITHUB_REPO", "GITHUB_REPO_URL", "GITHUB_URL", "GITHUB_REPOSITORY"]
    )
    branch = _read_first_secret_or_env(["GITHUB_BRANCH", "GLOBAL_GITHUB_BRANCH"])
    excel_path = _read_first_secret_or_env(
        ["GITHUB_EXCEL_PATH", "GLOBAL_GITHUB_EXCEL_PATH", "GITHUB_FILE_PATH", "GITHUB_PATH"]
    )
    token = _read_first_secret_or_env(
        ["GITHUB_TOKEN", "GH_TOKEN", "GITHUB_PAT", "GITHUB_ACCESS_TOKEN", "GH_PAT"]
    )
    sync_enabled = _read_first_secret_or_env(["GITHUB_SYNC_ENABLED", "GLOBAL_GITHUB_SYNC_ENABLED"])
    sync_on_change = _read_first_secret_or_env(["GITHUB_SYNC_ON_CHANGE", "GLOBAL_GITHUB_SYNC_ON_CHANGE"])

    result["repo"] = _normalize_github_repo_value(repo)
    result["branch"] = str(branch or "").strip()
    result["excel_path"] = str(excel_path or "").strip()
    result["token"] = str(token or "").strip()
    result["sync_enabled"] = str(sync_enabled or "").strip()
    result["sync_on_change"] = str(sync_on_change or "").strip()

    # 2) 중첩 키 [github] / [GITHUB] 지원
    for block_key in ["github", "GITHUB"]:
        block = {}
        try:
            maybe = st.secrets.get(block_key, {})
            if isinstance(maybe, dict):
                block = maybe
            else:
                try:
                    block = dict(maybe)
                except Exception:
                    block = {}
        except Exception:
            block = {}
        if not block:
            continue

        def _pick(*keys: str) -> str:
            for k in keys:
                val = str(block.get(k, "") or "").strip()
                if val:
                    return val
                up = str(k).upper()
                low = str(k).lower()
                val = str(block.get(up, "") or block.get(low, "") or "").strip()
                if val:
                    return val
            return ""

        if not result["repo"]:
            owner = _pick("owner")
            name = _pick("repo", "repository", "name")
            repo_url = _pick("repo_url", "url")
            result["repo"] = _normalize_github_repo_value(repo_url or (f"{owner}/{name}" if owner and name else name))
        if not result["branch"]:
            result["branch"] = _pick("branch")
        if not result["excel_path"]:
            result["excel_path"] = _pick("excel_path", "path", "file_path", "excel")
        if not result["token"]:
            result["token"] = _pick("token", "pat", "access_token", "github_token")
        if not result["sync_enabled"]:
            result["sync_enabled"] = _pick("sync_enabled", "enabled")
        if not result["sync_on_change"]:
            result["sync_on_change"] = _pick("sync_on_change", "on_change")

    return result


def get_github_sync_settings() -> dict[str, str | bool]:
    return {
        "enabled": _to_bool_flag(st.session_state.get("github_sync_enabled", False)),
        "repo": str(st.session_state.get("github_repo", "") or "").strip(),
        "branch": str(st.session_state.get("github_branch", "main") or "main").strip(),
        "excel_path": str(st.session_state.get("github_excel_path", "portfolio_auto.xlsx") or "").strip(),
        "token": str(st.session_state.get("github_token", "") or "").strip(),
    }


def fetch_excel_bytes_from_github(
    repo: str,
    path: str,
    branch: str = "main",
    token: str = "",
) -> tuple[bytes, str]:
    repo_text = str(repo or "").strip()
    path_text = str(path or "").strip()
    branch_text = str(branch or "main").strip() or "main"
    token_text = str(token or "").strip()
    if not repo_text or "/" not in repo_text:
        return b"", "GitHub repo 형식이 올바르지 않습니다. (예: owner/repo)"
    if not path_text:
        return b"", "GitHub 엑셀 경로를 입력해 주세요."

    headers = {"Accept": "application/vnd.github+json"}
    if token_text:
        headers["Authorization"] = f"Bearer {token_text}"

    url = f"https://api.github.com/repos/{repo_text}/contents/{path_text}"
    try:
        resp = requests.get(url, headers=headers, params={"ref": branch_text}, timeout=18)
        if resp.status_code == 404:
            return b"", "GitHub에 엑셀 파일이 아직 없습니다."
        resp.raise_for_status()
        payload = resp.json() or {}
        encoded = str(payload.get("content") or "").strip().replace("\n", "")
        if not encoded:
            return b"", "GitHub 파일 내용을 읽지 못했습니다."
        return base64.b64decode(encoded), ""
    except Exception as exc:
        return b"", f"GitHub 엑셀 다운로드 실패: {exc}"


def upload_excel_bytes_to_github(
    repo: str,
    path: str,
    branch: str,
    token: str,
    excel_bytes: bytes,
    commit_message: str,
) -> tuple[bool, str]:
    repo_text = str(repo or "").strip()
    path_text = str(path or "").strip()
    branch_text = str(branch or "main").strip() or "main"
    token_text = str(token or "").strip()
    if not repo_text or "/" not in repo_text:
        return False, "GitHub repo 형식이 올바르지 않습니다. (예: owner/repo)"
    if not path_text:
        return False, "GitHub 엑셀 경로를 입력해 주세요."
    if not token_text:
        return False, "GitHub Token이 필요합니다. (repo 권한)"
    if not excel_bytes:
        return False, "업로드할 엑셀 데이터가 비어 있습니다."

    headers = {
        "Accept": "application/vnd.github+json",
        "Authorization": f"Bearer {token_text}",
    }
    url = f"https://api.github.com/repos/{repo_text}/contents/{path_text}"

    existing_sha = ""
    try:
        existing_resp = requests.get(url, headers=headers, params={"ref": branch_text}, timeout=18)
        if existing_resp.status_code == 200:
            existing_payload = existing_resp.json() or {}
            existing_sha = str(existing_payload.get("sha") or "").strip()
        elif existing_resp.status_code not in {404}:
            existing_resp.raise_for_status()
    except Exception as exc:
        return False, f"GitHub 기존 파일 조회 실패: {exc}"

    body = {
        "message": commit_message,
        "content": base64.b64encode(excel_bytes).decode("utf-8"),
        "branch": branch_text,
    }
    if existing_sha:
        body["sha"] = existing_sha

    try:
        put_resp = requests.put(url, headers=headers, json=body, timeout=22)
        if put_resp.status_code not in {200, 201}:
            try:
                err_payload = put_resp.json() or {}
                err_msg = err_payload.get("message") or put_resp.text
            except Exception:
                err_msg = put_resp.text
            return False, f"GitHub 업로드 실패: {err_msg}"
        return True, "GitHub 엑셀 자동 저장 완료"
    except Exception as exc:
        return False, f"GitHub 업로드 실패: {exc}"


def load_all_snapshots_for_export() -> pd.DataFrame:
    conn = get_conn()
    try:
        df = pd.read_sql_query(
            """
            SELECT
                snapshot_date AS snapshot_date,
                stock_name AS stock_name,
                quantity AS quantity,
                COALESCE(currency, 'KRW') AS currency,
                COALESCE(fx_rate, 1) AS fx_rate,
                market_value AS market_value,
                pnl_value AS pnl_value,
                pnl_pct AS pnl_pct,
                created_at AS created_at
            FROM snapshots
            ORDER BY snapshot_date, market_value DESC, stock_name
            """,
            conn,
        )
    finally:
        conn.close()
    if df is None or df.empty:
        return pd.DataFrame(
            columns=[
                "snapshot_date",
                "stock_name",
                "quantity",
                "currency",
                "fx_rate",
                "market_value",
                "pnl_value",
                "pnl_pct",
                "created_at",
            ]
        )
    return df


def load_all_snapshot_cash_for_export() -> pd.DataFrame:
    conn = get_conn()
    try:
        df = pd.read_sql_query(
            """
            SELECT
                snapshot_date AS snapshot_date,
                COALESCE(cash_krw, 0) AS cash_krw,
                COALESCE(cash_usd, 0) AS cash_usd,
                created_at AS created_at,
                updated_at AS updated_at
            FROM snapshot_cash
            ORDER BY snapshot_date
            """,
            conn,
        )
    finally:
        conn.close()
    if df is None or df.empty:
        return pd.DataFrame(columns=["snapshot_date", "cash_krw", "cash_usd", "created_at", "updated_at"])
    return df


def build_portfolio_excel_bytes(snapshot_date: date, holdings_df: pd.DataFrame) -> bytes:
    usd_krw_rate = float(get_usd_krw_rate_for_date(snapshot_date)[0])
    view = ensure_portfolio_columns(holdings_df, usd_krw_rate, force_usd_rate=True).copy()
    for col in COLUMNS:
        if col not in view.columns:
            view[col] = pd.NA
    export_df = view[COLUMNS].copy()
    cash_krw, cash_usd = load_snapshot_cash(snapshot_date)
    meta_df = pd.DataFrame(
        [
            {
                "snapshot_date": snapshot_date.isoformat(),
                "cash_krw": float(cash_krw),
                "cash_usd": float(cash_usd),
                "saved_at": datetime.now().isoformat(timespec="seconds"),
                "usd_krw": usd_krw_rate,
            }
        ]
    )
    snapshot_rows = load_all_snapshots_for_export()
    cash_rows = load_all_snapshot_cash_for_export()

    out = BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        export_df.to_excel(writer, sheet_name="보유현황", index=False)
        meta_df.to_excel(writer, sheet_name="메타", index=False)
        snapshot_rows.to_excel(writer, sheet_name="스냅샷_보유현황", index=False)
        cash_rows.to_excel(writer, sheet_name="스냅샷_예수금", index=False)
    return out.getvalue()


def sync_snapshot_to_github_excel(snapshot_date: date, holdings_df: pd.DataFrame) -> tuple[bool, str]:
    cfg = get_github_sync_settings()
    if not bool(cfg["enabled"]):
        return False, ""

    repo = str(cfg["repo"] or "").strip()
    branch = str(cfg["branch"] or "main").strip() or "main"
    excel_path = str(cfg["excel_path"] or "").strip()
    token = str(cfg["token"] or "").strip()
    if not repo or not excel_path:
        return False, "GitHub 동기화가 켜져 있지만 repo/path 설정이 비어 있습니다."

    excel_bytes = build_portfolio_excel_bytes(snapshot_date, holdings_df)
    msg = (
        f"auto: portfolio snapshot {snapshot_date.isoformat()} "
        f"({datetime.now().isoformat(timespec='seconds')})"
    )
    ok, msg = upload_excel_bytes_to_github(
        repo=repo,
        path=excel_path,
        branch=branch,
        token=token,
        excel_bytes=excel_bytes,
        commit_message=msg,
    )
    if ok:
        try:
            st.session_state["uploaded_portfolio_excel_bytes"] = excel_bytes
            file_name = excel_path.split("/")[-1] if "/" in excel_path else excel_path
            st.session_state["uploaded_portfolio_excel_name"] = f"github:{file_name}"
            st.session_state["uploaded_portfolio_excel_sig"] = hashlib.sha256(excel_bytes).hexdigest()
        except Exception:
            pass
    return ok, msg


def bootstrap_excel_from_github_if_needed() -> None:
    cfg = get_github_sync_settings()
    if not bool(cfg["enabled"]):
        return

    try:
        has_uploaded = bool(st.session_state.get("uploaded_portfolio_excel_bytes", b""))
    except Exception:
        has_uploaded = False
    if has_uploaded:
        return

    repo = str(cfg["repo"] or "").strip()
    branch = str(cfg["branch"] or "main").strip() or "main"
    excel_path = str(cfg["excel_path"] or "").strip()
    token = str(cfg["token"] or "").strip()
    if not repo or not excel_path:
        return

    sync_sig = f"{repo}|{branch}|{excel_path}|{int(bool(token))}"
    if st.session_state.get("github_bootstrap_sig", "") == sync_sig:
        return
    st.session_state["github_bootstrap_sig"] = sync_sig

    excel_bytes, err = fetch_excel_bytes_from_github(
        repo=repo,
        path=excel_path,
        branch=branch,
        token=token,
    )
    if err:
        st.session_state["github_sync_notice"] = err
        return

    file_name = excel_path.split("/")[-1] if "/" in excel_path else excel_path
    st.session_state["uploaded_portfolio_excel_bytes"] = excel_bytes
    st.session_state["uploaded_portfolio_excel_name"] = f"github:{file_name}"
    st.session_state["uploaded_portfolio_excel_sig"] = hashlib.sha256(excel_bytes).hexdigest()
    st.session_state["editing_df_date"] = ""
    st.session_state["github_sync_notice"] = f"GitHub 엑셀 자동 불러오기 완료: {excel_path}"


def load_app_settings() -> dict[str, str]:
    conn = get_conn()
    try:
        rows = conn.execute("SELECT setting_key, setting_value FROM app_settings").fetchall()
    finally:
        conn.close()
    return {str(k): ("" if v is None else str(v)) for k, v in rows}


def save_app_settings(settings: dict[str, str]) -> None:
    now_str = datetime.now().isoformat(timespec="seconds")
    conn = get_conn()
    try:
        rows = [(str(k), str(v), now_str) for k, v in settings.items()]
        conn.executemany(
            """
            INSERT INTO app_settings (setting_key, setting_value, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(setting_key) DO UPDATE SET
                setting_value=excluded.setting_value,
                updated_at=excluded.updated_at
            """,
            rows,
        )
        conn.commit()
    finally:
        conn.close()


def save_app_settings_partial(settings: dict[str, str]) -> None:
    if not settings:
        return
    now_str = datetime.now().isoformat(timespec="seconds")
    conn = get_conn()
    try:
        rows = [(str(k), str(v), now_str) for k, v in settings.items()]
        conn.executemany(
            """
            INSERT INTO app_settings (setting_key, setting_value, updated_at)
            VALUES (?, ?, ?)
            ON CONFLICT(setting_key) DO UPDATE SET
                setting_value=excluded.setting_value,
                updated_at=excluded.updated_at
            """,
            rows,
        )
        conn.commit()
    finally:
        conn.close()


def initialize_api_settings(force: bool = False) -> None:
    settings = load_app_settings()
    store_sensitive = _to_bool_flag(settings.get("store_sensitive_keys", "false"))
    global_provider = normalize_ai_provider(settings.get("ai_provider", "claude"))
    global_openai_key = settings.get("openai_api_key", "") if store_sensitive else ""
    global_claude_key = settings.get("claude_api_key", "") if store_sensitive else ""
    global_alpha_key = settings.get("alpha_vantage_api_key", "") if store_sensitive else ""
    global_finnhub_key = settings.get("finnhub_api_key", "") if store_sensitive else ""
    global_openai_model = settings.get("openai_model", DEFAULT_OPENAI_MODEL) or DEFAULT_OPENAI_MODEL
    global_claude_model = settings.get("claude_model", DEFAULT_CLAUDE_MODEL) or DEFAULT_CLAUDE_MODEL
    github_sync_enabled_raw = str(settings.get("github_sync_enabled", "false") or "false").strip()
    github_sync_enabled = _to_bool_flag(github_sync_enabled_raw)
    github_sync_on_change = _to_bool_flag(settings.get("github_sync_on_change", "true"))
    github_repo = settings.get("github_repo", "")
    github_branch = settings.get("github_branch", "main") or "main"
    github_excel_path = settings.get("github_excel_path", "portfolio_auto.xlsx") or "portfolio_auto.xlsx"
    github_token = settings.get("github_token", "") if store_sensitive else ""
    daily_auto_enabled = _to_bool_flag(settings.get("daily_auto_snapshot_enabled", "false"))
    daily_auto_use_ai_ticker = _to_bool_flag(settings.get("daily_auto_snapshot_use_ai_ticker", "true"))
    daily_auto_hour_raw = settings.get("daily_auto_snapshot_hour", "18")
    try:
        daily_auto_hour = int(str(daily_auto_hour_raw).strip())
    except Exception:
        daily_auto_hour = 18
    daily_auto_hour = max(0, min(23, daily_auto_hour))
    daily_auto_last_run_date = str(settings.get("daily_auto_snapshot_last_run_date", "") or "").strip()
    daily_auto_last_run_at = str(settings.get("daily_auto_snapshot_last_run_at", "") or "").strip()
    daily_auto_last_attempt_date = str(settings.get("daily_auto_snapshot_last_attempt_date", "") or "").strip()
    daily_auto_last_summary = str(settings.get("daily_auto_snapshot_last_summary", "") or "").strip()

    # Secure source priority: secrets/env > DB
    global_openai_key = _read_first_secret_or_env(["OPENAI_API_KEY", "GLOBAL_OPENAI_API_KEY"]) or global_openai_key
    global_claude_key = _read_first_secret_or_env(["CLAUDE_API_KEY", "GLOBAL_CLAUDE_API_KEY"]) or global_claude_key
    global_alpha_key = _read_first_secret_or_env(["ALPHA_VANTAGE_API_KEY", "GLOBAL_ALPHA_VANTAGE_API_KEY"]) or global_alpha_key
    global_finnhub_key = _read_first_secret_or_env(["FINNHUB_API_KEY", "GLOBAL_FINNHUB_API_KEY"]) or global_finnhub_key
    gh_secret = _load_github_settings_from_secrets()
    github_token = (gh_secret.get("token") or "").strip() or github_token
    github_repo = _normalize_github_repo_value((gh_secret.get("repo") or "").strip()) or github_repo
    github_branch = (gh_secret.get("branch") or "").strip() or github_branch
    github_excel_path = (gh_secret.get("excel_path") or "").strip() or github_excel_path

    github_sync_enabled_secret = (gh_secret.get("sync_enabled") or "").strip()
    github_sync_on_change_secret = (gh_secret.get("sync_on_change") or "").strip()
    if github_sync_enabled_secret:
        github_sync_enabled = _to_bool_flag(github_sync_enabled_secret)
    elif github_repo and github_excel_path and github_token:
        # Secrets에 GitHub 핵심값이 있으면 자동 동기화를 기본 활성화한다.
        github_sync_enabled = True
    elif not github_sync_enabled_raw:
        # 설정 DB가 비어 있는 새 환경에서는 repo/path/token이 있으면 자동으로 활성화한다.
        github_sync_enabled = bool(github_repo and github_excel_path and github_token)

    if github_sync_on_change_secret:
        github_sync_on_change = _to_bool_flag(github_sync_on_change_secret)

    global_map = {
        "store_sensitive_keys": store_sensitive,
        "global_ai_provider": global_provider,
        "global_openai_api_key": global_openai_key,
        "global_claude_api_key": global_claude_key,
        "global_alpha_vantage_api_key": global_alpha_key,
        "global_finnhub_api_key": global_finnhub_key,
        "global_openai_model": global_openai_model,
        "global_claude_model": global_claude_model,
        "github_sync_enabled": github_sync_enabled,
        "github_sync_on_change": github_sync_on_change,
        "github_repo": github_repo,
        "github_branch": github_branch,
        "github_excel_path": github_excel_path,
        "github_token": github_token,
        "daily_auto_snapshot_enabled": daily_auto_enabled,
        "daily_auto_snapshot_use_ai_ticker": daily_auto_use_ai_ticker,
        "daily_auto_snapshot_hour": int(daily_auto_hour),
        "daily_auto_snapshot_last_run_date": daily_auto_last_run_date,
        "daily_auto_snapshot_last_run_at": daily_auto_last_run_at,
        "daily_auto_snapshot_last_attempt_date": daily_auto_last_attempt_date,
        "daily_auto_snapshot_last_summary": daily_auto_last_summary,
    }
    for k, v in global_map.items():
        if force or k not in st.session_state:
            st.session_state[k] = v

    # GitHub 관련은 Secrets를 항상 우선 반영한다(기존 세션 값이 있어도 덮어씀).
    if github_repo or github_excel_path or github_token or github_sync_enabled_secret or github_sync_on_change_secret:
        st.session_state["github_repo"] = github_repo
        st.session_state["github_branch"] = github_branch
        st.session_state["github_excel_path"] = github_excel_path
        st.session_state["github_token"] = github_token
        st.session_state["github_sync_enabled"] = github_sync_enabled
        st.session_state["github_sync_on_change"] = github_sync_on_change

    for prefix in ["analysis", "score", "compare"]:
        scoped_map = {
            f"{prefix}_ai_provider": global_provider,
            f"{prefix}_openai_api_key": global_openai_key,
            f"{prefix}_claude_api_key": global_claude_key,
            f"{prefix}_openai_model": global_openai_model,
            f"{prefix}_claude_model": global_claude_model,
        }
        for k, v in scoped_map.items():
            if force or k not in st.session_state:
                st.session_state[k] = v

    if force or "score_ai_api_key" not in st.session_state:
        st.session_state["score_ai_api_key"] = global_claude_key
    if force or "analysis_ai_api_key" not in st.session_state:
        st.session_state["analysis_ai_api_key"] = global_claude_key
    if force or "score_ai_model" not in st.session_state:
        st.session_state["score_ai_model"] = global_claude_model
    if force or "analysis_ai_model" not in st.session_state:
        st.session_state["analysis_ai_model"] = global_claude_model


@st.cache_data(ttl=600, show_spinner=False)
def fetch_openai_available_models(api_key: str) -> tuple[list[str], str]:
    key = (api_key or "").strip()
    if not key:
        return [], "OpenAI API Key를 입력해 주세요."
    try:
        resp = requests.get(
            "https://api.openai.com/v1/models",
            headers={"Authorization": f"Bearer {key}"},
            timeout=12,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        return [], f"OpenAI 모델 조회 실패: {exc}"

    model_ids = []
    for item in data.get("data", []):
        mid = str(item.get("id", "")).strip()
        if mid:
            model_ids.append(mid)
    model_ids = sorted(set(model_ids))
    if not model_ids:
        return [], "OpenAI에서 사용 가능한 모델 목록을 받지 못했습니다."
    return model_ids, ""


@st.cache_data(ttl=600, show_spinner=False)
def fetch_claude_available_models(api_key: str) -> tuple[list[str], str]:
    key = (api_key or "").strip()
    if not key:
        return [], "Claude API Key를 입력해 주세요."
    try:
        resp = requests.get(
            "https://api.anthropic.com/v1/models",
            headers={
                "x-api-key": key,
                "anthropic-version": "2023-06-01",
            },
            timeout=12,
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        return [], f"Claude 모델 조회 실패: {exc}"

    model_ids = []
    for item in data.get("data", []):
        mid = str(item.get("id", "")).strip()
        if mid:
            model_ids.append(mid)
    model_ids = sorted(set(model_ids))
    if not model_ids:
        return [], "Claude에서 사용 가능한 모델 목록을 받지 못했습니다."
    return model_ids, ""


def ensure_portfolio_columns(df: pd.DataFrame, usd_krw_rate: float, force_usd_rate: bool = False) -> pd.DataFrame:
    base = df.copy()
    if COL_CURRENCY not in base.columns:
        base[COL_CURRENCY] = "KRW"
    if COL_FX_RATE not in base.columns:
        base[COL_FX_RATE] = 1.0

    base[COL_CURRENCY] = (
        base[COL_CURRENCY]
        .astype(str)
        .str.strip()
        .str.upper()
        .replace({"": "KRW", "NAN": "KRW", "NONE": "KRW"})
    )
    base[COL_FX_RATE] = pd.to_numeric(base[COL_FX_RATE], errors="coerce")
    base.loc[base[COL_CURRENCY] == "KRW", COL_FX_RATE] = 1.0
    if force_usd_rate:
        base.loc[base[COL_CURRENCY] == "USD", COL_FX_RATE] = float(usd_krw_rate)
    else:
        base.loc[
            (base[COL_CURRENCY] == "USD")
            & (base[COL_FX_RATE].isna() | (base[COL_FX_RATE] <= 0)),
            COL_FX_RATE,
        ] = float(usd_krw_rate)
    base[COL_FX_RATE] = base[COL_FX_RATE].fillna(1.0)
    return base


def _batch_fetch_prices_krw(
    ticker_pairs: list[tuple[str, str]],
    target_date: date,
) -> dict[str, tuple[float | None, str]]:
    """yfinance.download()로 복수 티커를 한 번에 조회해 KRW 환산 가격을 반환한다.

    반환: {stock_name: (price_krw, source_label)}
    티커별 개별 API 호출 대신 단일 배치 호출로 성능을 개선한다.
    """
    try:
        import yfinance as yf
    except Exception:
        return {}

    valid_pairs = [(name, tkr) for name, tkr in ticker_pairs if tkr]
    if not valid_pairs:
        return {}

    tickers = [tkr for _, tkr in valid_pairs]
    try:
        batch_df = yf.download(tickers, period="5d", progress=False, auto_adjust=True)
    except Exception:
        return {}

    if batch_df is None or batch_df.empty:
        return {}

    close_data = batch_df.get("Close")
    if close_data is None:
        return {}

    result: dict[str, tuple[float | None, str]] = {}
    for stock_name, tkr in valid_pairs:
        try:
            if isinstance(close_data, pd.Series):
                # 단일 티커인 경우 Series 그대로
                prices = pd.to_numeric(close_data, errors="coerce").dropna()
            else:
                if tkr not in close_data.columns:
                    continue
                prices = pd.to_numeric(close_data[tkr], errors="coerce").dropna()

            if prices.empty:
                continue

            price_native = float(prices.iloc[-1])
            if price_native <= 0:
                continue

            # 통화 판별: 한국 거래소 종목은 KRW, 그 외는 USD로 기본 가정
            if tkr.upper().endswith(".KS") or tkr.upper().endswith(".KQ"):
                currency = "KRW"
            else:
                currency = "USD"

            fx_rate, fx_src = _fetch_currency_to_krw_rate(currency, target_date)
            if fx_rate is None or fx_rate <= 0:
                continue

            price_krw = price_native * fx_rate
            result[stock_name] = (price_krw, f"batch/{currency}→KRW({fx_src})")
        except Exception:
            continue

    return result


def _refresh_company_prices_for_portfolio(
    portfolio_df: pd.DataFrame,
    target_date: date,
    ai_provider: str,
    ai_api_key: str,
    ai_model: str,
    use_ai_ticker: bool,
) -> tuple[int, int, list[str]]:
    if portfolio_df is None or portfolio_df.empty:
        return 0, 0, []

    market_pref_map = build_market_preference_map(portfolio_df)
    unique_names = (
        portfolio_df[COL_NAME]
        .astype(str)
        .str.strip()
        .replace({"": pd.NA})
        .dropna()
        .drop_duplicates()
        .tolist()
    )
    updated_count = 0
    resolved_ticker_count = 0
    failures: list[str] = []

    # ── 1단계: 모든 종목의 티커 확보 ──────────────────────────────────
    ticker_map: dict[str, str] = {}  # stock_name → ticker
    for stock_name in unique_names:
        ticker = get_company_list_ticker(stock_name)
        if not ticker:
            ticker, _ = resolve_ticker_auto_with_retry(
                stock_name,
                use_ai=bool(use_ai_ticker and ai_api_key),
                api_key=ai_api_key,
                model=ai_model,
                provider=ai_provider,
                market_preference=market_pref_map.get(stock_name, ""),
            )
            ticker = clean_valid_ticker(ticker)
            if ticker:
                resolved_ticker_count += 1
        if ticker:
            ticker_map[stock_name] = ticker
        else:
            failures.append(f"{stock_name}(티커 미확보)")

    if not ticker_map:
        return updated_count, resolved_ticker_count, failures

    # ── 2단계: 배치 주가 조회 ────────────────────────────────────────
    ticker_pairs = list(ticker_map.items())
    batch_prices = _batch_fetch_prices_krw(ticker_pairs, target_date)

    # ── 3단계: 배치 실패 종목은 개별 조회로 fallback ──────────────────
    for stock_name, ticker in ticker_pairs:
        if stock_name in batch_prices:
            price_krw, price_src = batch_prices[stock_name]
        else:
            price_krw, price_src = fetch_current_price_krw_from_ticker(ticker, target_date)

        if price_krw is None or float(price_krw) <= 0:
            failures.append(f"{stock_name}({price_src or '주가 조회 실패'})")
            continue

        upsert_company_list_entry(
            stock_name,
            ticker=ticker,
            source="daily_auto_price",
            price_krw=float(price_krw),
            price_source=price_src or "daily_auto_price",
        )
        updated_count += 1

    return updated_count, resolved_ticker_count, failures


def run_daily_auto_snapshot(force: bool = False, target_date: date | None = None) -> tuple[bool, str]:
    run_date = target_date or date.today()
    run_date_text = run_date.isoformat()
    base_date = run_date - timedelta(days=1)
    base_date_text = base_date.isoformat()

    enabled = bool(st.session_state.get("daily_auto_snapshot_enabled", False))
    if not enabled and not force:
        return False, ""

    try:
        run_hour = int(st.session_state.get("daily_auto_snapshot_hour", 18))
    except Exception:
        run_hour = 18
    run_hour = max(0, min(23, run_hour))

    now_local = datetime.now()
    last_run_date = str(st.session_state.get("daily_auto_snapshot_last_run_date", "") or "").strip()
    last_attempt_date = str(st.session_state.get("daily_auto_snapshot_last_attempt_date", "") or "").strip()
    if not force:
        if run_date == date.today() and now_local.hour < run_hour:
            return False, ""
        if last_run_date == run_date_text:
            return False, ""
        if last_attempt_date == run_date_text:
            return False, ""

    st.session_state["daily_auto_snapshot_last_attempt_date"] = run_date_text
    save_app_settings_partial({"daily_auto_snapshot_last_attempt_date": run_date_text})

    try:
        bootstrap_excel_from_github_if_needed()
    except Exception:
        pass

    usd_krw_rate, _ = get_usd_krw_rate_for_date(run_date)
    base_df = ensure_portfolio_columns(load_snapshot(base_date), usd_krw_rate, force_usd_rate=True)
    if base_df.empty:
        fail_msg = "일일 자동 저장 실패: 전일 기준 보유 자산 데이터가 없어 스냅샷을 생성하지 못했습니다."
        st.session_state["daily_auto_snapshot_last_summary"] = fail_msg
        save_app_settings_partial({"daily_auto_snapshot_last_summary": fail_msg})
        return False, fail_msg

    ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
    use_ai_ticker = bool(st.session_state.get("daily_auto_snapshot_use_ai_ticker", True))

    try:
        updated_price_count, resolved_ticker_count, failures = _refresh_company_prices_for_portfolio(
            portfolio_df=base_df,
            target_date=run_date,
            ai_provider=ai_provider,
            ai_api_key=ai_api_key,
            ai_model=ai_model,
            use_ai_ticker=use_ai_ticker,
        )

        company_list_df = load_company_list()
        company_price_exact, company_price_norm = build_company_price_krw_maps(company_list_df)
        final_df = recalculate_portfolio_from_price_and_avg_buy(
            df=base_df,
            usd_krw_rate=usd_krw_rate,
            company_price_exact=company_price_exact,
            company_price_norm=company_price_norm,
        )
        final_df = ensure_numeric(final_df, usd_krw_rate)

        # 자동 실행은 전일 예수금을 그대로 승계한다.
        base_cash_krw, base_cash_usd = load_snapshot_cash(base_date)
        save_snapshot_cash(run_date, base_cash_krw, base_cash_usd)

        sync_ok, sync_msg = save_snapshot(
            snapshot_date=run_date,
            df=final_df,
            sync_to_github=True,
            sync_reason="daily_auto",
        )

        view_krw = to_krw_view(final_df, usd_krw_rate, force_usd_rate=True)
        stock_total_krw = float(pd.to_numeric(view_krw.get(COL_VALUE_KRW), errors="coerce").fillna(0).sum())
        stock_pnl_krw = float(pd.to_numeric(view_krw.get(COL_PNL_KRW), errors="coerce").fillna(0).sum())
        cash_total_krw, _, _ = get_snapshot_cash_krw(run_date, usd_krw_rate)
        total_asset_krw = stock_total_krw + float(cash_total_krw)

        sync_note = ""
        if sync_msg:
            sync_note = f" / GitHub {'완료' if sync_ok else '경고'}: {sync_msg}"
        summary = (
            f"일일 자동 저장 완료({run_date_text}) "
            f"[기준: {base_date_text} 보유/예수금] "
            f"- 총자산 {total_asset_krw:,.0f}원, 총손익 {stock_pnl_krw:,.0f}원, "
            f"주가갱신 {updated_price_count}개, 티커보강 {resolved_ticker_count}개"
            f"{sync_note}"
        )
        if failures:
            preview = ", ".join(failures[:3])
            remain = len(failures) - min(3, len(failures))
            tail = f" 외 {remain}개" if remain > 0 else ""
            summary += f" / 일부 실패: {preview}{tail}"

        now_text = datetime.now().isoformat(timespec="seconds")
        st.session_state["daily_auto_snapshot_last_run_date"] = run_date_text
        st.session_state["daily_auto_snapshot_last_run_at"] = now_text
        st.session_state["daily_auto_snapshot_last_summary"] = summary
        save_app_settings_partial(
            {
                "daily_auto_snapshot_last_run_date": run_date_text,
                "daily_auto_snapshot_last_run_at": now_text,
                "daily_auto_snapshot_last_summary": summary,
                "daily_auto_snapshot_last_attempt_date": run_date_text,
            }
        )
        return True, summary
    except Exception as exc:
        fail_msg = f"일일 자동 저장 실패: {exc}"
        st.session_state["daily_auto_snapshot_last_summary"] = fail_msg
        save_app_settings_partial({"daily_auto_snapshot_last_summary": fail_msg})
        return False, fail_msg


def to_krw_view(df: pd.DataFrame, usd_krw_rate: float, force_usd_rate: bool = False) -> pd.DataFrame:
    converted = ensure_portfolio_columns(df, usd_krw_rate, force_usd_rate=force_usd_rate)
    converted[COL_VALUE] = pd.to_numeric(converted[COL_VALUE], errors="coerce")
    converted[COL_PNL] = pd.to_numeric(converted[COL_PNL], errors="coerce")
    converted[COL_VALUE_KRW] = converted[COL_VALUE] * converted[COL_FX_RATE]
    converted[COL_PNL_KRW] = converted[COL_PNL] * converted[COL_FX_RATE]
    return converted


def ensure_numeric(df: pd.DataFrame, usd_krw_rate: float) -> pd.DataFrame:
    cleaned = df.copy()
    for col in [COL_QTY, COL_FX_RATE, COL_VALUE, COL_PNL, COL_RETURN]:
        cleaned[col] = pd.to_numeric(cleaned[col], errors="coerce")
    cleaned = ensure_portfolio_columns(cleaned, usd_krw_rate, force_usd_rate=True)
    cleaned[COL_NAME] = cleaned[COL_NAME].astype(str).str.strip()
    cleaned = cleaned.dropna(subset=[COL_NAME, COL_QTY, COL_VALUE, COL_PNL, COL_RETURN, COL_CURRENCY, COL_FX_RATE])
    cleaned = cleaned[cleaned[COL_NAME] != ""]
    cleaned = cleaned[COLUMNS]
    return cleaned


def get_holding_stock_names(current_df: pd.DataFrame) -> list[str]:
    names = set(current_df[COL_NAME].dropna().astype(str).tolist()) if not current_df.empty else set()
    conn = get_conn()
    try:
        rows = conn.execute("SELECT DISTINCT stock_name FROM snapshots ORDER BY stock_name").fetchall()
        names.update([row[0] for row in rows if row and row[0]])
    finally:
        conn.close()
    return sorted(names)


def get_all_stock_names(current_df: pd.DataFrame) -> list[str]:
    names = set(get_holding_stock_names(current_df))
    conn = get_conn()
    try:
        for table_name in ["company_scores", "company_analysis", "company_list"]:
            rows = conn.execute(
                f"SELECT DISTINCT stock_name FROM {table_name} WHERE stock_name IS NOT NULL AND stock_name != ''"
            ).fetchall()
            names.update([row[0] for row in rows if row and row[0]])
    finally:
        conn.close()
    return sorted(names)


def infer_market_preference_from_row(stock_name: str, currency: str = "", ticker: str = "") -> str:
    name = str(stock_name or "").strip()
    tkr = clean_valid_ticker(ticker)
    curr = str(currency or "").strip().upper()
    upper_name = name.upper()
    builtin_hint = get_builtin_ticker_hint(name)
    if builtin_hint:
        return "domestic" if (builtin_hint.endswith(".KS") or builtin_hint.endswith(".KQ")) else "foreign"
    if tkr.endswith(".KS") or tkr.endswith(".KQ"):
        return "domestic"
    if curr and curr in {"USD", "EUR", "JPY", "CNY", "GBP", "AUD", "CAD", "CHF"}:
        return "foreign"
    if "ADR" in upper_name:
        return "foreign"
    if _looks_foreign_hangul_name_hint(name):
        return "foreign"
    if _company_name_has_hangul(name):
        if _looks_explicit_foreign_company_name(name):
            return "foreign"
        if _looks_domestic_company_name_hint(name):
            return "domestic"
        return ""
    if tkr:
        return "foreign"
    if re.search(r"[A-Z]{2,}", upper_name):
        return "foreign"
    # 한글명이 없고 단서도 없으면 미분류.
    return ""


def build_market_preference_map(current_df: pd.DataFrame) -> dict[str, str]:
    pref_map: dict[str, str] = {}
    if current_df is None or current_df.empty:
        return pref_map

    list_ticker_map: dict[str, str] = {}
    company_list_df = load_company_list()
    if not company_list_df.empty:
        for _, row in company_list_df.iterrows():
            nm = str(row.get("stock_name") or "").strip()
            if not nm:
                continue
            list_ticker_map[nm] = clean_valid_ticker(str(row.get("ticker") or ""))

    for _, row in current_df.iterrows():
        name = str(row.get(COL_NAME) or "").strip()
        if not name:
            continue
        ticker = clean_valid_ticker(str(row.get("ticker") or "")) or list_ticker_map.get(name, "")
        pref = infer_market_preference_from_row(
            stock_name=name,
            currency=str(row.get(COL_CURRENCY) or ""),
            ticker=ticker,
        )
        if pref:
            pref_map[name] = pref
    return pref_map


def format_won(value: float) -> str:
    return f"{value:,.0f}원"


def format_usd(value: float) -> str:
    return f"${value:,.2f}"


def format_signed_won(value: float) -> str:
    sign = "+" if value > 0 else ""
    return f"{sign}{value:,.0f}원"


def format_signed_pct(value: float) -> str:
    sign = "+" if value > 0 else ""
    return f"{sign}{value:,.0f}%"


def value_class(value: float) -> str:
    if value > 0:
        return "positive"
    if value < 0:
        return "negative"
    return "neutral"


def render_summary_card(label: str, value: str, note: str, note_class: str = "neutral") -> None:
    st.markdown(
        f"""
        <div class="summary-card">
            <div class="summary-label">{label}</div>
            <div class="summary-value">{value}</div>
            <div class="summary-note {note_class}">{note}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def format_table_numbers(df: pd.DataFrame, percent_cols: set[str] | None = None) -> pd.DataFrame:
    view = df.copy()
    percent_cols = percent_cols or set()
    for col in view.columns:
        if not pd.api.types.is_numeric_dtype(view[col]):
            continue
        if col in percent_cols:
            view[col] = view[col].apply(lambda x: "" if pd.isna(x) else f"{x:,.0f}%")
        else:
            view[col] = view[col].apply(lambda x: "" if pd.isna(x) else f"{x:,.0f}")
    return view


def _signed_color_style(value) -> str:
    if value is None or pd.isna(value):
        return ""
    v = float(value)
    if v > 0:
        return "color: #d92d20; font-weight: 700;"
    if v < 0:
        return "color: #1570ef; font-weight: 700;"
    return "color: #64748b;"


def style_market_detail_table(df: pd.DataFrame):
    if df is None or df.empty:
        return df

    view = df.copy()
    numeric_cols = [
        COL_QTY,
        COL_AVG_BUY_KRW,
        COL_PRICE_KRW,
        "투자금액(원)",
        COL_VALUE_KRW,
        COL_PNL_KRW,
        COL_RETURN,
        "비중(%)",
    ]
    for col in numeric_cols:
        if col in view.columns:
            view[col] = pd.to_numeric(view[col], errors="coerce")

    fmt_map = {}
    for col in [COL_QTY, COL_AVG_BUY_KRW, COL_PRICE_KRW, "투자금액(원)", COL_VALUE_KRW, COL_PNL_KRW]:
        if col in view.columns:
            fmt_map[col] = "{:,.0f}"
    if COL_RETURN in view.columns:
        fmt_map[COL_RETURN] = "{:,.0f}%"
    if "비중(%)" in view.columns:
        fmt_map["비중(%)"] = "{:,.0f}%"

    styler = view.style.format(fmt_map, na_rep="")

    sign_cols = [col for col in [COL_PNL_KRW, COL_RETURN] if col in view.columns]
    if sign_cols:
        styler = styler.applymap(_signed_color_style, subset=sign_cols)

    if "시장구분" in view.columns:
        def _market_row_band(row: pd.Series):
            group = str(row.get("시장구분") or "")
            if group == "국내주식":
                bg = "background-color: rgba(29, 78, 216, 0.08);"
            elif group == "해외주식":
                bg = "background-color: rgba(15, 118, 110, 0.08);"
            else:
                bg = "background-color: rgba(148, 163, 184, 0.10);"
            return [bg] * len(row)

        styler = styler.apply(_market_row_band, axis=1)

    if "비중(%)" in view.columns and HAS_MATPLOTLIB:
        styler = styler.background_gradient(subset=["비중(%)"], cmap="Blues")

    try:
        styler = styler.hide(axis="index")
    except Exception:
        pass

    return styler


def style_market_summary_table(df: pd.DataFrame):
    if df is None or df.empty:
        return df

    view = df.copy()
    numeric_cols = ["종목수", "투자금액", "평가금액", "손익금액", "비중(%)"]
    for col in numeric_cols:
        if col in view.columns:
            view[col] = pd.to_numeric(view[col], errors="coerce")

    fmt_map = {}
    for col in ["종목수", "투자금액", "평가금액", "손익금액"]:
        if col in view.columns:
            fmt_map[col] = "{:,.0f}"
    if "비중(%)" in view.columns:
        fmt_map["비중(%)"] = "{:,.0f}%"

    styler = view.style.format(fmt_map, na_rep="")

    if "손익금액" in view.columns:
        styler = styler.applymap(_signed_color_style, subset=["손익금액"])
        if HAS_MATPLOTLIB:
            styler = styler.background_gradient(
                subset=["손익금액"],
                cmap=["#dbeafe", "#eef2f7", "#ffd9dc"],
            )

    if "비중(%)" in view.columns and HAS_MATPLOTLIB:
        styler = styler.background_gradient(subset=["비중(%)"], cmap="Blues")

    if "시장구분" in view.columns:
        def _market_row_band(row: pd.Series):
            group = str(row.get("시장구분") or "")
            if group == "국내주식":
                bg = "background-color: rgba(29, 78, 216, 0.08);"
            elif group == "해외주식":
                bg = "background-color: rgba(15, 118, 110, 0.08);"
            else:
                bg = "background-color: rgba(148, 163, 184, 0.10);"
            return [bg] * len(row)

        styler = styler.apply(_market_row_band, axis=1)

    try:
        styler = styler.hide(axis="index")
    except Exception:
        pass

    return styler


def auto_balance_yaxis(fig):
    # 이미 명시 범위가 지정된 그래프는 자동 스케일 재조정을 건너뛴다.
    yaxis = getattr(fig.layout, "yaxis", None)
    if yaxis is not None:
        fixed_range = getattr(yaxis, "range", None)
        if fixed_range and len(fixed_range) == 2 and fixed_range[0] is not None and fixed_range[1] is not None:
            return fig

    y_series = []
    for trace in fig.data:
        y_vals = getattr(trace, "y", None)
        if y_vals is None:
            continue
        numeric = pd.to_numeric(pd.Series(list(y_vals)), errors="coerce").dropna()
        if numeric.empty:
            continue
        y_series.append(numeric)

    if not y_series:
        return fig

    all_y = pd.concat(y_series, ignore_index=True)
    y_min = float(all_y.min())
    y_max = float(all_y.max())
    if pd.isna(y_min) or pd.isna(y_max):
        return fig

    spread = y_max - y_min
    if spread == 0:
        # 단일 값 그래프는 값의 크기에 비례해 여유 구간을 준다.
        pad = max(abs(y_max) * 0.10, 0.5)
    else:
        pad = max(spread * 0.15, 0.5)

    low = y_min - pad
    high = y_max + pad
    if low == high:
        high = low + 1.0

    fig.update_yaxes(range=[low, high])
    return fig


def build_zero_based_y_range(values, headroom_ratio: float = 0.12) -> list[float] | None:
    numeric = pd.to_numeric(pd.Series(values), errors="coerce").dropna()
    if numeric.empty:
        return None

    y_min = float(numeric.min())
    y_max = float(numeric.max())

    if y_min >= 0:
        upper = y_max * (1.0 + float(headroom_ratio))
        if upper <= 0:
            upper = 1.0
        return [0.0, float(upper)]

    if y_max <= 0:
        lower = y_min * (1.0 + float(headroom_ratio))
        if lower >= 0:
            lower = -1.0
        return [float(lower), 0.0]

    span = y_max - y_min
    pad = max(span * float(headroom_ratio), 0.5)
    return [float(y_min - pad), float(y_max + pad)]


def style_figure(fig):
    fig.update_layout(
        template="plotly_white",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(255,255,255,0)",
        legend_title_text="",
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        margin=dict(l=16, r=16, t=56, b=20),
        font=dict(family="Noto Sans KR", color="#0f172a"),
    )
    fig.update_xaxes(
        showgrid=True,
        gridcolor="rgba(148,163,184,0.24)",
        zeroline=False,
        tickformat="%Y-%m-%d",
        hoverformat="%Y-%m-%d",
    )
    fig.update_yaxes(showgrid=True, gridcolor="rgba(148,163,184,0.24)", zeroline=False, tickformat=",")
    return auto_balance_yaxis(fig)


def apply_daily_date_axis(fig):
    fig.update_xaxes(type="date", tickformat="%Y-%m-%d", hoverformat="%Y-%m-%d")

    x_all = []
    for trace in fig.data:
        x_vals = getattr(trace, "x", None)
        if x_vals is None:
            continue
        try:
            x_all.extend(list(x_vals))
        except Exception:
            continue
    if not x_all:
        return fig

    x_ser = pd.to_datetime(pd.Series(x_all), errors="coerce").dropna()
    if x_ser.empty:
        return fig
    x_ser = x_ser.dt.tz_localize(None) if getattr(x_ser.dt, "tz", None) is not None else x_ser
    unique_dates = sorted(pd.Series(x_ser.dt.date).dropna().unique().tolist())
    if not unique_dates:
        return fig

    if len(unique_dates) == 1:
        center = pd.Timestamp(unique_dates[0])
        fig.update_xaxes(
            range=[center - pd.Timedelta(days=2), center + pd.Timedelta(days=2)],
            tickmode="array",
            tickvals=[center],
            ticktext=[center.strftime("%Y-%m-%d")],
        )
        return fig

    min_dt = pd.Timestamp(unique_dates[0])
    max_dt = pd.Timestamp(unique_dates[-1])
    span_days = max(1, int((max_dt - min_dt).days))
    pad_days = max(1, min(14, int(round(span_days * 0.06))))

    xaxis_cfg = {"range": [min_dt - pd.Timedelta(days=pad_days), max_dt + pd.Timedelta(days=pad_days)]}
    if len(unique_dates) <= 12:
        xaxis_cfg["tickmode"] = "array"
        xaxis_cfg["tickvals"] = [pd.Timestamp(d) for d in unique_dates]
        xaxis_cfg["ticktext"] = [pd.Timestamp(d).strftime("%Y-%m-%d") for d in unique_dates]
    elif span_days <= 90:
        xaxis_cfg["dtick"] = "D7"
    elif span_days <= 365:
        xaxis_cfg["dtick"] = "M1"
    else:
        xaxis_cfg["dtick"] = "M3"

    fig.update_xaxes(**xaxis_cfg)
    return fig


def _label_text(value, pct: bool = False) -> str:
    if value is None or pd.isna(value):
        return ""
    value = float(value)
    if pct:
        return f"{value:,.0f}%"
    return f"{value:,.0f}"


def add_line_labels(fig, pct: bool = False, last_only: bool = False, max_labels: int = 6):
    line_traces = [t for t in fig.data if t.type in ("scatter", "scatterpolar")]
    for t_idx, trace in enumerate(line_traces):
        if trace.type not in ("scatter", "scatterpolar"):
            continue
        values = []
        if getattr(trace, "y", None) is not None:
            values = list(trace.y)
        elif getattr(trace, "r", None) is not None:
            values = list(trace.r)
        if not values:
            continue

        force_last_only = last_only or len(values) > max_labels or len(line_traces) > 1
        if force_last_only:
            labels = [""] * len(values)
            labels[-1] = _label_text(values[-1], pct=pct)
            positions = ["top right", "bottom right", "top left", "bottom left"]
            pos = positions[t_idx % len(positions)]
        else:
            step = max(1, len(values) // max_labels)
            labels = []
            for idx, value in enumerate(values):
                labels.append(_label_text(value, pct=pct) if (idx % step == 0 or idx == len(values) - 1) else "")
            pos = "top center"
        trace.update(mode="lines+markers+text", text=labels, textposition=pos)
    return fig


def estimate_textarea_height(
    value,
    min_height: int = 90,
    max_height: int = 640,
    chars_per_line: int = 60,
) -> int:
    text = str(value or "")
    lines = text.splitlines() or [""]
    virtual_lines = 0
    for line in lines:
        line_text = str(line or "")
        cjk_bonus = int(len(re.findall(r"[가-힣\u1100-\u11FF\u3130-\u318F\u4E00-\u9FFF]", line_text)) * 0.45)
        visual_len = len(line_text) + cjk_bonus
        wrapped = max(1, (visual_len + max(1, chars_per_line) - 1) // max(1, chars_per_line))
        virtual_lines += wrapped
    height = 44 + virtual_lines * 22
    return int(max(min_height, min(max_height, height)))


def estimate_dataframe_height(df: pd.DataFrame, min_height: int = 180, max_height: int = 3600) -> int:
    row_count = len(df.index) if isinstance(df, pd.DataFrame) else 0
    height = 42 + max(1, row_count) * 35
    return int(max(min_height, min(max_height, height)))


def apply_analysis_history_to_editor(latest_row: pd.Series) -> None:
    # Streamlit 제약: 이미 렌더된 위젯 키는 같은 실행 흐름에서 직접 재할당할 수 없다.
    # 따라서 pending 버퍼에 적재하고, 다음 rerun 시작 시점(위젯 렌더 전)에 반영한다.
    st.session_state["analysis_history_apply_pending"] = {
        "analysis_company_overview": str(latest_row.get("company_overview") or ""),
        "analysis_products_services": str(latest_row.get("products_services") or ""),
        "analysis_raw_materials": str(latest_row.get("raw_materials") or ""),
        "analysis_profit_up_factors": str(latest_row.get("profit_up_factors") or ""),
        "analysis_profit_down_factors": str(latest_row.get("profit_down_factors") or ""),
        "analysis_key_takeaway": str(latest_row.get("note") or latest_row.get("key_takeaway") or ""),
    }
    st.session_state["analysis_history_apply_notice"] = "선택한 이력 내용을 상단 '기업 분석 내용'에 반영했습니다."


def add_bar_labels(fig, pct: bool = False, max_labels: int = 10):
    for trace in fig.data:
        if trace.type != "bar":
            continue
        orientation = getattr(trace, "orientation", "v")
        values = list(trace.x) if orientation == "h" else list(trace.y)
        if len(values) <= max_labels:
            labels = [_label_text(v, pct=pct) for v in values]
        else:
            labels = [""] * len(values)
            ranked = sorted(
                range(len(values)),
                key=lambda i: abs(float(values[i])) if values[i] is not None and not pd.isna(values[i]) else -1,
                reverse=True,
            )
            for i in ranked[:max_labels]:
                labels[i] = _label_text(values[i], pct=pct)
        trace.update(text=labels, textposition="auto", cliponaxis=False)
    return fig


def make_top_scatter_text(df: pd.DataFrame, label_col: str, size_col: str, top_n: int = 8) -> pd.Series:
    text = pd.Series([""] * len(df), index=df.index, dtype="object")
    if df.empty or size_col not in df.columns or label_col not in df.columns:
        return text

    size_series = pd.to_numeric(df[size_col], errors="coerce").abs().fillna(0)
    top_idx = size_series.nlargest(min(top_n, len(df))).index
    text.loc[top_idx] = df.loc[top_idx, label_col].astype(str)
    return text


def compact_pie_df(df: pd.DataFrame, name_col: str, value_col: str, top_n: int = 10) -> pd.DataFrame:
    if df.empty:
        return df.copy()
    src = df[[name_col, value_col]].copy().sort_values(value_col, ascending=False)
    if len(src) <= top_n:
        return src

    top = src.head(top_n).copy()
    etc_value = src.iloc[top_n:][value_col].sum()
    if etc_value > 0:
        top = pd.concat([top, pd.DataFrame([{name_col: "기타", value_col: etc_value}])], ignore_index=True)
    return top


def compute_totals(
    df: pd.DataFrame,
    usd_krw_rate: float,
    snapshot_date: date | None = None,
) -> tuple[float, float, float, float]:
    total_value = 0.0
    total_pnl = 0.0
    if not df.empty:
        converted = to_krw_view(df, usd_krw_rate)
        total_value = float(converted[COL_VALUE_KRW].sum())
        total_pnl = float(converted[COL_PNL_KRW].sum())
    if snapshot_date is not None:
        cash_total_krw, _, _ = get_snapshot_cash_krw(snapshot_date, None)
        total_value += float(cash_total_krw)
    total_principal = total_value - total_pnl
    total_return = (total_pnl / total_principal * 100) if total_principal else 0.0
    return total_value, total_pnl, total_principal, total_return


def infer_stock_market_group(stock_name: str, currency: str, ticker: str = "") -> str:
    name = str(stock_name or "").strip()
    curr = str(currency or "").strip().upper()
    tkr = clean_valid_ticker(ticker)
    upper_name = name.upper()

    if tkr.endswith(".KS") or tkr.endswith(".KQ"):
        return "국내주식"
    if tkr:
        return "해외주식"
    if curr == "USD":
        return "해외주식"

    # 이름 기반 보조 힌트
    if "ADR" in upper_name:
        return "해외주식"
    if re.search(r"[A-Z]{2,}", upper_name):
        return "해외주식"
    if re.search(r"[가-힣]", name):
        return "국내주식"
    return "미분류"


def build_holdings_market_view(df: pd.DataFrame, usd_krw_rate: float) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame(
            columns=[
                "시장구분",
                COL_NAME,
                "티커",
                COL_CURRENCY,
                COL_QTY,
                COL_VALUE_KRW,
                COL_PNL_KRW,
                COL_RETURN,
                "비중(%)",
            ]
        )

    view = to_krw_view(df, usd_krw_rate).copy()
    view[COL_NAME] = view[COL_NAME].astype(str).str.strip()
    view[COL_CURRENCY] = view[COL_CURRENCY].astype(str).str.upper()

    ticker_map = {}
    company_list_df = load_company_list()
    if not company_list_df.empty:
        for _, row in company_list_df.iterrows():
            nm = str(row.get("stock_name") or "").strip()
            if not nm:
                continue
            ticker_map[nm] = clean_valid_ticker(str(row.get("ticker") or ""))

    view["티커"] = view[COL_NAME].map(ticker_map).fillna("")
    view["시장구분"] = view.apply(
        lambda r: infer_stock_market_group(
            stock_name=str(r.get(COL_NAME) or ""),
            currency=str(r.get(COL_CURRENCY) or ""),
            ticker=str(r.get("티커") or ""),
        ),
        axis=1,
    )

    total_value = float(view[COL_VALUE_KRW].sum())
    view["비중(%)"] = (view[COL_VALUE_KRW] / total_value * 100) if total_value else 0.0
    view = view.sort_values([COL_VALUE_KRW, COL_NAME], ascending=[False, True])

    return view[
        [
            "시장구분",
            COL_NAME,
            "티커",
            COL_CURRENCY,
            COL_QTY,
            COL_VALUE_KRW,
            COL_PNL_KRW,
            COL_RETURN,
            "비중(%)",
        ]
    ]


def get_pnl_color_config(values: pd.Series) -> dict:
    series = pd.to_numeric(values, errors="coerce").dropna()
    if series.empty:
        return {"scale": [(0.0, "#e2e8f0"), (1.0, "#94a3b8")], "midpoint": None}

    min_v = float(series.min())
    max_v = float(series.max())
    if min_v >= 0:
        # 수익만 있는 구간: 맑지만 이전보다 한 단계 진한 빨강 계열
        return {
            "scale": [(0.0, "#ffd9dc"), (0.45, "#ff8d94"), (1.0, "#ef4444")],
            "midpoint": None,
        }
    if max_v <= 0:
        # 손실만 있는 구간: 맑지만 이전보다 한 단계 진한 파랑 계열
        return {
            "scale": [(0.0, "#1d4ed8"), (0.55, "#6ea8ff"), (1.0, "#dbeafe")],
            "midpoint": None,
        }
    # 혼재 구간: 맑은 파랑-중립-맑은 빨강(명도 소폭 하향)
    return {
        "scale": [(0.0, "#2563eb"), (0.5, "#dbe4ef"), (1.0, "#ef4444")],
        "midpoint": 0.0,
    }


def filter_history_by_period(hist_df: pd.DataFrame, period: str) -> pd.DataFrame:
    if hist_df.empty or period == "전체":
        return hist_df

    end_date = hist_df["snapshot_date"].max()
    if period == "YTD":
        start_date = pd.Timestamp(year=end_date.year, month=1, day=1)
    elif period == "1개월":
        start_date = end_date - pd.DateOffset(months=1)
    elif period == "3개월":
        start_date = end_date - pd.DateOffset(months=3)
    elif period == "6개월":
        start_date = end_date - pd.DateOffset(months=6)
    elif period == "1년":
        start_date = end_date - pd.DateOffset(years=1)
    elif period == "2년":
        start_date = end_date - pd.DateOffset(years=2)
    elif period == "3년":
        start_date = end_date - pd.DateOffset(years=3)
    elif period == "5년":
        start_date = end_date - pd.DateOffset(years=5)
    elif period == "10년":
        start_date = end_date - pd.DateOffset(years=10)
    else:
        return hist_df

    return hist_df[hist_df["snapshot_date"] >= start_date].copy()


def add_history_features(hist_df: pd.DataFrame) -> pd.DataFrame:
    if hist_df.empty:
        return hist_df

    df = hist_df.copy()
    df["value_change"] = df["total_value"].diff().fillna(0)
    df["pnl_change"] = df["total_pnl"].diff().fillna(0)
    df["rolling_peak"] = df["total_value"].cummax()
    df["drawdown_pct"] = (df["total_value"] / df["rolling_peak"] - 1) * 100
    return df


def get_monthly_return_table(hist_df: pd.DataFrame) -> pd.DataFrame:
    if hist_df.empty:
        return pd.DataFrame()

    monthly = (
        hist_df.set_index("snapshot_date")
        .resample("M")
        .last()[["total_value"]]
        .dropna()
        .reset_index()
    )
    if monthly.empty:
        return pd.DataFrame()

    monthly["monthly_return_pct"] = monthly["total_value"].pct_change() * 100
    monthly["year"] = monthly["snapshot_date"].dt.year.astype(str)
    monthly["month"] = monthly["snapshot_date"].dt.month

    heat = monthly.pivot(index="year", columns="month", values="monthly_return_pct")
    heat = heat.reindex(columns=list(range(1, 13)))
    heat.columns = [f"{m}월" for m in range(1, 13)]
    return heat


def _render_allocation_target(
    total_value: float,
    domestic_holding_value: float,
    foreign_holding_value: float,
    cash_total_krw: float,
    unknown_holding_value: float,
) -> None:
    """자산 배분 목표 대비 현재 비중을 그룹 bar 차트로 렌더링한다.

    목표 비중은 app_settings의 'target_allocation_json' 키에 저장된다.
    형식: {"국내주식": 50, "해외주식": 30, "현금": 20}
    """
    import json as _json

    if total_value <= 0:
        return

    ALLOC_KEY = "target_allocation_json"
    settings = load_app_settings()
    raw = settings.get(ALLOC_KEY, "")
    try:
        targets: dict = _json.loads(raw) if raw else {}
    except Exception:
        targets = {}

    categories = ["국내주식", "해외주식", "현금", "미분류"]
    current_vals = {
        "국내주식": domestic_holding_value,
        "해외주식": foreign_holding_value,
        "현금": cash_total_krw,
        "미분류": unknown_holding_value,
    }
    current_pcts = {k: round(v / total_value * 100, 1) for k, v in current_vals.items()}

    # 목표 비중 슬라이더 UI
    with st.expander("자산 배분 목표 설정", expanded=False):
        st.caption("목표 비중(%)을 설정하면 현재 비중과 비교해 괴리를 시각화합니다.")
        slider_cols = st.columns(4)
        new_targets: dict[str, int] = {}
        for idx, cat in enumerate(["국내주식", "해외주식", "현금"]):
            with slider_cols[idx]:
                new_targets[cat] = st.slider(
                    cat,
                    min_value=0,
                    max_value=100,
                    value=int(targets.get(cat, 0)),
                    step=5,
                    key=f"alloc_target_{cat}",
                )
        with slider_cols[3]:
            total_set = sum(new_targets.values())
            st.metric("합계(%)", f"{total_set}%", delta=f"{total_set - 100:+d}%" if total_set != 100 else "✓ 100%")
        if st.button("목표 저장", key="alloc_target_save_btn"):
            save_app_settings({ALLOC_KEY: _json.dumps(new_targets, ensure_ascii=False)})
            st.success("목표 비중을 저장했습니다.")
            st.rerun()
        targets = new_targets  # 슬라이더 값을 즉시 차트에 반영

    has_targets = any(v > 0 for v in targets.values())
    if not has_targets:
        return

    bar_data: list[dict] = []
    for cat in categories:
        cur = current_pcts.get(cat, 0.0)
        tgt = float(targets.get(cat, 0))
        bar_data.append({"자산군": cat, "현재 비중(%)": cur, "목표 비중(%)": tgt})

    bar_df = pd.DataFrame(bar_data)
    alloc_fig = px.bar(
        bar_df,
        x="자산군",
        y=["현재 비중(%)", "목표 비중(%)"],
        barmode="group",
        title="자산 배분 현재 vs 목표",
        labels={"value": "비중(%)", "variable": "구분"},
        color_discrete_map={"현재 비중(%)": "#0f766e", "목표 비중(%)": "#cbd5e1"},
    )
    alloc_fig.update_yaxes(ticksuffix="%", range=[0, 105])
    st.plotly_chart(style_figure(alloc_fig), use_container_width=True)


def _render_stock_month_heatmap() -> None:
    """종목별 × 월별 수익률(pnl_pct) 히트맵을 렌더링한다.

    각 셀 = 해당 월의 마지막 스냅샷 기준 pnl_pct(누적 수익률 %).
    """
    all_df = load_all_snapshots_for_export()
    if all_df is None or all_df.empty:
        return

    all_df["snapshot_date"] = pd.to_datetime(all_df["snapshot_date"], errors="coerce")
    all_df = all_df.dropna(subset=["snapshot_date"])
    all_df["ym"] = all_df["snapshot_date"].dt.to_period("M").astype(str)  # 'YYYY-MM'

    # 각 (종목, 월) 에서 마지막 날짜의 pnl_pct 사용
    last_per_month = (
        all_df.sort_values("snapshot_date")
        .groupby(["stock_name", "ym"], sort=False)
        .last()
        .reset_index()
    )

    if last_per_month.empty:
        return

    pivot = last_per_month.pivot_table(
        index="stock_name",
        columns="ym",
        values="pnl_pct",
        aggfunc="last",
    )
    pivot.columns.name = None
    pivot.index.name = None

    if pivot.empty:
        return

    # 종목을 최근 pnl_pct 기준 내림차순 정렬
    try:
        last_col = pivot.columns[-1]
        pivot = pivot.sort_values(last_col, ascending=False)
    except Exception:
        pass

    heat_z = pivot.values.tolist()
    heat_y = pivot.index.tolist()
    heat_x = pivot.columns.tolist()

    hmap = go.Figure(
        go.Heatmap(
            z=heat_z,
            x=heat_x,
            y=heat_y,
            colorscale=[(0.0, "#1570ef"), (0.5, "#f1f5f9"), (1.0, "#d92d20")],
            zmid=0,
            text=[[f"{v:.1f}%" if v is not None and v == v else "" for v in row] for row in heat_z],
            texttemplate="%{text}",
            hovertemplate="종목: %{y}<br>월: %{x}<br>수익률: %{z:.1f}%<extra></extra>",
            colorbar=dict(title="수익률(%)"),
        )
    )
    hmap.update_layout(
        title="종목별 월간 수익률 히트맵 (누적 pnl_pct)",
        margin=dict(l=10, r=10, t=52, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        font=dict(family="Noto Sans KR", color="#0f172a"),
        xaxis=dict(tickangle=-45),
    )
    st.markdown("#### 종목별 월간 수익률 히트맵")
    st.caption("각 셀 = 해당 월 마지막 스냅샷 기준 누적 수익률(%). 빨강=양수, 파랑=음수.")
    st.plotly_chart(style_figure(hmap), use_container_width=True)


def _render_benchmark_comparison(featured_hist: pd.DataFrame) -> None:
    """포트폴리오 수익률과 KOSPI / S&P 500을 기준일=100으로 정규화해 비교 차트를 렌더링한다."""
    if featured_hist.empty:
        return

    try:
        dates = pd.to_datetime(featured_hist["snapshot_date"], errors="coerce").dropna()
        if dates.empty:
            return
        start_str = dates.min().strftime("%Y-%m-%d")
        end_str = (dates.max() + pd.Timedelta(days=2)).strftime("%Y-%m-%d")
    except Exception:
        return

    bench_df = fetch_benchmark_returns(start_str, end_str)

    # 포트폴리오 정규화 (기준일 총평가금액=100)
    port_series = (
        featured_hist[["snapshot_date", "total_value"]]
        .copy()
        .dropna(subset=["total_value"])
        .sort_values("snapshot_date")
    )
    port_series["snapshot_date"] = pd.to_datetime(port_series["snapshot_date"], errors="coerce")
    base_val = float(port_series["total_value"].iloc[0])
    if base_val <= 0:
        return
    port_series["내 포트폴리오"] = (port_series["total_value"] / base_val * 100).round(2)
    port_series = port_series.rename(columns={"snapshot_date": "date"})
    port_series["date"] = port_series["date"].dt.date

    # 벤치마크 병합
    if not bench_df.empty:
        merged = pd.merge(
            port_series[["date", "내 포트폴리오"]],
            bench_df,
            on="date",
            how="outer",
        ).sort_values("date")
    else:
        merged = port_series[["date", "내 포트폴리오"]].copy()

    bench_cols = [c for c in ["KOSPI", "S&P500"] if c in merged.columns]
    all_cols = ["내 포트폴리오"] + bench_cols
    merged = merged.dropna(subset=["내 포트폴리오"])
    if merged.empty or len(merged) < 2:
        return

    color_map = {"내 포트폴리오": "#dc2626", "KOSPI": "#1d4ed8", "S&P500": "#059669"}
    bench_fig = px.line(
        merged,
        x="date",
        y=all_cols,
        title="벤치마크 비교 (기준일=100)",
        labels={"date": "날짜", "value": "지수(기준=100)", "variable": "지표"},
        color_discrete_map=color_map,
        markers=False,
    )
    bench_fig.add_hline(y=100, line_dash="dot", line_color="#94a3b8", line_width=1)
    bench_fig.update_traces(line_width=2)
    bench_fig.update_yaxes(tickformat=".0f")
    st.markdown("#### 벤치마크 비교")
    if bench_cols:
        st.caption(
            "내 포트폴리오(빨강) vs KOSPI(파랑) vs S&P500(초록). "
            "기준일(조회 기간 첫 스냅샷)을 100으로 정규화."
        )
    else:
        st.caption("벤치마크 데이터를 불러오지 못했습니다. yfinance 연결을 확인해 주세요.")
    st.plotly_chart(style_figure(bench_fig), use_container_width=True)


def render_dashboard(current_df: pd.DataFrame, usd_krw_rate: float, selected_date: date) -> None:
    hist_df = load_history(as_of_date=selected_date)
    if hist_df.empty and current_df is not None and not current_df.empty:
        current_view = to_krw_view(current_df, usd_krw_rate)
        stock_value = float(current_view[COL_VALUE_KRW].sum())
        stock_pnl = float(current_view[COL_PNL_KRW].sum())
        cash_total_krw, cash_krw, cash_usd = get_snapshot_cash_krw(selected_date, usd_krw_rate)
        total_value = stock_value + float(cash_total_krw)
        total_pnl = stock_pnl
        total_principal = total_value - total_pnl
        total_return_pct = (total_pnl / total_principal * 100.0) if total_principal else 0.0
        hist_df = pd.DataFrame(
            [
                {
                    "snapshot_date": pd.Timestamp(selected_date),
                    "total_value": total_value,
                    "total_pnl": total_pnl,
                    "cash_krw": float(cash_krw),
                    "cash_usd": float(cash_usd),
                    "is_carry_forward": True,
                    "total_principal": total_principal,
                    "total_return_pct": total_return_pct,
                }
            ]
        )

    featured_hist = pd.DataFrame()

    asof_latest_date = get_latest_snapshot_date_on_or_before(selected_date)
    if asof_latest_date is not None:
        latest_df = load_snapshot(asof_latest_date)
    else:
        latest_df = empty_portfolio_df()
    source_df = current_df if not current_df.empty else latest_df
    summary_available = not source_df.empty
    exact_snapshot_exists = has_snapshot_on_date(selected_date)

    total_value = 0.0
    total_pnl = 0.0
    total_principal = 0.0
    total_return = 0.0
    cash_total_krw = 0.0
    cash_krw = 0.0
    cash_usd = 0.0
    total_holding_value = 0.0
    domestic_holding_value = 0.0
    foreign_holding_value = 0.0
    unknown_holding_value = 0.0
    foreign_holding_usd = 0.0
    foreign_cash_value = 0.0
    fx_for_cash = float(get_usd_krw_rate_for_date(selected_date)[0]) if selected_date else float(usd_krw_rate)
    if summary_available:
        source_df = to_krw_view(source_df, usd_krw_rate)
        base_date = selected_date
        total_value, total_pnl, total_principal, total_return = compute_totals(source_df, usd_krw_rate, base_date)
        cash_total_krw, cash_krw, cash_usd = get_snapshot_cash_krw(base_date, None)
        total_holding_value = float(pd.to_numeric(source_df[COL_VALUE_KRW], errors="coerce").fillna(0).sum())
        market_view_for_summary = build_holdings_market_view(source_df, usd_krw_rate)
        if not market_view_for_summary.empty:
            grouped = (
                market_view_for_summary.groupby("시장구분", as_index=False)[COL_VALUE_KRW]
                .sum()
                .set_index("시장구분")[COL_VALUE_KRW]
            )
            domestic_holding_value = float(grouped.get("국내주식", 0.0))
            foreign_holding_value = float(grouped.get("해외주식", 0.0))
            unknown_holding_value = float(grouped.get("미분류", 0.0))
        else:
            unknown_holding_value = total_holding_value
        foreign_cash_value = float(cash_usd) * fx_for_cash
        foreign_holding_usd = (foreign_holding_value / fx_for_cash) if fx_for_cash else 0.0

    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">전체 자산 요약</div>', unsafe_allow_html=True)
    if not summary_available:
        st.info("저장된 데이터가 없습니다. 기록 입력 탭에서 먼저 스냅샷을 저장해 주세요.")
    else:
        c1, c2, c3 = st.columns(3)
        with c1:
            total_note = f"주식 {format_won(total_holding_value)} + 예수금 {format_won(cash_total_krw)}<br>보유 종목 {len(source_df)}개"
            render_summary_card("총 평가금액", format_won(total_value), total_note)
        with c2:
            render_summary_card("총 손익", format_signed_won(total_pnl), "기준: 평가 - 원금", value_class(total_pnl))
        with c3:
            render_summary_card("총 수익률", format_signed_pct(total_return), f"총원금 {format_won(total_principal)}", value_class(total_return))

        c4, c5, c6, c7 = st.columns(4)
        with c4:
            dom_note = "국내 종목 평가금액 합계"
            if unknown_holding_value > 0:
                dom_note = f"{dom_note}<br>미분류 {format_won(unknown_holding_value)}"
            render_summary_card("국내주식 자산", format_won(domestic_holding_value), dom_note)
        with c5:
            render_summary_card(
                "해외주식 자산",
                format_won(foreign_holding_value),
                f"원화환산 기준 / 달러환산 {format_usd(foreign_holding_usd)}",
            )
        with c6:
            render_summary_card("원화 예수금", format_won(cash_krw), "KRW 현금")
        with c7:
            render_summary_card(
                "달러 예수금",
                format_usd(cash_usd),
                f"원화환산 {format_won(foreign_cash_value)} (USD/KRW {fx_for_cash:,.0f})",
            )
        if not exact_snapshot_exists:
            if asof_latest_date is not None:
                st.caption(
                    f"선택일 {selected_date.isoformat()} 스냅샷이 없어 "
                    f"{asof_latest_date.isoformat()} 기준 데이터를 승계 표시했습니다."
                )
            elif selected_date >= date.today() and not current_df.empty:
                st.caption("선택일 스냅샷이 없어 현재 엑셀 입력값으로 임시 표시 중입니다.")
    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">총 자산 추이 (평가금액/원금)</div>', unsafe_allow_html=True)
    if hist_df.empty:
        st.info("아직 저장된 스냅샷이 없어 자산 추이 그래프를 표시할 수 없습니다. 기록 입력 탭에서 먼저 저장해 주세요.")
    else:
        period_options = ["1개월", "3개월", "6개월", "YTD", "1년", "2년", "3년", "5년", "10년", "전체"]
        period = st.radio(
            "조회 기간",
            options=period_options,
            horizontal=True,
            index=period_options.index("전체"),
            key="dashboard_asset_period",
        )
        filtered_hist = filter_history_by_period(hist_df, period)
        if filtered_hist.empty:
            st.warning("선택한 기간에 데이터가 없습니다.")
        else:
            featured_hist = add_history_features(filtered_hist)
            core_line_fig = px.line(
                featured_hist,
                x="snapshot_date",
                y=["total_value", "total_principal"],
                markers=True,
                title="총 자산 추이",
                labels={"snapshot_date": "날짜", "value": "금액(원)", "variable": "지표"},
                color_discrete_sequence=["#0f766e", "#334155"],
            )
            y_max_raw = float(
                max(
                    pd.to_numeric(featured_hist["total_value"], errors="coerce").max(),
                    pd.to_numeric(featured_hist["total_principal"], errors="coerce").max(),
                    0.0,
                )
            )
            y_axis_max = y_max_raw * 1.12 if y_max_raw > 0 else 1.0

            for trace in core_line_fig.data:
                if trace.name == "total_value":
                    trace.update(
                        name="총평가금액",
                        line={"width": 3, "color": "#dc2626"},
                        marker={"size": 8, "color": "#dc2626"},
                        fill="tozeroy",
                        fillcolor="rgba(239,68,68,0.22)",
                    )
                else:
                    trace.update(
                        name="총원금",
                        line={"width": 3, "color": "#1d4ed8"},
                        marker={"size": 8, "color": "#1d4ed8"},
                        fill="tozeroy",
                        fillcolor="rgba(59,130,246,0.14)",
                    )
            core_line_fig.update_yaxes(range=[0, y_axis_max], tickformat=",")
            add_line_labels(core_line_fig, pct=False, last_only=False)
            st.plotly_chart(style_figure(apply_daily_date_axis(core_line_fig)), use_container_width=True)

        if "is_carry_forward" in hist_df.columns and bool(hist_df.iloc[-1].get("is_carry_forward", False)):
            st.caption("오늘 스냅샷이 없어 최근 저장 자산값을 오늘 날짜로 동일 반영했습니다.")
    st.markdown("</div>", unsafe_allow_html=True)

    if summary_available:
        st.markdown('<div class="section-shell">', unsafe_allow_html=True)
        st.markdown('<div class="section-title">자산 구성 분석</div>', unsafe_allow_html=True)

        dist_df = source_df.sort_values(COL_VALUE_KRW, ascending=False).copy()
        dist_df["비중(%)"] = (dist_df[COL_VALUE_KRW] / total_value * 100).round(2) if total_value else 0
        pnl_color_cfg = get_pnl_color_config(dist_df[COL_PNL_KRW])
        pnl_cont_kwargs = {"color_continuous_scale": pnl_color_cfg["scale"]}
        if pnl_color_cfg["midpoint"] is not None:
            pnl_cont_kwargs["color_continuous_midpoint"] = pnl_color_cfg["midpoint"]

        row1_col1, row1_col2 = st.columns([1, 1])
        with row1_col1:
            pie_df = compact_pie_df(dist_df, COL_NAME, COL_VALUE_KRW, top_n=10)
            pie_df["비중(%)"] = (pie_df[COL_VALUE_KRW] / pie_df[COL_VALUE_KRW].sum() * 100).round(2)
            pie_df["라벨"] = pie_df["비중(%)"].apply(lambda x: f"{x:,.0f}%" if x >= 3 else "")
            pie_fig = px.pie(
                pie_df,
                names=COL_NAME,
                values=COL_VALUE_KRW,
                title="현재 자산 비중",
                hole=0.50,
                color_discrete_sequence=px.colors.qualitative.Safe,
                custom_data=["라벨"],
            )
            pie_fig.update_traces(
                texttemplate="%{customdata[0]}",
                textposition="inside",
                insidetextorientation="horizontal",
                hovertemplate="%{label}<br>%{value:,.0f}원<br>%{percent:.1%}<extra></extra>",
            )
            pie_fig.update_layout(
                margin=dict(l=10, r=10, t=56, b=10),
                paper_bgcolor="rgba(0,0,0,0)",
                font=dict(family="Noto Sans KR", color="#0f172a"),
            )
            st.plotly_chart(pie_fig, use_container_width=True)

        with row1_col2:
            top_bar_fig = px.bar(
                dist_df.head(12),
                x=COL_VALUE_KRW,
                y=COL_NAME,
                color=COL_PNL_KRW,
                orientation="h",
                title="종목별 평가금액 (Top 12)",
                labels={COL_VALUE_KRW: "평가금액(원)", COL_NAME: "종목명", COL_PNL_KRW: "손익금액(원)"},
                **pnl_cont_kwargs,
            )
            top_bar_fig.update_coloraxes(showscale=False)
            top_bar_fig.update_layout(yaxis={"categoryorder": "total ascending"})
            add_bar_labels(top_bar_fig, pct=False)
            st.plotly_chart(style_figure(top_bar_fig), use_container_width=True)

        row2_col1, row2_col2 = st.columns([1, 1.1])
        with row2_col1:
            tree_fig = px.treemap(
                dist_df,
                path=[COL_NAME],
                values=COL_VALUE_KRW,
                color=COL_PNL_KRW,
                title="자산 트리맵",
                **pnl_cont_kwargs,
            )
            tree_fig.update_traces(textinfo="label+value+percent root")
            tree_fig.update_coloraxes(showscale=False)
            tree_fig.update_layout(
                margin=dict(l=6, r=6, t=48, b=8),
                paper_bgcolor="rgba(0,0,0,0)",
                font=dict(family="Noto Sans KR", color="#0f172a"),
            )
            st.plotly_chart(tree_fig, use_container_width=True)

        with row2_col2:
            risk_fig = px.scatter(
                dist_df,
                x=COL_RETURN,
                y="비중(%)",
                size=COL_VALUE_KRW,
                color=COL_PNL_KRW,
                hover_name=COL_NAME,
                text=make_top_scatter_text(dist_df, COL_NAME, COL_VALUE_KRW, top_n=7),
                title="수익률-비중 버블맵",
                labels={COL_RETURN: "수익률(%)", "비중(%)": "포트폴리오 비중(%)", COL_PNL_KRW: "손익금액(원)"},
                **pnl_cont_kwargs,
            )
            risk_fig.update_coloraxes(showscale=False)
            risk_fig.add_hline(y=dist_df["비중(%)"].mean(), line_dash="dot", line_color="#334155")
            risk_fig.add_vline(x=0, line_dash="dot", line_color="#334155")
            risk_fig.update_traces(textposition="top center")
            st.plotly_chart(style_figure(risk_fig), use_container_width=True)

        _render_stock_month_heatmap()
        _render_allocation_target(
            total_value=total_value,
            domestic_holding_value=domestic_holding_value,
            foreign_holding_value=foreign_holding_value,
            cash_total_krw=float(cash_total_krw),
            unknown_holding_value=unknown_holding_value,
        )

        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">시기별 자산 흐름</div>', unsafe_allow_html=True)

    if hist_df.empty:
        st.info("자산 흐름은 스냅샷이 1개 이상 저장되면 표시됩니다.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    if featured_hist.empty:
        active_period = str(st.session_state.get("dashboard_asset_period", "전체") or "전체")
        fallback_hist = filter_history_by_period(hist_df, active_period)
        if fallback_hist.empty:
            st.warning("선택한 기간에 데이터가 없습니다.")
            st.markdown("</div>", unsafe_allow_html=True)
            return
        featured_hist = add_history_features(fallback_hist)

    flow_col1, flow_col2 = st.columns([1.5, 1])
    with flow_col1:
        pnl_fig = px.bar(
            featured_hist,
            x="snapshot_date",
            y="total_pnl",
            title="날짜별 총 손익",
            color="total_pnl",
            labels={"snapshot_date": "날짜", "total_pnl": "손익(원)"},
            color_continuous_scale=[(0.0, "#1570ef"), (0.5, "#94a3b8"), (1.0, "#d92d20")],
        )
        pnl_fig.update_coloraxes(showscale=False)
        pnl_range = build_zero_based_y_range(featured_hist["total_pnl"], headroom_ratio=0.15)
        if pnl_range:
            pnl_fig.update_yaxes(range=pnl_range, tickformat=",")
        add_bar_labels(pnl_fig, pct=False)
        st.plotly_chart(style_figure(apply_daily_date_axis(pnl_fig)), use_container_width=True)

    with flow_col2:
        return_fig = px.line(
            featured_hist,
            x="snapshot_date",
            y="total_return_pct",
            markers=True,
            title="총 수익률 추이",
            labels={"snapshot_date": "날짜", "total_return_pct": "수익률(%)"},
            color_discrete_sequence=["#1d4ed8"],
        )
        ret_range = build_zero_based_y_range(featured_hist["total_return_pct"], headroom_ratio=0.18)
        if ret_range:
            return_fig.update_yaxes(range=ret_range)
        return_fig.update_yaxes(tickformat=",.0f", ticksuffix="%")
        add_line_labels(return_fig, pct=True, last_only=False)
        st.plotly_chart(style_figure(apply_daily_date_axis(return_fig)), use_container_width=True)

    extra_col1, extra_col2 = st.columns([1, 1])
    with extra_col1:
        change_fig = px.bar(
            featured_hist,
            x="snapshot_date",
            y="value_change",
            color="value_change",
            title="기간별 자산 증감",
            labels={"snapshot_date": "날짜", "value_change": "증감(원)"},
            color_continuous_scale=[(0.0, "#1570ef"), (0.5, "#94a3b8"), (1.0, "#d92d20")],
        )
        change_fig.update_coloraxes(showscale=False)
        add_bar_labels(change_fig, pct=False)
        st.plotly_chart(style_figure(apply_daily_date_axis(change_fig)), use_container_width=True)

    with extra_col2:
        dd_fig = px.line(
            featured_hist,
            x="snapshot_date",
            y="drawdown_pct",
            markers=True,
            title="드로우다운 추이",
            labels={"snapshot_date": "날짜", "drawdown_pct": "드로우다운(%)"},
            color_discrete_sequence=["#7c3aed"],
        )
        dd_fig.add_hline(y=0, line_dash="dot", line_color="#334155")
        dd_fig.update_yaxes(tickformat=",.0f", ticksuffix="%")
        add_line_labels(dd_fig, pct=True, last_only=False)
        st.plotly_chart(style_figure(apply_daily_date_axis(dd_fig)), use_container_width=True)

    heat_df = get_monthly_return_table(featured_hist)
    if not heat_df.empty and heat_df.notna().sum().sum() > 0:
        heat_fig = px.imshow(
            heat_df,
            text_auto=".0f",
            aspect="auto",
            color_continuous_scale=[(0.0, "#1570ef"), (0.5, "#e2e8f0"), (1.0, "#d92d20")],
            title="월간 수익률 히트맵 (%)",
            labels={"x": "월", "y": "연도", "color": "월수익률(%)"},
        )
        heat_fig.update_layout(
            margin=dict(l=12, r=12, t=52, b=10),
            paper_bgcolor="rgba(0,0,0,0)",
            font=dict(family="Noto Sans KR", color="#0f172a"),
        )
        st.plotly_chart(heat_fig, use_container_width=True)

    recent_hist = featured_hist.sort_values("snapshot_date", ascending=False).head(10).copy()
    recent_hist = recent_hist.rename(
        columns={
            "snapshot_date": "날짜",
            "total_value": "총평가금액",
            "total_principal": "총원금",
            "total_pnl": "총손익금액",
            "total_return_pct": "총수익률(%)",
            "value_change": "자산증감",
            "drawdown_pct": "드로우다운(%)",
            "cash_krw": "예수금(원화)",
            "cash_usd": "예수금(달러)",
        }
    )
    recent_hist["날짜"] = recent_hist["날짜"].dt.date
    st.caption("최근 기록")
    st.dataframe(
        format_table_numbers(recent_hist, percent_cols={"총수익률(%)", "드로우다운(%)"}),
        use_container_width=True,
        hide_index=True,
    )

    # ── 벤치마크 비교 차트 ──────────────────────────────────────────
    _render_benchmark_comparison(featured_hist)

    all_snapshot_dates = sorted(
        pd.to_datetime(hist_df["snapshot_date"], errors="coerce").dropna().dt.date.unique().tolist(),
        reverse=True,
    )
    if all_snapshot_dates:
        if st.session_state.pop("dashboard_delete_snapshot_confirm_reset", False):
            st.session_state["dashboard_delete_snapshot_confirm"] = False
        del_col1, del_col2, del_col3 = st.columns([1.2, 0.8, 1.4])
        with del_col1:
            st.selectbox(
                "삭제할 기록 날짜",
                options=all_snapshot_dates,
                key="dashboard_delete_snapshot_date",
                format_func=lambda d: d.isoformat() if hasattr(d, "isoformat") else str(d),
            )
        with del_col2:
            st.checkbox("삭제 확인", key="dashboard_delete_snapshot_confirm")
        with del_col3:
            delete_snapshot_btn = st.button("선택 날짜 기록 삭제", key="dashboard_delete_snapshot_btn")

        if delete_snapshot_btn:
            target_date = st.session_state.get("dashboard_delete_snapshot_date")
            confirmed = bool(st.session_state.get("dashboard_delete_snapshot_confirm", False))
            parsed_date = target_date if isinstance(target_date, date) else _safe_parse_date(target_date)
            if not parsed_date:
                st.warning("삭제할 날짜를 선택해 주세요.")
            elif not confirmed:
                st.warning("삭제 확인을 체크해 주세요.")
            else:
                ok, msg = delete_snapshot_by_date(
                    parsed_date,
                    delete_cash=True,
                    sync_to_github=True,
                )
                st.session_state["dashboard_delete_snapshot_confirm_reset"] = True
                if ok:
                    st.session_state["github_sync_notice"] = msg
                    st.rerun()
                st.warning(msg)
    st.markdown("</div>", unsafe_allow_html=True)


def render_input_tab(selected_date: date, edited_df: pd.DataFrame, usd_krw_rate: float) -> pd.DataFrame:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">포트폴리오 입력 및 저장</div>', unsafe_allow_html=True)
    st.caption(f"{selected_date} 기준 USD/KRW {usd_krw_rate:,.0f} 자동 적용")
    saved_cash_krw, saved_cash_usd = load_snapshot_cash(selected_date)
    st.caption(f"현재 저장 예수금: 원화 {saved_cash_krw:,.0f}원 / 달러 {format_usd(saved_cash_usd)}")

    editor_state = st.session_state.get("portfolio_editor")
    if isinstance(editor_state, pd.DataFrame):
        source_df = editor_state
    elif isinstance(editor_state, list):
        source_df = pd.DataFrame(editor_state)
    elif isinstance(editor_state, dict):
        try:
            source_df = pd.DataFrame(editor_state)
        except Exception:
            source_df = edited_df
    else:
        source_df = edited_df

    cleaned_df = ensure_numeric(
        ensure_portfolio_columns(source_df, usd_krw_rate, force_usd_rate=True),
        usd_krw_rate,
    )
    company_list_df = load_company_list()
    company_price_exact, company_price_norm = build_company_price_krw_maps(company_list_df)
    cleaned_df = recalculate_portfolio_from_price_and_avg_buy(
        cleaned_df,
        usd_krw_rate,
        company_price_exact=company_price_exact,
        company_price_norm=company_price_norm,
    )

    total_value, total_pnl, total_principal, total_return = compute_totals(cleaned_df, usd_krw_rate, selected_date)
    cash_total_krw, cash_krw, cash_usd = get_snapshot_cash_krw(selected_date, None)
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("총 평가금액", format_won(total_value))
    k2.metric("총 손익", format_signed_won(total_pnl))
    k3.metric("총 수익률", format_signed_pct(total_return))
    k4.metric("현재 예수금", format_won(cash_total_krw), f"원화 {cash_krw:,.0f} / 달러 {format_usd(cash_usd)}")

    if not cleaned_df.empty:
        st.markdown("#### 국내/해외 보유 리스트 한눈에 보기")
        market_view_df = build_holdings_market_view(cleaned_df, usd_krw_rate)
        market_view_df["투자금액(원)"] = market_view_df[COL_VALUE_KRW] - market_view_df[COL_PNL_KRW]
        market_view_df[COL_AVG_BUY_KRW] = (
            pd.to_numeric(market_view_df["투자금액(원)"], errors="coerce")
            / pd.to_numeric(market_view_df[COL_QTY], errors="coerce").replace(0, pd.NA)
        )
        market_view_df[COL_PRICE_KRW] = build_price_series_with_company_fallback(
            names=market_view_df[COL_NAME],
            qty=market_view_df[COL_QTY],
            value_krw=market_view_df[COL_VALUE_KRW],
            company_price_exact=company_price_exact,
            company_price_norm=company_price_norm,
        )
        group_summary = (
            market_view_df.groupby("시장구분", as_index=False)
            .agg(
                종목수=(COL_NAME, "count"),
                투자금액=("투자금액(원)", "sum"),
                평가금액=(COL_VALUE_KRW, "sum"),
                손익금액=(COL_PNL_KRW, "sum"),
            )
            .sort_values("평가금액", ascending=False)
        )
        group_total = float(group_summary["평가금액"].sum()) if not group_summary.empty else 0.0
        if group_total > 0:
            group_summary["비중(%)"] = group_summary["평가금액"] / group_total * 100.0
        else:
            group_summary["비중(%)"] = 0.0

        domestic_row = group_summary[group_summary["시장구분"] == "국내주식"]
        foreign_row = group_summary[group_summary["시장구분"] == "해외주식"]
        unknown_row = group_summary[group_summary["시장구분"] == "미분류"]

        d_val = float(domestic_row.iloc[0]["평가금액"]) if not domestic_row.empty else 0.0
        d_inv = float(domestic_row.iloc[0]["투자금액"]) if not domestic_row.empty else 0.0
        d_cnt = int(domestic_row.iloc[0]["종목수"]) if not domestic_row.empty else 0
        f_val = float(foreign_row.iloc[0]["평가금액"]) if not foreign_row.empty else 0.0
        f_inv = float(foreign_row.iloc[0]["투자금액"]) if not foreign_row.empty else 0.0
        f_cnt = int(foreign_row.iloc[0]["종목수"]) if not foreign_row.empty else 0
        u_cnt = int(unknown_row.iloc[0]["종목수"]) if not unknown_row.empty else 0
        u_inv = float(unknown_row.iloc[0]["투자금액"]) if not unknown_row.empty else 0.0
        d_ratio = (d_val / group_total * 100.0) if group_total else 0.0
        f_ratio = (f_val / group_total * 100.0) if group_total else 0.0
        total_dom_for_value = d_val + f_val

        g1, g2, g3, g4 = st.columns([1, 1, 1, 1.2])
        g1.metric("국내주식", format_won(d_val), f"투자금액 {d_inv:,.0f}원 | 비중 {d_ratio:,.0f}%")
        g2.metric("해외주식", format_won(f_val), f"투자금액 {f_inv:,.0f}원 | 비중 {f_ratio:,.0f}%")
        g3.metric("미분류", f"{u_cnt:,}", f"투자금액 {u_inv:,.0f}원")
        g4.metric("국내외 합계 평가금액", format_won(total_dom_for_value), "국내+해외 평가금액 합계")

        split_col1, split_col2 = st.columns([1.25, 1])
        with split_col1:
            split_fig = px.pie(
                group_summary,
                names="시장구분",
                values="평가금액",
                hole=0.5,
                title="국내/해외 자산 비중",
                color="시장구분",
                color_discrete_map={"국내주식": "#1d4ed8", "해외주식": "#0f766e", "미분류": "#94a3b8"},
            )
            split_fig.update_traces(
                texttemplate="%{value:,.0f}원<br>%{percent:.1%}",
                textposition="inside",
                hovertemplate="%{label}<br>%{value:,.0f}원<br>%{percent:.1%}<extra></extra>",
            )
            st.plotly_chart(style_figure(split_fig), use_container_width=True)
        with split_col2:
            st.caption("시장구분 요약")
            st.dataframe(
                style_market_summary_table(group_summary),
                use_container_width=True,
                hide_index=True,
            )

        t_all, t_dom, t_for, t_unk = st.tabs(["전체", "국내주식", "해외주식", "미분류"])
        view_percent_cols = {COL_RETURN, "비중(%)"}
        display_cols = [
            "시장구분",
            COL_NAME,
            "티커",
            COL_CURRENCY,
            COL_QTY,
            COL_AVG_BUY_KRW,
            COL_PRICE_KRW,
            "투자금액(원)",
            COL_VALUE_KRW,
            COL_PNL_KRW,
            COL_RETURN,
            "비중(%)",
        ]
        st.caption("아래 상세 리스트는 시장구분/손익 상태를 색상으로 시각화했습니다.")
        with t_all:
            st.dataframe(style_market_detail_table(market_view_df[display_cols]), use_container_width=True)
        with t_dom:
            dom_df = market_view_df[market_view_df["시장구분"] == "국내주식"]
            st.dataframe(style_market_detail_table(dom_df[display_cols]), use_container_width=True)
        with t_for:
            for_df = market_view_df[market_view_df["시장구분"] == "해외주식"]
            st.dataframe(style_market_detail_table(for_df[display_cols]), use_container_width=True)
        with t_unk:
            unk_df = market_view_df[market_view_df["시장구분"] == "미분류"]
            st.dataframe(style_market_detail_table(unk_df[display_cols]), use_container_width=True)
        st.caption("분류 기준: 티커(.KS/.KQ=국내) 우선, 이후 통화/종목명 힌트로 자동 분류합니다.")

        chart_df = to_krw_view(cleaned_df, usd_krw_rate).sort_values(COL_VALUE_KRW, ascending=False)
        chart_df["비중(%)"] = (chart_df[COL_VALUE_KRW] / total_value * 100).round(2) if total_value else 0

        c1, c2 = st.columns([1, 1])
        with c1:
            value_fig = px.bar(
                chart_df.head(12),
                x=COL_VALUE_KRW,
                y=COL_NAME,
                orientation="h",
                color=COL_VALUE_KRW,
                title="입력 데이터 기준 평가금액 Top 12",
                labels={COL_VALUE_KRW: "평가금액(원)", COL_NAME: "종목명"},
                color_continuous_scale="Tealgrn",
            )
            value_fig.update_coloraxes(showscale=False)
            value_fig.update_layout(yaxis={"categoryorder": "total ascending"})
            add_bar_labels(value_fig, pct=False)
            st.plotly_chart(style_figure(value_fig), use_container_width=True)

        with c2:
            pnl_plot_df = chart_df.copy()
            pnl_plot_df["손익구분"] = pnl_plot_df[COL_PNL_KRW].apply(
                lambda v: "수익(+)" if float(v) > 0 else ("손실(-)" if float(v) < 0 else "보합")
            )
            pnl_fig = px.bar(
                pnl_plot_df,
                x=COL_NAME,
                y=COL_PNL_KRW,
                color="손익구분",
                title="입력 데이터 기준 종목별 손익",
                labels={COL_NAME: "종목명", COL_PNL_KRW: "손익금액(원)", "손익구분": "구분"},
                color_discrete_map={
                    "수익(+)": "#d92d20",
                    "손실(-)": "#1570ef",
                    "보합": "#94a3b8",
                },
                category_orders={"손익구분": ["수익(+)", "손실(-)", "보합"]},
            )
            add_bar_labels(pnl_fig, pct=False)
            st.plotly_chart(style_figure(pnl_fig), use_container_width=True)

        bubble_fig = px.scatter(
            chart_df,
            x=COL_QTY,
            y=COL_RETURN,
            size=COL_VALUE_KRW,
            color=COL_PNL_KRW,
            hover_name=COL_NAME,
            text=make_top_scatter_text(chart_df, COL_NAME, COL_VALUE_KRW, top_n=7),
            title="수량-수익률 버블맵",
            labels={COL_QTY: "보유수량", COL_RETURN: "수익률(%)", COL_PNL_KRW: "손익금액(원)"},
            color_continuous_scale=[(0.0, "#1570ef"), (0.5, "#94a3b8"), (1.0, "#d92d20")],
        )
        bubble_fig.update_coloraxes(showscale=False)
        bubble_fig.add_hline(y=0, line_dash="dot", line_color="#334155")
        bubble_fig.update_traces(textposition="top center")
        st.plotly_chart(style_figure(bubble_fig), use_container_width=True)

    st.markdown("---")
    st.markdown("#### 이미지 붙여넣기 자동 등록 (AI)")
    st.caption("업로드 영역 클릭 후 Ctrl+V로 클립보드 이미지를 붙여넣거나 파일을 선택하세요.")

    if "portfolio_image_uploader_nonce" not in st.session_state:
        st.session_state["portfolio_image_uploader_nonce"] = 0
    uploader_key = f"portfolio_image_uploader_{st.session_state['portfolio_image_uploader_nonce']}"

    upload_col1, upload_col2 = st.columns([1.4, 1])
    with upload_col1:
        pasted_image = st.file_uploader(
            "보유현황 이미지",
            type=["png", "jpg", "jpeg", "webp"],
            key=uploader_key,
        )
    with upload_col2:
        st.checkbox(
            "이미지 인식 날짜 우선 사용",
            key="portfolio_use_ai_date",
            help="AI가 날짜를 읽으면 해당 날짜로 저장하고, 없으면 사이드바 선택 날짜로 저장합니다.",
        )
        auto_register_btn = st.button("이미지로 보유현황 자동 등록/저장", type="primary", key="portfolio_image_register_btn")

    if pasted_image is not None:
        st.image(pasted_image, caption="업로드된 이미지 미리보기", use_container_width=True)

    if auto_register_btn:
        if pasted_image is None:
            st.warning("먼저 이미지를 붙여넣거나 업로드해 주세요.")
        else:
            ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("global")
            parsed_payload, parse_err = extract_holdings_from_image_with_ai(
                image_bytes=pasted_image.getvalue(),
                mime_type=str(getattr(pasted_image, "type", "") or "image/png"),
                provider=ai_provider,
                api_key=ai_api_key,
                model=ai_model,
            )
            if parse_err:
                st.error(parse_err)
            else:
                holdings_rows = parsed_payload.get("holdings") if isinstance(parsed_payload, dict) else []
                incoming_df = build_holdings_df_from_ai_rows(holdings_rows if isinstance(holdings_rows, list) else [], usd_krw_rate)
                ai_cash_krw, ai_cash_usd = extract_cash_from_ai_payload(parsed_payload)
                if incoming_df.empty and ai_cash_krw is None and ai_cash_usd is None:
                    st.error("이미지에서 보유 종목/예수금을 추출하지 못했습니다. 더 선명한 이미지를 사용해 주세요.")
                else:
                    target_date = selected_date
                    ai_date = _safe_parse_date(parsed_payload.get("as_of_date")) if isinstance(parsed_payload, dict) else None
                    if st.session_state.get("portfolio_use_ai_date", False) and ai_date is not None:
                        target_date = ai_date

                    reflected_count = 0
                    if not incoming_df.empty:
                        # 저장 대상 날짜의 기존 스냅샷과 병합해, 같은 날짜/같은 기업은 새 이미지 값으로 덮어쓴다.
                        if target_date == selected_date:
                            base_df_for_target = cleaned_df
                        else:
                            base_df_for_target = ensure_portfolio_columns(
                                load_snapshot(target_date), usd_krw_rate, force_usd_rate=True
                            )
                        merged_df = merge_holdings_overwrite(base_df_for_target, incoming_df, usd_krw_rate)
                        target_usd_krw = float(get_usd_krw_rate_for_date(target_date)[0])
                        company_list_for_calc = load_company_list()
                        company_price_exact2, company_price_norm2 = build_company_price_krw_maps(company_list_for_calc)
                        merged_df = recalculate_portfolio_from_price_and_avg_buy(
                            merged_df,
                            target_usd_krw,
                            company_price_exact=company_price_exact2,
                            company_price_norm=company_price_norm2,
                        )
                        save_snapshot(target_date, merged_df, sync_to_github=False)
                        reflected_count = len(incoming_df)
                        if target_date == selected_date:
                            st.session_state["editing_df"] = merged_df

                    existing_cash_krw, existing_cash_usd = load_snapshot_cash(target_date)
                    final_cash_krw = ai_cash_krw if ai_cash_krw is not None else existing_cash_krw
                    final_cash_usd = ai_cash_usd if ai_cash_usd is not None else existing_cash_usd
                    if (
                        ai_cash_krw is not None
                        or ai_cash_usd is not None
                        or existing_cash_krw != 0
                        or existing_cash_usd != 0
                    ):
                        save_snapshot_cash(target_date, final_cash_krw, final_cash_usd)

                    sync_target_df = (
                        merged_df
                        if not incoming_df.empty
                        else ensure_portfolio_columns(load_snapshot(target_date), usd_krw_rate, force_usd_rate=True)
                    )
                    sync_ok, sync_msg = sync_snapshot_to_github_excel(target_date, sync_target_df)
                    sync_note = ""
                    if sync_msg:
                        sync_note = f" / GitHub {'저장 완료' if sync_ok else '저장 실패'}: {sync_msg}"

                    holdings_meta = holdings_rows if isinstance(holdings_rows, list) else []
                    for row in holdings_meta:
                        if not isinstance(row, dict):
                            continue
                        stock_name = str(row.get("stock_name") or row.get("name") or "").strip()
                        if not stock_name:
                            continue
                        if _classify_cash_bucket(stock_name, str(row.get("currency") or "")):
                            continue
                        ticker = clean_valid_ticker(str(row.get("ticker") or ""))
                        sector = str(row.get("sector") or "").strip()
                        upsert_company_list_entry(stock_name, ticker, sector=sector, source="image_ai")

                    # 성공 시 업로더를 초기화해 미리보기 이미지를 제거한다.
                    st.session_state["portfolio_image_uploader_nonce"] += 1
                    st.session_state["portfolio_image_notice"] = (
                        f"이미지 자동 등록 완료: {reflected_count}개 종목 반영, 저장일 {target_date.isoformat()} / "
                        f"예수금 원화 {final_cash_krw:,.0f}원, 달러 {format_usd(final_cash_usd)} "
                        f"(업로드 이미지 초기화 완료){sync_note}"
                    )
                    st.rerun()

    if "portfolio_image_notice" in st.session_state:
        st.success(st.session_state.pop("portfolio_image_notice"))

    st.markdown("---")
    st.markdown("#### 수동 입력 테이블")
    prepared_df = ensure_portfolio_columns(cleaned_df, usd_krw_rate, force_usd_rate=True)
    prepared_df[COL_PRICE_KRW] = pd.to_numeric(prepared_df.get(COL_PRICE_KRW), errors="coerce")
    prepared_df[COL_AVG_BUY_KRW] = pd.to_numeric(prepared_df.get(COL_AVG_BUY_KRW), errors="coerce")
    table_df = st.data_editor(
        prepared_df,
        num_rows="dynamic",
        use_container_width=True,
        key="portfolio_editor",
        column_config={
            COL_NAME: st.column_config.TextColumn(COL_NAME, required=True),
            COL_QTY: st.column_config.NumberColumn(COL_QTY, min_value=0, format="localized"),
            COL_CURRENCY: st.column_config.SelectboxColumn(COL_CURRENCY, options=["KRW", "USD"]),
            COL_FX_RATE: st.column_config.NumberColumn(COL_FX_RATE, min_value=0.0, format="localized"),
            COL_VALUE: st.column_config.NumberColumn(COL_VALUE, min_value=0, format="localized"),
            COL_PNL: st.column_config.NumberColumn(COL_PNL, format="localized"),
            COL_RETURN: st.column_config.NumberColumn(COL_RETURN, format="localized"),
            COL_AVG_BUY_KRW: st.column_config.NumberColumn(COL_AVG_BUY_KRW, format="localized"),
            COL_PRICE_KRW: st.column_config.NumberColumn(COL_PRICE_KRW, format="localized"),
        },
        column_order=COLUMNS + [COL_AVG_BUY_KRW, COL_PRICE_KRW],
        disabled=[COL_FX_RATE, COL_AVG_BUY_KRW, COL_PRICE_KRW],
    )
    final_df_base = ensure_numeric(table_df, usd_krw_rate)
    final_df_calc = recalculate_portfolio_from_price_and_avg_buy(
        final_df_base,
        usd_krw_rate,
        company_price_exact=company_price_exact,
        company_price_norm=company_price_norm,
    )
    final_df = ensure_numeric(final_df_calc, usd_krw_rate)

    # GitHub 동기화가 켜진 경우, 입력 테이블 변경을 즉시 스냅샷/원격으로 반영한다.
    auto_sync_on_change = bool(st.session_state.get("github_sync_on_change", True))
    if auto_sync_on_change and not final_df.empty:
        autosave_key = f"portfolio_autosave_hash::{selected_date.isoformat()}"
        hash_df = final_df.copy()
        hash_df[COL_NAME] = hash_df[COL_NAME].astype(str).str.strip()
        hash_df = hash_df.sort_values([COL_NAME, COL_CURRENCY], ascending=[True, True]).reset_index(drop=True)
        autosave_hash = hashlib.sha256(hash_df.to_csv(index=False).encode("utf-8")).hexdigest()
        prev_hash = st.session_state.get(autosave_key, None)
        if prev_hash is None:
            st.session_state[autosave_key] = autosave_hash
        elif prev_hash != autosave_hash:
            sync_ok, sync_msg = save_snapshot(
                selected_date,
                final_df,
                sync_to_github=True,
                sync_reason="input_auto_save",
            )
            st.session_state["editing_df"] = final_df
            st.session_state[autosave_key] = autosave_hash
            if sync_msg:
                if sync_ok:
                    st.info(f"자동 저장됨 ({selected_date}) / {sync_msg}")
                else:
                    st.warning(f"자동 저장됨 ({selected_date}) / GitHub 동기화 경고: {sync_msg}")
            else:
                st.info(f"자동 저장됨 ({selected_date})")

    save_col, dl_col = st.columns([1, 1])
    with save_col:
        if st.button("현재 날짜로 저장", type="primary", key="save_snapshot_btn"):
            if final_df.empty:
                st.error("저장할 데이터가 없습니다. 종목 정보를 입력하세요.")
            else:
                sync_ok, sync_msg = save_snapshot(selected_date, final_df, sync_to_github=True, sync_reason="manual_save")
                st.session_state["editing_df"] = final_df
                autosave_key = f"portfolio_autosave_hash::{selected_date.isoformat()}"
                hash_df = final_df.copy()
                hash_df[COL_NAME] = hash_df[COL_NAME].astype(str).str.strip()
                hash_df = hash_df.sort_values([COL_NAME, COL_CURRENCY], ascending=[True, True]).reset_index(drop=True)
                st.session_state[autosave_key] = hashlib.sha256(hash_df.to_csv(index=False).encode("utf-8")).hexdigest()
                if sync_msg:
                    if sync_ok:
                        st.success(f"{selected_date} 스냅샷 저장 완료 / {sync_msg}")
                    else:
                        st.warning(f"{selected_date} 스냅샷 저장 완료 / GitHub 동기화 경고: {sync_msg}")
                else:
                    st.success(f"{selected_date} 스냅샷 저장 완료")

    with dl_col:
        csv = final_df.to_csv(index=False).encode("utf-8-sig")
        st.download_button(
            label="현재 데이터 CSV 다운로드",
            data=csv,
            file_name=f"portfolio_{selected_date}.csv",
            mime="text/csv",
            disabled=final_df.empty,
        )

    st.markdown("</div>", unsafe_allow_html=True)
    return final_df


def render_company_analysis_tab(current_df: pd.DataFrame) -> None:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">기업정보</div>', unsafe_allow_html=True)

    corrected_builtin = reconcile_builtin_ticker_overrides()
    if corrected_builtin > 0:
        st.session_state["analysis_builtin_reconcile_notice"] = (
            f"내장 힌트 기준 티커 자동 교정: {corrected_builtin}개"
        )
        st.rerun()

    analysis_all = load_company_analysis_history()
    company_list_df = load_company_list()
    stock_names = get_holding_stock_names(current_df)
    market_pref_map = build_market_preference_map(current_df)
    analyzed_names = analysis_all["stock_name"].dropna().astype(str).unique().tolist() if not analysis_all.empty else []
    listed_names = company_list_df["stock_name"].dropna().astype(str).unique().tolist() if not company_list_df.empty else []
    options = sorted(set(stock_names + analyzed_names + listed_names))

    if "analysis_date" not in st.session_state:
        st.session_state["analysis_date"] = date.today()
    if "analysis_company_name" not in st.session_state:
        st.session_state["analysis_company_name"] = options[0] if options else ""
    if "analysis_ticker" not in st.session_state:
        st.session_state["analysis_ticker"] = ""
    if "analysis_company_name_input" not in st.session_state:
        st.session_state["analysis_company_name_input"] = st.session_state.get("analysis_company_name", "")
    if "analysis_ticker_input" not in st.session_state:
        st.session_state["analysis_ticker_input"] = st.session_state.get("analysis_ticker", "")
    if "analysis_note" not in st.session_state:
        st.session_state["analysis_note"] = ""
    if "analysis_use_ai_ticker" not in st.session_state:
        st.session_state["analysis_use_ai_ticker"] = True
    legacy_analysis_model = str(
        st.session_state.get(
            "analysis_ai_model",
            st.session_state.get("score_openai_model", st.session_state.get("score_ai_model", DEFAULT_OPENAI_MODEL)),
        )
        or ""
    )
    if "analysis_ai_provider" not in st.session_state:
        st.session_state["analysis_ai_provider"] = normalize_ai_provider(
            st.session_state.get("global_ai_provider", "claude")
        )
    if "analysis_openai_api_key" not in st.session_state:
        st.session_state["analysis_openai_api_key"] = st.session_state.get(
            "analysis_ai_api_key",
            st.session_state.get("score_openai_api_key", st.session_state.get("score_ai_api_key", "")),
        )
    if "analysis_claude_api_key" not in st.session_state:
        st.session_state["analysis_claude_api_key"] = ""
    if "analysis_openai_model" not in st.session_state:
        st.session_state["analysis_openai_model"] = (
            legacy_analysis_model if legacy_analysis_model and "claude" not in legacy_analysis_model.lower() else DEFAULT_OPENAI_MODEL
        )
    if "analysis_claude_model" not in st.session_state:
        st.session_state["analysis_claude_model"] = (
            legacy_analysis_model if "claude" in legacy_analysis_model.lower() else DEFAULT_CLAUDE_MODEL
        )
    if "analysis_ticker_source" not in st.session_state:
        st.session_state["analysis_ticker_source"] = ""
    if "analysis_prev_company" not in st.session_state:
        st.session_state["analysis_prev_company"] = ""
    if "analysis_financial_summary_cache" not in st.session_state:
        st.session_state["analysis_financial_summary_cache"] = {}
    if "analysis_company_name_pending" in st.session_state:
        next_name = (st.session_state.pop("analysis_company_name_pending") or "").strip()
        st.session_state["analysis_company_name_input"] = next_name
        st.session_state["analysis_company_name"] = next_name
    if "analysis_ticker_pending" in st.session_state:
        next_ticker = clean_valid_ticker(st.session_state.pop("analysis_ticker_pending") or "")
        st.session_state["analysis_ticker_input"] = next_ticker
        st.session_state["analysis_ticker"] = next_ticker
    if "analysis_ticker_autofill_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_ticker_autofill_notice"))
    if "analysis_history_apply_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_history_apply_notice"))
    if "analysis_new_company_ticker_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_new_company_ticker_notice"))
    if "analysis_builtin_reconcile_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_builtin_reconcile_notice"))
    for key in [
        "analysis_company_overview",
        "analysis_products_services",
        "analysis_raw_materials",
        "analysis_profit_up_factors",
        "analysis_profit_down_factors",
        "analysis_key_takeaway",
    ]:
        if key not in st.session_state:
            st.session_state[key] = ""
    pending_history_apply = st.session_state.pop("analysis_history_apply_pending", None)
    if isinstance(pending_history_apply, dict):
        for key in [
            "analysis_company_overview",
            "analysis_products_services",
            "analysis_raw_materials",
            "analysis_profit_up_factors",
            "analysis_profit_down_factors",
            "analysis_key_takeaway",
        ]:
            if key in pending_history_apply:
                st.session_state[key] = str(pending_history_apply.get(key) or "")

    # 위젯 상태가 배열/객체로 오염되면 레이아웃 겹침이 발생할 수 있어 렌더 전에 정규화한다.
    st.session_state["analysis_ai_provider"] = _coerce_choice(
        st.session_state.get("analysis_ai_provider"),
        {"openai", "claude"},
        "claude",
    )
    st.session_state["analysis_use_ai_ticker"] = _to_bool_flag(st.session_state.get("analysis_use_ai_ticker", False))
    st.session_state["analysis_watch_image_enrich_meta"] = _to_bool_flag(
        st.session_state.get("analysis_watch_image_enrich_meta", True)
    )
    for key in [
        "analysis_openai_api_key",
        "analysis_claude_api_key",
        "analysis_openai_model",
        "analysis_claude_model",
        "analysis_company_name_input",
        "analysis_ticker_input",
        "analysis_ticker_source",
        "analysis_selected_overview_company",
        "analysis_selected_overview_ticker_input",
        "analysis_selected_overview_sector_input",
    ]:
        st.session_state[key] = _sanitize_widget_text(st.session_state.get(key), "")
    st.session_state["analysis_company_hint"] = _sanitize_widget_text(
        st.session_state.get("analysis_company_hint"),
        "직접입력",
    )

    st.markdown("#### 기업 리스트 관리")
    if "analysis_new_company_name" not in st.session_state:
        st.session_state["analysis_new_company_name"] = ""
    if "analysis_new_company_ticker" not in st.session_state:
        st.session_state["analysis_new_company_ticker"] = ""
    if "analysis_new_company_sector" not in st.session_state:
        st.session_state["analysis_new_company_sector"] = ""
    if "analysis_watch_image_uploader_nonce" not in st.session_state:
        st.session_state["analysis_watch_image_uploader_nonce"] = 0
    if "analysis_watch_image_enrich_meta" not in st.session_state:
        st.session_state["analysis_watch_image_enrich_meta"] = True
    if "analysis_selected_overview_company" not in st.session_state:
        st.session_state["analysis_selected_overview_company"] = ""
    if "analysis_selected_overview_ticker_input" not in st.session_state:
        st.session_state["analysis_selected_overview_ticker_input"] = ""
    if "analysis_selected_overview_sector_input" not in st.session_state:
        st.session_state["analysis_selected_overview_sector_input"] = ""

    add_col1, add_col2, add_col3, add_col4, add_col5 = st.columns([1.3, 1.0, 1.0, 0.9, 0.7])
    with add_col1:
        st.text_input("추가 기업명", key="analysis_new_company_name", placeholder="예: 애플")
    with add_col2:
        st.text_input("추가 티커(선택)", key="analysis_new_company_ticker", placeholder="예: AAPL")
    with add_col3:
        st.text_input("산업섹터(선택)", key="analysis_new_company_sector", placeholder="예: 철강, 금융")
    with add_col4:
        lookup_new_company_ticker_btn = st.button("AI 티커 조회", key="analysis_lookup_new_company_ticker_btn")
    with add_col5:
        add_company_btn = st.button("기업 추가", key="analysis_add_company_btn")

    if lookup_new_company_ticker_btn:
        lookup_name = (st.session_state.get("analysis_new_company_name") or "").strip()
        if not lookup_name:
            st.warning("티커를 조회할 기업명을 먼저 입력해 주세요.")
        else:
            analysis_ai_provider, analysis_ai_api_key, analysis_ai_model = get_ai_settings_from_session("analysis")
            auto_ticker, auto_src = resolve_ticker_auto_with_retry(
                lookup_name,
                use_ai=bool((analysis_ai_api_key or "").strip()),
                api_key=analysis_ai_api_key,
                model=analysis_ai_model,
                provider=analysis_ai_provider,
                market_preference=market_pref_map.get(lookup_name, ""),
            )
            auto_ticker = clean_valid_ticker(auto_ticker)
            if auto_ticker:
                st.session_state["analysis_new_company_ticker"] = auto_ticker
                st.session_state["analysis_new_company_ticker_notice"] = (
                    f"{lookup_name} 티커 자동 입력: {auto_ticker} ({auto_src or '자동 탐색'})"
                )
                st.rerun()
            st.warning(auto_src or "티커를 찾지 못했습니다.")

    if add_company_btn:
        new_name = (st.session_state.get("analysis_new_company_name") or "").strip()
        new_ticker = clean_valid_ticker(st.session_state.get("analysis_new_company_ticker") or "")
        new_sector = (st.session_state.get("analysis_new_company_sector") or "").strip()
        analysis_ai_provider, analysis_ai_api_key, analysis_ai_model = get_ai_settings_from_session("analysis")
        if not new_name:
            st.warning("추가할 기업명을 입력해 주세요.")
        else:
            resolved_sector = new_sector
            if not resolved_sector:
                resolved_sector, sector_src = resolve_sector_auto(
                    new_name,
                    new_ticker,
                    ai_api_key=analysis_ai_api_key,
                    ai_model=analysis_ai_model,
                    ai_provider=analysis_ai_provider,
                )
                if not resolved_sector and sector_src:
                    st.caption(f"섹터 자동 조회 실패: {sector_src}")
            if not resolved_sector:
                resolved_sector = infer_sector_from_name_heuristic(new_name, new_ticker)
            upsert_company_list_entry(new_name, new_ticker, sector=resolved_sector, source="manual")
            st.session_state["analysis_new_company_name"] = ""
            st.session_state["analysis_new_company_ticker"] = ""
            st.session_state["analysis_new_company_sector"] = ""
            st.session_state["analysis_company_name_pending"] = new_name
            if new_ticker:
                st.session_state["analysis_ticker_pending"] = new_ticker
            st.success(f"{new_name} 기업을 리스트에 추가했습니다.")
            st.rerun()

    saved_list_names = sorted(set(listed_names))
    remove_options = ["선택안함"] + saved_list_names
    if (
        "analysis_remove_company_name" in st.session_state
        and st.session_state["analysis_remove_company_name"] not in remove_options
    ):
        st.session_state["analysis_remove_company_name"] = "선택안함"
    del_col1, del_col2 = st.columns([1.5, 0.8])
    with del_col1:
        remove_name = st.selectbox("추가기업 삭제", remove_options, key="analysis_remove_company_name")
    with del_col2:
        remove_btn = st.button("선택 삭제", key="analysis_remove_company_btn")
    if remove_btn:
        if remove_name == "선택안함":
            st.warning("삭제할 기업을 선택해 주세요.")
        else:
            delete_company_list_entry(remove_name)
            if (st.session_state.get("analysis_company_name_input") or "").strip() == remove_name:
                st.session_state["analysis_company_name_pending"] = ""
                st.session_state["analysis_ticker_pending"] = ""
            st.success(f"{remove_name} 기업을 추가 리스트에서 삭제했습니다.")
            st.rerun()

    st.markdown("##### 이미지로 관심종목 리스트 추가 (AI)")
    st.caption("여기서 추가되는 항목은 보유현황이 아니라 기업 리스트(관심종목)만 업데이트됩니다.")
    watch_uploader_key = f"analysis_watch_image_uploader_{st.session_state['analysis_watch_image_uploader_nonce']}"
    watch_col1, watch_col2 = st.columns([1.35, 1.0])
    with watch_col1:
        watch_image = st.file_uploader(
            "관심종목 이미지",
            type=["png", "jpg", "jpeg", "webp"],
            key=watch_uploader_key,
        )
    with watch_col2:
        st.checkbox(
            "티커/산업섹터 자동 보강",
            key="analysis_watch_image_enrich_meta",
            help="이미지에서 빈 값이면 API/AI로 티커와 섹터를 추가 탐색합니다.",
        )
        watch_import_btn = st.button(
            "이미지로 관심종목 추가",
            key="analysis_watch_image_import_btn",
            type="primary",
        )

    if watch_image is not None:
        st.image(watch_image, caption="업로드된 관심종목 이미지", use_container_width=True)

    if watch_import_btn:
        if watch_image is None:
            st.warning("먼저 이미지를 업로드해 주세요.")
        else:
            ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
            parsed_rows, parse_err = extract_company_watchlist_from_image_with_ai(
                image_bytes=watch_image.getvalue(),
                mime_type=str(getattr(watch_image, "type", "") or "image/png"),
                provider=ai_provider,
                api_key=ai_api_key,
                model=ai_model,
            )
            if parse_err:
                st.error(parse_err)
            elif not parsed_rows:
                st.warning("이미지에서 기업 목록을 찾지 못했습니다. 더 선명한 표/목록 이미지를 사용해 주세요.")
            else:
                enrich_meta = bool(st.session_state.get("analysis_watch_image_enrich_meta", True))
                inserted_count = 0
                unresolved = []
                first_added_name = ""
                with st.spinner("관심종목 리스트에 반영 중입니다..."):
                    for item in parsed_rows:
                        if not isinstance(item, dict):
                            continue
                        stock_name = str(item.get("stock_name") or "").strip()
                        if not stock_name:
                            continue
                        ticker = clean_valid_ticker(str(item.get("ticker") or ""))
                        sector = str(item.get("sector") or "").strip()
                        market_pref = market_pref_map.get(stock_name, "")

                        if enrich_meta and not ticker:
                            auto_ticker, _ = resolve_ticker_auto_with_retry(
                                stock_name,
                                use_ai=bool(ai_api_key),
                                api_key=ai_api_key,
                                model=ai_model,
                                provider=ai_provider,
                                market_preference=market_pref,
                            )
                            if auto_ticker:
                                ticker = auto_ticker

                        if enrich_meta and not sector:
                            auto_sector, _ = resolve_sector_auto(
                                stock_name,
                                ticker,
                                ai_api_key=ai_api_key,
                                ai_model=ai_model,
                                ai_provider=ai_provider,
                            )
                            if auto_sector:
                                sector = auto_sector
                        if not sector:
                            sector = infer_sector_from_name_heuristic(stock_name, ticker)

                        upsert_company_list_entry(stock_name, ticker, sector=sector, source="image_watch_ai")
                        inserted_count += 1
                        if not first_added_name:
                            first_added_name = stock_name
                        if not ticker or not sector:
                            unresolved.append(stock_name)

                st.session_state["analysis_watch_image_uploader_nonce"] += 1
                if first_added_name:
                    st.session_state["analysis_company_name_pending"] = first_added_name
                unresolved_count = len(set(unresolved))
                msg = f"이미지 관심종목 추가 완료: {inserted_count}개 반영"
                if unresolved_count > 0:
                    msg += f" / 티커·섹터 일부 미완성 {unresolved_count}개"
                st.session_state["analysis_watch_image_notice"] = msg
                st.rerun()

    if "analysis_watch_image_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_watch_image_notice"))

    holding_set = set(stock_names)
    listed_set = set(listed_names)
    ticker_map = {}
    sector_map = {}
    price_map = {}
    source_map = {}
    registered_at_map = {}
    if not company_list_df.empty:
        for _, row in company_list_df.iterrows():
            nm = str(row.get("stock_name") or "").strip()
            if not nm:
                continue
            ticker_map[nm] = clean_valid_ticker(str(row.get("ticker") or ""))
            sector_map[nm] = str(row.get("sector") or "").strip()
            price_val = _safe_to_float(row.get("price_krw"))
            if price_val is not None and price_val > 0:
                price_map[nm] = float(price_val)
            source_map[nm] = str(row.get("source") or "").strip() or "manual"
            created_text = str(row.get("created_at") or "").strip()
            registered_at_map[nm] = created_text
    overview_rows = []
    for nm in sorted(holding_set | listed_set):
        tags = []
        if nm in holding_set:
            tags.append("보유종목")
        if nm in listed_set:
            tags.append("추가리스트")
        overview_rows.append(
            {
                "기업명": nm,
                "티커": ticker_map.get(nm, ""),
                "산업섹터": sector_map.get(nm, ""),
                "현재주가(원화환산)": price_map.get(nm),
                "등록일시": registered_at_map.get(nm, ""),
                "구분": ", ".join(tags),
                "리스트소스": source_map.get(nm, ""),
            }
        )
    def _bulk_fill_company_meta(
        target_rows: list[dict],
        ai_provider: str,
        ai_api_key: str,
        ai_model: str,
        force_refresh: bool = False,
    ) -> tuple[int, int, list[str]]:
        updated_count = 0
        skipped_count = 0
        unresolved: list[str] = []

        for row in target_rows:
            company_name = str(row.get("기업명") or "").strip()
            row_ticker = clean_valid_ticker(str(row.get("티커") or ""))
            row_sector = str(row.get("산업섹터") or "").strip()
            if not company_name:
                skipped_count += 1
                continue

            current_ticker = "" if force_refresh else row_ticker
            current_sector = "" if force_refresh else row_sector
            next_ticker = current_ticker
            next_sector = current_sector
            market_pref = market_pref_map.get(company_name, "")
            if _company_name_has_hangul(company_name) and _looks_domestic_company_name_hint(company_name):
                market_pref = "domestic"
            builtin_hint = get_builtin_ticker_hint(company_name)
            if builtin_hint:
                market_pref = "domestic" if (builtin_hint.endswith(".KS") or builtin_hint.endswith(".KQ")) else "foreign"
            if not market_pref:
                if row_ticker:
                    if not force_refresh:
                        market_pref = "domestic" if (row_ticker.endswith(".KS") or row_ticker.endswith(".KQ")) else "foreign"
                elif _company_name_has_hangul(company_name) and _looks_domestic_company_name_hint(company_name):
                    market_pref = "domestic"

            if (not force_refresh) and current_ticker:
                suspicious = False
                is_kr_ticker = current_ticker.endswith(".KS") or current_ticker.endswith(".KQ")
                if market_pref == "domestic" and not is_kr_ticker:
                    suspicious = True
                elif market_pref == "foreign" and is_kr_ticker:
                    suspicious = True
                elif (
                    _company_name_has_hangul(company_name)
                    and not _looks_explicit_foreign_company_name(company_name)
                    and not is_kr_ticker
                ):
                    suspicious = True
                if suspicious:
                    current_ticker = ""
                    next_ticker = ""

            if builtin_hint and _ticker_matches_market_preference(builtin_hint, market_pref):
                next_ticker = builtin_hint

            if not next_ticker:
                auto_ticker, _ = resolve_ticker_auto_with_retry(
                    company_name,
                    use_ai=bool(ai_api_key),
                    api_key=ai_api_key,
                    model=ai_model,
                    provider=ai_provider,
                    market_preference=market_pref,
                )
                if auto_ticker:
                    next_ticker = auto_ticker

            if next_ticker and market_pref == "domestic" and not (
                next_ticker.endswith(".KS") or next_ticker.endswith(".KQ")
            ):
                next_ticker = ""
            if next_ticker and builtin_hint and (builtin_hint.endswith(".KS") or builtin_hint.endswith(".KQ")):
                if not (next_ticker.endswith(".KS") or next_ticker.endswith(".KQ")):
                    next_ticker = ""

            if not next_sector:
                auto_sector, _ = resolve_sector_auto(
                    company_name,
                    next_ticker,
                    ai_api_key=ai_api_key,
                    ai_model=ai_model,
                    ai_provider=ai_provider,
                )
                if auto_sector:
                    next_sector = auto_sector

            changed = (next_ticker != current_ticker) or (next_sector != current_sector)
            has_any_new = bool(next_ticker or next_sector)
            if changed and has_any_new:
                source_tag = "meta_refill_ai" if (force_refresh and ai_api_key) else "meta_refill"
                if not force_refresh:
                    source_tag = "auto_fill_ai" if ai_api_key else "auto_fill"
                upsert_company_list_entry(
                    company_name,
                    next_ticker,
                    sector=next_sector,
                    source=source_tag,
                )
                updated_count += 1
            else:
                skipped_count += 1
                if (not next_ticker) or (not next_sector):
                    unresolved.append(company_name)
        return updated_count, skipped_count, unresolved

    meta_btn_col1, meta_btn_col2, meta_btn_col3 = st.columns([1.1, 1.5, 1.1])
    with meta_btn_col1:
        auto_fill_missing_btn = st.button(
            "빈 티커/산업섹터 일괄 채우기 (API+AI)",
            key="analysis_fill_missing_company_meta_btn",
        )
    with meta_btn_col2:
        reset_and_refill_btn = st.button(
            "티커/산업섹터 전체 초기화 후 재탐색 (API+AI)",
            key="analysis_reset_refill_company_meta_btn",
        )
    with meta_btn_col3:
        refresh_price_btn = st.button(
            "현재 주가 일괄 불러오기 + DB저장 (API+AI)",
            key="analysis_fill_company_price_btn",
        )

    if auto_fill_missing_btn:
        ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
        targets = [row for row in overview_rows if not str(row.get("티커") or "").strip() or not str(row.get("산업섹터") or "").strip()]
        if not targets:
            st.info("이미 모든 기업에 티커/산업섹터 정보가 있습니다.")
        else:
            with st.spinner("빈 정보만 자동으로 찾는 중입니다..."):
                updated_count, skipped_count, unresolved = _bulk_fill_company_meta(
                    targets,
                    ai_provider=ai_provider,
                    ai_api_key=ai_api_key,
                    ai_model=ai_model,
                    force_refresh=False,
                )

            if updated_count > 0:
                st.session_state["analysis_bulk_fill_notice"] = (
                    f"일괄 채우기 완료: 업데이트 {updated_count}개, 유지/실패 {skipped_count}개"
                )
            elif skipped_count > 0:
                st.session_state["analysis_bulk_fill_notice"] = "일괄 채우기 결과: 새로 업데이트된 항목이 없습니다."
            if unresolved:
                preview = ", ".join(unresolved[:8])
                remain = len(unresolved) - min(8, len(unresolved))
                tail = f" 외 {remain}개" if remain > 0 else ""
                st.session_state["analysis_bulk_fill_warning"] = f"일부 항목은 여전히 빈 값입니다: {preview}{tail}"
            st.rerun()

    if reset_and_refill_btn:
        ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
        targets = [
            {"기업명": str(row.get("기업명") or "").strip(), "티커": "", "산업섹터": ""}
            for row in overview_rows
            if str(row.get("기업명") or "").strip()
        ]
        if not targets:
            st.info("재탐색할 기업이 없습니다.")
        else:
            cleared_count = clear_company_list_meta_all(source="manual_meta_reset")
            with st.spinner("티커/산업섹터 전체 초기화 후 재탐색 중입니다..."):
                updated_count, skipped_count, unresolved = _bulk_fill_company_meta(
                    targets,
                    ai_provider=ai_provider,
                    ai_api_key=ai_api_key,
                    ai_model=ai_model,
                    force_refresh=True,
                )
            st.session_state["analysis_bulk_reset_notice"] = (
                f"초기화+재탐색 완료: 초기화 {cleared_count}개, 업데이트 {updated_count}개, 유지/실패 {skipped_count}개"
            )
            if unresolved:
                preview = ", ".join(unresolved[:8])
                remain = len(unresolved) - min(8, len(unresolved))
                tail = f" 외 {remain}개" if remain > 0 else ""
                st.session_state["analysis_bulk_reset_warning"] = f"일부 항목은 여전히 미완성입니다: {preview}{tail}"
            st.rerun()

    if refresh_price_btn:
        ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
        targets = [row for row in overview_rows if str(row.get("기업명") or "").strip()]
        if not targets:
            st.info("주가를 갱신할 기업이 없습니다.")
        else:
            updated_count = 0
            failed_details: list[str] = []
            with st.spinner("기업 리스트 현재 주가를 일괄 불러오는 중입니다..."):
                for row in targets:
                    company_name = str(row.get("기업명") or "").strip()
                    current_ticker = clean_valid_ticker(str(row.get("티커") or ""))
                    if not company_name:
                        continue

                    ticker = current_ticker
                    if not ticker:
                        ticker, _ = resolve_ticker_auto_with_retry(
                            company_name,
                            use_ai=bool(ai_api_key),
                            api_key=ai_api_key,
                            model=ai_model,
                            provider=ai_provider,
                            market_preference=market_pref_map.get(company_name, ""),
                        )
                    if not ticker:
                        failed_details.append(f"{company_name}(티커 없음)")
                        continue

                    price_krw, price_src = fetch_current_price_krw_from_ticker(ticker, date.today())
                    if price_krw is None or float(price_krw) <= 0:
                        fail_reason = (price_src or "주가 조회 실패").strip()
                        failed_details.append(f"{company_name}({fail_reason})")
                        continue

                    upsert_company_list_entry(
                        company_name,
                        ticker=ticker,
                        source=None,
                        price_krw=float(price_krw),
                        price_source=price_src or "api",
                    )
                    updated_count += 1

            auto_save_suffix = ""
            auto_save_warn = ""
            if updated_count > 0:
                try:
                    target_date_text = str(st.session_state.get("editing_df_date", "") or "").strip()
                    target_date = _safe_parse_date(target_date_text) or date.today()
                    session_df = st.session_state.get("editing_df", pd.DataFrame())
                    if isinstance(session_df, pd.DataFrame) and not session_df.empty:
                        usd_rate, _ = get_usd_krw_rate_for_date(target_date)
                        refreshed_company_df = load_company_list()
                        price_exact, price_norm = build_company_price_krw_maps(refreshed_company_df)
                        recalced_df = recalculate_portfolio_from_price_and_avg_buy(
                            ensure_portfolio_columns(session_df.copy(), usd_rate, force_usd_rate=True),
                            usd_rate,
                            company_price_exact=price_exact,
                            company_price_norm=price_norm,
                        )
                        sync_ok, sync_msg = save_snapshot(
                            target_date,
                            recalced_df,
                            sync_to_github=True,
                            sync_reason="price_refresh_auto_save",
                        )
                        st.session_state["editing_df"] = recalced_df
                        st.session_state["editing_df_date"] = target_date.isoformat()
                        if sync_msg:
                            auto_save_suffix = f" / 스냅샷 저장·GitHub 동기화 {'완료' if sync_ok else '경고'}"
                        else:
                            auto_save_suffix = " / 스냅샷 자동 저장 완료"
                    else:
                        latest_date_text, latest_df = load_latest_snapshot()
                        if latest_df is not None and not latest_df.empty:
                            target_date2 = _safe_parse_date(latest_date_text) or date.today()
                            sync_ok, sync_msg = sync_snapshot_to_github_excel(target_date2, latest_df)
                            if sync_msg:
                                auto_save_suffix = f" / GitHub 동기화 {'완료' if sync_ok else '경고'}"
                except Exception as exc:
                    auto_save_warn = f"주가 갱신 후 자동 저장 실패: {exc}"

            if updated_count > 0:
                st.session_state["analysis_bulk_price_notice"] = f"현재 주가 업데이트 완료: {updated_count}개{auto_save_suffix}"
            else:
                st.session_state["analysis_bulk_price_notice"] = "현재 주가를 새로 업데이트하지 못했습니다."
            if auto_save_warn:
                failed_details.append(auto_save_warn)
            if failed_details:
                preview = ", ".join(failed_details[:5])
                remain = len(failed_details) - min(5, len(failed_details))
                tail = f" 외 {remain}개" if remain > 0 else ""
                st.session_state["analysis_bulk_price_warning"] = f"일부 실패: {preview}{tail}"
            st.rerun()

    if "analysis_bulk_fill_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_bulk_fill_notice"))
    if "analysis_bulk_fill_warning" in st.session_state:
        st.warning(st.session_state.pop("analysis_bulk_fill_warning"))
    if "analysis_bulk_price_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_bulk_price_notice"))
    if "analysis_bulk_price_warning" in st.session_state:
        st.warning(st.session_state.pop("analysis_bulk_price_warning"))
    if "analysis_bulk_reset_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_bulk_reset_notice"))
    if "analysis_bulk_reset_warning" in st.session_state:
        st.warning(st.session_state.pop("analysis_bulk_reset_warning"))
    if "analysis_bulk_selected_reset_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_bulk_selected_reset_notice"))
    if "analysis_bulk_selected_reset_warning" in st.session_state:
        st.warning(st.session_state.pop("analysis_bulk_selected_reset_warning"))
    if "analysis_bulk_selected_save_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_bulk_selected_save_notice"))
    if "analysis_bulk_selected_save_warning" in st.session_state:
        st.warning(st.session_state.pop("analysis_bulk_selected_save_warning"))
    if "analysis_bulk_selected_delete_notice" in st.session_state:
        st.success(st.session_state.pop("analysis_bulk_selected_delete_notice"))
    if "analysis_bulk_selected_delete_warning" in st.session_state:
        st.warning(st.session_state.pop("analysis_bulk_selected_delete_warning"))

    if overview_rows:
        st.caption("목록에서 기업 행을 선택하면 아래 기업명/티커 입력이 자동 선택됩니다.")
        search_keyword = st.text_input(
            "기업 종합검색",
            key="analysis_overview_search_keyword",
            placeholder="기업명/티커/산업섹터/구분/소스/등록일시/주가 포함 검색",
        )
        overview_df = pd.DataFrame(overview_rows)
        overview_df["현재주가(원화환산)"] = pd.to_numeric(overview_df["현재주가(원화환산)"], errors="coerce")
        overview_df["등록일시"] = pd.to_datetime(overview_df.get("등록일시"), errors="coerce")
        overview_df = overview_df.sort_values(
            ["등록일시", "기업명"],
            ascending=[False, True],
            na_position="last",
        ).reset_index(drop=True)
        search_text = str(search_keyword or "").strip()
        if search_text:
            needle = search_text.casefold()
            mask = pd.Series(False, index=overview_df.index)
            for col in ["기업명", "티커", "산업섹터", "구분", "리스트소스", "등록일시", "현재주가(원화환산)"]:
                series_text = overview_df[col].astype(str).str.casefold()
                mask = mask | series_text.str.contains(re.escape(needle), regex=True, na=False)
            overview_view_df = overview_df[mask].copy()
            st.caption(f"검색 결과: {len(overview_view_df):,} / 전체 {len(overview_df):,}")
        else:
            overview_view_df = overview_df.copy()
        current_input_name = (st.session_state.get("analysis_company_name_input") or "").strip()
        current_input_ticker = clean_valid_ticker(st.session_state.get("analysis_ticker_input") or "")
        selected_rows = []
        # 기업 리스트 표는 화면 길이 과확장을 막기 위해 고정 높이 스크롤을 사용한다.
        overview_height = 560
        try:
            table_event = st.dataframe(
                overview_view_df,
                height=overview_height,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "현재주가(원화환산)": st.column_config.NumberColumn("현재주가(원화환산)", format="localized"),
                    "등록일시": st.column_config.DatetimeColumn("등록일시", format="YYYY-MM-DD HH:mm:ss"),
                },
                on_select="rerun",
                selection_mode="multi-row",
                key="analysis_company_overview_table",
            )
            try:
                selected_rows = list(table_event.selection.rows)
            except Exception:
                selected_rows = []
        except TypeError:
            st.dataframe(overview_view_df, height=overview_height, use_container_width=True, hide_index=True)

        selected_names = []
        if selected_rows:
            for idx in selected_rows:
                try:
                    row_idx = int(idx)
                except Exception:
                    continue
                if 0 <= row_idx < len(overview_view_df):
                    nm = str(overview_view_df.iloc[row_idx].get("기업명") or "").strip()
                    if nm:
                        selected_names.append(nm)
        selected_names = list(dict.fromkeys(selected_names))
        st.session_state["analysis_selected_overview_names"] = selected_names

        sel_col1, sel_col2, sel_col3, sel_col4 = st.columns([1.1, 1.0, 0.95, 1.45])
        with sel_col1:
            reset_selected_btn = st.button(
                "선택 항목 초기화 후 재탐색 (API+AI)",
                key="analysis_reset_refill_selected_company_meta_btn",
            )
        with sel_col2:
            save_selected_btn = st.button(
                "선택 항목 수동 저장",
                key="analysis_save_selected_company_meta_btn",
            )
        with sel_col3:
            delete_selected_btn = st.button(
                "선택 관심목록 삭제",
                key="analysis_delete_selected_company_list_btn",
            )
        with sel_col4:
            st.caption(f"현재 선택: {len(selected_names):,}개")

        if reset_selected_btn:
            targets = [{"기업명": nm, "티커": "", "산업섹터": ""} for nm in selected_names]
            if not targets:
                st.warning("먼저 표에서 기업을 1개 이상 체크해 주세요.")
            else:
                ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
                cleared_count = clear_company_list_meta_by_names(
                    selected_names,
                    source="manual_meta_reset_selected",
                )
                with st.spinner("선택 기업 티커/산업섹터 초기화 후 재탐색 중입니다..."):
                    updated_count, skipped_count, unresolved = _bulk_fill_company_meta(
                        targets,
                        ai_provider=ai_provider,
                        ai_api_key=ai_api_key,
                        ai_model=ai_model,
                        force_refresh=True,
                    )
                st.session_state["analysis_bulk_selected_reset_notice"] = (
                    f"선택 초기화+재탐색 완료: 초기화 {cleared_count}개, 업데이트 {updated_count}개, 유지/실패 {skipped_count}개"
                )
                if unresolved:
                    preview = ", ".join(unresolved[:8])
                    remain = len(unresolved) - min(8, len(unresolved))
                    tail = f" 외 {remain}개" if remain > 0 else ""
                    st.session_state["analysis_bulk_selected_reset_warning"] = (
                        f"선택 항목 중 일부는 여전히 미완성입니다: {preview}{tail}"
                    )
                st.rerun()

        if save_selected_btn:
            if not selected_rows:
                st.warning("먼저 표에서 기업을 1개 이상 체크해 주세요.")
            else:
                saved_count = 0
                skipped: list[str] = []
                for idx in selected_rows:
                    try:
                        row_idx = int(idx)
                    except Exception:
                        continue
                    if not (0 <= row_idx < len(overview_view_df)):
                        continue
                    row_view = overview_view_df.iloc[row_idx]
                    company_name = str(row_view.get("기업명") or "").strip()
                    ticker = clean_valid_ticker(str(row_view.get("티커") or ""))
                    sector = str(row_view.get("산업섹터") or "").strip()
                    if not company_name:
                        continue
                    if not ticker and not sector:
                        skipped.append(company_name)
                        continue
                    upsert_company_list_entry(
                        company_name,
                        ticker=ticker,
                        sector=sector,
                        source="manual_save_selected",
                    )
                    saved_count += 1
                st.session_state["analysis_bulk_selected_save_notice"] = (
                    f"선택 항목 저장 완료: {saved_count}개"
                )
                if skipped:
                    preview = ", ".join(skipped[:6])
                    remain = len(skipped) - min(6, len(skipped))
                    tail = f" 외 {remain}개" if remain > 0 else ""
                    st.session_state["analysis_bulk_selected_save_warning"] = (
                        f"티커/섹터가 비어 저장하지 못한 항목: {preview}{tail}"
                    )
                st.rerun()

        if delete_selected_btn:
            if not selected_names:
                st.warning("먼저 표에서 기업을 1개 이상 체크해 주세요.")
            else:
                deleted_count = 0
                skipped_non_watch: list[str] = []
                for name in selected_names:
                    nm = str(name or "").strip()
                    if not nm:
                        continue
                    if nm not in listed_set:
                        skipped_non_watch.append(nm)
                        continue
                    delete_company_list_entry(nm)
                    deleted_count += 1
                    if (st.session_state.get("analysis_company_name_input") or "").strip() == nm:
                        st.session_state["analysis_company_name_pending"] = ""
                        st.session_state["analysis_ticker_pending"] = ""

                if deleted_count > 0:
                    st.session_state["analysis_bulk_selected_delete_notice"] = (
                        f"선택 관심목록 삭제 완료: {deleted_count}개"
                    )
                if skipped_non_watch:
                    preview = ", ".join(skipped_non_watch[:6])
                    remain = len(skipped_non_watch) - min(6, len(skipped_non_watch))
                    tail = f" 외 {remain}개" if remain > 0 else ""
                    st.session_state["analysis_bulk_selected_delete_warning"] = (
                        f"보유종목만 존재해 관심목록 삭제를 건너뜀: {preview}{tail}"
                    )
                if deleted_count == 0 and not skipped_non_watch:
                    st.session_state["analysis_bulk_selected_delete_warning"] = "삭제 가능한 관심목록 항목이 없습니다."
                st.rerun()

        if selected_rows:
            row_idx = int(selected_rows[0])
            if 0 <= row_idx < len(overview_view_df):
                picked_name = str(overview_view_df.iloc[row_idx].get("기업명") or "").strip()
                picked_ticker = clean_valid_ticker(str(overview_view_df.iloc[row_idx].get("티커") or ""))
                picked_sector = str(overview_view_df.iloc[row_idx].get("산업섹터") or "").strip()
                prev_selected_name = _sanitize_widget_text(st.session_state.get("analysis_selected_overview_company"), "")
                if picked_name and picked_name != prev_selected_name:
                    st.session_state["analysis_selected_overview_company"] = picked_name
                    st.session_state["analysis_selected_overview_ticker_input"] = picked_ticker
                    st.session_state["analysis_selected_overview_sector_input"] = picked_sector
                need_apply = False
                if picked_name and picked_name != current_input_name:
                    st.session_state["analysis_company_name_pending"] = picked_name
                    need_apply = True
                if picked_ticker and picked_ticker != current_input_ticker:
                    st.session_state["analysis_ticker_pending"] = picked_ticker
                    need_apply = True
                if need_apply:
                    st.session_state["analysis_company_hint"] = "직접입력"
                    st.rerun()

    selected_company_for_edit = _sanitize_widget_text(st.session_state.get("analysis_selected_overview_company"), "")
    if selected_company_for_edit:
        st.markdown("##### 선택 기업 티커/섹터 수정")
        edit_col1, edit_col2, edit_col3, edit_col4, edit_col5 = st.columns([1.0, 1.0, 1.1, 0.75, 0.85])
        with edit_col1:
            st.caption(f"선택 기업: **{selected_company_for_edit}**")
        with edit_col2:
            st.text_input(
                "수정 티커",
                key="analysis_selected_overview_ticker_input",
                placeholder="예: NOV / 005930.KS",
            )
        with edit_col3:
            st.text_input(
                "수정 산업섹터(선택)",
                key="analysis_selected_overview_sector_input",
                placeholder="예: Energy / 반도체",
            )
        with edit_col4:
            save_selected_meta_btn = st.button("선택값 저장", key="analysis_save_selected_meta_btn")
        with edit_col5:
            clear_selected_ticker_btn = st.button("티커 비우기", key="analysis_clear_selected_ticker_btn")

        if save_selected_meta_btn:
            ticker_raw = _sanitize_widget_text(st.session_state.get("analysis_selected_overview_ticker_input"), "")
            sector_new = _sanitize_widget_text(st.session_state.get("analysis_selected_overview_sector_input"), "")
            ticker_new = clean_valid_ticker(ticker_raw)
            if ticker_raw and not ticker_new:
                st.warning("티커 형식이 올바르지 않습니다. 예: NOV, AAPL, 005930.KS")
            elif not ticker_new and not sector_new:
                st.warning("수정할 티커 또는 산업섹터를 입력해 주세요.")
            else:
                upsert_company_list_entry(
                    selected_company_for_edit,
                    ticker=ticker_new,
                    sector=sector_new,
                    source="manual_edit",
                )
                if ticker_new:
                    st.session_state["analysis_ticker_pending"] = ticker_new
                    st.session_state["analysis_ticker_source"] = "기업 리스트 수동 수정"
                st.session_state["analysis_company_name_pending"] = selected_company_for_edit
                st.session_state["analysis_company_hint"] = "직접입력"
                st.session_state["analysis_ticker_autofill_notice"] = (
                    f"{selected_company_for_edit} 수정 저장 완료"
                    + (f" (티커 {ticker_new})" if ticker_new else "")
                )
                st.rerun()

        if clear_selected_ticker_btn:
            cleared = clear_company_list_ticker(selected_company_for_edit, source="manual_edit_clear")
            if not cleared:
                st.warning("티커를 비울 기업을 찾지 못했습니다.")
            else:
                st.session_state["analysis_selected_overview_ticker_input"] = ""
                st.session_state["analysis_ticker_pending"] = ""
                st.session_state["analysis_ticker_source"] = "기업 리스트 수동 초기화"
                st.session_state["analysis_company_name_pending"] = selected_company_for_edit
                st.session_state["analysis_company_hint"] = "직접입력"
                st.session_state["analysis_ticker_autofill_notice"] = (
                    f"{selected_company_for_edit} 티커를 비우고 주가 캐시를 초기화했습니다."
                )
                st.rerun()

    st.markdown("##### AI 옵션")
    st.checkbox("yfinance 티커 검색 실패 시 AI 티커 추론 사용", key="analysis_use_ai_ticker")
    st.caption("AI 키/모델/제공자는 API 설정 탭의 공통값을 사용합니다.")
    st.caption("기업분석 탐색원: (국내) Naver -> yfinance -> Google -> Toss -> Alpha/Finnhub, (해외) yfinance -> Naver -> Google -> Toss -> Alpha/Finnhub")

    c1, c2, c3, c4 = st.columns([1, 1.3, 1.1, 1.2])
    with c1:
        st.date_input("분석일", key="analysis_date")
    with c2:
        st.text_input("기업명", key="analysis_company_name_input")
    with c3:
        st.text_input("티커", key="analysis_ticker_input", placeholder="005930.KS / AAPL")
    with c4:
        hint_options = ["직접입력"] + options
        if (
            "analysis_company_hint" in st.session_state
            and st.session_state["analysis_company_hint"] not in hint_options
        ):
            st.session_state["analysis_company_hint"] = "직접입력"
        selected = st.selectbox("기업 리스트", hint_options, index=0, key="analysis_company_hint")
        if selected != "직접입력":
            current_input_name = (st.session_state.get("analysis_company_name_input") or "").strip()
            if selected != current_input_name:
                st.session_state["analysis_company_name_pending"] = selected
                st.rerun()

    analysis_company_name_value = (st.session_state.get("analysis_company_name_input") or "").strip()
    analysis_ticker_raw = _sanitize_widget_text(st.session_state.get("analysis_ticker_input"), "")
    analysis_ticker_value = clean_valid_ticker(analysis_ticker_raw)
    if analysis_ticker_raw and analysis_ticker_value and analysis_ticker_raw != analysis_ticker_value:
        st.session_state["analysis_ticker_input"] = analysis_ticker_value
    st.session_state["analysis_company_name"] = analysis_company_name_value
    st.session_state["analysis_ticker"] = analysis_ticker_value

    analysis_ai_provider, analysis_ai_api_key, analysis_ai_model = get_ai_settings_from_session("analysis")

    company_name = analysis_company_name_value
    if company_name and (company_name != st.session_state.get("analysis_prev_company") or not analysis_ticker_value):
        list_ticker = get_company_list_ticker(company_name)
        market_pref = market_pref_map.get(company_name, "")
        list_is_kr = bool(list_ticker.endswith(".KS") or list_ticker.endswith(".KQ"))
        q_norm = normalize_company_name_for_match(company_name)
        short_hangul_name = bool(re.search(r"[가-힣]", company_name)) and len(q_norm) <= 4
        need_recheck = (not list_ticker) or (
            short_hangul_name and list_is_kr and _market_pref_normalized(market_pref) != "domestic"
        )

        chosen_ticker = list_ticker
        chosen_source = "기업 리스트 저장값" if list_ticker else ""
        if need_recheck:
            tkr, src = resolve_ticker_auto_with_retry(
                company_name,
                use_ai=bool(st.session_state.get("analysis_use_ai_ticker", False)),
                api_key=analysis_ai_api_key,
                model=analysis_ai_model,
                provider=analysis_ai_provider,
                market_preference=market_pref,
            )
            tkr_is_kr = bool(tkr.endswith(".KS") or tkr.endswith(".KQ")) if tkr else False
            prefer_auto = bool(tkr) and (
                not list_ticker
                or (list_is_kr and not tkr_is_kr)
                or str(src or "").startswith("웹검색")
            )
            if prefer_auto:
                chosen_ticker = tkr
                chosen_source = src

        if chosen_ticker:
            if chosen_ticker != analysis_ticker_value:
                st.session_state["analysis_ticker_pending"] = chosen_ticker
                st.session_state["analysis_ticker_autofill_notice"] = (
                    f"티커 자동 입력: {chosen_ticker} ({chosen_source or '자동 탐색'})"
                )
            st.session_state["analysis_ticker_source"] = chosen_source or "자동 탐색"
        st.session_state["analysis_prev_company"] = company_name

        latest_df = analysis_all[analysis_all["stock_name"] == company_name] if not analysis_all.empty else pd.DataFrame()
        if not latest_df.empty:
            latest = latest_df.sort_values(["analysis_date", "updated_at"], ascending=False).iloc[0]
            if not analysis_ticker_value:
                latest_ticker = clean_valid_ticker(latest.get("ticker") or "")
                if latest_ticker:
                    st.session_state["analysis_ticker_pending"] = latest_ticker
            latest_financial = parse_financial_summary_json(latest.get("financial_summary_json"))
            if latest_financial:
                st.session_state["analysis_financial_summary_cache"] = latest_financial
            st.session_state["analysis_company_overview"] = latest.get("company_overview") or ""
            st.session_state["analysis_products_services"] = latest.get("products_services") or ""
            st.session_state["analysis_raw_materials"] = latest.get("raw_materials") or ""
            st.session_state["analysis_profit_up_factors"] = latest.get("profit_up_factors") or ""
            st.session_state["analysis_profit_down_factors"] = latest.get("profit_down_factors") or ""
            st.session_state["analysis_key_takeaway"] = latest.get("note") or ""
        if "analysis_ticker_pending" in st.session_state:
            st.rerun()

    action_col1, action_col2, action_col3 = st.columns([1, 1.4, 1])
    with action_col1:
        auto_ticker_btn = st.button("티커 자동 입력", key="analysis_auto_ticker_btn")
    with action_col2:
        generate_btn = st.button("AI로 기업 분석 생성 및 저장", type="primary", key="analysis_generate_btn")
    with action_col3:
        manual_save_btn = st.button("현재 내용 수동 저장", key="analysis_manual_save_btn")

    src = st.session_state.get("analysis_ticker_source", "")
    if src:
        st.caption(f"티커 소스: {src}")

    if auto_ticker_btn:
        tkr, src = resolve_ticker_auto_with_retry(
            analysis_company_name_value,
            use_ai=bool(st.session_state.get("analysis_use_ai_ticker", False)),
            api_key=analysis_ai_api_key,
            model=analysis_ai_model,
            provider=analysis_ai_provider,
            market_preference=market_pref_map.get(analysis_company_name_value, ""),
        )
        if tkr:
            st.session_state["analysis_ticker_pending"] = tkr
            st.session_state["analysis_ticker_source"] = src
            st.session_state["analysis_ticker_autofill_notice"] = f"티커 자동 입력: {tkr} ({src})"
            st.rerun()
        else:
            st.warning(src or "티커를 찾지 못했습니다.")

    financial_summary = {}
    financial_error = ""

    if generate_btn:
        company_name = analysis_company_name_value
        if not company_name:
            st.error("기업명을 입력해 주세요.")
        else:
            ticker, financial_summary, financial_source, financial_error = fetch_company_financial_summary_for_analysis(
                company_name=company_name,
                ticker_input=analysis_ticker_value,
                use_ai_ticker=bool(st.session_state.get("analysis_use_ai_ticker", False)),
                ai_provider=analysis_ai_provider,
                ai_api_key=analysis_ai_api_key,
                ai_model=analysis_ai_model,
                market_preference=market_pref_map.get(company_name, ""),
            )
            ticker = clean_valid_ticker(ticker)
            if ticker and ticker != analysis_ticker_value:
                st.session_state["analysis_ticker_pending"] = ticker
                st.session_state["analysis_ticker_source"] = "기업분석 재탐색"

            if financial_error:
                st.warning(financial_error)
            if financial_summary:
                st.session_state["analysis_financial_summary_cache"] = financial_summary
                st.caption(f"재무 데이터 소스: {financial_source}")
            google_context_text, google_research_note = fetch_google_company_research_context(
                company_name=company_name,
                ticker=ticker,
                max_items=8,
            )
            if google_research_note:
                st.caption(f"기업 실체 정보 소스: {google_research_note}")

            used_ai_provider = normalize_ai_provider(analysis_ai_provider)
            used_ai_model = analysis_ai_model
            analysis = {}
            ai_err = ""

            if (analysis_ai_api_key or "").strip():
                analysis, ai_err, research_note = generate_company_analysis_with_ai(
                    company_name=company_name,
                    ticker=ticker,
                    financial_summary=financial_summary,
                    api_key=analysis_ai_api_key,
                    model=analysis_ai_model,
                    provider=analysis_ai_provider,
                    google_context_text=google_context_text,
                    google_research_note=google_research_note,
                )
                if research_note:
                    st.caption(f"기업 실체 정보 소스: {research_note}")
            else:
                ai_err = "AI API 키가 없어 템플릿 기반으로 분석을 생성합니다."

            if ai_err and (analysis_ai_api_key or "").strip():
                fallback_provider = "openai" if normalize_ai_provider(analysis_ai_provider) == "claude" else "claude"
                fallback_key = (
                    st.session_state.get("global_openai_api_key", "")
                    if fallback_provider == "openai"
                    else st.session_state.get("global_claude_api_key", "")
                )
                fallback_model = (
                    st.session_state.get("global_openai_model", DEFAULT_OPENAI_MODEL)
                    if fallback_provider == "openai"
                    else st.session_state.get("global_claude_model", DEFAULT_CLAUDE_MODEL)
                )
                fallback_key = (fallback_key or "").strip()
                fallback_model = (fallback_model or "").strip()
                if fallback_key:
                    st.warning(
                        f"{ai_provider_label(analysis_ai_provider)} 호출 실패로 "
                        f"{ai_provider_label(fallback_provider)} 모델로 1회 재시도합니다."
                    )
                    fallback_analysis, fallback_err, fallback_research_note = generate_company_analysis_with_ai(
                        company_name=company_name,
                        ticker=ticker,
                        financial_summary=financial_summary,
                        api_key=fallback_key,
                        model=fallback_model,
                        provider=fallback_provider,
                        google_context_text=google_context_text,
                        google_research_note=google_research_note,
                    )
                    if fallback_research_note:
                        st.caption(f"기업 실체 정보 소스: {fallback_research_note}")
                    if not fallback_err and fallback_analysis:
                        analysis = fallback_analysis
                        ai_err = ""
                        used_ai_provider = fallback_provider
                        used_ai_model = fallback_model
                        st.caption(f"AI 생성 소스: {ai_provider_label(used_ai_provider)}")
                    else:
                        ai_err = f"{ai_err} | 대체 호출 실패: {fallback_err}"

            if ai_err:
                st.warning(ai_err)
                analysis = generate_company_analysis_template(
                    company_name=company_name,
                    ticker=ticker,
                    financial_summary=financial_summary,
                    google_context_text=google_context_text,
                )
                used_ai_provider = "template"
                used_ai_model = "rule-based"
                st.caption("AI 실패/미설정으로 템플릿 기반 분석을 생성했습니다.")

            st.session_state["analysis_company_overview"] = analysis.get("company_overview", "")
            st.session_state["analysis_products_services"] = analysis.get("products_services", "")
            st.session_state["analysis_raw_materials"] = analysis.get("raw_materials", "")
            st.session_state["analysis_profit_up_factors"] = analysis.get("profit_up_factors", "")
            st.session_state["analysis_profit_down_factors"] = analysis.get("profit_down_factors", "")
            st.session_state["analysis_key_takeaway"] = analysis.get("key_takeaway", "")
            save_company_analysis(
                analysis_date=st.session_state["analysis_date"],
                stock_name=company_name,
                ticker=ticker,
                financial_summary=financial_summary,
                analysis=analysis,
                source=f"ai:{used_ai_provider}",
                ai_model=f"{ai_provider_label(used_ai_provider)}:{used_ai_model}",
                note=analysis.get("key_takeaway", ""),
            )
            upsert_company_list_entry(
                company_name,
                ticker,
                sector=(financial_summary.get("sector") or "") if isinstance(financial_summary, dict) else "",
                source="analysis",
            )
            st.success(f"{company_name} 기업 분석 생성/저장을 완료했습니다.")

    st.markdown("#### 기업 분석 내용")
    st.text_area(
        "기업 개요",
        key="analysis_company_overview",
        height=estimate_textarea_height(
            st.session_state.get("analysis_company_overview", ""),
            min_height=96,
            max_height=340,
            chars_per_line=62,
        ),
    )
    st.text_area(
        "핵심 제품/서비스·돈 버는 방식",
        key="analysis_products_services",
        height=estimate_textarea_height(
            st.session_state.get("analysis_products_services", ""),
            min_height=96,
            max_height=420,
            chars_per_line=62,
        ),
    )
    st.text_area(
        "핵심 원재료/투입요소",
        key="analysis_raw_materials",
        height=estimate_textarea_height(
            st.session_state.get("analysis_raw_materials", ""),
            min_height=96,
            max_height=380,
            chars_per_line=62,
        ),
    )
    st.text_area(
        "이익 증가 요인·좋은 변화(촉매)",
        key="analysis_profit_up_factors",
        height=estimate_textarea_height(
            st.session_state.get("analysis_profit_up_factors", ""),
            min_height=96,
            max_height=430,
            chars_per_line=62,
        ),
    )
    st.text_area(
        "이익 감소 요인(리스크)",
        key="analysis_profit_down_factors",
        height=estimate_textarea_height(
            st.session_state.get("analysis_profit_down_factors", ""),
            min_height=96,
            max_height=430,
            chars_per_line=62,
        ),
    )
    st.text_area(
        "요약 메모",
        key="analysis_key_takeaway",
        height=estimate_textarea_height(
            st.session_state.get("analysis_key_takeaway", ""),
            min_height=90,
            max_height=300,
            chars_per_line=62,
        ),
    )

    report_analysis = {
        "company_overview": st.session_state.get("analysis_company_overview", ""),
        "products_services": st.session_state.get("analysis_products_services", ""),
        "raw_materials": st.session_state.get("analysis_raw_materials", ""),
        "profit_up_factors": st.session_state.get("analysis_profit_up_factors", ""),
        "profit_down_factors": st.session_state.get("analysis_profit_down_factors", ""),
        "key_takeaway": st.session_state.get("analysis_key_takeaway", ""),
    }
    report_has_content = any(str(v or "").strip() for v in report_analysis.values())
    report_company_name = (analysis_company_name_value or "").strip()
    report_ticker = clean_valid_ticker(analysis_ticker_value)
    report_file_name = _safe_report_filename(
        company_name=report_company_name or "company",
        ticker=report_ticker or "ticker",
        analysis_date_value=st.session_state.get("analysis_date"),
    )
    if report_has_content:
        if HAS_PYTHON_DOCX:
            try:
                report_bytes = build_company_analysis_docx_bytes(
                    company_name=report_company_name,
                    ticker=report_ticker,
                    analysis_date_value=st.session_state.get("analysis_date"),
                    analysis_fields=report_analysis,
                    financial_summary=st.session_state.get("analysis_financial_summary_cache", {}),
                )
                st.download_button(
                    "기업 분석 보고서 다운로드 (Word)",
                    data=report_bytes,
                    file_name=report_file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="analysis_download_word_report_btn",
                    use_container_width=False,
                )
            except Exception as exc:
                st.warning(f"워드 보고서 생성 실패: {exc}")
        else:
            st.warning("워드 다운로드를 위해 `python-docx` 설치가 필요합니다.")

    if manual_save_btn:
        company_name = analysis_company_name_value
        ticker = analysis_ticker_value
        financial_summary = st.session_state.get("analysis_financial_summary_cache", {})
        if not company_name:
            st.error("기업명을 입력해 주세요.")
        else:
            resolved_ticker, fetched_summary, fetched_source, fetched_err = fetch_company_financial_summary_for_analysis(
                company_name=company_name,
                ticker_input=ticker,
                use_ai_ticker=bool(st.session_state.get("analysis_use_ai_ticker", False)),
                ai_provider=analysis_ai_provider,
                ai_api_key=analysis_ai_api_key,
                ai_model=analysis_ai_model,
                market_preference=market_pref_map.get(company_name, ""),
            )
            if resolved_ticker:
                ticker = resolved_ticker
                if ticker != analysis_ticker_value:
                    st.session_state["analysis_ticker_pending"] = ticker
                    st.session_state["analysis_ticker_source"] = "기업분석 재탐색"
            if fetched_err:
                st.warning(f"재무제표 조회 경고: {fetched_err}")
            if fetched_summary:
                financial_summary = fetched_summary
                st.session_state["analysis_financial_summary_cache"] = fetched_summary
                st.caption(f"재무 데이터 소스: {fetched_source}")
            save_company_analysis(
                analysis_date=st.session_state["analysis_date"],
                stock_name=company_name,
                ticker=ticker,
                financial_summary=financial_summary,
                analysis={
                    "company_overview": st.session_state.get("analysis_company_overview", ""),
                    "products_services": st.session_state.get("analysis_products_services", ""),
                    "raw_materials": st.session_state.get("analysis_raw_materials", ""),
                    "profit_up_factors": st.session_state.get("analysis_profit_up_factors", ""),
                    "profit_down_factors": st.session_state.get("analysis_profit_down_factors", ""),
                    "key_takeaway": st.session_state.get("analysis_key_takeaway", ""),
                },
                source="manual",
                ai_model=f"{ai_provider_label(analysis_ai_provider)}:{analysis_ai_model}",
                note=st.session_state.get("analysis_key_takeaway", ""),
            )
            upsert_company_list_entry(
                company_name,
                ticker,
                sector=(financial_summary.get("sector") or "") if isinstance(financial_summary, dict) else "",
                source="analysis",
            )
            st.success(f"{company_name} 분석 내용을 수동 저장했습니다.")

    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">기업별 재무제표/분석 이력</div>', unsafe_allow_html=True)
    analysis_all = load_company_analysis_history()
    if analysis_all.empty:
        st.info("저장된 기업 분석 이력이 없습니다.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    latest_by_company = (
        analysis_all.sort_values(["analysis_date", "updated_at"], ascending=[False, False])
        .drop_duplicates(subset=["stock_name"], keep="first")
        .copy()
    )
    latest_by_company["analysis_date"] = pd.to_datetime(latest_by_company["analysis_date"]).dt.date
    company_list_view = latest_by_company.rename(
        columns={
            "analysis_date": "최근분석일",
            "stock_name": "기업명",
            "ticker": "티커",
            "source": "생성방식",
            "ai_model": "모델",
        }
    )
    st.caption("저장된 기업 리스트 (최신 분석 기준)")
    company_list_view_df = company_list_view[["기업명", "티커", "최근분석일", "생성방식", "모델"]]
    st.dataframe(
        company_list_view_df,
        height=estimate_dataframe_height(company_list_view_df, min_height=220, max_height=2200),
        use_container_width=True,
        hide_index=True,
    )

    target_name = (st.session_state.get("analysis_company_name") or "").strip()
    if target_name:
        target_df = analysis_all[analysis_all["stock_name"] == target_name].copy()
        if target_df.empty:
            st.info(f"'{target_name}' 기업의 저장된 분석 이력이 없습니다.")
            st.markdown("</div>", unsafe_allow_html=True)
            return
    else:
        target_df = analysis_all.copy()

    target_df = target_df.sort_values(["analysis_date", "updated_at"], ascending=[False, False]).copy()
    target_df["analysis_date"] = pd.to_datetime(target_df["analysis_date"])
    pick_labels = []
    for _, row in target_df.iterrows():
        date_label = row["analysis_date"].date().isoformat()
        src_label = str(row.get("source") or "-")
        tkr_label = str(row.get("ticker") or "-")
        pick_labels.append(f"{date_label} | {src_label} | {tkr_label}")
    selected_idx = st.selectbox(
        "상세 조회 이력",
        options=list(range(len(pick_labels))),
        format_func=lambda i: pick_labels[i],
        key=f"analysis_history_pick_{target_name or 'all'}",
    )

    latest = target_df.iloc[int(selected_idx)]
    financial = parse_financial_summary_json(latest.get("financial_summary_json"))
    hist_info_col1, hist_info_col2 = st.columns([1.4, 1])
    with hist_info_col1:
        st.caption(
            "선택 이력: "
            f"{pd.to_datetime(latest.get('analysis_date')).date().isoformat()} | "
            f"{str(latest.get('source') or '-')} | "
            f"{str(latest.get('ticker') or '-')}"
        )
    with hist_info_col2:
        if st.button("선택 이력 내용을 상단 분석 내용에 불러오기", key=f"analysis_apply_history_btn_{target_name or 'all'}"):
            apply_analysis_history_to_editor(latest)
            st.success("선택한 이력 내용을 상단 '기업 분석 내용'에 반영했습니다.")
            st.rerun()

    metric_keys = [
        ("market_cap", "시가총액"),
        ("enterprise_value", "기업가치(EV)"),
        ("total_revenue", "매출"),
        ("ebitda", "EBITDA"),
        ("net_income_to_common", "당기순이익"),
        ("operating_cashflow", "영업현금흐름"),
        ("free_cashflow", "잉여현금흐름"),
        ("trailing_pe", "PER"),
        ("price_to_book", "PBR"),
        ("roe_pct", "ROE(%)"),
        ("operating_margin_pct", "영업이익률(%)"),
        ("revenue_growth_pct", "매출성장률(%)"),
        ("earnings_growth_pct", "이익성장률(%)"),
        ("debt_to_equity", "부채비율"),
        ("current_ratio", "유동비율"),
    ]
    metric_rows = []
    for key, label in metric_keys:
        if key in financial and financial.get(key) is not None:
            metric_rows.append({"항목": label, "값": financial.get(key)})
    if metric_rows:
        st.caption("핵심 재무 지표")
        metric_df = format_table_numbers(pd.DataFrame(metric_rows))
        st.dataframe(
            metric_df,
            height=estimate_dataframe_height(metric_df, min_height=180, max_height=1400),
            use_container_width=True,
            hide_index=True,
        )

    income_rows = financial.get("income_statement_annual") if isinstance(financial, dict) else None
    if income_rows:
        st.caption("연간 손익계산서 요약")
        income_df = format_table_numbers(pd.DataFrame(income_rows))
        st.dataframe(
            income_df,
            height=estimate_dataframe_height(income_df, min_height=180, max_height=2400),
            use_container_width=True,
            hide_index=True,
        )

    balance_rows = financial.get("balance_sheet_annual") if isinstance(financial, dict) else None
    if balance_rows:
        st.caption("연간 재무상태표 요약")
        balance_df = format_table_numbers(pd.DataFrame(balance_rows))
        st.dataframe(
            balance_df,
            height=estimate_dataframe_height(balance_df, min_height=180, max_height=2400),
            use_container_width=True,
            hide_index=True,
        )

    cash_rows = financial.get("cashflow_annual") if isinstance(financial, dict) else None
    if cash_rows:
        st.caption("연간 현금흐름표 요약")
        cash_df = format_table_numbers(pd.DataFrame(cash_rows))
        st.dataframe(
            cash_df,
            height=estimate_dataframe_height(cash_df, min_height=180, max_height=2400),
            use_container_width=True,
            hide_index=True,
        )

    hist_view = target_df.sort_values(["analysis_date", "updated_at"], ascending=[False, False]).copy()
    hist_view["analysis_date"] = pd.to_datetime(hist_view["analysis_date"]).dt.date
    hist_view = hist_view.rename(
        columns={
            "analysis_date": "분석일",
            "stock_name": "기업명",
            "ticker": "티커",
            "source": "생성방식",
            "ai_model": "모델",
            "updated_at": "수정시각",
        }
    )
    st.caption("분석 이력")
    hist_view_df = hist_view[["분석일", "기업명", "티커", "생성방식", "모델", "수정시각"]]
    st.dataframe(
        hist_view_df,
        height=estimate_dataframe_height(hist_view_df, min_height=220, max_height=2400),
        use_container_width=True,
        hide_index=True,
    )
    st.markdown("</div>", unsafe_allow_html=True)


def render_value_chain_tab() -> None:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">밸류체인</div>', unsafe_allow_html=True)
    st.caption(
        "밸류체인 이미지를 올리면 단계/세부공정/기업명을 추출한 뒤 "
        "관심기업(기업정보 리스트)과 자동 매칭해서 시각화합니다."
    )

    if "value_chain_uploader_nonce" not in st.session_state:
        st.session_state["value_chain_uploader_nonce"] = 0
    if "value_chain_result" not in st.session_state:
        st.session_state["value_chain_result"] = {}
    if "value_chain_notice" not in st.session_state:
        st.session_state["value_chain_notice"] = ""
    if "value_chain_warning" not in st.session_state:
        st.session_state["value_chain_warning"] = ""
    if "value_chain_summary_notice" not in st.session_state:
        st.session_state["value_chain_summary_notice"] = ""
    if "value_chain_summary_warning" not in st.session_state:
        st.session_state["value_chain_summary_warning"] = ""
    if "value_chain_save_name" not in st.session_state:
        st.session_state["value_chain_save_name"] = ""

    # ── 저장된 밸류체인 불러오기 ────────────────────────────────────────
    saved_chains = load_value_chains()
    if saved_chains:
        st.markdown("##### 저장된 밸류체인 불러오기")
        saved_options = {f"[{c['id']}] {c['chain_name']} ({c['created_at'][:10]})": c["id"] for c in saved_chains}
        load_col, del_col = st.columns([3, 1])
        with load_col:
            selected_label = st.selectbox(
                "저장 목록",
                options=["(새로 분석)"] + list(saved_options.keys()),
                key="value_chain_select_saved",
                label_visibility="collapsed",
            )
        with del_col:
            if selected_label != "(새로 분석)":
                if st.button("삭제", key="value_chain_delete_saved_btn", type="secondary"):
                    delete_value_chain(saved_options[selected_label])
                    st.session_state["value_chain_result"] = {}
                    st.rerun()
        if selected_label != "(새로 분석)":
            selected_id = saved_options[selected_label]
            if st.button("불러오기", key="value_chain_load_btn", type="primary"):
                loaded = load_value_chain_by_id(selected_id)
                if loaded:
                    chain_data = loaded["chain_json"]
                    chain_data["matched_rows"] = []
                    chain_data["created_at"] = loaded["created_at"]
                    chain_data["overview_text"] = chain_data.get("overview_text", "")
                    chain_data["overview_source"] = chain_data.get("overview_source", "")
                    company_list_df2 = load_company_list()
                    chain_data["matched_rows"] = _build_value_chain_match_rows(chain_data, company_list_df2)
                    st.session_state["value_chain_result"] = chain_data
                    st.session_state["value_chain_notice"] = f"'{loaded['chain_name']}' 밸류체인을 불러왔습니다."
                    st.rerun()
        st.divider()

    uploader_key = f"value_chain_image_uploader_{st.session_state['value_chain_uploader_nonce']}"
    up_col, opt_col = st.columns([1.8, 1])
    with up_col:
        vc_file = st.file_uploader(
            "밸류체인 이미지",
            type=["png", "jpg", "jpeg", "webp"],
            key=uploader_key,
        )
    with opt_col:
        show_only_matched = st.checkbox("관심기업 매칭 결과만 보기", value=False, key="value_chain_show_only_matched")
        run_btn = st.button("이미지로 밸류체인 분석", type="primary", key="value_chain_run_btn")

    if vc_file is not None:
        st.image(vc_file, caption="업로드된 밸류체인 이미지", use_container_width=True)

    if run_btn:
        if vc_file is None:
            st.warning("먼저 밸류체인 이미지를 업로드해 주세요.")
        else:
            img_bytes = vc_file.getvalue()
            img_hash = _compute_image_hash(img_bytes)
            # 동일 이미지 중복 체크
            existing_chains = load_value_chains()
            dup_chain = next((c for c in existing_chains if c.get("source_image_hash") == img_hash), None)
            if dup_chain:
                st.warning(
                    f"동일한 이미지로 이미 저장된 밸류체인이 있습니다: "
                    f"[{dup_chain['id']}] {dup_chain['chain_name']} ({dup_chain['created_at'][:10]}). "
                    "저장 목록에서 불러오거나 다른 이미지를 사용하세요."
                )
            else:
                ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
                with st.spinner("밸류체인 이미지에서 단계/기업을 추출하고 매칭 중입니다..."):
                    parsed, parse_err = extract_value_chain_from_image_with_ai(
                        image_bytes=img_bytes,
                        mime_type=getattr(vc_file, "type", "image/png"),
                        provider=ai_provider,
                        api_key=ai_api_key,
                        model=ai_model,
                    )
                    if parse_err:
                        st.error(parse_err)
                    else:
                        company_list_df = load_company_list()
                        match_rows = _build_value_chain_match_rows(parsed, company_list_df)
                        parsed["matched_rows"] = match_rows
                        parsed["created_at"] = datetime.now().isoformat(timespec="seconds")
                        parsed["overview_text"] = ""
                        parsed["overview_source"] = ""
                        parsed["_image_hash"] = img_hash
                        st.session_state["value_chain_result"] = parsed
                        st.session_state["value_chain_uploader_nonce"] += 1
                        total_nodes = len(match_rows)
                        matched_cnt = sum(1 for r in match_rows if bool(r.get("matched")))
                        st.session_state["value_chain_notice"] = (
                            f"밸류체인 분석 완료: 기업 {total_nodes}개 중 관심기업 매칭 {matched_cnt}개"
                        )
                        st.rerun()

    if st.session_state.get("value_chain_notice"):
        st.success(st.session_state.pop("value_chain_notice"))
    if st.session_state.get("value_chain_warning"):
        st.warning(st.session_state.pop("value_chain_warning"))
    if st.session_state.get("value_chain_summary_notice"):
        st.success(st.session_state.pop("value_chain_summary_notice"))
    if st.session_state.get("value_chain_summary_warning"):
        st.warning(st.session_state.pop("value_chain_summary_warning"))

    result = st.session_state.get("value_chain_result") if isinstance(st.session_state.get("value_chain_result"), dict) else {}
    if not result or not result.get("rows"):
        st.info("아직 밸류체인 결과가 없습니다. 이미지를 업로드해 분석해 주세요.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    chain_name = str(result.get("chain_name") or "밸류체인").strip()
    match_rows = result.get("matched_rows") if isinstance(result.get("matched_rows"), list) else []
    if not match_rows:
        company_list_df = load_company_list()
        match_rows = _build_value_chain_match_rows(result, company_list_df)
        result["matched_rows"] = match_rows
        st.session_state["value_chain_result"] = result

    df = pd.DataFrame(match_rows)
    if df.empty:
        st.warning("밸류체인 결과를 표시할 행이 없습니다.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    if bool(show_only_matched):
        view_df = df[df["matched"] == True].copy()  # noqa: E712
    else:
        view_df = df.copy()

    total_count = len(df)
    matched_count = int(df["matched"].sum()) if "matched" in df.columns else 0
    unmatched_count = max(0, total_count - matched_count)
    stage_count = len(df["stage"].dropna().unique())
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("체인명", chain_name)
    m2.metric("체인 기업수", f"{total_count:,}")
    m3.metric("관심기업 매칭", f"{matched_count:,}")
    m4.metric("미매칭", f"{unmatched_count:,}")
    st.caption(f"단계 수: {stage_count} / 생성시각: {str(result.get('created_at') or '-')}")

    summary_col1, summary_col2 = st.columns([1, 1.8])
    with summary_col1:
        make_summary_btn = st.button("밸류체인 총 설명 생성 (AI+검색)", key="value_chain_make_summary_btn")
    with summary_col2:
        overview_source = str(result.get("overview_source") or "").strip()
        if overview_source:
            st.caption(f"설명 생성 검색소스: {overview_source}")

    if make_summary_btn:
        ai_provider, ai_api_key, ai_model = get_ai_settings_from_session("analysis")
        with st.spinner("구글 검색 기반으로 밸류체인 총 설명을 생성 중입니다..."):
            summary_text, summary_err, summary_source = generate_value_chain_overview_with_ai(
                chain_name=chain_name,
                match_rows=df.to_dict("records"),
                provider=ai_provider,
                api_key=ai_api_key,
                model=ai_model,
            )
            if summary_err:
                st.session_state["value_chain_summary_warning"] = summary_err
            else:
                result["overview_text"] = summary_text
                result["overview_source"] = summary_source
                st.session_state["value_chain_result"] = result
                st.session_state["value_chain_summary_notice"] = "밸류체인 총 설명 생성을 완료했습니다."
        st.rerun()

    overview_text = str(result.get("overview_text") or "").strip()
    if overview_text:
        st.markdown("#### 밸류체인 총 설명")
        st.markdown(overview_text)

    stage_order = ["업스트림", "미드스트림", "다운스트림"]
    view_df["stage_order"] = view_df["stage"].apply(lambda x: stage_order.index(x) if x in stage_order else 99)
    view_df = view_df.sort_values(["stage_order", "segment", "input_company"]).copy()

    company_list_df = load_company_list()
    analysis_all_df = load_company_analysis_history()
    latest_analysis_map: dict[str, dict] = {}
    if isinstance(analysis_all_df, pd.DataFrame) and not analysis_all_df.empty:
        tmp = analysis_all_df.copy()
        tmp["analysis_date"] = pd.to_datetime(tmp["analysis_date"], errors="coerce")
        tmp = tmp.sort_values(["analysis_date", "updated_at"], ascending=[False, False])
        for _, row in tmp.iterrows():
            nm = str(row.get("stock_name") or "").strip()
            if not nm or nm in latest_analysis_map:
                continue
            latest_analysis_map[nm] = row.to_dict()

    popover_seq = 0
    for stage_name in list(dict.fromkeys(view_df["stage"].tolist())):
        stage_df = view_df[view_df["stage"] == stage_name].copy()
        if stage_df.empty:
            continue
        st.markdown(f"#### {stage_name}")
        segments = stage_df["segment"].dropna().astype(str).unique().tolist()
        cols = st.columns(max(1, min(4, len(segments))))
        for idx, seg in enumerate(segments):
            seg_df = stage_df[stage_df["segment"] == seg].copy()
            with cols[idx % len(cols)]:
                st.markdown(f"**{seg}**")
                for _, rr in seg_df.iterrows():
                    popover_seq += 1
                    in_name = str(rr.get("input_company") or "").strip()
                    m_name = str(rr.get("matched_company") or "").strip()
                    matched = bool(rr.get("matched"))
                    if matched:
                        head_text = f"매칭 | {in_name} -> {m_name}"
                    else:
                        head_text = f"미매칭 | {in_name}"
                    st.caption(head_text)
                    with st.popover(f"{in_name or '기업'} 상세보기 ({popover_seq})"):
                        _render_value_chain_company_detail(
                            input_company=in_name,
                            matched_company=m_name,
                            company_list_df=company_list_df,
                            latest_analysis_map=latest_analysis_map,
                        )

    sankey_fig = _build_value_chain_sankey_figure(view_df.to_dict("records"), chain_name)
    if sankey_fig is not None:
        st.plotly_chart(style_figure(sankey_fig), use_container_width=True)

    out_df = view_df.rename(
        columns={
            "stage": "단계",
            "segment": "세부공정",
            "input_company": "이미지기업명",
            "matched_company": "매칭기업명",
            "matched": "매칭여부",
            "match_score": "매칭점수",
            "ticker": "티커",
            "sector": "산업섹터",
        }
    )
    out_df["매칭점수"] = pd.to_numeric(out_df["매칭점수"], errors="coerce").round(3)
    st.markdown("#### 매칭 상세")
    st.dataframe(
        out_df[["단계", "세부공정", "이미지기업명", "매칭기업명", "티커", "산업섹터", "매칭여부", "매칭점수"]],
        height=estimate_dataframe_height(out_df, min_height=240, max_height=1200),
        use_container_width=True,
        hide_index=True,
    )

    result_bytes = json.dumps(result, ensure_ascii=False, indent=2).encode("utf-8")
    st.download_button(
        "밸류체인 결과 JSON 다운로드",
        data=result_bytes,
        file_name=f"value_chain_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        mime="application/json",
        key="value_chain_download_json_btn",
    )

    # ── 밸류체인 저장 ────────────────────────────────────────────────
    st.divider()
    st.markdown("##### 이 밸류체인 저장")
    save_col1, save_col2 = st.columns([2, 1])
    with save_col1:
        default_save_name = str(result.get("chain_name") or "밸류체인").strip()
        save_input_name = st.text_input(
            "저장 이름",
            value=st.session_state.get("value_chain_save_name") or default_save_name,
            key="value_chain_save_name_input",
            placeholder="예: 반도체 밸류체인 2026-03",
        )
    with save_col2:
        st.write("")  # 레이블 높이 맞춤
        st.write("")
        if st.button("DB에 저장", key="value_chain_save_btn", type="primary"):
            if not save_input_name.strip():
                st.warning("저장 이름을 입력해 주세요.")
            else:
                matched_names = [
                    str(r.get("matched_company") or "")
                    for r in match_rows if bool(r.get("matched"))
                ]
                saved_id = save_value_chain(
                    chain_name=save_input_name.strip(),
                    chain_json={k: v for k, v in result.items() if k not in ("matched_rows", "_image_hash")},
                    matched_companies=matched_names,
                    image_hash=result.get("_image_hash"),
                )
                st.session_state["value_chain_save_name"] = ""
                st.session_state["value_chain_notice"] = f"밸류체인 '{save_input_name.strip()}'을 저장했습니다 (id={saved_id})."
                st.rerun()

    st.markdown("</div>", unsafe_allow_html=True)


def render_stock_tab(current_df: pd.DataFrame) -> None:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">종목별 추적</div>', unsafe_allow_html=True)

    stock_names = get_all_stock_names(current_df)
    if not stock_names:
        st.info("종목 데이터가 없습니다. 기록 입력 탭에서 먼저 저장해 주세요.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    selected_stock = st.selectbox("종목 선택", stock_names)
    stock_hist = load_stock_history(selected_stock)

    if stock_hist.empty:
        st.caption("해당 종목의 저장 기록이 아직 없습니다.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    stock_hist = stock_hist.copy()
    stock_hist["value_change"] = stock_hist["market_value_krw"].diff().fillna(0)
    stock_hist["pnl_change"] = stock_hist["pnl_value_krw"].diff().fillna(0)

    chart_col, info_col = st.columns([1.6, 1])
    with chart_col:
        stock_fig = px.line(
            stock_hist,
            x="snapshot_date",
            y=["market_value_krw", "pnl_value_krw"],
            markers=True,
            color_discrete_sequence=["#1d4ed8", "#d92d20"],
            title=f"{selected_stock} 금액 추이",
            labels={"snapshot_date": "날짜", "value": "금액(원)", "variable": "지표"},
        )
        stock_fig.for_each_trace(lambda t: t.update(name="평가금액" if t.name == "market_value_krw" else "손익금액"))
        add_line_labels(stock_fig, pct=False, last_only=False)
        st.plotly_chart(style_figure(stock_fig), use_container_width=True)

    with info_col:
        latest = stock_hist.iloc[-1]
        st.metric("최근 평가금액", format_won(float(latest["market_value_krw"])))
        st.metric("최근 손익금액", format_signed_won(float(latest["pnl_value_krw"])))
        st.metric("최근 수익률", format_signed_pct(float(latest["pnl_pct"])))
        st.metric("최근 보유수량", f"{float(latest['quantity']):,.0f}주")

    row2_col1, row2_col2 = st.columns([1, 1])
    with row2_col1:
        qty_fig = px.area(
            stock_hist,
            x="snapshot_date",
            y="quantity",
            title="보유수량 추이",
            labels={"snapshot_date": "날짜", "quantity": "보유수량(주)"},
            color_discrete_sequence=["#0f766e"],
        )
        add_line_labels(qty_fig, pct=False, last_only=False)
        st.plotly_chart(style_figure(qty_fig), use_container_width=True)

    with row2_col2:
        pct_fig = px.line(
            stock_hist,
            x="snapshot_date",
            y="pnl_pct",
            markers=True,
            title="수익률(%) 추이",
            labels={"snapshot_date": "날짜", "pnl_pct": "수익률(%)"},
            color_discrete_sequence=["#1d4ed8"],
        )
        pct_fig.add_hline(y=0, line_dash="dot", line_color="#334155")
        pct_fig.update_yaxes(tickformat=",.0f", ticksuffix="%")
        add_line_labels(pct_fig, pct=True, last_only=False)
        st.plotly_chart(style_figure(pct_fig), use_container_width=True)

    row3_col1, row3_col2 = st.columns([1, 1])
    with row3_col1:
        value_change_fig = px.bar(
            stock_hist,
            x="snapshot_date",
            y="value_change",
            color="value_change",
            title="평가금액 증감",
            labels={"snapshot_date": "날짜", "value_change": "증감(원)"},
            color_continuous_scale=[(0.0, "#1570ef"), (0.5, "#94a3b8"), (1.0, "#d92d20")],
        )
        value_change_fig.update_coloraxes(showscale=False)
        add_bar_labels(value_change_fig, pct=False)
        st.plotly_chart(style_figure(value_change_fig), use_container_width=True)

    with row3_col2:
        pnl_change_fig = px.bar(
            stock_hist,
            x="snapshot_date",
            y="pnl_change",
            color="pnl_change",
            title="손익금액 증감",
            labels={"snapshot_date": "날짜", "pnl_change": "증감(원)"},
            color_continuous_scale=[(0.0, "#1570ef"), (0.5, "#94a3b8"), (1.0, "#d92d20")],
        )
        pnl_change_fig.update_coloraxes(showscale=False)
        add_bar_labels(pnl_change_fig, pct=False)
        st.plotly_chart(style_figure(pnl_change_fig), use_container_width=True)

    table_view = stock_hist.rename(
        columns={
            "snapshot_date": "날짜",
            "stock_name": "종목명",
            "quantity": "보유수량",
            "currency": "통화",
            "fx_effective": "환율(원화기준)",
            "market_value_krw": "평가금액",
            "pnl_value_krw": "손익금액",
            "pnl_pct": "수익률(%)",
        }
    ).sort_values("날짜", ascending=False)
    table_view["날짜"] = table_view["날짜"].dt.date
    st.dataframe(
        format_table_numbers(table_view, percent_cols={"수익률(%)"}),
        use_container_width=True,
        hide_index=True,
    )
    st.markdown("</div>", unsafe_allow_html=True)


def render_fx_tab() -> None:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">주요 환율 대시보드</div>', unsafe_allow_html=True)

    period_map = {
        "1개월": 45,
        "3개월": 120,
        "6개월": 220,
        "1년": 400,
        "3년": 1200,
    }
    selected_period = st.selectbox("조회 기간", list(period_map.keys()), index=2)
    days = period_map[selected_period]
    end_date = date.today() + timedelta(days=1)
    start_date = date.today() - timedelta(days=days)

    series_map: dict[str, pd.DataFrame] = {}
    for row in FX_TRACKERS:
        series = fetch_fx_series(row["ticker"], start_date.isoformat(), end_date.isoformat())
        if not series.empty:
            series_map[row["pair"]] = series

    if not series_map:
        st.warning("환율 데이터를 불러오지 못했습니다. 네트워크 또는 데이터 소스를 확인해 주세요.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    summary_df = get_fx_tracker_summary(series_map)
    if summary_df.empty:
        st.warning("환율 요약 데이터를 만들 수 없습니다.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    summary_df = summary_df.sort_values("통화쌍")
    latest_update_date = pd.to_datetime(summary_df["기준일"], errors="coerce").max()
    if pd.notna(latest_update_date):
        st.markdown(
            (
                '<div class="fx-reference-banner">'
                '<span class="fx-label">환율 마지막 업데이트</span>'
                f'<span class="fx-date">{latest_update_date.date().isoformat()}</span>'
                "</div>"
            ),
            unsafe_allow_html=True,
        )

    major_pairs = ["USD/KRW", "EUR/KRW", "JPY/KRW", "CNY/KRW"]
    card_cols = st.columns(4)
    for idx, pair in enumerate(major_pairs):
        row_df = summary_df[summary_df["통화쌍"] == pair]
        with card_cols[idx]:
            if row_df.empty:
                render_summary_card(pair, "-", "데이터 없음", "neutral")
            else:
                r = row_df.iloc[0]
                delta = _label_text(r["1일변동(%)"], pct=True)
                cls = value_class(r["1일변동(%)"] if pd.notna(r["1일변동(%)"]) else 0)
                render_summary_card(pair, _label_text(r["현재환율"]), f"1일 변동 {delta}", cls)

    available_pairs = sorted(series_map.keys())
    focus_index = available_pairs.index("USD/KRW") if "USD/KRW" in available_pairs else 0
    focus_pair = st.selectbox("상세 통화쌍", available_pairs, index=focus_index)
    focus_meta = summary_df[summary_df["통화쌍"] == focus_pair]
    if not focus_meta.empty:
        focus_date = pd.to_datetime(focus_meta.iloc[0]["기준일"], errors="coerce")
        if pd.notna(focus_date):
            st.markdown(
                (
                    '<div class="fx-reference-banner">'
                    f'<span class="fx-label">{focus_pair} 기준일</span>'
                    f'<span class="fx-date">{focus_date.date().isoformat()}</span>'
                    "</div>"
                ),
                unsafe_allow_html=True,
            )
    focus_df = series_map[focus_pair].copy().sort_values("date")
    focus_fig = px.line(
        focus_df,
        x="date",
        y="rate",
        markers=True,
        title=f"{focus_pair} 환율 추이",
        labels={"date": "날짜", "rate": "환율"},
        color_discrete_sequence=["#1d4ed8"],
    )
    add_line_labels(focus_fig, pct=False, last_only=True)
    st.plotly_chart(style_figure(focus_fig), use_container_width=True)

    preferred_default = ["USD/KRW", "EUR/KRW", "JPY/KRW", "CNY/KRW"]
    default_pairs = [p for p in preferred_default if p in available_pairs]
    if len(default_pairs) < min(4, len(available_pairs)):
        default_pairs += [p for p in available_pairs if p not in default_pairs]
    default_pairs = default_pairs[: min(4, len(available_pairs))]
    selected_pairs = st.multiselect(
        "비교 통화쌍",
        available_pairs,
        default=default_pairs,
    )
    if not selected_pairs:
        selected_pairs = available_pairs[:1]

    compare_rows = []
    for pair in selected_pairs:
        df = series_map[pair].copy().sort_values("date")
        if df.empty:
            continue
        base = float(df.iloc[0]["rate"])
        if base == 0:
            continue
        df["pair"] = pair
        df["index100"] = df["rate"] / base * 100.0
        compare_rows.append(df[["date", "pair", "index100"]])

    if compare_rows:
        compare_df = pd.concat(compare_rows, ignore_index=True)
        compare_fig = px.line(
            compare_df,
            x="date",
            y="index100",
            color="pair",
            markers=True,
            title="환율 지수 비교 (첫날=100)",
            labels={"date": "날짜", "index100": "지수", "pair": "통화쌍"},
        )
        add_line_labels(compare_fig, pct=False, last_only=True)
        st.plotly_chart(style_figure(compare_fig), use_container_width=True)

    change_df = summary_df.melt(
        id_vars=["통화쌍"],
        value_vars=["1일변동(%)", "1주변동(%)", "1개월변동(%)"],
        var_name="구간",
        value_name="변동률",
    ).dropna(subset=["변동률"])
    if not change_df.empty:
        change_fig = px.bar(
            change_df,
            x="통화쌍",
            y="변동률",
            color="구간",
            barmode="group",
            title="통화쌍별 변동률 비교",
            labels={"통화쌍": "통화쌍", "변동률": "변동률(%)", "구간": "비교구간"},
        )
        add_bar_labels(change_fig, pct=True)
        change_fig.update_yaxes(tickformat=",.0f", ticksuffix="%")
        st.plotly_chart(style_figure(change_fig), use_container_width=True)

    st.caption("주요 환율 요약")
    st.dataframe(
        format_table_numbers(summary_df, percent_cols={"1일변동(%)", "1주변동(%)", "1개월변동(%)"}),
        use_container_width=True,
        hide_index=True,
    )
    st.caption("데이터 소스: yfinance (시장 휴일에는 직전 거래일 값이 반영될 수 있습니다.)")
    st.markdown("</div>", unsafe_allow_html=True)


def render_company_score_tab(current_df: pd.DataFrame) -> None:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">기업 장기투자 종합 점수</div>', unsafe_allow_html=True)

    suggestion_names = get_all_stock_names(current_df)
    market_pref_map = build_market_preference_map(current_df)
    if "score_date" not in st.session_state:
        st.session_state["score_date"] = date.today()
    if "score_stock_name" not in st.session_state:
        st.session_state["score_stock_name"] = suggestion_names[0] if suggestion_names else ""
    if "score_ticker" not in st.session_state:
        st.session_state["score_ticker"] = ""
    if "score_stock_name_input" not in st.session_state:
        st.session_state["score_stock_name_input"] = st.session_state.get("score_stock_name", "")
    if "score_ticker_input" not in st.session_state:
        st.session_state["score_ticker_input"] = st.session_state.get("score_ticker", "")
    if "score_note" not in st.session_state:
        st.session_state["score_note"] = ""
    if "score_source" not in st.session_state:
        st.session_state["score_source"] = "manual"
    if "score_ticker_source" not in st.session_state:
        st.session_state["score_ticker_source"] = ""
    if "score_use_ai" not in st.session_state:
        st.session_state["score_use_ai"] = True
    legacy_score_model = str(st.session_state.get("score_ai_model", DEFAULT_OPENAI_MODEL) or "")
    if "score_ai_provider" not in st.session_state:
        st.session_state["score_ai_provider"] = normalize_ai_provider(
            st.session_state.get("global_ai_provider", "claude")
        )
    if "score_openai_api_key" not in st.session_state:
        st.session_state["score_openai_api_key"] = st.session_state.get(
            "score_ai_api_key",
            st.session_state.get("analysis_openai_api_key", ""),
        )
    if "score_claude_api_key" not in st.session_state:
        st.session_state["score_claude_api_key"] = st.session_state.get("analysis_claude_api_key", "")
    if "score_openai_model" not in st.session_state:
        st.session_state["score_openai_model"] = (
            legacy_score_model if legacy_score_model and "claude" not in legacy_score_model.lower() else DEFAULT_OPENAI_MODEL
        )
    if "score_claude_model" not in st.session_state:
        st.session_state["score_claude_model"] = (
            legacy_score_model if "claude" in legacy_score_model.lower() else DEFAULT_CLAUDE_MODEL
        )
    if "score_prev_stock_name" not in st.session_state:
        st.session_state["score_prev_stock_name"] = ""

    if "score_stock_name_pending" in st.session_state:
        next_name = (st.session_state.pop("score_stock_name_pending") or "").strip()
        st.session_state["score_stock_name_input"] = next_name
    if "score_ticker_pending" in st.session_state:
        next_ticker = clean_valid_ticker(st.session_state.pop("score_ticker_pending") or "")
        st.session_state["score_ticker_input"] = next_ticker
    if "score_autofill_notice" in st.session_state:
        st.success(st.session_state.pop("score_autofill_notice"))

    for key in SCORE_METRIC_CONFIG:
        ss_key = f"score_{key}"
        if ss_key not in st.session_state:
            st.session_state[ss_key] = 0.0

    for w_key, w_value in DEFAULT_SCORE_WEIGHTS.items():
        ss_key = f"weight_{w_key}"
        if ss_key not in st.session_state:
            st.session_state[ss_key] = int(w_value)

    with st.expander("AI 티커 추론 설정 (선택)", expanded=False):
        st.checkbox("yfinance 검색 실패 시 AI로 티커 추론", key="score_use_ai")
        st.caption("AI 키/모델/제공자는 API 설정 탭의 공통값을 사용합니다.")

    top_col1, top_col2, top_col3, top_col4 = st.columns([1, 1.2, 1.1, 1.2])
    with top_col1:
        st.date_input("평가일", key="score_date")
    with top_col2:
        st.text_input("기업명", key="score_stock_name_input")
    with top_col3:
        st.text_input("티커", key="score_ticker_input", placeholder="005930.KS / AAPL")
    with top_col4:
        selected_hint = st.selectbox("보유 종목 불러오기", ["직접입력"] + suggestion_names, index=0, key="score_name_hint")
        if selected_hint != "직접입력":
            current_input_name = (st.session_state.get("score_stock_name_input") or "").strip()
            if selected_hint != current_input_name:
                st.session_state["score_stock_name_pending"] = selected_hint
                st.rerun()

    score_name_value = (st.session_state.get("score_stock_name_input") or "").strip()
    score_ticker_value = clean_valid_ticker(st.session_state.get("score_ticker_input") or "")

    score_ai_provider, score_ai_api_key, score_ai_model = get_ai_settings_from_session("score")

    current_name = score_name_value
    prev_name = st.session_state.get("score_prev_stock_name", "")
    if current_name and (current_name != prev_name or not score_ticker_value):
        auto_ticker, auto_src = resolve_ticker_auto_with_retry(
            current_name,
            use_ai=bool(st.session_state.get("score_use_ai", False)),
            api_key=score_ai_api_key,
            model=score_ai_model,
            provider=score_ai_provider,
            market_preference=market_pref_map.get(current_name, ""),
        )
        if auto_ticker and (not score_ticker_value or selected_hint != "직접입력"):
            st.session_state["score_ticker_pending"] = auto_ticker
            st.session_state["score_ticker_source"] = auto_src
            st.session_state["score_autofill_notice"] = f"티커 자동 입력 완료: {auto_ticker} ({auto_src})"
            st.session_state["score_prev_stock_name"] = current_name
            st.rerun()
        st.session_state["score_prev_stock_name"] = current_name

    auto_col1, auto_col2 = st.columns([1, 2])
    with auto_col1:
        auto_ticker_clicked = st.button("티커 자동 입력", key="score_auto_ticker_btn")
    with auto_col2:
        ticker_source = st.session_state.get("score_ticker_source", "")
        if ticker_source:
            st.caption(f"티커 소스: {ticker_source}")
        else:
            st.caption("기업명 기반으로 티커를 자동 추천합니다.")

    if auto_ticker_clicked:
        auto_ticker, auto_src = resolve_ticker_auto_with_retry(
            score_name_value,
            use_ai=bool(st.session_state.get("score_use_ai", False)),
            api_key=score_ai_api_key,
            model=score_ai_model,
            provider=score_ai_provider,
            market_preference=market_pref_map.get(score_name_value, ""),
        )
        if auto_ticker:
            st.session_state["score_ticker_pending"] = auto_ticker
            st.session_state["score_ticker_source"] = auto_src
            st.session_state["score_autofill_notice"] = f"티커 자동 입력 완료: {auto_ticker} ({auto_src})"
            st.rerun()
        else:
            st.warning(auto_src or "티커를 찾지 못했습니다.")

    fetch_col1, fetch_col2 = st.columns([1, 2])
    with fetch_col1:
        fetch_clicked = st.button("지표 자동 불러오기 (멀티소스)", key="score_fetch_metrics_btn")
    with fetch_col2:
        st.caption("티커 자동 입력 후 클릭하면 yfinance → Alpha Vantage → Finnhub 순으로 지표를 조회합니다.")

    if fetch_clicked:
        fetched, fetched_name, err_msg, metric_source = fetch_company_metrics_multi_source(score_ticker_value)
        if err_msg:
            st.warning(err_msg)
        else:
            for metric_key, metric_value in fetched.items():
                if metric_value is not None:
                    st.session_state[f"score_{metric_key}"] = float(metric_value)
            if fetched_name and not score_name_value:
                st.session_state["score_stock_name_pending"] = fetched_name
            st.session_state["score_source"] = metric_source or "multi_source"
            st.success(f"지표 자동 불러오기를 완료했습니다. (소스: {metric_source or 'multi_source'})")
            if fetched_name and not score_name_value:
                st.rerun()

    st.markdown("#### 지표 입력")
    metric_cols = st.columns(3)
    metric_order = list(SCORE_METRIC_CONFIG.keys())
    for idx, metric_key in enumerate(metric_order):
        with metric_cols[idx % 3]:
            cfg = SCORE_METRIC_CONFIG[metric_key]
            st.number_input(
                cfg["label"],
                key=f"score_{metric_key}",
                step=1.0,
                format="%.0f",
                min_value=-1000.0,
                max_value=1000.0,
            )

    st.markdown("#### 카테고리 가중치")
    w1, w2, w3, w4 = st.columns(4)
    with w1:
        st.slider("배당", min_value=0, max_value=100, key="weight_dividend")
    with w2:
        st.slider("성장", min_value=0, max_value=100, key="weight_growth")
    with w3:
        st.slider("안정성", min_value=0, max_value=100, key="weight_stability")
    with w4:
        st.slider("밸류", min_value=0, max_value=100, key="weight_valuation")

    metrics = {key: float(st.session_state[f"score_{key}"]) for key in SCORE_METRIC_CONFIG}
    weights = {
        "dividend": float(st.session_state["weight_dividend"]),
        "growth": float(st.session_state["weight_growth"]),
        "stability": float(st.session_state["weight_stability"]),
        "valuation": float(st.session_state["weight_valuation"]),
    }
    scores = compute_company_scores(metrics, weights)
    normalized_weights = scores["weights"]

    if abs(sum(weights.values()) - 100.0) > 0.01:
        st.info(
            "가중치 합계가 100이 아니어서 자동 정규화했습니다. "
            f"배당 {normalized_weights['dividend']:,.0f} / 성장 {normalized_weights['growth']:,.0f} / "
            f"안정성 {normalized_weights['stability']:,.0f} / 밸류 {normalized_weights['valuation']:,.0f}"
        )

    score_cards = st.columns(5)
    with score_cards[0]:
        render_summary_card("총점", f"{scores['total_score']:,.0f}", "장기투자 종합 점수", "neutral")
    with score_cards[1]:
        render_summary_card("배당 점수", f"{scores['dividend_score']:,.0f}", "배당수익률 기반", "neutral")
    with score_cards[2]:
        render_summary_card("성장 점수", f"{scores['growth_score']:,.0f}", "매출/EPS/ROE 기반", "neutral")
    with score_cards[3]:
        render_summary_card("안정성 점수", f"{scores['stability_score']:,.0f}", "부채/유동/마진 기반", "neutral")
    with score_cards[4]:
        render_summary_card("밸류 점수", f"{scores['valuation_score']:,.0f}", "PER/PBR 기반", "neutral")

    radar_df = pd.DataFrame(
        {
            "영역": ["배당", "성장", "안정성", "밸류"],
            "점수": [
                scores["dividend_score"],
                scores["growth_score"],
                scores["stability_score"],
                scores["valuation_score"],
            ],
        }
    )
    metric_score_df = pd.DataFrame(
        {
            "지표": [cfg["label"] for cfg in SCORE_METRIC_CONFIG.values()],
            "점수": [
                score_linear(metrics[key], cfg["min"], cfg["max"], cfg["reverse"])
                for key, cfg in SCORE_METRIC_CONFIG.items()
            ],
        }
    ).sort_values("점수", ascending=True)

    viz_col1, viz_col2 = st.columns([1, 1])
    with viz_col1:
        radar_fig = px.line_polar(
            radar_df,
            r="점수",
            theta="영역",
            line_close=True,
            range_r=[0, 100],
            title="카테고리 점수 레이더",
            color_discrete_sequence=["#1d4ed8"],
        )
        radar_fig.update_traces(
            fill="toself",
            text=[_label_text(v, pct=False) for v in radar_df["점수"]],
            mode="lines+markers+text",
            textposition="top center",
        )
        radar_fig.update_layout(
            margin=dict(l=20, r=20, t=56, b=20),
            paper_bgcolor="rgba(0,0,0,0)",
            font=dict(family="Noto Sans KR", color="#0f172a"),
        )
        st.plotly_chart(radar_fig, use_container_width=True)

    with viz_col2:
        metric_fig = px.bar(
            metric_score_df,
            x="점수",
            y="지표",
            orientation="h",
            color="점수",
            title="지표별 환산 점수(0~100)",
            color_continuous_scale="Viridis",
            range_x=[0, 100],
        )
        metric_fig.update_coloraxes(showscale=False)
        add_bar_labels(metric_fig, pct=False)
        st.plotly_chart(style_figure(metric_fig), use_container_width=True)

    gauge = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=scores["total_score"],
            number={"suffix": "점", "valueformat": ",.0f"},
            title={"text": "종합 점수 게이지"},
            gauge={
                "axis": {"range": [0, 100]},
                "bar": {"color": "#1d4ed8"},
                "steps": [
                    {"range": [0, 40], "color": "#fee2e2"},
                    {"range": [40, 70], "color": "#fef3c7"},
                    {"range": [70, 100], "color": "#dcfce7"},
                ],
            },
        )
    )
    gauge.update_layout(
        margin=dict(l=20, r=20, t=40, b=20),
        paper_bgcolor="rgba(0,0,0,0)",
        font=dict(family="Noto Sans KR", color="#0f172a"),
    )
    st.plotly_chart(gauge, use_container_width=True)

    st.text_area("분석 메모", key="score_note", height=80, placeholder="투자 아이디어, 리스크, 체크포인트...")
    if st.button("기업 점수 저장", type="primary", key="score_save_btn"):
        stock_name = score_name_value
        if not stock_name:
            st.error("기업명을 입력해 주세요.")
        else:
            save_company_score(
                score_date=st.session_state["score_date"],
                stock_name=stock_name,
                ticker=score_ticker_value,
                metrics=metrics,
                scores=scores,
                note=st.session_state["score_note"],
                source=st.session_state.get("score_source", "manual"),
            )
            upsert_company_list_entry(
                stock_name,
                score_ticker_value,
                source="score",
            )
            st.success(f"{stock_name} 점수를 저장했습니다.")
            st.session_state["score_source"] = "manual"

    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">기업 점수 히스토리</div>', unsafe_allow_html=True)
    hist_df = load_company_score_history()

    if hist_df.empty:
        st.info("저장된 기업 점수 이력이 없습니다.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    companies = sorted(hist_df["stock_name"].dropna().unique().tolist())
    filter_col1, filter_col2 = st.columns([1.5, 1])
    with filter_col1:
        selected_companies = st.multiselect("기업 필터", companies, default=companies[: min(3, len(companies))])
    with filter_col2:
        focus_company = st.selectbox("상세 추적 기업", companies)

    filtered_hist = hist_df.copy()
    if selected_companies:
        filtered_hist = filtered_hist[filtered_hist["stock_name"].isin(selected_companies)]

    line_fig = px.line(
        filtered_hist,
        x="score_date",
        y="total_score",
        color="stock_name",
        markers=True,
        title="기업별 종합 점수 추이",
        labels={"score_date": "날짜", "total_score": "종합 점수", "stock_name": "기업"},
    )
    add_line_labels(line_fig, pct=False, last_only=False)
    st.plotly_chart(style_figure(line_fig), use_container_width=True)

    latest_rank = (
        filtered_hist.sort_values("score_date")
        .groupby("stock_name", as_index=False)
        .tail(1)
        .sort_values("total_score", ascending=False)
    )
    rank_fig = px.bar(
        latest_rank,
        x="stock_name",
        y="total_score",
        color="total_score",
        title="최신 시점 기업 점수 랭킹",
        labels={"stock_name": "기업", "total_score": "종합 점수"},
        color_continuous_scale="Tealgrn",
    )
    rank_fig.update_coloraxes(showscale=False)
    add_bar_labels(rank_fig, pct=False)
    st.plotly_chart(style_figure(rank_fig), use_container_width=True)

    focus_df = hist_df[hist_df["stock_name"] == focus_company].sort_values("score_date")
    if not focus_df.empty:
        cat_df = focus_df.melt(
            id_vars=["score_date"],
            value_vars=["dividend_score", "growth_score", "stability_score", "valuation_score", "total_score"],
            var_name="score_type",
            value_name="score",
        )
        cat_df["score_type"] = cat_df["score_type"].map(
            {
                "dividend_score": "배당",
                "growth_score": "성장",
                "stability_score": "안정성",
                "valuation_score": "밸류",
                "total_score": "종합",
            }
        )
        cat_fig = px.line(
            cat_df,
            x="score_date",
            y="score",
            color="score_type",
            markers=True,
            title=f"{focus_company} 카테고리 점수 추이",
            labels={"score_date": "날짜", "score": "점수", "score_type": "구분"},
        )
        add_line_labels(cat_fig, pct=False, last_only=False)
        st.plotly_chart(style_figure(cat_fig), use_container_width=True)

        metric_trend = focus_df[[
            "score_date",
            "dividend_yield",
            "revenue_growth",
            "eps_growth",
            "roe",
            "operating_margin",
            "debt_ratio",
            "current_ratio",
            "per",
            "pbr",
        ]].copy()
        metric_trend = metric_trend.melt(
            id_vars=["score_date"], var_name="metric", value_name="value"
        )
        metric_trend = metric_trend.dropna(subset=["value"])
        metric_trend["metric"] = metric_trend["metric"].map(
            {
                "dividend_yield": "배당수익률",
                "revenue_growth": "매출성장률",
                "eps_growth": "EPS성장률",
                "roe": "ROE",
                "operating_margin": "영업이익률",
                "debt_ratio": "부채비율",
                "current_ratio": "유동비율",
                "per": "PER",
                "pbr": "PBR",
            }
        )
        if not metric_trend.empty:
            metric_fig2 = px.line(
                metric_trend,
                x="score_date",
                y="value",
                color="metric",
                markers=True,
                title=f"{focus_company} 핵심 지표 변화",
                labels={"score_date": "날짜", "value": "지표값", "metric": "지표"},
            )
            add_line_labels(metric_fig2, pct=False, last_only=False)
            st.plotly_chart(style_figure(metric_fig2), use_container_width=True)

    hist_table = filtered_hist.sort_values(["score_date", "stock_name"], ascending=[False, True]).copy()
    hist_table["score_date"] = hist_table["score_date"].dt.date
    hist_table = hist_table.rename(
        columns={
            "score_date": "평가일",
            "stock_name": "기업명",
            "ticker": "티커",
            "total_score": "종합점수",
            "dividend_score": "배당점수",
            "growth_score": "성장점수",
            "stability_score": "안정성점수",
            "valuation_score": "밸류점수",
            "source": "입력소스",
        }
    )
    st.dataframe(
        format_table_numbers(
            hist_table[
                [
                    "평가일",
                    "기업명",
                    "티커",
                    "종합점수",
                    "배당점수",
                    "성장점수",
                    "안정성점수",
                    "밸류점수",
                    "입력소스",
                    "note",
                ]
            ]
        ),
        use_container_width=True,
        hide_index=True,
    )
    st.markdown("</div>", unsafe_allow_html=True)


def render_company_compare_tab(current_df: pd.DataFrame) -> None:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">기업분석</div>', unsafe_allow_html=True)

    all_companies = _coerce_string_list(get_all_stock_names(current_df))
    holding_companies = _coerce_string_list(get_holding_stock_names(current_df))
    company_list_df = load_company_list()
    if not all_companies:
        st.info("비교할 기업이 없습니다. 기업정보 탭에서 기업을 추가하거나 기록을 저장해 주세요.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    sector_map = {name: "" for name in all_companies}
    if not company_list_df.empty:
        for _, row in company_list_df.iterrows():
            nm = str(row.get("stock_name") or "").strip()
            if not nm:
                continue
            sector_map[nm] = str(row.get("sector") or "").strip()
    sector_values = sorted({(sector_map.get(name, "").strip() or "미분류") for name in all_companies})
    sector_options = ["전체"] + sector_values

    default_metric_keys = [m for m in ["per", "pbr", "roe", "dividend_yield"] if m in SCORE_METRIC_CONFIG]
    default_companies = (
        holding_companies[: min(12, len(holding_companies))] if holding_companies else all_companies[: min(12, len(all_companies))]
    )

    if "compare_companies" not in st.session_state:
        st.session_state["compare_companies"] = default_companies
    else:
        st.session_state["compare_companies"] = _coerce_string_list(st.session_state.get("compare_companies"))
    if "compare_metrics" not in st.session_state:
        st.session_state["compare_metrics"] = default_metric_keys
    else:
        st.session_state["compare_metrics"] = _coerce_string_list(st.session_state.get("compare_metrics"))
    if "compare_custom_weights" not in st.session_state:
        st.session_state["compare_custom_weights"] = False
    if "compare_use_ai_ticker" not in st.session_state:
        st.session_state["compare_use_ai_ticker"] = False
    if "compare_ai_provider" not in st.session_state:
        st.session_state["compare_ai_provider"] = normalize_ai_provider(
            st.session_state.get("global_ai_provider", "claude")
        )
    if "compare_openai_api_key" not in st.session_state:
        st.session_state["compare_openai_api_key"] = st.session_state.get(
            "score_openai_api_key",
            st.session_state.get("analysis_openai_api_key", ""),
        )
    if "compare_claude_api_key" not in st.session_state:
        st.session_state["compare_claude_api_key"] = st.session_state.get(
            "score_claude_api_key",
            st.session_state.get("analysis_claude_api_key", ""),
        )
    if "compare_openai_model" not in st.session_state:
        st.session_state["compare_openai_model"] = st.session_state.get("score_openai_model", DEFAULT_OPENAI_MODEL)
    if "compare_claude_model" not in st.session_state:
        st.session_state["compare_claude_model"] = st.session_state.get("score_claude_model", DEFAULT_CLAUDE_MODEL)
    if "compare_sector_filter" not in st.session_state:
        st.session_state["compare_sector_filter"] = "전체"
    if st.session_state.get("compare_sector_filter") not in sector_options:
        st.session_state["compare_sector_filter"] = "전체"
    if "compare_set_name" not in st.session_state:
        st.session_state["compare_set_name"] = ""
    if "compare_set_note" not in st.session_state:
        st.session_state["compare_set_note"] = ""
    if "compare_selected_set_name" not in st.session_state:
        st.session_state["compare_selected_set_name"] = "선택안함"

    if "compare_pending_companies" in st.session_state:
        pending_companies = _coerce_string_list(st.session_state.pop("compare_pending_companies"))
        st.session_state["compare_companies"] = [c for c in pending_companies if c in all_companies]
    if "compare_pending_metrics" in st.session_state:
        pending_metrics = _coerce_string_list(st.session_state.pop("compare_pending_metrics"))
        st.session_state["compare_metrics"] = [m for m in pending_metrics if m in SCORE_METRIC_CONFIG]
    if "compare_pending_custom_weights" in st.session_state:
        st.session_state["compare_custom_weights"] = bool(st.session_state.pop("compare_pending_custom_weights"))
    if "compare_pending_weights" in st.session_state:
        pending_weights = st.session_state.pop("compare_pending_weights") or {}
        for metric_key, weight_value in pending_weights.items():
            if metric_key in SCORE_METRIC_CONFIG:
                try:
                    st.session_state[f"compare_weight_{metric_key}"] = float(weight_value)
                except Exception:
                    st.session_state[f"compare_weight_{metric_key}"] = 0.0
    if "compare_pending_sector_filter" in st.session_state:
        pending_sector = str(st.session_state.pop("compare_pending_sector_filter") or "전체")
        st.session_state["compare_sector_filter"] = pending_sector if pending_sector in sector_options else "전체"
    if "compare_set_name_pending" in st.session_state:
        st.session_state["compare_set_name"] = str(st.session_state.pop("compare_set_name_pending") or "")
    if "compare_set_note_pending" in st.session_state:
        st.session_state["compare_set_note"] = str(st.session_state.pop("compare_set_note_pending") or "")
    if "compare_selected_set_pending" in st.session_state:
        st.session_state["compare_selected_set_name"] = str(st.session_state.pop("compare_selected_set_pending") or "선택안함")
    if "compare_set_notice" in st.session_state:
        st.success(st.session_state.pop("compare_set_notice"))

    # 위젯 상태가 배열/객체로 오염되면 레이아웃 겹침이 발생할 수 있어 렌더 전에 정규화한다.
    st.session_state["compare_ai_provider"] = _coerce_choice(
        st.session_state.get("compare_ai_provider"),
        {"openai", "claude"},
        "claude",
    )
    st.session_state["compare_custom_weights"] = _to_bool_flag(st.session_state.get("compare_custom_weights", False))
    st.session_state["compare_use_ai_ticker"] = _to_bool_flag(st.session_state.get("compare_use_ai_ticker", False))
    for key in [
        "compare_openai_api_key",
        "compare_claude_api_key",
        "compare_openai_model",
        "compare_claude_model",
        "compare_set_name",
        "compare_set_note",
        "compare_selected_set_name",
    ]:
        st.session_state[key] = _sanitize_widget_text(st.session_state.get(key), "")
    st.session_state["compare_sector_filter"] = _sanitize_widget_text(
        st.session_state.get("compare_sector_filter"),
        "전체",
    )
    if st.session_state.get("compare_sector_filter") not in sector_options:
        st.session_state["compare_sector_filter"] = "전체"

    compare_sets_df = load_company_compare_sets()
    saved_set_names = compare_sets_df["set_name"].dropna().astype(str).tolist() if not compare_sets_df.empty else []
    saved_set_options = ["선택안함"] + saved_set_names
    if st.session_state.get("compare_selected_set_name", "선택안함") not in saved_set_options:
        st.session_state["compare_selected_set_name"] = "선택안함"

    def parse_json_list(value) -> list:
        if isinstance(value, list):
            return value
        if not value:
            return []
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, list) else []
        except Exception:
            return []

    def parse_json_dict(value) -> dict:
        if isinstance(value, dict):
            return value
        if not value:
            return {}
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except Exception:
            return {}

    st.markdown("#### 비교 세트 관리")
    set_col1, set_col2 = st.columns([1.2, 1.2])
    with set_col1:
        selected_set_name = st.selectbox("저장된 세트", saved_set_options, key="compare_selected_set_name")
    with set_col2:
        st.text_input("세트 이름", key="compare_set_name", placeholder="예: 철강/조선 고배당")
    st.text_input("세트 메모(선택)", key="compare_set_note", placeholder="비교 목적/조건 메모")

    set_btn_col1, set_btn_col2, set_btn_col3 = st.columns([1, 1.2, 1])
    with set_btn_col1:
        load_set_btn = st.button("세트 불러오기", key="compare_load_set_btn")
    with set_btn_col2:
        save_set_btn = st.button("현재 조건 세트 저장", key="compare_save_set_btn")
    with set_btn_col3:
        delete_set_btn = st.button("세트 삭제", key="compare_delete_set_btn")

    st.markdown("#### 비교 조건")
    st.selectbox("산업섹터 필터", options=sector_options, key="compare_sector_filter")
    active_sector_filter = st.session_state.get("compare_sector_filter", "전체")
    if active_sector_filter == "전체":
        filtered_companies = list(all_companies)
    else:
        filtered_companies = [
            name for name in all_companies if (sector_map.get(name, "").strip() or "미분류") == active_sector_filter
        ]

    if not filtered_companies:
        st.info("선택한 산업섹터에 해당하는 기업이 없습니다. 기업정보 탭에서 섹터를 입력해 주세요.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    current_selected_companies = st.session_state.get("compare_companies", [])
    pruned_companies = [c for c in current_selected_companies if c in filtered_companies]
    if pruned_companies != current_selected_companies:
        st.session_state["compare_companies"] = pruned_companies

    selected_companies = st.multiselect("비교 기업 선택", options=filtered_companies, key="compare_companies")
    selected_metrics = st.multiselect(
        "비교 지표 선택",
        options=list(SCORE_METRIC_CONFIG.keys()),
        format_func=lambda k: SCORE_METRIC_CONFIG[k]["label"],
        key="compare_metrics",
    )

    if load_set_btn:
        if selected_set_name == "선택안함":
            st.warning("불러올 세트를 선택해 주세요.")
        else:
            picked = compare_sets_df[compare_sets_df["set_name"] == selected_set_name]
            if picked.empty:
                st.warning("선택한 세트를 찾지 못했습니다.")
            else:
                row = picked.iloc[0]
                loaded_companies = [str(v).strip() for v in parse_json_list(row.get("companies_json")) if str(v).strip()]
                loaded_metrics = [str(v).strip() for v in parse_json_list(row.get("metrics_json")) if str(v).strip()]
                loaded_weights = parse_json_dict(row.get("weights_json"))
                st.session_state["compare_pending_companies"] = loaded_companies
                st.session_state["compare_pending_metrics"] = loaded_metrics
                st.session_state["compare_pending_custom_weights"] = bool(loaded_weights)
                st.session_state["compare_pending_weights"] = loaded_weights
                st.session_state["compare_pending_sector_filter"] = str(row.get("sector_filter") or "전체")
                st.session_state["compare_set_name_pending"] = selected_set_name
                st.session_state["compare_set_note_pending"] = str(row.get("note") or "")
                st.session_state["compare_set_notice"] = f"세트 `{selected_set_name}`을 불러왔습니다."
                st.rerun()

    if not selected_metrics:
        st.warning("최소 1개 이상의 지표를 선택해 주세요.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    st.checkbox("지표 가중치 직접 설정", key="compare_custom_weights")
    metric_weights = {}
    if st.session_state.get("compare_custom_weights", False):
        cols = st.columns(min(3, max(1, len(selected_metrics))))
        default_weight = round(100 / max(1, len(selected_metrics)))
        for idx, metric_key in enumerate(selected_metrics):
            label = SCORE_METRIC_CONFIG[metric_key]["label"]
            with cols[idx % len(cols)]:
                metric_weights[metric_key] = float(
                    st.number_input(
                        f"{label} 가중치",
                        min_value=0.0,
                        max_value=100.0,
                        step=1.0,
                        value=float(st.session_state.get(f"compare_weight_{metric_key}", default_weight)),
                        key=f"compare_weight_{metric_key}",
                        format="%.0f",
                    )
                )
    else:
        metric_weights = {m: 1.0 for m in selected_metrics}
        st.caption("선택 지표를 동일 가중치로 계산합니다.")

    if save_set_btn:
        target_set_name = (st.session_state.get("compare_set_name") or "").strip()
        if not target_set_name and selected_set_name != "선택안함":
            target_set_name = selected_set_name
        if not target_set_name:
            st.warning("저장할 세트 이름을 입력해 주세요.")
        elif not selected_companies:
            st.warning("세트에 저장할 기업을 1개 이상 선택해 주세요.")
        else:
            save_company_compare_set(
                set_name=target_set_name,
                companies=selected_companies,
                metrics=selected_metrics,
                weights=metric_weights,
                sector_filter=st.session_state.get("compare_sector_filter", "전체"),
                note=st.session_state.get("compare_set_note", ""),
            )
            st.session_state["compare_set_notice"] = f"세트 `{target_set_name}`을 저장했습니다."
            st.session_state["compare_selected_set_pending"] = target_set_name
            st.session_state["compare_set_name_pending"] = target_set_name
            st.rerun()

    if delete_set_btn:
        target_set_name = selected_set_name if selected_set_name != "선택안함" else (st.session_state.get("compare_set_name") or "").strip()
        if not target_set_name:
            st.warning("삭제할 세트를 선택해 주세요.")
        else:
            delete_company_compare_set(target_set_name)
            st.session_state["compare_set_notice"] = f"세트 `{target_set_name}`을 삭제했습니다."
            st.session_state["compare_selected_set_pending"] = "선택안함"
            if (st.session_state.get("compare_set_name") or "").strip() == target_set_name:
                st.session_state["compare_set_name_pending"] = ""
                st.session_state["compare_set_note_pending"] = ""
            st.rerun()

    total_w = sum(max(0.0, float(metric_weights.get(m, 0.0))) for m in selected_metrics)
    if total_w <= 0:
        normalized = {m: 100.0 / len(selected_metrics) for m in selected_metrics}
    else:
        normalized = {m: max(0.0, float(metric_weights.get(m, 0.0))) / total_w * 100.0 for m in selected_metrics}
    st.caption(
        "적용 가중치: "
        + " / ".join([f"{SCORE_METRIC_CONFIG[m]['label']} {normalized[m]:,.0f}" for m in selected_metrics])
    )

    with st.expander("AI 티커 추론 설정 (선택)", expanded=False):
        st.checkbox("티커 자동 탐색 실패 시 AI 티커 추론 사용", key="compare_use_ai_ticker")
        st.caption("AI 키/모델/제공자는 API 설정 탭의 공통값을 사용합니다.")

    compute_btn = st.button("선택 지표로 기업 점수 계산", type="primary", key="compare_compute_btn")
    if compute_btn:
        if not selected_companies:
            st.error("비교할 기업을 1개 이상 선택해 주세요.")
        else:
            comp_provider, comp_api_key, comp_model = get_ai_settings_from_session("compare")
            with st.spinner("기업 지표를 불러와 점수를 계산하는 중입니다..."):
                result_df, error_df = compute_company_metric_ranking(
                    companies=selected_companies,
                    metric_keys=selected_metrics,
                    metric_weights=metric_weights,
                    use_ai_ticker=bool(st.session_state.get("compare_use_ai_ticker", False)),
                    ai_provider=comp_provider,
                    ai_api_key=comp_api_key,
                    ai_model=comp_model,
                )
            st.session_state["compare_result_df"] = result_df
            st.session_state["compare_error_df"] = error_df
            st.session_state["compare_metric_snapshot"] = list(selected_metrics)
            if result_df.empty:
                st.warning("점수 계산 결과가 없습니다. 티커/데이터 조회 상태를 확인해 주세요.")
            else:
                normal_count = int((result_df["상태"] == "정상").sum())
                missing_count = int((result_df["상태"] == "지표 부족").sum())
                data_error_count = int((result_df["상태"] == "데이터 오류").sum())
                no_ticker_count = int((result_df["상태"] == "티커 없음").sum())
                st.success(
                    "기업 비교 점수 계산 완료 "
                    f"(정상 {normal_count:,}개 / 지표 부족 {missing_count:,}개 / "
                    f"데이터 오류 {data_error_count:,}개 / 티커 없음 {no_ticker_count:,}개)"
                )

    result_df = st.session_state.get("compare_result_df", pd.DataFrame())
    if result_df is None or result_df.empty:
        st.info("조건을 선택하고 `선택 지표로 기업 점수 계산` 버튼을 눌러 결과를 확인하세요.")
        st.markdown("</div>", unsafe_allow_html=True)
        return

    metric_snapshot = st.session_state.get("compare_metric_snapshot", selected_metrics)
    result_df = sanitize_compare_result_df(result_df, metric_snapshot)
    st.session_state["compare_result_df"] = result_df

    status_series = result_df["상태"].astype(str) if "상태" in result_df.columns else pd.Series(dtype=str)
    normal_count = int((status_series == "정상").sum())
    missing_count = int((status_series == "지표 부족").sum())
    data_error_count = int((status_series == "데이터 오류").sum())
    no_ticker_count = int((status_series == "티커 없음").sum())
    st.caption(
        f"결과 요약: 정상 {normal_count:,} / 지표 부족 {missing_count:,} / "
        f"데이터 오류 {data_error_count:,} / 티커 없음 {no_ticker_count:,}"
    )

    score_cols = [f"{SCORE_METRIC_CONFIG[m]['label']} 점수" for m in metric_snapshot if f"{SCORE_METRIC_CONFIG[m]['label']} 점수" in result_df.columns]
    value_cols = [f"{SCORE_METRIC_CONFIG[m]['label']} 값" for m in metric_snapshot if f"{SCORE_METRIC_CONFIG[m]['label']} 값" in result_df.columns]

    normal_df = result_df[result_df["상태"] == "정상"].copy()
    if not normal_df.empty:
        rank_fig = px.bar(
            normal_df.sort_values("총점", ascending=False),
            x="기업명",
            y="총점",
            color="총점",
            title="기업 종합 점수 랭킹 (높은 순)",
            labels={"기업명": "기업", "총점": "종합 점수"},
            color_continuous_scale="Tealgrn",
        )
        rank_fig.update_coloraxes(showscale=False)
        add_bar_labels(rank_fig, pct=False)
        st.plotly_chart(style_figure(rank_fig), use_container_width=True)

        if "산업섹터" in normal_df.columns and normal_df["산업섹터"].nunique() > 1:
            sector_rank = (
                normal_df.groupby("산업섹터", as_index=False)["총점"]
                .mean()
                .sort_values("총점", ascending=False)
            )
            sector_fig = px.bar(
                sector_rank,
                x="산업섹터",
                y="총점",
                color="총점",
                title="산업섹터별 평균 종합 점수",
                labels={"산업섹터": "산업섹터", "총점": "평균 종합 점수"},
                color_continuous_scale="Tealgrn",
            )
            sector_fig.update_coloraxes(showscale=False)
            add_bar_labels(sector_fig, pct=False)
            st.plotly_chart(style_figure(sector_fig), use_container_width=True)

        if score_cols:
            melted = normal_df.melt(
                id_vars=["기업명"],
                value_vars=score_cols,
                var_name="지표",
                value_name="점수",
            )
            metric_fig = px.bar(
                melted,
                x="기업명",
                y="점수",
                color="지표",
                barmode="group",
                title="선택 지표별 점수 비교",
                labels={"기업명": "기업", "점수": "점수", "지표": "지표"},
            )
            st.plotly_chart(style_figure(metric_fig), use_container_width=True)

    table_cols = ["순위", "기업명", "산업섹터", "티커", "총점"] + value_cols + score_cols + ["상태", "티커소스", "데이터소스"]
    table_cols = [c for c in table_cols if c in result_df.columns]
    percent_cols = {c for c in table_cols if "(%)" in c}
    st.caption("기업 점수 계산 결과")
    st.dataframe(
        format_table_numbers(result_df[table_cols], percent_cols=percent_cols),
        use_container_width=True,
        hide_index=True,
    )

    err_df = st.session_state.get("compare_error_df", pd.DataFrame())
    if err_df is not None and not err_df.empty:
        st.caption("조회 실패 목록")
        st.dataframe(format_table_numbers(err_df), use_container_width=True, hide_index=True)

    if not compare_sets_df.empty:
        set_view = compare_sets_df.rename(
            columns={
                "set_name": "세트명",
                "sector_filter": "섹터필터",
                "note": "메모",
                "updated_at": "수정시각",
            }
        ).copy()
        set_view["기업수"] = set_view["companies_json"].apply(lambda v: len(parse_json_list(v)))
        set_view["지표수"] = set_view["metrics_json"].apply(lambda v: len(parse_json_list(v)))
        st.caption("저장된 비교 세트")
        st.dataframe(
            format_table_numbers(set_view[["세트명", "섹터필터", "기업수", "지표수", "메모", "수정시각"]]),
            use_container_width=True,
            hide_index=True,
        )

    st.markdown("</div>", unsafe_allow_html=True)


def render_api_settings_tab() -> None:
    st.markdown('<div class="section-shell">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">API 설정</div>', unsafe_allow_html=True)
    st.caption(
        "OpenAI/Claude API 키와 기본 모델을 저장하면 기업정보/기업분석/기업 점수 탭에서 공통으로 사용합니다. "
        "추가로 Alpha Vantage/Finnhub 키를 넣으면 야후 외 경로로 기업 데이터를 보조 수집합니다. "
        "GitHub 동기화를 켜면 엑셀을 원격 저장소에서 자동 불러오고 저장 시 자동 커밋합니다."
    )

    if "api_settings_saved_notice" in st.session_state:
        st.success(st.session_state.pop("api_settings_saved_notice"))

    if "global_openai_model_options" not in st.session_state:
        st.session_state["global_openai_model_options"] = []
    if "global_claude_model_options" not in st.session_state:
        st.session_state["global_claude_model_options"] = []
    if "store_sensitive_keys" not in st.session_state:
        st.session_state["store_sensitive_keys"] = False

    st.selectbox(
        "기본 AI 제공자",
        options=["openai", "claude"],
        format_func=lambda x: ai_provider_label(x),
        key="global_ai_provider",
    )
    st.text_input("OpenAI API Key", key="global_openai_api_key", type="password", placeholder="sk-...")
    st.text_input("Claude API Key", key="global_claude_api_key", type="password", placeholder="sk-ant-...")
    st.text_input(
        "Alpha Vantage API Key (선택)",
        key="global_alpha_vantage_api_key",
        type="password",
        placeholder="알파밴티지 키",
    )
    st.text_input(
        "Finnhub API Key (선택)",
        key="global_finnhub_api_key",
        type="password",
        placeholder="finnhub 키",
    )
    st.checkbox(
        "민감키를 로컬 DB에 저장 (비권장)",
        key="store_sensitive_keys",
        help="권장: 해제 상태 유지 후 Streamlit Secrets/환경변수로 관리",
    )
    if not bool(st.session_state.get("store_sensitive_keys", False)):
        st.caption("권장 모드: API 키/GitHub Token은 DB에 저장하지 않고 현재 세션 + Secrets/환경변수만 사용")

    fetch_col1, fetch_col2, fetch_col3 = st.columns([1, 1, 1.2])
    with fetch_col1:
        fetch_openai_btn = st.button("OpenAI 모델 조회", key="api_fetch_openai_models_btn")
    with fetch_col2:
        fetch_claude_btn = st.button("Claude 모델 조회", key="api_fetch_claude_models_btn")
    with fetch_col3:
        fetch_all_btn = st.button("모델 전체 조회", key="api_fetch_all_models_btn")

    if fetch_openai_btn or fetch_all_btn:
        models, err = fetch_openai_available_models(st.session_state.get("global_openai_api_key", ""))
        if err:
            st.warning(err)
        else:
            st.session_state["global_openai_model_options"] = models
            current = str(st.session_state.get("global_openai_model", "") or "").strip()
            if not current or current not in models:
                st.session_state["global_openai_model"] = models[0]
            st.success(f"OpenAI 사용 가능 모델 {len(models):,.0f}개를 불러왔습니다.")

    if fetch_claude_btn or fetch_all_btn:
        models, err = fetch_claude_available_models(st.session_state.get("global_claude_api_key", ""))
        if err:
            st.warning(err)
        else:
            st.session_state["global_claude_model_options"] = models
            current = str(st.session_state.get("global_claude_model", "") or "").strip()
            if not current or current not in models:
                st.session_state["global_claude_model"] = models[0]
            st.success(f"Claude 사용 가능 모델 {len(models):,.0f}개를 불러왔습니다.")

    openai_options = list(st.session_state.get("global_openai_model_options", []))
    current_openai = str(st.session_state.get("global_openai_model", DEFAULT_OPENAI_MODEL) or DEFAULT_OPENAI_MODEL).strip()
    if current_openai and current_openai not in openai_options:
        openai_options = [current_openai] + openai_options
    if openai_options:
        st.selectbox("OpenAI 모델", options=openai_options, key="global_openai_model")
    else:
        st.text_input("OpenAI 모델", key="global_openai_model")

    claude_options = list(st.session_state.get("global_claude_model_options", []))
    current_claude = str(st.session_state.get("global_claude_model", DEFAULT_CLAUDE_MODEL) or DEFAULT_CLAUDE_MODEL).strip()
    if current_claude and current_claude not in claude_options:
        claude_options = [current_claude] + claude_options
    if claude_options:
        st.selectbox("Claude 모델", options=claude_options, key="global_claude_model")
    else:
        st.text_input("Claude 모델", key="global_claude_model")

    st.markdown("#### GitHub 엑셀 자동 동기화")
    st.checkbox("GitHub 동기화 사용", key="github_sync_enabled")
    st.checkbox(
        "포트폴리오 입력 변경 시 즉시 GitHub 저장",
        key="github_sync_on_change",
        help="기록 입력 탭 테이블 값이 바뀌면 현재 선택 날짜 스냅샷을 자동 저장/동기화합니다.",
    )
    st.text_input("GitHub Repo", key="github_repo", placeholder="owner/repo")
    st.text_input("GitHub Branch", key="github_branch", placeholder="main")
    st.text_input("GitHub Excel Path", key="github_excel_path", placeholder="portfolio_auto.xlsx")
    st.text_input("GitHub Token", key="github_token", type="password", placeholder="ghp_... (repo 권한)")
    gh_secret_cfg = _load_github_settings_from_secrets()
    if any(
        str(gh_secret_cfg.get(k, "") or "").strip()
        for k in ["repo", "branch", "excel_path", "token", "sync_enabled", "sync_on_change"]
    ):
        st.caption("GitHub 설정은 Streamlit Secrets/환경변수 값을 우선 사용 중입니다.")

    gh_col1, gh_col2 = st.columns([1, 1.2])
    with gh_col1:
        test_pull_btn = st.button("GitHub에서 엑셀 불러오기 테스트", key="api_test_pull_github_excel_btn")
    with gh_col2:
        test_push_btn = st.button("현재 데이터 GitHub에 즉시 저장", key="api_test_push_github_excel_btn")

    if test_pull_btn:
        cfg = get_github_sync_settings()
        excel_bytes, err = fetch_excel_bytes_from_github(
            repo=str(cfg["repo"] or ""),
            path=str(cfg["excel_path"] or ""),
            branch=str(cfg["branch"] or "main"),
            token=str(cfg["token"] or ""),
        )
        if err:
            st.warning(err)
        else:
            st.session_state["uploaded_portfolio_excel_bytes"] = excel_bytes
            path_text = str(cfg["excel_path"] or "").strip()
            file_name = path_text.split("/")[-1] if "/" in path_text else path_text
            st.session_state["uploaded_portfolio_excel_name"] = f"github:{file_name}"
            st.session_state["uploaded_portfolio_excel_sig"] = hashlib.sha256(excel_bytes).hexdigest()
            st.session_state["editing_df_date"] = ""
            st.success("GitHub 엑셀 불러오기 테스트 성공")
            st.rerun()

    if test_push_btn:
        cfg = get_github_sync_settings()
        if not bool(cfg["enabled"]):
            st.warning("GitHub 동기화 사용을 먼저 켜 주세요.")
        else:
            latest_date, latest_df = load_latest_snapshot()
            if latest_df.empty:
                latest_df = load_snapshot(DEFAULT_DATE)
                target_date = DEFAULT_DATE
            else:
                target_date = _safe_parse_date(latest_date) or DEFAULT_DATE

            if latest_df.empty:
                st.warning("GitHub로 저장할 포트폴리오 데이터가 없습니다.")
            else:
                ok, msg = sync_snapshot_to_github_excel(target_date, latest_df)
                if ok:
                    st.success(msg)
                elif msg:
                    st.warning(msg)

    st.markdown("#### 일일 자동 주가 갱신/자산 저장")
    st.checkbox(
        "하루 1회 자동 실행 사용",
        key="daily_auto_snapshot_enabled",
        help="기준 시각 이후 첫 실행에서 전일 보유/예수금을 기준으로 주가 갱신→재계산→오늘 스냅샷 저장을 수행합니다.",
    )
    st.number_input(
        "자동 실행 기준 시각(한국시간, 시)",
        min_value=0,
        max_value=23,
        step=1,
        key="daily_auto_snapshot_hour",
    )
    st.checkbox(
        "티커 비어있을 때 AI 티커 보강 사용",
        key="daily_auto_snapshot_use_ai_ticker",
        help="자동 실행 시 티커가 없는 종목만 AI/API 경로로 티커를 보강합니다.",
    )
    st.caption(
        "자동 실행 기준: 전일(어제) 스냅샷의 보유종목/예수금을 고정으로 가져와, "
        "당일 주가만 갱신해 총자산/총손익을 재계산 후 저장합니다."
    )
    st.caption(
        "주의: Streamlit 앱은 요청이 있을 때 실행됩니다. "
        "완전 무인 실행이 필요하면 OS 작업 스케줄러/cron으로 `daily_auto_snapshot.py`를 하루 1회 실행하세요."
    )

    last_run_date = str(st.session_state.get("daily_auto_snapshot_last_run_date", "") or "").strip()
    last_run_at = str(st.session_state.get("daily_auto_snapshot_last_run_at", "") or "").strip()
    last_summary = str(st.session_state.get("daily_auto_snapshot_last_summary", "") or "").strip()
    if last_run_date:
        st.caption(f"최근 자동 저장일: {last_run_date}" + (f" ({last_run_at})" if last_run_at else ""))
    if last_summary:
        st.caption(f"최근 자동 저장 결과: {last_summary}")

    if st.button("일일 자동 실행 지금 테스트", key="api_daily_auto_snapshot_test_btn"):
        ok, msg = run_daily_auto_snapshot(force=True, target_date=date.today())
        if ok:
            st.success(msg)
        elif msg:
            st.warning(msg)
        else:
            st.info("실행할 항목이 없습니다.")

    action_col1, action_col2 = st.columns([1, 1.2])
    with action_col1:
        submit_save = st.button("API 설정 저장", type="primary", key="api_save_settings_btn")
    with action_col2:
        reload_btn = st.button("저장값 다시 불러오기", key="api_reload_settings_btn")

    if submit_save:
        persist_sensitive = bool(st.session_state.get("store_sensitive_keys", False))
        save_app_settings(
            {
                "store_sensitive_keys": "true" if persist_sensitive else "false",
                "ai_provider": normalize_ai_provider(st.session_state.get("global_ai_provider", "claude")),
                "openai_api_key": st.session_state.get("global_openai_api_key", "") if persist_sensitive else "",
                "claude_api_key": st.session_state.get("global_claude_api_key", "") if persist_sensitive else "",
                "alpha_vantage_api_key": st.session_state.get("global_alpha_vantage_api_key", "") if persist_sensitive else "",
                "finnhub_api_key": st.session_state.get("global_finnhub_api_key", "") if persist_sensitive else "",
                "openai_model": st.session_state.get("global_openai_model", DEFAULT_OPENAI_MODEL),
                "claude_model": st.session_state.get("global_claude_model", DEFAULT_CLAUDE_MODEL),
                "github_sync_enabled": "true" if bool(st.session_state.get("github_sync_enabled", False)) else "false",
                "github_sync_on_change": "true" if bool(st.session_state.get("github_sync_on_change", True)) else "false",
                "github_repo": st.session_state.get("github_repo", ""),
                "github_branch": st.session_state.get("github_branch", "main"),
                "github_excel_path": st.session_state.get("github_excel_path", "portfolio_auto.xlsx"),
                "github_token": st.session_state.get("github_token", "") if persist_sensitive else "",
                "daily_auto_snapshot_enabled": "true" if bool(st.session_state.get("daily_auto_snapshot_enabled", False)) else "false",
                "daily_auto_snapshot_use_ai_ticker": "true" if bool(st.session_state.get("daily_auto_snapshot_use_ai_ticker", True)) else "false",
                "daily_auto_snapshot_hour": str(int(st.session_state.get("daily_auto_snapshot_hour", 18) or 18)),
                "daily_auto_snapshot_last_run_date": st.session_state.get("daily_auto_snapshot_last_run_date", ""),
                "daily_auto_snapshot_last_run_at": st.session_state.get("daily_auto_snapshot_last_run_at", ""),
                "daily_auto_snapshot_last_attempt_date": st.session_state.get("daily_auto_snapshot_last_attempt_date", ""),
                "daily_auto_snapshot_last_summary": st.session_state.get("daily_auto_snapshot_last_summary", ""),
            }
        )
        st.session_state["api_settings_saved_notice"] = (
            "API 설정 저장 완료 (민감키 DB 저장 ON)"
            if persist_sensitive
            else "API 설정 저장 완료 (민감키 DB 미저장, Secrets/환경변수 권장)"
        )
        st.rerun()

    if reload_btn:
        st.session_state["force_reload_api_settings"] = True
        st.rerun()

    st.caption(
        "설정은 로컬 DB(`portfolio.db`)에 저장됩니다. "
        "민감키 DB 저장을 끄면 키는 비워 저장되며, Streamlit Secrets/환경변수를 우선 사용합니다."
    )
    st.markdown("</div>", unsafe_allow_html=True)


def _read_secret_or_env(key: str) -> str:
    value = ""
    try:
        value = str(st.secrets.get(key, "") or "").strip()
    except Exception:
        value = ""
    if value:
        return value
    return str(os.getenv(key, "") or "").strip()


def _get_access_password_config() -> tuple[str, str]:
    # Preferred: sha256 hash. Fallback: plain password.
    raw_hash = _read_secret_or_env("APP_PASSWORD_HASH")
    if raw_hash:
        h = raw_hash.strip().lower()
        if h.startswith("sha256:"):
            h = h.split(":", 1)[1].strip()
        if re.fullmatch(r"[0-9a-f]{64}", h):
            return "hash", h

    raw_plain = _read_secret_or_env("APP_PASSWORD")
    if raw_plain:
        return "plain", raw_plain
    return "", ""


def _verify_access_password(input_password: str) -> bool:
    mode, stored = _get_access_password_config()
    if not mode:
        return False
    candidate = str(input_password or "")
    if mode == "hash":
        digest = hashlib.sha256(candidate.encode("utf-8")).hexdigest()
        return hmac.compare_digest(digest, stored)
    return hmac.compare_digest(candidate, stored)


def require_password_gate() -> None:
    mode, _ = _get_access_password_config()
    if not mode:
        st.error("접속 비밀번호가 설정되지 않았습니다. Secrets 또는 환경변수 설정 후 다시 실행해 주세요.")
        st.markdown("다음 중 하나를 설정하세요.")
        st.code(
            "APP_PASSWORD = \"원하는비밀번호\"\n"
            "# 또는\n"
            "APP_PASSWORD_HASH = \"sha256:64자리해시\"",
            language="toml",
        )
        st.stop()

    if st.session_state.get("auth_ok", False):
        return

    st.markdown(
        """
        <div class="hero">
            <h1>접속 비밀번호 확인</h1>
            <p>이 앱은 비밀번호 인증 후에만 사용할 수 있습니다.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    with st.form("access_password_form", clear_on_submit=True):
        input_password = st.text_input("비밀번호", type="password", placeholder="접속 비밀번호 입력")
        submitted = st.form_submit_button("로그인", type="primary")
    if submitted:
        if _verify_access_password(input_password):
            st.session_state["auth_ok"] = True
            st.session_state.pop("auth_error", None)
            st.rerun()
        st.session_state["auth_error"] = "비밀번호가 올바르지 않습니다."

    auth_err = str(st.session_state.get("auth_error", "") or "").strip()
    if auth_err:
        st.error(auth_err)
    st.stop()


def main() -> None:
    st.set_page_config(page_title="투자 포트폴리오 기록장", page_icon=":chart_with_upwards_trend:", layout="wide")
    inject_styles()
    require_password_gate()
    force_reload = bool(st.session_state.pop("force_reload_api_settings", False))
    initialize_api_settings(force=force_reload)
    bootstrap_excel_from_github_if_needed()
    auto_ok, auto_msg = run_daily_auto_snapshot(force=False, target_date=date.today())
    if auto_msg:
        if auto_ok:
            st.session_state["daily_auto_snapshot_notice_success"] = auto_msg
        else:
            st.session_state["daily_auto_snapshot_notice_warning"] = auto_msg

    st.markdown(
        """
        <div class="hero">
            <h1>투자 포트폴리오 기록장</h1>
            <p>대시보드에서 전체 자산을 한눈에 보고, 기록 탭에서 날짜별 스냅샷을 저장하세요.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        if st.button("로그아웃", key="sidebar_logout_btn"):
            st.session_state["auth_ok"] = False
            st.rerun()
        st.subheader("기록 설정")
        selected_date = st.date_input("기록 날짜", value=DEFAULT_DATE)
        selected_date_key = selected_date.isoformat()
        usd_krw_rate, fx_source = get_usd_krw_rate_for_date(selected_date)
        st.metric("해당일 USD/KRW", f"{usd_krw_rate:,.0f}")
        st.caption(f"환율 소스: {fx_source}")
        uploaded_excel = st.file_uploader(
            "포트폴리오 엑셀 업로드 (.xlsx)",
            type=["xlsx"],
            key="sidebar_portfolio_excel_upload",
            help="Cloud/원격 실행 환경에서는 이 업로드 파일을 우선 사용합니다.",
        )
        if uploaded_excel is not None:
            uploaded_bytes = uploaded_excel.getvalue()
            uploaded_sig = hashlib.sha256(uploaded_bytes).hexdigest() if uploaded_bytes else ""
            prev_sig = str(st.session_state.get("uploaded_portfolio_excel_sig", "") or "")
            if uploaded_sig and uploaded_sig != prev_sig:
                st.session_state["uploaded_portfolio_excel_bytes"] = uploaded_bytes
                st.session_state["uploaded_portfolio_excel_name"] = str(uploaded_excel.name or "").strip()
                st.session_state["uploaded_portfolio_excel_sig"] = uploaded_sig
                st.session_state["editing_df_date"] = ""
                st.session_state["portfolio_excel_notice"] = (
                    f"엑셀 업로드 반영: {st.session_state['uploaded_portfolio_excel_name']}"
                )
                st.rerun()

        uploaded_name = str(st.session_state.get("uploaded_portfolio_excel_name", "") or "").strip()
        excel_path = resolve_excel_path()
        if uploaded_name:
            st.caption(f"엑셀 소스: 업로드 파일 ({uploaded_name})")
        elif excel_path:
            st.caption(f"엑셀 자동 불러오기: {excel_path.name}")
        else:
            st.caption("엑셀 소스 없음: 파일 업로드 또는 서버 경로 설정 필요")
        if uploaded_name and st.button("업로드 엑셀 해제", key="sidebar_clear_uploaded_excel_btn"):
            for key in [
                "uploaded_portfolio_excel_bytes",
                "uploaded_portfolio_excel_name",
                "uploaded_portfolio_excel_sig",
            ]:
                st.session_state.pop(key, None)
            st.session_state["editing_df_date"] = ""
            st.session_state["portfolio_excel_notice"] = "업로드 엑셀 연결을 해제했습니다."
            st.rerun()
        use_sample = st.button("샘플 데이터로 시작", key="sidebar_use_sample_btn")

        if st.button("선택 날짜 데이터 불러오기", key="sidebar_load_date_btn"):
            st.session_state["editing_df"] = ensure_portfolio_columns(
                load_snapshot(selected_date), usd_krw_rate, force_usd_rate=True
            )
            st.session_state["editing_df_date"] = selected_date_key
            st.success(f"{selected_date} 데이터 불러오기 완료")

    if "portfolio_excel_notice" in st.session_state:
        st.success(st.session_state.pop("portfolio_excel_notice"))
    if "daily_auto_snapshot_notice_success" in st.session_state:
        st.success(st.session_state.pop("daily_auto_snapshot_notice_success"))
    if "daily_auto_snapshot_notice_warning" in st.session_state:
        st.warning(st.session_state.pop("daily_auto_snapshot_notice_warning"))
    if "github_sync_notice" in st.session_state:
        gh_notice = str(st.session_state.pop("github_sync_notice") or "").strip()
        if gh_notice:
            if "완료" in gh_notice:
                st.success(gh_notice)
            else:
                st.warning(gh_notice)

    if use_sample:
        st.session_state["editing_df"] = ensure_portfolio_columns(
            pd.DataFrame(DEFAULT_HOLDINGS), usd_krw_rate, force_usd_rate=True
        )
        st.session_state["editing_df_date"] = selected_date_key

    should_reload_for_date = st.session_state.get("editing_df_date", "") != selected_date_key
    if "editing_df" not in st.session_state or should_reload_for_date:
        st.session_state["editing_df"] = ensure_portfolio_columns(
            load_snapshot(selected_date), usd_krw_rate, force_usd_rate=True
        )
        st.session_state["editing_df_date"] = selected_date_key
    else:
        st.session_state["editing_df"] = ensure_portfolio_columns(
            st.session_state["editing_df"], usd_krw_rate, force_usd_rate=True
        )

    corrected_builtin_global = reconcile_builtin_ticker_overrides()
    if corrected_builtin_global > 0:
        st.session_state["github_sync_notice"] = f"내장 티커 자동 교정 완료: {corrected_builtin_global}개"
        st.rerun()

    tab_options = ["대시보드", "기록 입력", "환율", "기업정보", "기업분석", "밸류체인", "기업 점수", "API 설정"]
    if "active_main_tab" not in st.session_state:
        st.session_state["active_main_tab"] = "대시보드"
    if st.session_state.get("active_main_tab") not in tab_options:
        st.session_state["active_main_tab"] = "대시보드"

    active_tab = st.radio(
        "메뉴",
        options=tab_options,
        horizontal=True,
        label_visibility="collapsed",
        key="active_main_tab",
    )

    if active_tab == "대시보드":
        render_dashboard(st.session_state["editing_df"], usd_krw_rate, selected_date)
    elif active_tab == "기록 입력":
        st.session_state["editing_df"] = render_input_tab(selected_date, st.session_state["editing_df"], usd_krw_rate)
    elif active_tab == "환율":
        render_fx_tab()
    elif active_tab == "기업정보":
        render_company_analysis_tab(st.session_state["editing_df"])
    elif active_tab == "기업분석":
        render_company_compare_tab(st.session_state["editing_df"])
    elif active_tab == "밸류체인":
        render_value_chain_tab()
    elif active_tab == "기업 점수":
        render_company_score_tab(st.session_state["editing_df"])
    elif active_tab == "API 설정":
        render_api_settings_tab()


if __name__ == "__main__":
    main()

